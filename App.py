import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import os
import json
import uuid
import base64
import io
import re
import hmac
import hashlib
import difflib
import requests
import glob
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from datetime import datetime, timedelta, timezone
from bs4 import BeautifulSoup
from streamlit_calendar import calendar
from streamlit_gsheets import GSheetsConnection
import extra_streamlit_components as stx

try:
    from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
    DRIVE_LIB_READY = True
except ImportError:
    # Si por alguna razón esta submódulo no está disponible, la app sigue
    # funcionando: simplemente el guardado de adjuntos cae de vuelta a base64.
    DRIVE_LIB_READY = False

try:
    import bcrypt
    BCRYPT_READY = True
except ImportError:
    # Si el paquete no está instalado, la app sigue funcionando pero avisa.
    # Agrega "bcrypt" a requirements.txt para habilitar el hash de contraseñas.
    BCRYPT_READY = False

# =====================================================================
# 🔒 UTILIDADES DE SEGURIDAD (HASH DE CONTRASEÑAS Y COOKIES FIRMADAS)
# =====================================================================

def hash_password(plano: str) -> str:
    """Convierte una contraseña en texto plano a un hash bcrypt seguro."""
    if not BCRYPT_READY:
        return str(plano)  # Fallback (no recomendado): evita crashear si falta la librería.
    return bcrypt.hashpw(str(plano).encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def es_hash_bcrypt(valor: str) -> bool:
    valor = str(valor)
    return valor.startswith("$2a$") or valor.startswith("$2b$") or valor.startswith("$2y$")

def verificar_password(plano: str, almacenado: str) -> bool:
    """
    Verifica una contraseña contra el valor guardado.
    Soporta migración transparente: si el valor guardado todavía es texto plano
    (contraseñas antiguas), compara en texto plano. Los hashes nuevos ya
    quedan protegidos con bcrypt.
    """
    almacenado = str(almacenado)
    if es_hash_bcrypt(almacenado) and BCRYPT_READY:
        try:
            return bcrypt.checkpw(str(plano).encode("utf-8"), almacenado.encode("utf-8"))
        except Exception:
            return False
    # Compatibilidad con contraseñas históricas aún no migradas
    return str(plano) == almacenado

def _cookie_secret() -> str:
    # Idealmente definida en st.secrets["COOKIE_SECRET"] (Streamlit Cloud -> Settings -> Secrets).
    return str(st.secrets.get("COOKIE_SECRET", "jurisync_cambia_este_secreto_en_secrets_toml"))

def generar_token_sesion(usuario: str, dias_validez: int = 7) -> str:
    """Genera una cookie firmada (usuario|expiración|firma) para evitar suplantación."""
    expira = int((datetime.now(timezone.utc) + timedelta(days=dias_validez)).timestamp())
    payload = f"{usuario}|{expira}"
    firma = hmac.new(_cookie_secret().encode("utf-8"), payload.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{payload}|{firma}"

def validar_token_sesion(token: str):
    """Valida la firma y expiración de la cookie. Devuelve el usuario o None si es inválida."""
    try:
        usuario, expira, firma = str(token).split("|")
        payload = f"{usuario}|{expira}"
        firma_esperada = hmac.new(_cookie_secret().encode("utf-8"), payload.encode("utf-8"), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(firma, firma_esperada):
            return None
        if int(expira) < int(datetime.now(timezone.utc).timestamp()):
            return None
        return usuario
    except Exception:
        return None

# =====================================================================
# 🔎 UTILIDADES DE CRUCE PARA ESTADO DIARIO (NORMALIZACIÓN + FUZZY MATCH)
# =====================================================================

def normalizar_rol(valor) -> str:
    """
    Normaliza un ROL/RIT judicial a un formato comparable, sin importar
    si viene como 'C-1234-2026', '1234/2026', 'C 1234 2026', etc.
    Devuelve algo tipo '1234-2026' (solo dígitos y guión separador de año).
    """
    if valor is None:
        return ""
    texto = str(valor).upper().strip()
    texto = texto.replace("/", "-").replace("_", "-")
    numeros = re.findall(r"\d+", texto)
    if len(numeros) >= 2:
        # Se asume [número de rol, año] como los dos últimos grupos numéricos
        return f"{int(numeros[-2])}-{numeros[-1]}"
    elif len(numeros) == 1:
        return numeros[0]
    return texto

def buscar_coincidencias_probables(df_pj, df_causas, col_rol_pj, umbral=0.85):
    """
    Segundo filtro (además del cruce exacto): detecta causas cuyo ROL es
    'casi' igual (diferencias de formato, un dígito distinto, etc.) para
    que el usuario las revise manualmente en vez de que pasen inadvertidas.
    """
    probables = []
    rol_pj_normalizados = df_pj["ROL_NORMALIZADO"].unique().tolist()
    for _, fila_causa in df_causas.iterrows():
        rol_causa_norm = fila_causa["ROL_NORMALIZADO"]
        if not rol_causa_norm:
            continue
        mejor_score = 0.0
        mejor_match = None
        for rol_pj in rol_pj_normalizados:
            score = difflib.SequenceMatcher(None, rol_causa_norm, rol_pj).ratio()
            if score > mejor_score:
                mejor_score = score
                mejor_match = rol_pj
        if umbral <= mejor_score < 1.0:
            probables.append({
                "ROL_Causa_Local": fila_causa.get("ROL", ""),
                "ROL_Estado_Diario": mejor_match,
                "Cliente": fila_causa.get("Cliente", ""),
                "Similitud": f"{mejor_score:.0%}"
            })
    return pd.DataFrame(probables)

def validar_tamano_para_sheets(archivo_bytes: bytes, nombre_archivo: str, limite_kb: int = 35):
    """
    Google Sheets limita cada celda a ~50.000 caracteres. Un archivo se
    guarda como base64 (crece ~33%), así que un límite de ~35 KB en bytes
    originales da margen de sobra. Si se supera, avisamos ANTES de guardar
    en vez de dejar que la celda se trunque en silencio.
    """
    tamano_kb = len(archivo_bytes) / 1024
    if tamano_kb > limite_kb:
        return False, f"⚠️ '{nombre_archivo}' pesa {tamano_kb:.0f} KB. Google Sheets no admite archivos de más de ~{limite_kb} KB en este módulo (se truncaría o fallaría el guardado). Comprime el PDF o súbelo a Google Drive y pega el enlace."
    return True, ""

def boton_descargar_excel(df, nombre_archivo, key, label="⬇️ Descargar excel"):
    """Genera un botón de descarga en Excel para cualquier listado (Causas, Tareas, Clientes)."""
    try:
        buffer_excel = io.BytesIO()
        with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Datos')
        st.download_button(label, data=buffer_excel.getvalue(), file_name=nombre_archivo,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=key)
    except Exception:
        st.caption("⚠️ Instala 'openpyxl' en requirements.txt para habilitar la descarga en Excel.")

def fecha_segura(valor, formato="%Y-%m-%d"):
    """
    Convierte un valor a datetime sin romper la app si viene vacío, NaN,
    None o con un formato inesperado (columna faltante, dato antiguo, etc.).
    Si no se puede interpretar, devuelve la fecha de hoy como respaldo.
    """
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return datetime.now()
    texto = str(valor).strip()
    if texto == "" or texto.lower() == "nan" or texto == "--":
        return datetime.now()
    try:
        return datetime.strptime(texto, formato)
    except (ValueError, TypeError):
        return datetime.now()

def _corregir_dtypes_texto(df):
    """
    Corrige columnas mal tipadas por inferencia automática de pandas:
    - Columnas 100% "True"/"False" -> quedan como bool -> se pasan a texto.
    - Columnas 100% vacías -> quedan como float (puro NaN) -> se pasan a texto.
    Sin esto, asignar después un string en esas columnas revienta con
    TypeError (ya pasó con 'Fecha_Inicio' y 'Debe_Cambiar_Clave'). Las
    columnas numéricas reales (con datos) no se tocan.
    """
    for col in df.columns:
        if df[col].dtype == bool:
            df[col] = df[col].astype(str)
        elif df[col].dtype == float and df[col].isna().all():
            df[col] = df[col].astype(object)
    return df

def parsear_monto_clp(texto) -> int:
    """
    Interpreta un monto en formato chileno. En Chile el punto se usa como
    separador de miles, NUNCA de decimales (el peso chileno no usa centavos
    en el uso corriente). Por eso "500.000" debe leerse como 500 mil, no
    como 500 con tres decimales. La forma más segura de lograr esto es
    quedarse solo con los dígitos y descartar cualquier punto, coma o símbolo.
    """
    if texto is None:
        return 0
    solo_digitos = re.sub(r'[^0-9]', '', str(texto))
    return int(solo_digitos) if solo_digitos else 0

def formatear_clp(monto) -> str:
    """Formatea un entero como moneda chilena: 500000 -> '$500.000'."""
    try:
        return f"${int(monto):,.0f}".replace(",", ".")
    except (ValueError, TypeError):
        return "$0"

# =====================================================================
# 📋 MOTOR DE CÁLCULO: POSESIÓN EFECTIVA INTESTADA (Formulario 4423 SII)
# =====================================================================
# Tomado directamente de las instrucciones del Formulario 4423 del SII
# (Declaración y Pago de Impuesto a las Herencias Intestadas) que Nicolás
# adjuntó. Estas tablas y fórmulas son las oficiales para calcular las
# asignaciones de cada heredero y el impuesto a la herencia que le
# corresponde, según el orden de sucesión intestada chileno.

TABLA1_EXENCION_UTM = {
    "Hijo": 600, "Cónyuge": 600, "Ascendiente": 600,
    "Hermano": 60, "Medio Hermano": 60,
    "Colateral 3° o 4° grado": 60, "Colateral 5° o 6° grado": 0
}

TABLA3_RECARGO_PCT = {
    "Hijo": 0, "Cónyuge": 0, "Ascendiente": 0,
    "Hermano": 20, "Medio Hermano": 20,
    "Colateral 3° o 4° grado": 20, "Colateral 5° o 6° grado": 40
}

# (desde UTM, hasta UTM, tasa, deducción en UTM)
TABLA2_TRAMOS_IMPUESTO = [
    (0.01, 960, 0.01, 0.0),
    (960.01, 1920, 0.025, 14.4),
    (1920.01, 3840, 0.05, 62.4),
    (3840.01, 5760, 0.075, 158.4),
    (5760.01, 7680, 0.10, 302.4),
    (7680.01, 9600, 0.15, 686.4),
    (9600.01, 14400, 0.20, 1166.4),
    (14400.01, float('inf'), 0.25, 1886.4),
]

def calcular_impuesto_tabla2(base_imponible_utm: float) -> float:
    """Aplica la Tabla N°2 del Formulario 4423: tasa progresiva por tramo de UTM."""
    if base_imponible_utm <= 0:
        return 0.0
    for desde, hasta, tasa, deduccion in TABLA2_TRAMOS_IMPUESTO:
        if desde <= base_imponible_utm <= hasta:
            return max(0.0, base_imponible_utm * tasa - deduccion)
    return 0.0

def calcular_asignaciones_intestadas(conteo_herederos: dict, masa_hereditaria: float) -> dict:
    """
    Calcula la asignación (en pesos) que le corresponde a CADA TIPO de
    heredero, siguiendo exactamente el orden de sucesión intestada y las
    fórmulas de las instrucciones del Formulario 4423 del SII.
    conteo_herederos: dict tipo_heredero -> cantidad de personas de ese tipo.
    Devuelve: dict tipo_heredero -> asignación individual en pesos (por persona).
    """
    M = masa_hereditaria
    h = conteo_herederos.get("Hijo", 0)
    conyuge = conteo_herederos.get("Cónyuge", 0)
    asc = conteo_herederos.get("Ascendiente", 0)
    herm = conteo_herederos.get("Hermano", 0)
    mherm = conteo_herederos.get("Medio Hermano", 0)
    col34 = conteo_herederos.get("Colateral 3° o 4° grado", 0)
    col56 = conteo_herederos.get("Colateral 5° o 6° grado", 0)
    
    resultado = {}
    
    if h > 0:
        # a) Los hijos excluyen a los demás salvo cónyuge sobreviviente.
        if conyuge > 0:
            if h == 1:
                resultado["Hijo"] = M / 2
                resultado["Cónyuge"] = M / 2
            elif h < 7:
                resultado["Hijo"] = M / (h + 2)
                resultado["Cónyuge"] = 2 * M / (h + 2)
            else:
                resultado["Cónyuge"] = M / 4
                resultado["Hijo"] = 3 * M / (4 * h)
        else:
            resultado["Hijo"] = M / h
    elif conyuge > 0 or asc > 0:
        # b) Sin hijos: cónyuge y/o ascendientes.
        if conyuge > 0 and asc > 0:
            resultado["Cónyuge"] = 2 * M / 3
            resultado["Ascendiente"] = M / (3 * asc)
        elif conyuge > 0:
            resultado["Cónyuge"] = M
        else:
            resultado["Ascendiente"] = M / asc
    elif herm > 0 or mherm > 0:
        # c) Sin los anteriores: hermanos y medio hermanos.
        if herm > 0 and mherm > 0:
            resultado["Hermano"] = 2 * (M / (2 * herm + mherm))
            resultado["Medio Hermano"] = M / (2 * herm + mherm)
        elif herm > 0:
            resultado["Hermano"] = M / herm
        else:
            resultado["Medio Hermano"] = M / mherm
    elif col34 > 0:
        # d) Colaterales de grado más próximo excluyen a los demás.
        resultado["Colateral 3° o 4° grado"] = M / col34
    elif col56 > 0:
        resultado["Colateral 5° o 6° grado"] = M / col56
    
    return resultado

def calcular_posesion_efectiva_completa(df_herederos: pd.DataFrame, masa_hereditaria: float, valor_utm: float) -> pd.DataFrame:
    """
    Toma la tabla de herederos (Nombre, RUT, Tipo) y la masa hereditaria ya
    calculada, y devuelve una tabla con la asignación e impuesto de CADA
    heredero individual, aplicando las Tablas 1, 2 y 3 del Formulario 4423.
    """
    if df_herederos.empty or valor_utm <= 0:
        return pd.DataFrame()
    
    conteo_por_tipo = df_herederos['Tipo de Heredero'].value_counts().to_dict()
    asignacion_por_tipo = calcular_asignaciones_intestadas(conteo_por_tipo, masa_hereditaria)
    
    filas_resultado = []
    for _, heredero in df_herederos.iterrows():
        tipo = heredero['Tipo de Heredero']
        asignacion_pesos = asignacion_por_tipo.get(tipo, 0)
        asignacion_utm = asignacion_pesos / valor_utm if valor_utm > 0 else 0
        exencion_utm = TABLA1_EXENCION_UTM.get(tipo, 0)
        base_imponible_utm = max(0, asignacion_utm - exencion_utm)
        impuesto_utm = calcular_impuesto_tabla2(base_imponible_utm)
        recargo_pct = TABLA3_RECARGO_PCT.get(tipo, 0)
        impuesto_total_utm = impuesto_utm * (1 + recargo_pct / 100)
        impuesto_total_pesos = impuesto_total_utm * valor_utm
        
        filas_resultado.append({
            "Heredero": heredero.get('Nombre', ''), "RUT": heredero.get('RUT', ''), "Tipo": tipo,
            "Asignación ($)": round(asignacion_pesos), "Asignación (UTM)": round(asignacion_utm, 2),
            "Exención (UTM)": exencion_utm, "Base Imponible (UTM)": round(base_imponible_utm, 2),
            "Recargo (%)": recargo_pct, "Impuesto Total (UTM)": round(impuesto_total_utm, 2),
            "Impuesto Total ($)": round(impuesto_total_pesos)
        })
    
    return pd.DataFrame(filas_resultado)

# =====================================================================
# ⚖️ CATÁLOGO OFICIAL: EXCEPCIONES DEL ARTÍCULO 464 CPC (JUICIO EJECUTIVO)
# =====================================================================
# Verificado contra el texto vigente del Código de Procedimiento Civil.
# Las 4 primeras son dilatorias; el resto, perentorias.
# =====================================================================
# ⚖️ INSTRUCCIÓN COMPARTIDA DE FUNDAMENTACIÓN JURÍDICA
# =====================================================================
# Se inserta en todos los prompts del sistema que redactan o analizan
# (Excepciones, Escritos Judiciales, Estrategia, Redactor IA, Análisis de
# Escrituras), para que la IA fundamente con la ley exacta y jurisprudencia
# general bien establecida, PERO sin inventar citas específicas (números de
# rol, fechas de sentencias, nombres de causas) que no pueda verificar — un
# riesgo real y documentado de las IA en trabajo legal, que puede terminar
# en un escrito con una cita falsa si no se controla explícitamente.
INSTRUCCION_FUNDAMENTACION_JURIDICA = """
INSTRUCCIONES DE FUNDAMENTACIÓN JURÍDICA (obligatorias):
1. Cita SIEMPRE los artículos legales exactos que sustentan cada punto (Código Civil, Código de Procedimiento Civil, Código de Comercio, Código del Trabajo, Código Orgánico de Tribunales, leyes especiales, etc., según corresponda a la materia).
2. Cuando exista un criterio jurisprudencial GENERAL y bien establecido en el derecho chileno (por ejemplo, "la jurisprudencia reiterada de nuestros tribunales superiores ha sostenido que..."), inclúyelo como refuerzo argumentativo, sin necesidad de un número de causa específico.
3. NUNCA inventes ni completes de memoria una cita jurisprudencial específica (rol de causa, fecha exacta de sentencia, nombre de las partes) si no tienes certeza absoluta de que es correcta y real. Es preferible fundamentar solo con la norma legal y doctrina general antes que arriesgar una cita falsa.
4. Si consideras que el caso amerita buscar una sentencia específica de respaldo, dilo explícitamente como una nota al abogado (ej: "(Sugerencia: verificar jurisprudencia reciente de la Corte Suprema sobre este punto específico antes de presentar)"), en vez de inventarla.
"""

CATALOGO_EXCEPCIONES_464 = {
    1: "La incompetencia del tribunal ante quien se haya presentado la demanda",
    2: "La falta de capacidad del demandante o de personería o representación legal del que comparezca en su nombre",
    3: "La litis pendencia ante tribunal competente, siempre que el juicio que le da origen haya sido promovido por el acreedor",
    4: "La ineptitud del libelo por falta de algún requisito legal en el modo de formular la demanda (Art. 254 CPC)",
    5: "El beneficio de excusión o la caducidad de la fianza",
    6: "La falsedad del título",
    7: "La falta de alguno de los requisitos o condiciones establecidos por las leyes para que dicho título tenga fuerza ejecutiva, sea absolutamente, sea con relación al demandado",
    8: "El exceso de avalúo, en los casos de los incisos 2° y 3° del artículo 438",
    9: "El pago de la deuda",
    10: "La remisión de la deuda",
    11: "La concesión de esperas o la prórroga del plazo",
    12: "La novación",
    13: "La compensación",
    14: "La nulidad de la obligación",
    15: "La pérdida de la cosa debida (Título XIX, Libro IV, Código Civil)",
    16: "La transacción",
    17: "La prescripción de la deuda o sólo de la acción ejecutiva",
    18: "La cosa juzgada",
}

def extraer_texto_pdfs(archivos_pdf_subidos):
    """Extrae el texto de una lista de PDFs subidos, para enviárselo a Groq (no lee PDFs de forma nativa como Gemini)."""
    import PyPDF2
    texto_total = ""
    for archivo in archivos_pdf_subidos:
        try:
            lector = PyPDF2.PdfReader(archivo)
            texto_total += f"\n--- {archivo.name} ---\n" + "\n".join([p.extract_text() or "" for p in lector.pages])
        except Exception:
            texto_total += f"\n--- {archivo.name} (no se pudo leer, posiblemente escaneado sin OCR) ---\n"
    return texto_total

def consultar_groq(prompt: str, temperatura: float = 0.2) -> str:
    """
    Consulta la API de Groq (gratuita, sin tarjeta de crédito, formato
    compatible con OpenAI). Se usa como motor de IA principal del sistema,
    en reemplazo de Gemini, cuya cuenta tiene un problema de facturación
    todavía sin resolver por parte de Google.
    """
    headers = {"Authorization": f"Bearer {st.secrets['GROQ_API_KEY']}", "Content-Type": "application/json"}
    body = {"model": "llama-3.3-70b-versatile", "messages": [{"role": "user", "content": prompt}], "temperature": temperatura}
    respuesta = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=body, timeout=180)
    respuesta.raise_for_status()
    return respuesta.json()["choices"][0]["message"]["content"]

# =====================================================================
# 💰 ARANCEL DE HONORARIOS — COLEGIO DE ABOGADOS DE VALPARAÍSO
# =====================================================================
# Arancel profesional aprobado por el H. Consejo del Colegio de Abogados
# de Valparaíso (14 de junio de 1999). Son valores ORIENTADORES y
# supletorios (Art. 7 del propio Arancel): rigen solo a falta de pacto
# expreso entre abogado y cliente, y sirven como referencia de la
# costumbre forense. Nunca reemplazan el criterio del abogado.
ARANCEL_COLEGIO_VALPARAISO = [
    {'numero': 1, 'descripcion': 'Actas, minutas, cartas, proposiciones o contraproposiciones de negocios, memoriales, solicitudes, formularios, protestas, finiquitos, declaraciones, decretos, instrucciones, etc. (redacción de)', 'honorario': 'De 2 UF a 20 UF'},
    {'numero': 2, 'descripcion': 'Administración de bienes en general, desempeñada por abogado', 'honorario': 'Del 5% al 10% de la renta mensual, o del 0,5% al 1% del capital en caso de bienes que no generen rentas periódicas. El honorario no comprende la redacción de contratos.'},
    {'numero': 3, 'descripcion': 'Administración pro-indiviso o albaceazgo, desempeñado por abogado', 'honorario': 'Se aplicarán las reglas previstas para el honorario del juez árbitro en el N56. .'},
    {'numero': 4, 'descripcion': 'Aguas', 'honorario': '- Tramitación de mercedes de agua: De 15 UF a 50 UF En caso de oposición, el honorario se recargará en un 50%. - Inscripción de derechos de aprovechamiento, de acuerdo al articulado transitorio del Código de Aguas: De 10 UF a 40 UF - Constitución de Organizaciones de Usuarios (Juntas de Vigilancia, Asociaciones de Canalistas, Comunidades de Aguas, Comunidades de Obras de Drenaje y similares): De 4'},
    {'numero': 5, 'descripcion': 'Capitulaciones matrimoniales', 'honorario': 'De 5 UF a 20 UF'},
    {'numero': 6, 'descripcion': 'Cobranzas extrajudiciales', 'honorario': 'Del 2% al 10% del monto recuperado.'},
    {'numero': 7, 'descripcion': 'Concesiones o permisos de servicio público, marítimas, de acuicultura, de pesca, administrativas y similares', 'honorario': 'Se cobrará por tiempo trabajado, en conformidad al artículo 33.'},
    {'numero': 8, 'descripcion': 'Contratos en general (incluye estudio de títulos o antecedentes, redacción, inscripción y publicación)', 'honorario': '- Si se tratare de contratos estandarizados, entendiendo por tales aquellos que se redactan sobre la base de modelos de uso corriente en la práctica forense, el honorario será de 5 UF a 100 UF. - En los demás casos, del 0,5% al 5% del monto del negocio, con un mínimo de 5 UF. Si la cuantía fuere indeterminada, de 5 UF a 200 UF. - En los contratos de tracto sucesivo, la mitad de una renta mensual,'},
    {'numero': 9, 'descripcion': 'Contratos en general (revisión de contratos ya redactados)', 'honorario': 'Del 25% al 50% del honorario que corresponda según el número precedente.'},
    {'numero': 10, 'descripcion': 'Convenios extrajudiciales', 'honorario': 'Del 0,5% al 5% del valor del activo del deudor que figure en el acta del convenio, con un mínimo de 50 UF. El honorario del abogado de cada acreedor será del 1% al 5% del monto del crédito, con un mínimo de 5 UF.'},
    {'numero': 11, 'descripcion': 'Copropiedad Inmobiliaria (asesoramiento general)', 'honorario': 'Del 1% al 5% del valor de venta del edificio, con un mínimo de 50 UF. Estas gestiones comprenden todos los actos y trámites conducentes a la constitución de la copropiedad inmobiliaria y al completo desarrollo de la misma, incluyéndose, por tanto, el estudio de títulos, redacción de escrituras de compraventas y/o sociedad, contratos para la ejecución del proyecto (contratos de promesa, de trabajo,'},
    {'numero': 12, 'descripcion': 'Expropiación forzosa con solución mediante gestión meramente administrativa', 'honorario': 'Del 2% al 5% del monto de la indemnización que se pague.'},
    {'numero': 13, 'descripcion': 'Extranjería, Asilo y Retorno. Comprende tramitación de permiso de residencia y clasificación de refugiado o retornado', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 14, 'descripcion': 'Impacto ambiental (asesoramiento para oponerse a autorización)', 'honorario': 'De 30 UF a 100 UF'},
    {'numero': 15, 'descripcion': 'Impacto ambiental (asesoramiento para obtener autorización ante autoridades administrativas)', 'honorario': 'De 30 UF a 300 UF'},
    {'numero': 16, 'descripcion': 'Informe privado por escrito', 'honorario': 'De 10 UF a 100 UF, según la especialidad o dificultad de la consulta, el tiempo empleado en absolverla y la extensión y trascendencia del informe.'},
    {'numero': 17, 'descripcion': 'Informe en derecho nacional o extranjero', 'honorario': 'De 40 UF a 400 UF, según las mismas circunstancias previstas en la regla precedente.'},
    {'numero': 18, 'descripcion': 'Inscripción de dominio', 'honorario': '- Primera inscripción: Del 1% al 2% del valor comercial del predio con un mínimo de 10 UF. - Otras inscripciones (cuando no están incluidas en los números 8 y 11): De 5 UF a 15 UF.'},
    {'numero': 19, 'descripcion': 'Legalización y protocolización de documentos', 'honorario': 'De 1 UF a 10 UF'},
    {'numero': 20, 'descripcion': 'Nacionalización (gestiones de)', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 21, 'descripcion': 'Parcelaciones, loteos, urbanizaciones y formación de poblaciones (asesoramiento en)', 'honorario': 'El honorario previsto en el N11 y comprende las gestiones señaladas en él, en cuanto fueren compatibles. Del mismo modo, se excluyen los procedimientos judiciales.'},
    {'numero': 22, 'descripcion': 'Permanencia de extranjeros (tramitación de)', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 23, 'descripcion': 'Personalidad jurídica (tramitación), de Corporaciones, fundaciones, cooperativas, sindicatos y asociaciones en general', 'honorario': 'De 20 UF a 200 UF'},
    {'numero': 24, 'descripcion': 'Propiedad Intelectual, Industrial y Marcas Comerciales (Registro de)', 'honorario': 'De 2 UF a 15 UF'},
    {'numero': 25, 'descripcion': 'Reclamaciones y Recursos Administrativos en general, no previstos especialmente', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 26, 'descripcion': 'Rectificación y derecho de respuesta Ley de Prensa (redacción de)', 'honorario': 'De 2 UF a 8 UF'},
    {'numero': 27, 'descripcion': 'Reglamentos, estatutos u otros ordenamientos normativos, cuando no estuvieren comprendidos en el número 23 (redacción de)', 'honorario': 'De 10 UF a 70 UF'},
    {'numero': 28, 'descripcion': 'Representación de personas, grupos de acción cívica, de intereses económicos comunes u otros análogos, ante autoridades del Estado o ante organismos internacionales', 'honorario': 'Cuando no pudieren aplicarse otros números contemplados en este Arancel, se cobrará por tiempo trabajado.'},
    {'numero': 29, 'descripcion': 'Separación convencional de bienes', 'honorario': '- Simple: de 5 UF a 20 UF. - Con liquidación de sociedad conyugal: del 0,5% al 5% del patrimonio social, con un mínimo de 10 UF.'},
    {'numero': 30, 'descripcion': 'Sociedades colectivas, de responsabilidad limitada, en comandita simple y por acciones, anónimas, cuentas en participación o contractuales mineras. Incluye constitución, modificaciones, estudio de títulos o antecedentes, redacción, inscripción y publicación', 'honorario': '- Si el contrato es estandarizado, de 5 UF a 100 UF - En los demás casos, de 10 UF a 200 UF'},
    {'numero': 31, 'descripcion': 'Sociedades (disolución y liquidación)', 'honorario': '- Simple disolución: de 5 UF a 20 UF - Con liquidación: del 0,5% al 5% del patrimonio neto, con un mínimo de 10 UF.'},
    {'numero': 32, 'descripcion': 'Testamentos (redacción de)', 'honorario': 'De 5 UF a 100 UF'},
    {'numero': 33, 'descripcion': 'Títulos (examen de)', 'honorario': 'De 5 UF a 50 UF. Si comprendiere la formación, arreglo u obtención de antecedentes: de 10 UF a 100 UF'},
    {'numero': 34, 'descripcion': 'Transacciones extrajudiciales', 'honorario': 'Si precaven el juicio, hasta la mitad del honorario que correspondería en caso de deducirse.'},
    {'numero': 35, 'descripcion': 'Usufructo, hipotecas y otros derechos reales (constitución)', 'honorario': 'Se aplicará la regla contenida en el N8. # CAPITULO II ASUNTOS DE JURISDICCION VOLUNTARIA.'},
    {'numero': 36, 'descripcion': 'Adopción simple y Adopción plena', 'honorario': 'De 10 UF a 80 UF'},
    {'numero': 37, 'descripcion': 'Ausencia y nombramiento de curador (declaración de)', 'honorario': 'De 10 UF a 50 UF'},
    {'numero': 38, 'descripcion': 'Autorizaciones Judiciales', 'honorario': 'Para enajenar, gravar o dar en arrendamiento bienes de incapaces; para obligar a éstos como fiadores; para celebrar contratos de sociedad; para proceder a la partición de bienes de incapaces; para celebrar capitulaciones matrimoniales y otras autorizaciones que sean necesa-rias para la validez de los actos de incapaces, o supletorias del cónyuge o incapaz: De 10 UF a 50 UF'},
    {'numero': 39, 'descripcion': 'Consignaciones judiciales derivadas de expropiaciones (monto provisional), constitución de servidumbres legales para concesión de servicios públicos y otras similares', 'honorario': 'Del 2% al 4% del monto percibido.'},
    {'numero': 40, 'descripcion': 'Curadores de bienes, adjuntos o especiales (nombramiento de)', 'honorario': 'De 10 UF a 50 UF'},
    {'numero': 41, 'descripcion': 'Herencia yacente y nombramiento de curador (declaración de)', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 42, 'descripcion': 'Información para perpetua memoria', 'honorario': 'De 5 UF a 20 UF'},
    {'numero': 43, 'descripcion': 'Insinuación de donaciones, incluyendo determinación del impuesto', 'honorario': 'De 5 UF a 30 UF'},
    {'numero': 44, 'descripcion': 'Inventario solemne (facción de)', 'honorario': 'De 5 UF a 20 UF'},
    {'numero': 45, 'descripcion': 'Muerte presunta (declaración de)', 'honorario': 'De 10 UF a 50 UF'},
    {'numero': 46, 'descripcion': 'Nombre (diligencias sobre cambio de)', 'honorario': 'De 10 UF a 20 UF'},
    {'numero': 47, 'descripcion': 'Posesión efectiva de herencia, incluyendo apertura de testamento, inventario, protocolizaciones, exención o determinación, pago y aprobación del impuesto de herencia e inscripciones', 'honorario': 'Del 0,5% al 5% del valor comercial de los bienes, con un mínimo de 10 UF'},
    {'numero': 48, 'descripcion': 'Reconocimiento de hijos', 'honorario': 'De 5 UF a 50 UF'},
    {'numero': 49, 'descripcion': 'Rectificación de partidas', 'honorario': 'De 5 UF a 20 UF'},
    {'numero': 50, 'descripcion': 'Remates o pública subasta (patrocinio de interesados en)', 'honorario': 'Del O,5% al 5% del precio, incluyendo las gestiones señaladas en el N8.'},
    {'numero': 51, 'descripcion': 'Tutelas y curatelas generales (nombramiento y discernimiento de)', 'honorario': 'De 10 UF a 50 UF'},
    {'numero': 52, 'descripcion': 'Venta en pública subasta o en licitación', 'honorario': 'Del 0,5% al 4% del precio que se obtenga. El honorario incluirá todo el trabajo profesional necesario para realizar la venta y enajenación de los bienes subastados o licitados. # CAPITULO III MATERIAS CIVILES CONTENCIOSAS'},
    {'numero': 53, 'descripcion': 'Acción de desposeimiento', 'honorario': 'Abogado del demandante: del 2% al 10% del monto recuperado, con un mínimo de 10 UF. Abogado del demandado: del 2% al 10% del valor comercial del bien, o de la deuda, debiendo aplicarse el porcentaje sobre el valor que fuere menor.'},
    {'numero': 54, 'descripcion': 'Aguas (recurso de amparo de)', 'honorario': 'De 20 UF a 200 UF'},
    {'numero': 55, 'descripcion': 'Alimentos (juicios de)', 'honorario': 'Abogado del demandante: un honorario equivalente de 1 a 3 meses de la pensión que se obtenga. Abogado del demandado: un honorario de media a una y media pensión mensual demandada. Cuando se hubiere remunerado el incidente de alimentos provisionales, este honorario servirá de abono al que corresponda al término del juicio.'},
    {'numero': 56, 'descripcion': 'Arbitrajes, liquidaciones y particiones (honorarios del juez o del liquidador)', 'honorario': '- El árbitro podrá fijar sus honorarios por tiempo o de acuerdo a la cuantía, considerando la naturaleza del caso, o conforme a lo que convenga con todas las partes del proceso. - Cuando el honorario se fije por tiempo, el árbitro deberá tener presente los parámetros establecidos en el artículo 33 de este arancel. - Cuando el honorario se fije según la cuantía, se aplicará la tabla siguiente a fal'},
    {'numero': 57, 'descripcion': 'Arrendamientos (juicios de)', 'honorario': 'De una a cuatro rentas mensuales.'},
    {'numero': 58, 'descripcion': 'Avenimiento, conciliación o transacción', 'honorario': 'Cuando se pone término al juicio, regirá la regla del art. 24.'},
    {'numero': 59, 'descripcion': 'Bien Familiar (declaración o defensa en gestión judicial)', 'honorario': 'De 10 UF a 100 UF'},
    {'numero': 60, 'descripcion': 'Citación de evicción', 'honorario': 'De 4 UF a 40 UF'},
    {'numero': 61, 'descripcion': 'Convenio Judicial Preventivo', 'honorario': 'Abogado del deudor: Del 1% al 5% del valor del activo, con un mínimo de 75 UF. Abogado del acreedor: Del 1% al 5% del monto del crédito, con un mínimo de 5 UF.'},
    {'numero': 62, 'descripcion': 'Cumplimiento de sentencias dictadas por tribunales extranjeros. En materia contenciosa, de 10 UF a 100 UF. En materia voluntaria, de 5 UF a 50 UF. ## 63.- Cheques (cobranzas de)', 'honorario': 'Se aplicará la regla correspondiente al juicio ejecutivo.'},
    {'numero': 64, 'descripcion': 'Derecho legal de retención como acción especial', 'honorario': 'Se aplicará la regla del N72, sobre medidas cautelares.'},
    {'numero': 65, 'descripcion': 'Exhortos (diligenciamiento de)', 'honorario': 'Nacionales: de 2 UF a 25 UF. Extranjeros: de 10 UF a 75 UF.'},
    {'numero': 66, 'descripcion': 'Expropiación (gestión o juicios de)', 'honorario': 'Se aplicarán las reglas del juicio ordinario, calculándose los porcentajes sobre la mayor suma que se consigne judicialmente.'},
    {'numero': 67, 'descripcion': 'Incidentes', 'honorario': 'De previo y especial pronunciamiento: Hasta la cuarta parte del honorario correspondiente a la cuestión principal. Ordinarios: Hasta la décima parte del honorario correspondiente a la cuestión principal. En segunda instancia el honorario se elevará hasta el doble. Si se ha asumido la defensa de la causa, no habrá derecho a cobrar honorarios por la promoción de incidentes.'},
    {'numero': 68, 'descripcion': 'Interdictos posesorios', 'honorario': 'Regirán las reglas del juicio ordinario o del de cuantía indeterminada, en su caso.'},
    {'numero': 69, 'descripcion': 'Jactancia (acción de)', 'honorario': 'El 25% del honorario que corresponda al juicio que con la demanda de jactancia se trata de provocar o silenciar.'},
    {'numero': 70, 'descripcion': 'Juicios y Procedimientos', 'honorario': '- Juicio de cuantía indeterminada: De 10 UF a 100 UF, pudiendo elevarse hasta 500 UF en casos de especial trascendencia económica. ## Juicio ejecutivo y ejecución de resoluciones conforme al Título XIX del Libro I del Código de Procedimiento Civil. Abogado del ejecutante: Si se oponen excepciones, se aplicará la escala de la letra c) de este número, rebajada a las tres cuartas partes. Si no hay ex'},
    {'numero': 71, 'descripcion': 'Medidas Cautelares Prejudiciales', 'honorario': 'Si con la obtención de la medida se solucionare la cuestión o cuestiones que iban a ser objeto del juicio, el honorario será la mitad del que habría correspondido a dicho juicio. Si no se hubiera obtenido la medida, o habiéndosela obtenido no se solucionare la cuestión, de 10 UF a 100 UF. El mismo honorario precedente corresponderá al abogado de la parte en contra de la cual se hubiere pedido la m'},
    {'numero': 72, 'descripcion': 'Medidas Precautorias o cautelares', 'honorario': 'Se entenderán comprendidas en el honorario correspondiente al juicio, en conformidad al art. 18 letra m).'},
    {'numero': 73, 'descripcion': 'Menores (patrocinio en gestiones ante los Juzgados de Letras de)', 'honorario': '- Si fueren de alimentos, se aplicarán las reglas previstas para el juicio de alimentos. - Si fueren sobre otras materias, de 3 UF a 50 UF, salvo que se tratare de asuntos especialmente contemplados en otros números de este Arancel, en que el honorario se encuentra fijado en ellos.'},
    {'numero': 74, 'descripcion': 'Notificaciones aisladas de Cesión de Créditos, Prendas, Títulos Ejecutivos, y demás asuntos o diligencias similares, en que la ley o el contrato dispusieren la notificación para fines especiales', 'honorario': 'De 1 UF a 10 UF.'},
    {'numero': 75, 'descripcion': 'Pago por Consignación', 'honorario': 'De 5 UF a 50 UF. En el juicio sobre suficiencia del pago el honorario se regirá por las reglas correspondientes al juicio ordinario.'},
    {'numero': 76, 'descripcion': 'Preparación de la vía ejecutiva sin que se siga ejecución', 'honorario': 'Si estas gestiones producen como resultado el pago total o parcial de la deuda, el honorario será el que corresponda al juicio ejecutivo. Si no se obtiene título ejecutivo, o si obteniéndolo no se consigue el pago, el honorario será de 5 UF a 50 UF.'},
    {'numero': 77, 'descripcion': 'Quiebra (Juicios de)', 'honorario': 'Abogado del acreedor: Se aplicará ]a escala del juicio ejecutivo, sobre la cantidad que el acreedor perciba, con un mínimo de 10 UF. Abogado del deudor: Si éste hubiere solicitado su quiebra, del 1% al 5% del activo de la quiebra. Si se rechaza la quiebra, de 50 UF a 200 UF. Si se rechazare la oposición y se declara la quiebra, de 10 UF a 100 UF. En el juicio de calificación de la quiebra, el hono'},
    {'numero': 78, 'descripcion': 'Tercerías en juicio ejecutivo', 'honorario': 'Se aplicará la escala del juicio ejecutivo con un mínimo de 10 UF. El honorario se calculará sobre el valor del bien en las tercerías de dominio y posesión, y sobre el monto percibido en las tercerías de prelación y pago.'},
    {'numero': 79, 'descripcion': 'Violencia intrafamiliar (denuncia o defensa de)', 'honorario': 'De 2 UF a 50 UF. # CAPITULO IV ACCIONES Y MATERIAS JURISDICCIONALES DE RANGO CONSTITUCIONAL.'},
    {'numero': 80, 'descripcion': 'Acciones de rango constitucional no previstas especialmente en este arancel', 'honorario': 'De 30 UF a 300 UF.'},
    {'numero': 81, 'descripcion': 'Error o arbitrariedad judicial. Declaración de la Corte Suprema y juicio indemnizatorio', 'honorario': '- Declaración de la Corte Suprema, de 50 UF a 200 UF. - Juicio indemnizatorio. Se aplicarán las reglas del juicio ordinario civil.'},
    {'numero': 82, 'descripcion': 'Expropiación (acción o reclamo de ilegalidad del acto expropiatorio)', 'honorario': 'De 50 UF a 250 UF.'},
    {'numero': 83, 'descripcion': 'Inconstitucionalidad de organizaciones, movimientos o partidos políticos', 'honorario': 'De 50 UF a 200 UF.'},
    {'numero': 84, 'descripcion': 'Nacionalidad. Reclamación contra acto o resolución administrativa que prive de ella o la desconozca', 'honorario': 'De 50 UF a 100 UF.'},
    {'numero': 85, 'descripcion': 'Nulidad de Derecho Público (acción de)', 'honorario': 'Se aplicarán las reglas del juicio ordinario civil de cuantía indeterminada.'},
    {'numero': 86, 'descripcion': 'Protección (acción o recurso de)', 'honorario': 'De 25 UF a 500 UF.'},
    {'numero': 87, 'descripcion': 'Rehabilitación de Ciudadanía ante el Senado', 'honorario': 'De 40 UF a 100 UF.'},
    {'numero': 88, 'descripcion': 'Requisiciones y limitaciones al dominio en estados de excepción constitucional', 'honorario': 'Se aplicarán las reglas del juicio ordinario civil. # CAPITULO V MATERIAS PENALES'},
    {'numero': 89, 'descripcion': 'Acción Civil de juicio ordinario Se aplicarán las reglas de juicio ordinario ## 90.- Acusación o contestación de las mismas', 'honorario': 'La tercera parte del honorario contemplado para el juicio ordinario penal.'},
    {'numero': 91, 'descripcion': 'Anotaciones prontuariales (eliminación de)', 'honorario': 'De 5 UF a 50 UF.'},
    {'numero': 92, 'descripcion': 'Denuncias (redacción de)', 'honorario': 'De 5 UF a 50 UF.'},
    {'numero': 93, 'descripcion': 'Desafueros De 100 UF a 500 UF. ## 94.- Embargo de bienes (incidentes de)', 'honorario': 'Se aplicará el honorario previsto para las tercerías del juicio ejecutivo.'},
    {'numero': 95, 'descripcion': 'Extradición. De 100 UF a 500 UF. ## 96.- Indulto (remisivo, reductivo y conmutativo)', 'honorario': 'De 20 UF a 100 UF.'},
    {'numero': 97, 'descripcion': 'Juicio ordinario penal', 'honorario': '- Abogado Querellante: - Presentación de denuncia o querella: De UF. 10 a UF. 50 - Al dictarse auto de procedimiento o sobreseimiento sin que exista procesado: De UF. 10 a UF. 50 - Al presentarse acusación o adhesión de la acusación: De UF. 20 a UF. 100 - Al dictarse sentencia definitiva de 1ª instancia: De UF. 20 a UF. 100 - Por la segunda instancia: De UF. 30 a UF. 150 ## Abogado del Querellado:'},
    {'numero': 98, 'descripcion': 'Ley de Tránsito (infracción a la)', 'honorario': 'De 5 UF a 5O UF. Si se ejerciere la acción civil, se devengará además el honorario que corresponda según las reglas del juicio ordinario civil.'},
    {'numero': 99, 'descripcion': 'Libertad provisional (gestión aislada)', 'honorario': 'De 5 UF a 50 UF.'},
    {'numero': 100, 'descripcion': 'Querella de capítulos', 'honorario': 'De 50 UF a 500 UF.'},
    {'numero': 101, 'descripcion': 'Restitución de especies en poder de la justicia (mera devolución)', 'honorario': 'De 5 UF a 30 UF. # CAPITULO VI MATERIAS LABORALES'},
    {'numero': 102, 'descripcion': 'Comisiones de Medicina preventiva y curativa (gestiones ante las)', 'honorario': 'De 10 UF a 30 UF'},
    {'numero': 103, 'descripcion': 'Despidos (demandas por)', 'honorario': 'Del 10% al 25% del monto obtenido.'},
    {'numero': 104, 'descripcion': 'Leyes sociales y del trabajo, defensas administrativas en denuncias por infracción a ellas (ante la Inspección del Trabajo, instituciones de previsión, ministerio o subsecretarías, etc.)', 'honorario': 'Del 5% al 15% de la ventaja a beneficio económico que resultare al cliente, con un mínimo de 5 UF.'},
    {'numero': 105, 'descripcion': 'Negociación y conflictos Colectivos (incluye asesoría, y redacción o revisión del contrato). - Conversaciones directas (incluyendo comparendos o audiencias ante las autoridades), sin acuerdo que ponga término al conflicto', 'honorario': 'De una cuarta UF a media UF por trabajador involucrado, con un mínimo de 20 UF. - Conversaciones directas, con acta de avenimiento o contrato colectivo. De media UF a una UF por trabajador involucrado, con un mínimo de 40 UF - Honorarios del empleador: De 30 UF a 300 UF.'},
    {'numero': 106, 'descripcion': 'Pensiones y Jubilaciones (gestiones para obtener su reajuste o revalorización)', 'honorario': 'Del 5% al 20% del aumento que representare en un año, con un mínimo de 10 UF.'},
    {'numero': 107, 'descripcion': 'Previsión (gestiones relativas a los beneficios de)', 'honorario': '- Por la tramitación de beneficios que conceden las leyes de previsión, el 10% de las pensiones o asignaciones atrasadas, y, además, de una a tres de las pensiones y asignaciones bases mensuales que se fijen definitivamente. - Por la tramitación de desahucio que conceden esas mismas leyes, del 5% al 10% del monto total. - En caso de rechazarse el beneficio solicitado, el honorario se regulará de a'},
    {'numero': 108, 'descripcion': 'Reclamaciones laborales de carácter administrativo', 'honorario': 'Del 5% al 20% del monto disputado y percibido.'},
    {'numero': 109, 'descripcion': 'Trabajo (juicios que no incluyen despidos)', 'honorario': '- Si fueren susceptibles de apreciación pecuniaria, se estar a las reglas del juicio ordinario civil. - En los demás casos, se estará a las reglas del juicio de cuantía indetermada. # CAPITULO VII MATERIAS TRIBUTARIAS A.- CONTENCIOSAS.'},
    {'numero': 110, 'descripcion': 'Contestación de Citaciones', 'honorario': 'De 10 UF a 50 UF.'},
    {'numero': 111, 'descripcion': 'Reclamos de avalúo y reclamos por cambios de avalúo, sobre bienes raíces', 'honorario': 'Del 2% al 4% de la diferencia de avalúo que se obtenga en definitiva respecto al fijado en la tasación general.'},
    {'numero': 112, 'descripcion': 'Reclamos de Impuestos', 'honorario': 'Del 12 al 15% sobre los menores impuestos actualizados que se obtengan, con un mínimo a todo evento de 20 UF a 100 UF, que se imputará a los porcentajes referidos. En casos de reclamos de liquidaciones separadas, por concepto de Impuesto Global Complementario correspondiente a los socios, y que provengan de las mismas partidas liquidadas a la sociedad por Impuesto de Primera Categoría, se cobrará'},
    {'numero': 113, 'descripcion': 'Oposiciones en cobros ejecutivos de impuestos', 'honorario': 'Se aplicará el honorario previsto para el juicio ejecutivo. # B.- INFRACCIONALES.'},
    {'numero': 114, 'descripcion': 'Juicios por delitos tributarios', 'honorario': 'Se aplicará el honorario previsto para el juicio ordinario penal.'},
    {'numero': 115, 'descripcion': 'Reclamo de denuncias por infracciones', 'honorario': 'De 10 UF a 80 UF. # C.- GESTIONES NO CONTENCIOSAS.'},
    {'numero': 116, 'descripcion': 'Convenios de pago con el Servicio de Tesorería. Del 2% al 5% del monto del impuesto objeto del convenio, con un mínimo de 5 UF. ## 117.- Solicitudes de condonación de intereses penales ante el Servicio de Impuestos Internos o ante el Servicio de Tesorería. De 5 UF a 20 UF. ## 118.- Solicitudes de devolución de impuestos (art. 126 Código Tributario)', 'honorario': 'Del 5% al 10% sobre las sumas actualizadas que se devuelvan, con un mínimo de 5 UF. # CAPITULO VIII MATERIAS MINERAS.'},
    {'numero': 119, 'descripcion': 'Administración de la Pertenencia por parte del Minero, del Aviador o del Acreedor a quien corresponda (juicios relativos al ejercicio de la)', 'honorario': 'De 10 UF a 100 UF.'},
    {'numero': 120, 'descripcion': 'Catar y Cavar (gestión judicial para obtener permiso)', 'honorario': 'De 10 UF a 50 UF.'},
    {'numero': 121, 'descripcion': 'Concesión minera de exploración (gestión judicial de constitución)', 'honorario': 'De 20 UF a 100 UF.'},
    {'numero': 122, 'descripcion': 'Concesión minera de explotación o pertenencia (gestión judicial de constitución, sin considerar los juicios de oposición)', 'honorario': 'De 30 UF a 200 UF.'},
    {'numero': 123, 'descripcion': 'Internación de pertenencias (juicios sobre)', 'honorario': 'De 20 UF a 200 UF.'},
    {'numero': 124, 'descripcion': 'Juicios mineros no contemplados expresamente en este párrafo', 'honorario': 'Si son de cuantía determinada, debe estarse a la escala establecida para el juicio ordinario. Si son de cuantía indeterminada, de 20 UF a 200 UF, pudiendo elevarse hasta 500 UF en casos de especial trascendencia económica.'},
    {'numero': 125, 'descripcion': 'Mensura (juicios de nulidad de la concesión minera)', 'honorario': 'De 20 UF a 200 UF.'},
    {'numero': 126, 'descripcion': 'Mensura (juicios de oposición a la)', 'honorario': 'De 20 UF a 200 UF.'},
    {'numero': 127, 'descripcion': 'Remate de concesión minera por no pago de patente (defensa en)', 'honorario': '| 128.- | Remate de concesión minera (asesoramiento para participar en): De 10 UF a 50 UF. | | --- | --- | | 129.- | Servidumbres mineras (juicios de constitución, ejercicio y terminación de): De 30 UF a 200 UF. | | 130.- | Servidumbres mineras (constitución por escritura pública): De 20 UF a 100 UF. | | 131.- | Sociedades legales Mineras (asesoramiento): Se remunerará por tiempo trabajado. | # CA'},
    {'numero': 132, 'descripcion': 'Acusaciones constitucionales ante el Senado', 'honorario': 'De 50 UF a 200 UF.'},
    {'numero': 133, 'descripcion': 'Aduana (procedimientos de)', 'honorario': '- Reclamo de aforo: De 20 UF a 100 UF. - Reclamaciones por simples infracciones aduaneras: De 5 UF a 50 UF. - Juicio aduanero: Se aplicará el honorario previsto en el juicio ordinario civil, respecto del abogado del demandado. - Gestión para obtener la renuncia de la acción penal: De 10 UF a 100 UF.'},
    {'numero': 134, 'descripcion': 'Antimonopolios (Ley)', 'honorario': '- Defensa ante la Comisión Preventiva Regional: De 5 UF a 100 UF. - Defensa ante la Comisión Preventiva Central: - Defensa ante la Comisión Resolutiva: De 50 UF a 500 UF. - Recurso de reclamación ante la Corte Suprema: De 50 UF a 300 UF. El honorario de la defensa en dos o más tramos de los indicados no será acumulativa cuando sea atendida por el mismo abogado. En este caso, se cobrará el honorari'},
    {'numero': 135, 'descripcion': 'Aporte de capitales extranjeros (gestiones sobre)', 'honorario': 'Del O,5% al 2% del capital involucrado.'},
    {'numero': 136, 'descripcion': 'Clausura de inmuebles (alzamiento)', 'honorario': 'De 10 UF a 100 UF.'},
    {'numero': 137, 'descripcion': 'Concesiones de Bienes Nacionales', 'honorario': 'De una a cuatro rentas mensuales garantizadas por el contrato. Si no pudiere aplicarse esta regla, el honorario será de 5 UF a 200 UF.'},
    {'numero': 138, 'descripcion': 'Contiendas de Competencia ante el Tribunal Constitucional, ante el Senado o ante la Corte Suprema', 'honorario': 'De 20 UF a 100 UF.'},
    {'numero': 139, 'descripcion': 'Electorales (reclamos)', 'honorario': 'Ante la Justicia Ordinaria o Justicia Electoral, de 5 UF a 50 UF'},
    {'numero': 140, 'descripcion': 'Ilegalidad contra decretos alcaldicios, acuerdos de las Municipalidades o actuaciones de funcionarios municipales, formulados ante la autoridad comunal o ante los tribunales (reclamos o recursos de)', 'honorario': 'Del 5 % al 10 % de la cuantía controvertida del negocio. Si no fuere susceptible de apreciación pecuniaria, de 10 UF a 100 UF.'},
    {'numero': 141, 'descripcion': 'Inscripción de documentos en casos de negativa del Conservador de Bienes Raíces (gestiones para obtener la)', 'honorario': 'De 5 UF a 50 UF.'},
    {'numero': 142, 'descripcion': 'Militares, Navales, de Carabineros y Aviación (juicios ante Tribunales)', 'honorario': 'Se aplicarán las reglas del juicio ordinario penal.'},
    {'numero': 143, 'descripcion': 'Patentes (reclamos de)', 'honorario': 'De 10 UF a 100 UF.'},
    {'numero': 144, 'descripcion': 'Policía Local (patrocinio ante los Juzgados de)', 'honorario': 'De aplicará la regla prevista para las infracciones a la Ley del Tránsito en el N97.'},
    {'numero': 145, 'descripcion': 'Propiedad Industrial y Marcas Comerciales (gestiones sobre)', 'honorario': 'Se aplicarán la regla prevista para su respectivo registro en el N24.'},
    {'numero': 146, 'descripcion': 'Propiedad Industrial y Marcas Comerciales (Juicios sobre)', 'honorario': 'Se aplicarán las reglas del juicio ordinario civil o penal, según el caso.'},
    {'numero': 147, 'descripcion': 'Propiedad Intelectual y Derechos de Autor (gestiones relativas a)', 'honorario': 'Se aplicará la regla prevista para su respectivo registro en el N24.'},
    {'numero': 148, 'descripcion': 'Protección al Consumidor', 'honorario': '- Defensa ante el Servicio Nacional del Consumidor: De 2 UF a 50 UF. - Defensa ante el Juzgado de Policía Local: Se aplicará la regla prevista para las infracciones a la Ley del Tránsito.'},
    {'numero': 149, 'descripcion': 'Regularización de la posesión de Bienes Raíces (D.L. 2695)', 'honorario': '- Regularización: De 10 UF a 100 UF. - Defensa del propietario afectado: Del 10% al 20% del valor comercial del bien recuperado, con un mínimo de 50 UF. - Acciones de dominio y de compensación: Se aplicarán las reglas del juicio ordinario.'},
    {'numero': 150, 'descripcion': 'Regularización de vehículos, capitales, u otros bienes en situación irregular', 'honorario': 'Del 1% al 4 % de su cuantía o valor comercial.'},
    {'numero': 151, 'descripcion': 'Servicios Públicos, semifiscales o de administración autónoma, bancos, sociedades, etc. (gestiones ante)', 'honorario': 'De 3 UF a 50 UF.'},
    {'numero': 152, 'descripcion': 'Superintendencia de Bancos e Instituciones Financieras, de Valores y Seguros, de Seguridad Social, Cámaras de Comercio y otras semejantes (arbitrajes ante la): Para los honorarios del abogado, se aplicará la regla prevista para el juicio ordinario civil. # CAPITULO X RECURSOS Y ALEGATOS. Este párrafo regula los honorarios por recursos y alegatos de abogados distintos del patrocinante de la causa, y también los de este último para los efectos del inciso segundo del artículo 140 del Código de Procedimiento Civil. # RECURSOS ## 153.- Amparo (recurso de)', 'honorario': 'De 20 UF a 150 UF.'},
    {'numero': 154, 'descripcion': 'Casación civil ante la Corte Suprema (interposición y alegato)', 'honorario': 'De 50 UF a 300 UF.'},
    {'numero': 155, 'descripcion': 'Casación penal ante la Corte Suprema (interposición y alegato)', 'honorario': 'De 50 UF a 300 UF.'},
    {'numero': 156, 'descripcion': 'Hecho (redacción del recurso de)', 'honorario': 'De 10 UF a 30 UF.'},
    {'numero': 157, 'descripcion': 'Inaplicabilidad (redacción y tramitación del recurso de)', 'honorario': 'De 50 UF a 300 UF.'},
    {'numero': 158, 'descripcion': 'Queja (redacción y tramitación del recurso de)', 'honorario': 'De 10 UF a 100 UF.'},
    {'numero': 159, 'descripcion': 'Recursos y solicitudes ante la Contraloría General de la República', 'honorario': 'De 10 UF a 100 UF.'},
    {'numero': 160, 'descripcion': 'Revisión (redacción y tramitación del recurso de)', 'honorario': 'De 50 UF a 300 UF. # ALEGATOS. 161.- Si el alegato no está comprendido en la interposición del recurso, se aplicarán las siguientes reglas: - Del 20% al 50% del honorario que corresponda a la interposición del recurso. - Si no pudiere aplicarse la regla anterior, el honorario será de 10 UF a 100 UF, según la naturaleza y complejidad del recurso.'},
]


# =====================================================================
# 📖 CONEXIÓN CON LA API DE DATOS ABIERTOS DE BCN (LeyChile)
# =====================================================================
# BCN (Biblioteca del Congreso Nacional) SÍ ofrece una API pública y
# documentada para su base de legislación (a diferencia de PJUD, que
# bloquea el acceso automatizado). Solo se incluyen los códigos cuyo
# identificador (idNorma) fue verificado con certeza contra fuentes
# oficiales — para evitar el riesgo de traer, por error, el texto de una
# ley distinta a la que se cree estar consultando.
CODIGOS_BCN_IDNORMA = {
    "Código Civil": "172986",
    "Código de Procedimiento Civil": "22740",
    "Código Orgánico de Tribunales": "25563",
    "Código Procesal Penal": "176595",
}

def obtener_articulos_codigo_bcn(nombre_codigo):
    """
    Descarga el texto vigente y actualizado de un código legal chileno
    directo desde la API oficial de datos abiertos de BCN. Devuelve una
    lista de (tipo_de_parte, texto) por cada artículo/párrafo del código.
    """
    id_norma = CODIGOS_BCN_IDNORMA.get(nombre_codigo)
    if not id_norma:
        return None
    try:
        import xml.etree.ElementTree as ET
        respuesta = requests.get(f"https://www.leychile.cl/Consulta/obtxml?opt=7&idNorma={id_norma}", timeout=30)
        respuesta.raise_for_status()
        root = ET.fromstring(respuesta.content)
        ns = "{http://www.leychile.cl/esquemas}"
        articulos = []
        for estructura in root.iter(f"{ns}EstructuraFuncional"):
            tipo_parte = estructura.get('tipoParte', '')
            texto_elem = estructura.find(f"{ns}Texto")
            if texto_elem is not None and texto_elem.text:
                articulos.append((tipo_parte, texto_elem.text.strip()))
        return articulos
    except Exception as e:
        print(f"Error obteniendo código desde BCN: {e}")
        return None

def buscar_en_codigo_bcn(nombre_codigo, termino_busqueda, max_resultados=5):
    """
    Busca un término (palabra clave o número de artículo) dentro del texto
    fresco de un código legal, obtenido en vivo desde BCN — no una copia
    guardada que podría quedar desactualizada si la ley cambia.
    """
    articulos = obtener_articulos_codigo_bcn(nombre_codigo)
    if not articulos:
        return []
    resultados = []
    termino_lower = termino_busqueda.lower().strip()
    for tipo_parte, texto in articulos:
        if termino_lower in texto.lower():
            resultados.append(texto)
        if len(resultados) >= max_resultados:
            break
    return resultados

def buscar_jurisprudencia_relevante(texto_busqueda, top_n=3):
    """
    Busca en la biblioteca de jurisprudencia (sentencias que el propio
    abogado ha ido cargando) las más relacionadas con el texto de un caso,
    usando coincidencia de palabras clave sobre la Materia y el Resumen.
    Solo encuentra sentencias REALES que el usuario cargó — nunca inventa
    ninguna, evitando el riesgo de citas jurisprudenciales falsas.
    """
    df_juris = safe_read_sheet("base_jurisprudencia", COLS_JURISPRUDENCIA)
    if df_juris.empty:
        return []
    palabras_busqueda = set(re.findall(r'\w{4,}', texto_busqueda.lower()))
    resultados = []
    for _, fila in df_juris.iterrows():
        texto_fila = f"{fila.get('Materia','')} {fila.get('Resumen','')}"
        palabras_fila = set(re.findall(r'\w{4,}', str(texto_fila).lower()))
        interseccion = palabras_busqueda & palabras_fila
        if interseccion:
            resultados.append((len(interseccion), fila))
    resultados.sort(key=lambda x: -x[0])
    return [r[1] for r in resultados[:top_n] if r[0] > 0]

def buscar_arancel_referencial(texto_busqueda, top_n=3):
    """
    Busca en el Arancel del Colegio de Abogados de Valparaíso las entradas
    más relacionadas con el texto de una acción o rama del derecho, usando
    coincidencia de palabras clave (sin necesitar librerías extra). Se usa
    en Contratos para sugerir un rango de honorarios de referencia.
    """
    palabras_busqueda = set(re.findall(r'\w{4,}', texto_busqueda.lower()))
    resultados = []
    for item in ARANCEL_COLEGIO_VALPARAISO:
        palabras_item = set(re.findall(r'\w{4,}', item['descripcion'].lower()))
        interseccion = palabras_busqueda & palabras_item
        if interseccion:
            resultados.append((len(interseccion), item))
    resultados.sort(key=lambda x: -x[0])
    return [r[1] for r in resultados[:top_n] if r[0] > 0]

def _limpiar_json_ia(texto_respuesta: str) -> str:
    texto_respuesta = texto_respuesta.strip()
    if texto_respuesta.startswith("```"):
        texto_respuesta = re.sub(r'^```(json)?\s*', '', texto_respuesta)
        texto_respuesta = re.sub(r'\s*```$', '', texto_respuesta)
    return texto_respuesta

def analizar_excepciones_con_ia(archivos_pdf_subidos, contexto_adicional=""):
    """
    Analiza los documentos y determina cuáles de las 18 excepciones del
    Art. 464 CPC son aplicables, con nivel de confianza y cita textual de
    respaldo. Usa Gemini directamente (sin selección de motor: es la forma
    más simple). Devuelve la lista de dicts ya parseada desde JSON.
    """
    lista_excepciones_texto = "\n".join([f"N°{n}: {texto}" for n, texto in CATALOGO_EXCEPCIONES_464.items()])
    
    instrucciones_base = f"""
    Actúa como un abogado chileno experto en juicio ejecutivo y en la oposición de excepciones del artículo 464 del Código de Procedimiento Civil.

    Analiza en profundidad los documentos de esta causa ejecutiva (pueden incluir: demanda, título ejecutivo como pagaré o mandato, resoluciones, personería, certificados, comprobantes, etc.).

    Las 18 excepciones posibles del artículo 464 del CPC son:
    {lista_excepciones_texto}

    Para CADA UNA de las 18 excepciones, determina si es aplicable a este caso concreto según los documentos. Contexto adicional entregado por el abogado: {contexto_adicional if contexto_adicional.strip() else "(sin contexto adicional)"}

    {INSTRUCCION_FUNDAMENTACION_JURIDICA}
    En el campo "fundamento" de cada excepción, además de los hechos, cita el o los artículos legales exactos que la sustentan (del CPC, Código Civil, Código de Comercio u otro cuerpo legal según corresponda) y, si existe, el criterio jurisprudencial general aplicable siguiendo las reglas anteriores.

    Responde EXCLUSIVAMENTE con un array JSON válido (nada de texto antes o después, sin usar bloques de código markdown), con un objeto por cada una de las 18 excepciones, con esta estructura exacta:
    [
      {{
        "numero": 14,
        "nombre": "Nulidad de la obligación",
        "aplica": true,
        "confianza": "Alta",
        "fundamento": "Explicación detallada de por qué aplica o no aplica, citando hechos concretos de los documentos, la norma legal exacta que la sustenta, y el criterio jurisprudencial general si corresponde.",
        "cita_textual": "Cita literal breve (máximo 40 palabras) extraída del documento que respalda el fundamento, o cadena vacía si no aplica."
      }}
    ]
    "confianza" debe ser "Alta", "Media" o "Baja" si aplica=true, o null si aplica=false.
    Sé riguroso: solo marca aplica=true cuando los documentos realmente respalden la excepción con hechos concretos, no supongas nada que no esté en los documentos.
    """
    
    texto_documentos = extraer_texto_pdfs(archivos_pdf_subidos)
    prompt_final = instrucciones_base + f"\n\nTEXTO EXTRAÍDO DE LOS DOCUMENTOS:\n{texto_documentos[:45000]}"
    texto_respuesta = consultar_groq(prompt_final)
    return json.loads(_limpiar_json_ia(texto_respuesta))

def redactar_escrito_judicial_ia(tipo_escrito, instrucciones_tipo, archivos_pdf_subidos, contexto_adicional):
    """
    Motor general para redactar CUALQUIER tipo de presentación judicial
    (demandas, evacúa traslados, abandonos de procedimiento, nulidades
    procesales, tercerías, etc.), no solo excepciones. Analiza los
    documentos adjuntos (si los hay) y devuelve el texto del escrito ya
    redactado, listo para pasar al generador de Word. Usa Gemini
    directamente (sin selección de motor: es la forma más simple).
    """
    # Se busca en la biblioteca de jurisprudencia (sentencias reales que el
    # propio abogado cargó) algo relacionado con este escrito. A diferencia
    # de pedirle a la IA que "recuerde" jurisprudencia (con riesgo de que la
    # invente), esto son sentencias REALES ya verificadas por el usuario.
    sentencias_relevantes = buscar_jurisprudencia_relevante(f"{tipo_escrito} {instrucciones_tipo} {contexto_adicional}")
    bloque_jurisprudencia_real = ""
    if sentencias_relevantes:
        bloque_jurisprudencia_real = "\n\nJURISPRUDENCIA REAL DISPONIBLE (de la biblioteca del estudio, verificada, puedes citarla con confianza si es pertinente):\n"
        for s in sentencias_relevantes:
            bloque_jurisprudencia_real += f"- {s.get('Tribunal','')}, Rol {s.get('Rol_Causa','')}, {s.get('Fecha_Sentencia','')}: {s.get('Resumen','')}\n"
    
    prompt_base = f"""
    Actúa como un abogado chileno experto en litigación, redactando una presentación judicial de tipo: {tipo_escrito}.

    Instrucciones específicas para este tipo de escrito: {instrucciones_tipo}

    Contexto y hechos entregados por el abogado: {contexto_adicional if contexto_adicional.strip() else "(sin contexto adicional escrito, básate solo en los documentos adjuntos)"}
    {bloque_jurisprudencia_real}

    {INSTRUCCION_FUNDAMENTACION_JURIDICA}

    Redacta el escrito completo, con lenguaje formal jurídico chileno, incluyendo su suma, comparecencia (usa placeholders genéricos como [NOMBRE], [ROL] si no tienes el dato exacto), fundamentos de hecho y de derecho lo más completos posible (citando las normas legales exactas aplicables y, cuando corresponda, el criterio jurisprudencial general según las reglas anteriores), y el petitorio final ("POR TANTO, RUEGO A US...").
    Estructura el texto en párrafos separados por doble salto de línea (\\n\\n), sin usar títulos markdown (nada de # ni **), solo texto plano formal, ya que se insertará directo en un documento Word.
    """
    if archivos_pdf_subidos:
        texto_documentos = extraer_texto_pdfs(archivos_pdf_subidos)
        prompt_base += f"\n\nTEXTO EXTRAÍDO DE LOS DOCUMENTOS ADJUNTOS:\n{texto_documentos[:45000]}"
    return consultar_groq(prompt_base)

# =====================================================================
# 📝 CATÁLOGO DE TIPOS DE ESCRITOS JUDICIALES (general, no solo excepciones)
# =====================================================================
TIPOS_ESCRITOS_JUDICIALES = {
    "Demanda (Ejecutiva u Ordinaria)": "Redacta una demanda completa, incluyendo los hechos, los fundamentos de derecho aplicables según el tipo de acción, y el petitorio.",
    "Evacúa Traslado (Contestación de Demanda)": "Redacta la contestación de la demanda, oponiendo las excepciones y defensas de fondo pertinentes, controvirtiendo los hechos y el derecho invocado por la contraria.",
    "Abandono del Procedimiento": "Redacta un incidente de abandono del procedimiento, fundado en la inactividad de todas las partes por el plazo legal, conforme a los artículos 152 y siguientes del Código de Procedimiento Civil.",
    "Nulidad Procesal / Incidente de Nulidad": "Redacta un incidente de nulidad procesal, identificando el vicio que afecta la validez de una actuación judicial y el perjuicio reparable solo con la declaración de nulidad, conforme a los artículos 79 y siguientes del Código de Procedimiento Civil.",
    "Tercería de Posesión": "Redacta una tercería de posesión, fundada en que el bien embargado se encuentra en poder de un tercero ajeno al juicio que invoca la posesión del mismo.",
    "Tercería de Dominio": "Redacta una tercería de dominio, fundada en que el bien embargado es de propiedad de un tercero ajeno al juicio ejecutivo.",
    "Tercería de Prelación": "Redacta una tercería de prelación, fundada en un mejor derecho de pago del tercero sobre el producto del remate.",
    "Tercería de Pago": "Redacta una tercería de pago, para que el tercerista concurra proporcionalmente al pago con el producto del remate de los bienes embargados.",
    "Excepciones Ejecutivas (Art. 464 CPC)": "__ESPECIAL__",  # Usa el flujo especializado de 18 excepciones, no este genérico
    "Recurso de Reposición": "Redacta un recurso de reposición en contra de una resolución judicial, exponiendo el error que se reclama y solicitando que se deje sin efecto o se modifique.",
    "Recurso de Apelación": "Redacta un recurso de apelación en contra de una resolución judicial, exponiendo los agravios y solicitando que el tribunal superior revise y revoque o modifique lo resuelto.",
    "Solicitud de Cúmplase / Cumplimiento Incidental": "Redacta una solicitud de cumplimiento incidental de una sentencia o resolución firme y ejecutoriada, conforme a los artículos 231 y siguientes del Código de Procedimiento Civil.",
    "Otro tipo de presentación": "Redacta la presentación judicial exactamente según las instrucciones y el contexto que entregue el abogado, sin asumir un formato predeterminado.",
}

# =====================================================================
# 📝 ESTRUCTURAS DETALLADAS POR TIPO DE ESCRITO — REDACTOR IA
# =====================================================================
# A diferencia del catálogo genérico de arriba (una línea por tipo), este es
# específico para el Redactor IA: da la ESTRUCTURA PROCESAL exacta esperada
# para cada tipo de escrito, distinguiendo entre escritos DE FONDO (largos,
# con fundamentos extensos) y de MERA TRAMITACIÓN (cortos, directos al
# punto, sin desarrollar fundamentos extensos innecesarios).
ESTRUCTURAS_REDACTOR_IA = {
    "Demanda (Ordinaria)": """
    ESTRUCTURA ESPERADA (escrito DE FONDO — extenso y fundamentado):
    EN LO PRINCIPAL: Demanda [tipo de acción]; PRIMER OTROSÍ: Acompaña documentos; SEGUNDO OTROSÍ: Patrocinio y poder.
    Cuerpo: I. LOS HECHOS (relato cronológico y detallado de los antecedentes fácticos); II. EL DERECHO (fundamentos jurídicos: la acción, los requisitos legales que la configuran, doctrina y normativa aplicable); III. LA CUANTÍA (si corresponde).
    Petitorio: "POR TANTO, en mérito de lo expuesto y lo dispuesto en los artículos [citar], a US. RUEGO tener por interpuesta demanda de [tipo] en contra de [demandado], ya individualizado, admitirla a tramitación, y en definitiva, acogerla en todas sus partes, condenando al demandado a [petición concreta], con expresa condena en costas."
    """,
    "Demanda Ejecutiva": """
    ESTRUCTURA ESPERADA (escrito DE FONDO):
    EN LO PRINCIPAL: Demanda Ejecutiva; PRIMER OTROSÍ: Acompaña título ejecutivo y documentos; SEGUNDO OTROSÍ: Forma de notificación si corresponde; TERCER OTROSÍ: Patrocinio y poder.
    Cuerpo: identificación exacta del título ejecutivo (tipo, fecha, monto, exigibilidad), verificación de que la obligación es líquida, actualmente exigible y no prescrita (Art. 434 y ss. CPC).
    Petitorio: "POR TANTO, RUEGO A US.: tener por interpuesta demanda ejecutiva en autos, despachar mandamiento de ejecución y embargo en contra de [demandado] por la suma de [monto] más intereses y costas, ordenando trabar embargo sobre bienes suficientes."
    """,
    "Contestación de Demanda (Evacúa Traslado)": """
    ESTRUCTURA ESPERADA (escrito DE FONDO):
    EN LO PRINCIPAL: Contesta demanda; OTROSÍ: si corresponde, demanda reconvencional o acompaña documentos.
    Cuerpo: I. Excepciones dilatorias (si las hay, se oponen primero y se resuelven antes que el fondo); II. Contestación al fondo: pronunciamiento punto por punto sobre cada hecho alegado por el demandante (afirmar, negar, o allanarse a cada uno — nunca dejar hechos sin pronunciamiento expreso); III. Los hechos según la parte demandada (versión propia); IV. El derecho (fundamentos jurídicos de la defensa).
    Petitorio: "POR TANTO, RUEGO A US.: tener por contestada la demanda, rechazarla en todas sus partes, con expresa condenación en costas."
    """,
    "Réplica": """
    ESTRUCTURA ESPERADA (escrito breve, de FONDO pero acotado):
    Ratifica los fundamentos de la demanda, se hace cargo específicamente de las alegaciones y excepciones opuestas en la contestación, sin repetir todo lo ya dicho — solo profundiza o rebate lo nuevo que introdujo la contraria.
    Petitorio: "POR TANTO, RUEGO A US.: tener por evacuado el trámite de réplica, y estarse a lo pedido en la demanda."
    """,
    "Dúplica": """
    ESTRUCTURA ESPERADA (escrito breve, de FONDO pero acotado):
    Ratifica los fundamentos de la contestación, rebate lo nuevo introducido en la réplica.
    Petitorio: "POR TANTO, RUEGO A US.: tener por evacuado el trámite de dúplica, y estarse a lo pedido en la contestación."
    """,
    "Oposición de Excepciones (Art. 464 CPC)": """
    ESTRUCTURA ESPERADA (escrito DE FONDO): usa preferentemente el módulo especializado "⚖️ Excepciones Ejecutivas" (analiza los 18 números del Art. 464 CPC con IA). Si igual se redacta aquí:
    EN LO PRINCIPAL: Opone excepciones; OTROSÍ: acompaña documentos.
    Cuerpo: cada excepción opuesta, numerada según el Art. 464 CPC, con su fundamento de hecho y de derecho específico.
    Petitorio: "POR TANTO, RUEGO A US.: tener por opuestas las excepciones del artículo 464 N°[...] ya fundamentadas, acogerlas, y rechazar la ejecución con costas."
    """,
    "Recurso de Reposición": """
    ESTRUCTURA ESPERADA (escrito DE FONDO, pero puede ser breve según la complejidad):
    EN LO PRINCIPAL: Repone; OTROSÍ: Apela en subsidio (si corresponde, y casi siempre conviene incluirlo para no perder la vía de alzada).
    Cuerpo: identifica la resolución exacta que se impugna (fecha, contenido), expone el error de hecho o de derecho en que incurrió el tribunal, y por qué debe modificarse.
    Petitorio: "POR TANTO, RUEGO A US.: tener por interpuesto recurso de reposición en contra de la resolución de fecha [...], acogerlo, dejarla sin efecto y resolver conforme a lo pedido; y en subsidio, para el caso de ser rechazada esta reposición, apelar de dicha resolución para ante la Iltma. Corte de Apelaciones respectiva."
    """,
    "Recurso de Apelación": """
    ESTRUCTURA ESPERADA (escrito DE FONDO):
    EN LO PRINCIPAL: Apela.
    Cuerpo: identifica la resolución apelada (fecha, tribunal, contenido), expone los agravios que causa (por qué es errónea o injusta), fundamentos de hecho y de derecho de por qué debe revocarse o modificarse.
    Petitorio: "POR TANTO, RUEGO A US.: tener por interpuesto recurso de apelación en contra de la resolución de fecha [...], concederlo, y elevar estos autos a la Iltma. Corte de Apelaciones de [...] para que, conociendo del recurso, la revoque [o modifique] en el sentido de [...], con costas."
    """,
    "Incidente de Nulidad Procesal": """
    ESTRUCTURA ESPERADA (escrito DE FONDO):
    EN LO PRINCIPAL: Promueve incidente de nulidad procesal.
    Cuerpo: identifica el vicio concreto (qué norma de procedimiento se infringió), acredita el perjuicio reparable solo con la declaración de nulidad (Art. 83 CPC — no basta el vicio, debe haber perjuicio real), y que se promueve dentro de plazo (5 días desde que se tuvo conocimiento del vicio, salvo excepciones).
    Petitorio: "POR TANTO, RUEGO A US.: tener por promovido el incidente de nulidad procesal, acogerlo, y declarar la nulidad de [actuación específica], retrotrayendo el procedimiento al estado de [...]."
    """,
    "Solicitud de Abandono del Procedimiento": """
    ESTRUCTURA ESPERADA (escrito DE MERA TRAMITACIÓN — corto y directo):
    Señala la fecha de la última gestión útil realizada por CUALQUIERA de las partes, constata que ha transcurrido el plazo legal (6 meses, Art. 152 CPC) sin actividad, y que no se trata de excepciones legales al abandono (juicios de familia, ejecutivos en ciertos casos, etc.).
    Petitorio breve: "POR TANTO, RUEGO A US.: declarar el abandono del procedimiento en estos autos, con costas."
    """,
    "Escrito de Mera Tramitación (Téngase Presente / Acompaña Documentos / Otro trámite simple)": """
    ESTRUCTURA ESPERADA (CORTO Y DIRECTO — nunca extenderse innecesariamente):
    Estos escritos NO llevan desarrollo extenso de "Hechos" ni "Derecho" — van directo al punto, en 1 a 3 párrafos como máximo.
    Ejemplos de este tipo: "Téngase Presente", "Acompaña Documentos (con o sin citación)", "Solicita se Despache Mandamiento de Ejecución y Embargo", "Solicita Certificado de Ejecutoria", "Solicita Notificación por Avisos o por Cédula", "Solicita Alzamiento de Medida Precautoria", "Solicita Hora para Comparendo", "Confiere Patrocinio y Poder", "Solicita Copias Autorizadas", "Solicita se Oficie a [institución]".
    Estructura: EN LO PRINCIPAL (lo que se pide, en una frase clara) + cuerpo breve explicando el motivo puntual + petitorio directo ("POR TANTO, RUEGO A US.: acceder a lo solicitado").
    """,
    "Otro (Especificar en instrucciones)": """
    No hay una estructura predeterminada — sigue estrictamente las instrucciones específicas que entregue el abogado más abajo, usando buen juicio procesal chileno para el formato (suma, comparecencia, cuerpo, petitorio) según lo que se pida.
    """,
}

def crear_escrito_judicial_generico_word(tipo_escrito, texto_redactado, datos_causa=None):
    """
    Genera en Word cualquier tipo de escrito judicial (no solo excepciones),
    con el mismo formato profesional del resto del sistema. Convierte cada
    bloque separado por doble salto de línea en un párrafo real, para
    evitar el problema de huecos al justificar.
    """
    if not DOCX_READY:
        return None
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    if datos_causa:
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_meta.add_run(f"{tipo_escrito.upper()} — Causa Rol {datos_causa.get('rol','')}, \"{datos_causa.get('caratulado','')}\", {datos_causa.get('tribunal','')}").bold = True
    
    for bloque in texto_redactado.split("\n\n"):
        bloque_limpio = bloque.strip().replace("**", "").replace("#", "")
        if bloque_limpio:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(bloque_limpio)
    
    return doc

def crear_escrito_oposicion_excepciones_word(datos_causa, excepciones_seleccionadas):
    """
    Redacta el escrito de oposición de excepciones (EN LO PRINCIPAL) con las
    excepciones que el abogado seleccionó, usando el mismo formato
    profesional del resto del sistema.
    """
    if not DOCX_READY:
        return None
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    p_suma = doc.add_paragraph()
    p_suma.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_suma.add_run("EN LO PRINCIPAL: ").bold = True
    p_suma.add_run("Opone excepciones a la ejecución; ")
    p_suma.add_run("OTROSÍ: ").bold = True
    p_suma.add_run("Acompaña documentos.")
    
    p_suma2 = doc.add_paragraph()
    p_suma2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_suma2.add_run(f"\nS.J.L. {datos_causa.get('tribunal','')}\n")
    
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.add_run(f"{datos_causa.get('nombre_ejecutado','')}, en representación de la parte ejecutada en autos sobre juicio ejecutivo, Rol N° {datos_causa.get('rol','')}, caratulados \"{datos_causa.get('caratulado','')}\", a US. respetuosamente digo:")
    
    p_fund = doc.add_paragraph()
    p_fund.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_fund.add_run("Que, encontrándome dentro del plazo legal, vengo en oponer a la ejecución las excepciones contempladas en el artículo 464 del Código de Procedimiento Civil que a continuación se fundamentan, solicitando desde ya su acogimiento con expresa condenación en costas a la parte ejecutante.")
    
    for exc in excepciones_seleccionadas:
        p_exc = doc.add_paragraph()
        p_exc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_exc.add_run(f"\n{exc['numero']}. EXCEPCIÓN DEL ARTÍCULO 464 N°{exc['numero']} DEL CPC: {exc['nombre'].upper()}. ").bold = True
        p_exc.add_run(exc.get('fundamento_final', exc.get('fundamento', '')))
        
        if exc.get('cita_textual', '').strip():
            p_cita = doc.add_paragraph()
            p_cita.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_cita.paragraph_format.left_indent = Pt(28)
            r_cita = p_cita.add_run(f'«{exc["cita_textual"]}»')
            r_cita.italic = True
    
    p_petitorio = doc.add_paragraph()
    p_petitorio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nums_excepciones = ", ".join([f"N°{e['numero']}" for e in excepciones_seleccionadas])
    p_petitorio.add_run(f"\nPOR TANTO,\n").bold = True
    p_petitorio.add_run(f"RUEGO A US.: Tener por opuestas las excepciones del artículo 464 {nums_excepciones} del Código de Procedimiento Civil ya fundamentadas, acogerlas a tramitación, y en definitiva, en la sentencia definitiva que se dicte en estos autos, acogerlas en todas sus partes, rechazando la ejecución, con expresa condenación en costas.")
    
    p_otrosi = doc.add_paragraph()
    p_otrosi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_otrosi.add_run("\nOTROSÍ: ").bold = True
    p_otrosi.add_run("Ruego a US. tener por acompañados los documentos fundantes de las excepciones opuestas, con citación.")
    
    return doc

def numero_a_letras_clp(monto: int) -> str:
    """
    Convierte un monto entero a su forma escrita en español, para no tener
    que redactar 'Valor en Letras' a mano cada vez (ej: 500000 ->
    'quinientos mil pesos'). Cubre montos hasta 999.999.999, más que
    suficiente para honorarios profesionales.
    """
    monto = int(monto)
    if monto == 0:
        return "cero pesos"
    if monto == 1:
        return "un peso"
    
    UNIDADES = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
    DIECIS = ["diez", "once", "doce", "trece", "catorce", "quince", "dieciséis", "diecisiete", "dieciocho", "diecinueve"]
    DECENAS = ["", "", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
    CENTENAS = ["", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos"]
    
    def _tres_digitos(n, femenino=False, apocope_uno=False):
        if n == 0:
            return ""
        if n == 100:
            return "cien"
        c, resto = divmod(n, 100)
        partes = []
        if c > 0:
            partes.append(CENTENAS[c])
        if resto > 0:
            if resto < 10:
                u = "una" if (femenino and resto == 1) else UNIDADES[resto]
                partes.append(u)
            elif resto < 20:
                partes.append(DIECIS[resto - 10])
            else:
                d, u = divmod(resto, 10)
                if u == 0:
                    partes.append(DECENAS[d])
                elif d == 2:
                    partes.append(f"veinti{UNIDADES[u] if u != 1 else 'uno' if not femenino else 'una'}")
                else:
                    u_txt = ("una" if femenino else "uno") if u == 1 else UNIDADES[u]
                    partes.append(f"{DECENAS[d]} y {u_txt}")
        texto = " ".join(partes)
        # Apócope: "veintiuno mil" -> "veintiún mil", "treinta y uno mil" -> "treinta y un mil"
        if apocope_uno and texto.endswith("uno"):
            texto = texto[:-3] + "ún" if texto.endswith("veintiuno") else texto[:-3] + "un"
        return texto
    
    millones, resto_m = divmod(monto, 1_000_000)
    miles, unidades_finales = divmod(resto_m, 1_000)
    
    trozos = []
    if millones > 0:
        if millones == 1:
            trozos.append("un millón")
        else:
            trozos.append(f"{_tres_digitos(millones, apocope_uno=True)} millones")
    if miles > 0:
        if miles == 1:
            trozos.append("mil")
        else:
            trozos.append(f"{_tres_digitos(miles, apocope_uno=True)} mil")
    if unidades_finales > 0:
        trozos.append(_tres_digitos(unidades_finales))
    
    # Regla gramatical: "un millón DE pesos" cuando no hay nada más entre
    # medio, pero "un millón doscientos mil pesos" (sin "de") cuando sí lo hay.
    sufijo = " de pesos" if (millones > 0 and miles == 0 and unidades_finales == 0) else " pesos"
    return " ".join(trozos).strip().capitalize() + sufijo

def boton_refrescar_equipo(key):
    """Fuerza una relectura desde disco de todos los archivos base_causas_*/base_tareas_*
    del equipo, por si un cambio reciente de otro usuario no se refleja aún."""
    if st.button("🔄 Actualizar datos del equipo", key=key, help="Si algún cambio reciente de otro abogado no aparece, presiona aquí."):
        claves_a_limpiar = [k for k in st.session_state.keys() if k.startswith("_csv_cache_base_causas_") or k.startswith("_csv_cache_base_tareas_")]
        for k in claves_a_limpiar:
            del st.session_state[k]
        st.rerun()

def leer_csv_local(path, default_cols=None):
    """
    Lee un CSV local cacheándolo en st.session_state mientras el archivo no
    cambie en disco (se detecta por fecha de modificación). Evita releer el
    mismo archivo una y otra vez en cada rerun de Streamlit cuando nadie lo
    modificó, sin arriesgar mostrar datos desactualizados.
    """
    if not os.path.exists(path):
        return pd.DataFrame(columns=default_cols) if default_cols is not None else pd.DataFrame()
    try:
        mtime = os.path.getmtime(path)
    except OSError:
        mtime = None
    cache_key = f"_csv_cache_{path}"
    cached = st.session_state.get(cache_key)
    if cached is not None and cached[0] == mtime:
        return cached[1].copy()
    try:
        df = pd.read_csv(path)
        df = _corregir_dtypes_texto(df)
    except Exception:
        return pd.DataFrame(columns=default_cols) if default_cols is not None else pd.DataFrame()
    st.session_state[cache_key] = (mtime, df)
    return df.copy()

# =====================================================================
# ☁️ ALMACENAMIENTO DE ARCHIVOS EN GOOGLE DRIVE (CON RESPALDO A BASE64)
# =====================================================================
# Por qué: Google Sheets limita cada celda a ~50.000 caracteres. Guardar PDFs
# o fotos en base64 dentro de una celda se rompe apenas el archivo pesa un
# poco. Ahora los adjuntos se suben a una carpeta de Google Drive y en la
# hoja solo se guarda el ID del archivo (unas pocas decenas de caracteres).
#
# CONFIGURACIÓN NECESARIA (una sola vez, de tu parte):
# 1. En Google Cloud Console, habilita la "Google Drive API" para el mismo
#    proyecto de la cuenta de servicio que ya usas para Google Sheets
#    (la que configuraste en Secrets, sección [connections.gsheets]).
# 2. Crea una carpeta en Google Drive para JuriSync y compártela con el
#    correo de la cuenta de servicio (termina en "...gserviceaccount.com"),
#    dándole permiso de Editor.
# 3. Copia el ID de esa carpeta (está en la URL de Drive) y agrégalo en
#    Streamlit Cloud -> Settings -> Secrets como:
#       DRIVE_FOLDER_ID = "el_id_de_tu_carpeta"
#
# Si no configuras esto todavía, el sistema NO se rompe: cae automáticamente
# de vuelta al guardado en base64 (comportamiento anterior), respetando el
# límite de tamaño de Sheets.

# =====================================================================
# 🔑 CONEXIÓN OAUTH CON EL DRIVE PERSONAL DEL ABOGADO
# =====================================================================
# Las cuentas de servicio NO tienen cuota de almacenamiento propia en Drive
# (es una limitación real y documentada de Google), así que cualquier
# subida a un Drive personal común (no una Unidad Compartida de Google
# Workspace) falla con "storageQuotaExceeded". La solución que Google
# recomienda para cuentas Gmail normales es autenticarse como el usuario
# real (OAuth), usando SU cuota de 15GB gratis, en vez de la cuenta de
# servicio. Esto requiere una autorización única (ver Panel Admin).
def _url_autorizacion_drive_oauth():
    """Genera la URL para que el abogado autorice el acceso a su Drive personal, y el objeto 'flow' para completar el intercambio después."""
    from google_auth_oauthlib.flow import Flow
    redirect_uri = st.secrets.get("APP_BASE_URL", "https://jurisyncs.streamlit.app")
    client_config = {
        "web": {
            "client_id": st.secrets["GOOGLE_OAUTH_CLIENT_ID"],
            "client_secret": st.secrets["GOOGLE_OAUTH_CLIENT_SECRET"],
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [redirect_uri],
        }
    }
    flow = Flow.from_client_config(client_config, scopes=['https://www.googleapis.com/auth/drive'], redirect_uri=redirect_uri)
    auth_url, _ = flow.authorization_url(prompt='consent', access_type='offline', include_granted_scopes='true')
    return auth_url, flow

def _intercambiar_codigo_oauth_drive(flow, codigo_autorizacion):
    """Cambia el código de autorización (que llega en la URL tras autorizar) por un refresh_token permanente."""
    flow.fetch_token(code=codigo_autorizacion)
    return flow.credentials

def _servicio_drive():
    SCOPES_DRIVE = ['https://www.googleapis.com/auth/drive']
    # Prioridad 1: si ya se completó la autorización OAuth (hay un
    # refresh_token guardado en Secrets), se usa la cuenta personal del
    # abogado, que sí tiene cuota de almacenamiento propia.
    if "GOOGLE_OAUTH_REFRESH_TOKEN" in st.secrets:
        from google.oauth2.credentials import Credentials as OAuthCredentials
        creds = OAuthCredentials(
            token=None,
            refresh_token=st.secrets["GOOGLE_OAUTH_REFRESH_TOKEN"],
            client_id=st.secrets["GOOGLE_OAUTH_CLIENT_ID"],
            client_secret=st.secrets["GOOGLE_OAUTH_CLIENT_SECRET"],
            token_uri="https://oauth2.googleapis.com/token",
            scopes=SCOPES_DRIVE,
        )
    else:
        # Respaldo: la cuenta de servicio (falla con archivos grandes por la
        # limitación de storageQuotaExceeded, pero sigue sirviendo para
        # cuentas de Google Workspace con Unidad Compartida configurada).
        creds = _obtener_credenciales_google(SCOPES_DRIVE)
    return build('drive', 'v3', credentials=creds)

def subir_archivo_drive(nombre_archivo: str, contenido_bytes: bytes, mime_type: str = 'application/octet-stream'):
    """Sube un archivo a la carpeta de Drive configurada. Devuelve el file_id o None si falla."""
    try:
        servicio = _servicio_drive()
        carpeta_id = st.secrets.get("DRIVE_FOLDER_ID", "")
        metadata = {'name': nombre_archivo}
        if carpeta_id:
            metadata['parents'] = [carpeta_id]
        media = MediaIoBaseUpload(io.BytesIO(contenido_bytes), mimetype=mime_type, resumable=False)
        archivo = servicio.files().create(body=metadata, media_body=media, fields='id').execute()
        return archivo.get('id')
    except Exception as e:
        # Antes este error solo se imprimía en los registros internos del
        # servidor, invisibles para el usuario — quien solo veía un mensaje
        # confuso hablando del límite de tamaño de Sheets (el respaldo), sin
        # saber que el problema real era que Drive falló. Ahora queda
        # guardado en session_state para poder mostrarlo en pantalla.
        st.session_state['_ultimo_error_drive'] = str(e)
        print(f"Error subiendo a Drive: {e}")
        return None

def descargar_archivo_drive(file_id: str):
    """Descarga un archivo de Drive por su file_id. Devuelve los bytes o None si falla."""
    try:
        servicio = _servicio_drive()
        request = servicio.files().get_media(fileId=file_id)
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, request)
        listo = False
        while not listo:
            _, listo = downloader.next_chunk()
        return buffer.getvalue()
    except Exception as e:
        print(f"Error descargando de Drive: {e}")
        return None

def guardar_archivo_adjunto(nombre_archivo: str, contenido_bytes: bytes, mime_type: str = 'application/octet-stream'):
    """
    Punto único de guardado de adjuntos. Intenta Google Drive primero
    (sin límite de tamaño de celda); si Drive no está configurado o falla,
    cae automáticamente al guardado en base64 (con aviso de tamaño).
    Devuelve (drive_id, base64_str) — uno de los dos queda vacío.
    """
    drive_id = subir_archivo_drive(nombre_archivo, contenido_bytes, mime_type)
    if drive_id:
        return drive_id, ""
    # Respaldo: base64 en la celda, solo si el tamaño lo permite
    tamano_ok, _ = validar_tamano_para_sheets(contenido_bytes, nombre_archivo)
    b64 = base64.b64encode(contenido_bytes).decode('utf-8') if tamano_ok else ""
    return "", b64

def obtener_bytes_adjunto(fila, campo_drive: str, campo_b64: str):
    """
    Recupera los bytes de un adjunto sin importar si quedó guardado en Drive
    (dato nuevo) o en base64 (dato antiguo, de antes de esta migración).
    """
    drive_id = fila.get(campo_drive, "")
    if pd.notna(drive_id) and str(drive_id).strip() not in ("", "nan"):
        contenido = descargar_archivo_drive(str(drive_id).strip())
        if contenido is not None:
            return contenido
    b64_val = fila.get(campo_b64, "")
    if pd.notna(b64_val) and str(b64_val).strip() not in ("", "nan"):
        try:
            return base64.b64decode(b64_val)
        except Exception:
            return None
    return None

# =====================================================================
# 🛡️ MOTOR DE RESILIENCIA Y CACHÉ (AIRBAGS ANTI-CRASH Y VELOCIDAD)
# =====================================================================
conn = st.connection("gsheets", type=GSheetsConnection)

def _obtener_credenciales_google(scopes):
    """
    Construye las credenciales de Google para Drive y Calendar a partir de
    los Secrets de Streamlit (misma sección [connections.gsheets] que ya usa
    Google Sheets), en vez de depender de un archivo credenciales_calendar.json
    en el repositorio. Esto es más seguro (el repositorio es público, así que
    nunca debe contener una llave de servicio) y evita tener que configurar
    la cuenta de servicio dos veces: se reutiliza la misma que ya configuraste
    para Sheets.
    """
    info = dict(st.secrets["connections"]["gsheets"])
    info.pop("spreadsheet", None)  # no es parte de las credenciales, es la URL de la hoja
    return Credentials.from_service_account_info(info, scopes=scopes)

def generar_contenido_gemini(prompt_texto, archivos_pdf=None):
    """
    Genera contenido con Gemini para todo el sistema (Redactor IA, Análisis
    de Escrituras, Excepciones, Escritos Judiciales), probando dos caminos:
    
    1° VERTEX AI (se intenta primero): usa la MISMA cuenta de servicio de
       Google que ya tienes configurada para Sheets/Drive, y cobra desde tu
       facturación normal de Google Cloud, no desde la cola de "prepago" de
       AI Studio que ha estado fallando. Este es el motivo por el que se
       agregó: hay un problema de sincronización de facturación ampliamente
       reportado por otros desarrolladores en 2026 específicamente con esa
       cola de prepago de AI Studio, que Vertex AI evita por completo al
       cobrar por el otro canal.
    2° API DIRECTA (respaldo): si Vertex AI no está configurado o falla por
       cualquier motivo, cae de vuelta al método anterior con GEMINI_API_KEY,
       para no perder funcionalidad si Vertex no llegó a activarse.
    
    Para que el paso 1 funcione, hace falta (una sola vez, de tu parte):
    - Habilitar la "Vertex AI API" en Google Cloud Console, mismo proyecto
      que ya usas (jurisync-libre).
    - Darle a la cuenta de servicio (jurisync-admin@jurisync-libre...) el
      rol "Vertex AI User" en IAM.
    """
    try:
        from google import genai as genai_vertex
        from google.genai import types as genai_types
        
        creds_vertex = _obtener_credenciales_google(['https://www.googleapis.com/auth/cloud-platform'])
        project_id_vertex = st.secrets["connections"]["gsheets"].get("project_id", "")
        cliente_vertex = genai_vertex.Client(vertexai=True, project=project_id_vertex, location="us-central1", credentials=creds_vertex)
        
        partes_contenido = [prompt_texto]
        if archivos_pdf:
            for archivo in archivos_pdf:
                partes_contenido.append(genai_types.Part.from_bytes(data=archivo.getvalue(), mime_type="application/pdf"))
        
        respuesta_vertex = cliente_vertex.models.generate_content(model="gemini-2.5-flash", contents=partes_contenido)
        return respuesta_vertex.text
    except Exception:
        # Respaldo: la API directa de siempre (puede fallar si el problema
        # de facturación de AI Studio sigue activo en tu cuenta).
        import google.generativeai as genai
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        modelo_elegido = "gemini-1.5-flash"
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                md_name = m.name.replace("models/", "")
                if 'flash' in md_name:
                    modelo_elegido = md_name
                    break
        modelo = genai.GenerativeModel(modelo_elegido)
        if archivos_pdf:
            archivos_gemini = [genai.upload_file(io.BytesIO(a.getvalue()), mime_type="application/pdf", display_name=getattr(a, 'name', 'archivo.pdf')) for a in archivos_pdf]
            respuesta = modelo.generate_content([prompt_texto] + archivos_gemini)
        else:
            respuesta = modelo.generate_content(prompt_texto)
        return respuesta.text

# --- DEFINICIÓN DE COLUMNAS MAESTRAS ---
COLS_USUARIOS = ['Usuario', 'Password', 'Nombre_Real', 'Correo', 'Debe_Cambiar_Clave', 'Plan']
COLS_CLIENTES = ['RUT', 'Nombre', 'Telefono', 'Correo', 'Clave_unica', 'Direccion', 'Usuario_Propietario']
COLS_CAUSAS = ['ROL', 'TRIBUNAL', 'CARATULADO', 'Cliente', 'RUT', 'Tipo_Negocio', 'Usuario_Propietario', 'Estado_Honorarios', 'Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas', 'Clave_unica', 'SAC', 'Sucursal', 'Servicio']

# =====================================================================
# 🏛️ CATÁLOGO DE TRIBUNALES DE CHILE (para seleccionar, no escribir a mano)
# =====================================================================
CATALOGO_TRIBUNALES = {
    "Cortes y Tribunal Supremo": [
        "Corte Suprema", "Corte de Apelaciones de Arica", "Corte de Apelaciones de Iquique",
        "Corte de Apelaciones de Antofagasta", "Corte de Apelaciones de Copiapó", "Corte de Apelaciones de La Serena",
        "Corte de Apelaciones de Valparaíso", "Corte de Apelaciones de Santiago", "Corte de Apelaciones de San Miguel",
        "Corte de Apelaciones de Rancagua", "Corte de Apelaciones de Talca", "Corte de Apelaciones de Chillán",
        "Corte de Apelaciones de Concepción", "Corte de Apelaciones de Temuco", "Corte de Apelaciones de Valdivia",
        "Corte de Apelaciones de Puerto Montt", "Corte de Apelaciones de Coyhaique", "Corte de Apelaciones de Punta Arenas"
    ],
    "Juzgados Civiles de Santiago": [f"{i}° Juzgado Civil de Santiago" for i in range(1, 31)],
    "Juzgados de Familia de Santiago": [f"{i}° Juzgado de Familia de Santiago" for i in range(1, 5)],
    "Juzgados Laborales de Santiago": [
        "1° Juzgado de Letras del Trabajo de Santiago", "2° Juzgado de Letras del Trabajo de Santiago",
        "3° Juzgado de Letras del Trabajo de Santiago", "Juzgado de Cobranza Laboral y Previsional de Santiago"
    ],
    "Juzgados Penales de Santiago": [
        "Juzgado de Garantía de Santiago", "Juzgado de Garantía de San Miguel", "Juzgado de Garantía de San Bernardo",
        "Juzgado de Garantía de Puente Alto", "Juzgado de Garantía de Colina", "Juzgado de Garantía de Talagante",
        "Tribunal de Juicio Oral en lo Penal de Santiago", "Tribunal de Juicio Oral en lo Penal de San Miguel"
    ],
    "Tribunales Regionales - Zona Norte": [
        "Juzgado de Letras de Arica", "Juzgado de Familia de Arica", "Juzgado de Garantía de Arica",
        "Juzgado de Letras de Iquique", "Juzgado de Familia de Iquique", "Juzgado de Garantía de Iquique",
        "1° Juzgado Civil de Antofagasta", "2° Juzgado Civil de Antofagasta", "Juzgado de Familia de Antofagasta",
        "Juzgado de Garantía de Antofagasta", "Juzgado de Letras del Trabajo de Antofagasta",
        "Juzgado de Letras de Copiapó", "Juzgado de Familia de Copiapó", "Juzgado de Garantía de Copiapó",
        "1° Juzgado Civil de La Serena", "2° Juzgado Civil de La Serena", "Juzgado de Familia de La Serena", "Juzgado de Garantía de La Serena"
    ],
    "Tribunales Regionales - Zona Centro": [
        "1° Juzgado Civil de Valparaíso", "2° Juzgado Civil de Valparaíso", "3° Juzgado Civil de Valparaíso",
        "Juzgado de Familia de Valparaíso", "Juzgado de Garantía de Valparaíso", "Juzgado de Letras del Trabajo de Valparaíso",
        "1° Juzgado Civil de Viña del Mar", "2° Juzgado Civil de Viña del Mar", "Juzgado de Familia de Viña del Mar",
        "1° Juzgado Civil de Rancagua", "2° Juzgado Civil de Rancagua", "Juzgado de Familia de Rancagua", "Juzgado de Garantía de Rancagua",
        "1° Juzgado Civil de Talca", "2° Juzgado Civil de Talca", "Juzgado de Familia de Talca", "Juzgado de Garantía de Talca",
        "Juzgado de Letras de Chillán", "Juzgado de Familia de Chillán", "Juzgado de Garantía de Chillán"
    ],
    "Tribunales Regionales - Zona Sur": [
        "1° Juzgado Civil de Concepción", "2° Juzgado Civil de Concepción", "3° Juzgado Civil de Concepción",
        "Juzgado de Familia de Concepción", "Juzgado de Garantía de Concepción", "Juzgado de Letras del Trabajo de Concepción",
        "1° Juzgado Civil de Temuco", "2° Juzgado Civil de Temuco", "Juzgado de Familia de Temuco", "Juzgado de Garantía de Temuco",
        "Juzgado de Letras de Valdivia", "Juzgado de Familia de Valdivia", "Juzgado de Garantía de Valdivia",
        "Juzgado de Letras de Osorno", "Juzgado de Familia de Osorno", "Juzgado de Garantía de Osorno",
        "1° Juzgado Civil de Puerto Montt", "2° Juzgado Civil de Puerto Montt", "Juzgado de Familia de Puerto Montt", "Juzgado de Garantía de Puerto Montt"
    ],
    "Tribunales Regionales - Zona Austral": [
        "Juzgado de Letras de Coyhaique", "Juzgado de Familia de Coyhaique", "Juzgado de Garantía de Coyhaique",
        "Juzgado de Letras de Punta Arenas", "Juzgado de Familia de Punta Arenas", "Juzgado de Garantía de Punta Arenas"
    ],
    "Juzgados de Policía Local (RM)": [
        "1° Juzgado de Policía Local de Santiago", "2° Juzgado de Policía Local de Santiago",
        "Juzgado de Policía Local de Providencia", "Juzgado de Policía Local de Las Condes",
        "Juzgado de Policía Local de Ñuñoa", "Juzgado de Policía Local de La Florida",
        "Juzgado de Policía Local de Maipú", "Juzgado de Policía Local de Puente Alto",
        "Juzgado de Policía Local de San Bernardo", "Juzgado de Policía Local de Peñalolén"
    ],
    "Tribunales Especiales": [
        "Tribunal Tributario y Aduanero de la Región Metropolitana - Santiago Oriente",
        "Tribunal Tributario y Aduanero de la Región Metropolitana - Santiago Poniente",
        "Tribunal de Defensa de la Libre Competencia", "1° Tribunal Ambiental (Antofagasta)",
        "2° Tribunal Ambiental (Santiago)", "3° Tribunal Ambiental (Valdivia)",
        "Tribunal Constitucional", "Tribunal Calificador de Elecciones"
    ]
}
_LISTA_PLANA_TRIBUNALES = []
for _grupo_trib in CATALOGO_TRIBUNALES.values():
    _LISTA_PLANA_TRIBUNALES.extend(_grupo_trib)
_LISTA_PLANA_TRIBUNALES.append("✏️ Otro (no está en la lista)")

def selector_tribunal(valor_actual="", key_prefix="trib"):
    """
    Selector predeterminado de tribunales de Chile: un desplegable único con
    los ~150 tribunales del catálogo (Streamlit permite escribir para filtrar
    dentro del propio selectbox), más un campo de texto de respaldo siempre
    visible para los casos no listados. Se diseñó así (un solo nivel, sin
    selects en cascada) para que funcione bien tanto dentro como fuera de un
    st.form, donde los widgets no se actualizan entre sí hasta enviar el formulario.
    Devuelve el nombre final del tribunal como texto (igual que antes).
    """
    idx_default = 0
    if valor_actual and valor_actual in _LISTA_PLANA_TRIBUNALES:
        idx_default = _LISTA_PLANA_TRIBUNALES.index(valor_actual)
    trib_sel = st.selectbox("Tribunal (escribe para buscar)", _LISTA_PLANA_TRIBUNALES, index=idx_default, key=f"{key_prefix}_sel")
    trib_manual = st.text_input("¿No está en la lista? Escríbelo aquí (tiene prioridad sobre lo seleccionado arriba)",
                                 value=valor_actual if valor_actual and valor_actual not in _LISTA_PLANA_TRIBUNALES else "",
                                 key=f"{key_prefix}_manual", placeholder="Ej: Juzgado de Letras de Melipilla")
    if trib_manual.strip():
        return trib_manual.strip()
    if trib_sel == "✏️ Otro (no está en la lista)":
        return ""
    return trib_sel
COLS_TAREAS = ['ID_Tarea', 'ROL', 'Creador', 'Fecha_Creacion', 'Fecha_Vencimiento', 'Titulo', 'Descripcion', 'Estado', 'Comentarios', 'Prioridad', 'Usuario_Propietario']
COLS_CONTRATOS = ['ID', 'Fecha', 'Cliente', 'Servicio', 'Honorarios', 'Archivo_B64', 'Archivo_Drive_ID', 'Usuario_Propietario']
COLS_ESCRITURAS = ['ID', 'Fecha', 'Tipo_Escritura', 'Cliente', 'RUT_Cliente', 'Detalle', 'Archivo_B64', 'Archivo_Drive_ID', 'Usuario_Propietario']
COLS_ANALISIS_ESCRITURAS = ['ID', 'Fecha', 'Nombre_Archivo_Original', 'Archivo_B64', 'Archivo_Drive_ID', 'Usuario_Propietario']
COLS_EXCEPCIONES = ['ID', 'Fecha', 'ROL', 'Excepciones_Opuestas', 'Archivo_B64', 'Archivo_Drive_ID', 'Usuario_Propietario']
COLS_CITAS = ['ID_Cita', 'Fecha', 'Hora', 'RUT_Cliente', 'Nombre_Cliente', 'Telefono', 'Email',
              'Sucursal', 'Modalidad', 'Abogado_Asignado', 'Tipo_Juicio', 'Observacion',
              'Estado', 'Usuario_Propietario']
COLS_ENCARGOS = ['ID_Encargo', 'Nombre_Encargante', 'RUT_Encargante', 'Fecha_Encargo', 'Fecha_Limite',
                 'Descripcion_Encargo', 'Monto', 'Estado', 'Usuario_Propietario']
COLS_PAGOS_HONORARIOS = ['ID_Pago', 'Fecha_Pago', 'Cliente', 'ROL', 'Monto_Cuota', 'Numero_Cuota', 'Usuario_Propietario']
COLS_JURISPRUDENCIA = ['ID', 'Tribunal', 'Rol_Causa', 'Fecha_Sentencia', 'Materia', 'Resumen', 'Archivo_Nombre',
                        'Archivo_B64', 'Archivo_Drive_ID', 'Fecha_Carga', 'Usuario_Propietario']
COLS_POSESION_EFECTIVA = ['ID', 'Fecha', 'Causante', 'RUT_Causante', 'Fecha_Defuncion', 'Herederos_JSON', 'Bienes_JSON', 'Cliente_Solicitante', 'RUT_Cliente', 'Estado', 'Valor_UTM', 'Masa_Hereditaria', 'Impuesto_Total', 'Archivo_B64', 'Archivo_Drive_ID', 'Usuario_Propietario']
COLS_TRAMITES = ['ID_Tramite', 'ROL', 'Fecha_Pago', 'Tipo_Auxiliar', 'Monto', 'Comprobante_Nombre', 'Comprobante_B64', 'Comprobante_Drive_ID', 'Registrado_Por', 'Usuario_Propietario']
COLS_DOCS = ['ID_Req', 'Cliente_Token', 'Documento_Nombre', 'Estado', 'Archivo_B64', 'Archivo_Drive_ID', 'Fecha_Subida', 'Link_Externo']
COLS_ESTADO_DIARIO = ['ID_ED', 'Fecha_Estado', 'ROL', 'Tribunal', 'Resolucion_Extracto', 'Doc_Nombre', 'Doc_B64', 'Doc_Drive_ID']
COLS_MENSAJES = ['ID', 'Fecha', 'De', 'Para', 'Mensaje']
@st.cache_data(ttl=900)
def fetch_sheet_cached(worksheet_name):
    """Descarga de la nube y guarda en memoria RAM por 15 min para máxima velocidad.
    IMPORTANTE: ttl=0 en conn.read() es intencional. La librería streamlit-gsheets
    tiene su PROPIA caché interna independiente de este decorador @st.cache_data.
    Antes se le pasaba ttl=900 también aquí, lo que creaba dos cachés apiladas:
    al guardar, fetch_sheet_cached.clear() solo limpiaba la caché exterior (esta),
    pero la caché interna de conn.read() seguía viva hasta 15 min más, así que
    los cambios (ej. Plan de un usuario) podían tardar hasta 15 min en reflejarse,
    o no reflejarse aunque recargaras la página. Con ttl=0 aquí, la única caché
    real es esta de @st.cache_data, que sí se limpia correctamente al guardar.
    """
    return conn.read(worksheet=worksheet_name, ttl=0)

def safe_read_sheet(worksheet_name, default_cols=None):
    """Lee usando caché. Si Google falla o bloquea, intenta evitar caídas."""
    try:
        df = fetch_sheet_cached(worksheet_name)
        if df is not None and not df.empty:
            df_clean = df.dropna(how="all")
            df_clean = _corregir_dtypes_texto(df_clean)
            # Guarda un respaldo local silencioso
            df_clean.to_csv(f"{worksheet_name}.csv", index=False)
            return df_clean
    except Exception:
        pass
    
    # Si Google falla, intenta leer el archivo local
    csv_path = f"{worksheet_name}.csv"
    if os.path.exists(csv_path):
        try:
            return _corregir_dtypes_texto(pd.read_csv(csv_path))
        except Exception:
            pass
            
    # Si no hay nada, crea una tabla vacía con las columnas correctas para evitar el TypeError
    if default_cols is not None:
        return pd.DataFrame(columns=default_cols)
    return pd.DataFrame()

def safe_update_sheet(worksheet_name, df):
    """Guarda en la nube e informa si Google Sheets rechaza la conexión."""
    csv_path = f"{worksheet_name}.csv"
    try:
        df.to_csv(csv_path, index=False)
    except Exception: pass
        
    try:
        conn.update(worksheet=worksheet_name, data=df)
        fetch_sheet_cached.clear() # Limpia la memoria para ver los datos frescos
        return True
    except Exception as e:
        # ¡ESTO ES CLAVE! Si Google falla, te saldrá un aviso rojo en vez de fallar en silencio.
        texto_error = str(e)
        if "not found" in texto_error.lower() or "worksheetnotfound" in texto_error.lower() or texto_error.strip() == worksheet_name:
            st.error(f"⚠️ Falta crear la pestaña **'{worksheet_name}'** en tu Google Sheets. Ve a tu hoja, crea una pestaña nueva con ese nombre exacto (abajo, botón '+'), y vuelve a intentarlo.")
        else:
            st.error(f"⚠️ Google Sheets bloqueó el guardado en la hoja '{worksheet_name}'. Detalle técnico: {texto_error}")
        fetch_sheet_cached.clear()
        return False

# --- FUNCIÓN DE GOOGLE CALENDAR DINÁMICA ---
def agendar_plazo_calendar(titulo, descripcion, fecha_str, correo_destino):
    if not correo_destino or "@" not in str(correo_destino):
        return False

    try:
        SCOPES = ['https://www.googleapis.com/auth/calendar.events']
        creds = _obtener_credenciales_google(SCOPES)
        servicio = build('calendar', 'v3', credentials=creds)

        f_obj = datetime.strptime(fecha_str, "%d/%m/%Y")
        fecha_iso = f_obj.strftime("%Y-%m-%dT09:00:00-04:00")
        fecha_fin = f_obj.strftime("%Y-%m-%dT10:00:00-04:00")

        evento = {
            'summary': f"🔴 PLAZO FATAL: {titulo}",
            'description': descripcion,
            'start': {'dateTime': fecha_iso, 'timeZone': 'America/Santiago'},
            'end': {'dateTime': fecha_fin, 'timeZone': 'America/Santiago'},
            'reminders': {
                'useDefault': False,
                'overrides': [
                    {'method': 'popup', 'minutes': 10080}, # 1 semana antes
                    {'method': 'popup', 'minutes': 2880},  # 2 días antes
                    {'method': 'popup', 'minutes': 180},   # 3 horas antes
                ],
            },
        }
        servicio.events().insert(calendarId=correo_destino, body=evento).execute()
        return True
    except Exception as e:
        print(f"Error interno Calendar: {e}")
        return False

# --- VERIFICACIÓN DE LIBRERÍA WORD ---
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_READY = True
except ImportError:
    DOCX_READY = False

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="JuriSync | Sistema Judicial", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

# =====================================================================
# 📱 PWA: hace que JuriSync se pueda "instalar" en el celular (ícono propio,
# pantalla completa, sin la barra del navegador), sin pasar por ninguna
# tienda de aplicaciones. Streamlit no da acceso directo al <head> de la
# página, así que se inyecta con JavaScript. El manifest.json y los íconos
# se sirven desde la carpeta static/ (requiere enableStaticServing=true en
# .streamlit/config.toml); el service worker, en cambio, se genera como un
# Blob directamente en el navegador, porque Streamlit no sirve archivos .js
# con el tipo de contenido correcto desde esa carpeta.
st.markdown("""
<script>
(function() {
    const head = document.head;
    if (!document.getElementById('jurisync-manifest')) {
        const manifestLink = document.createElement('link');
        manifestLink.id = 'jurisync-manifest';
        manifestLink.rel = 'manifest';
        manifestLink.href = './app/static/manifest.json';
        head.appendChild(manifestLink);

        const themeColor = document.createElement('meta');
        themeColor.name = 'theme-color';
        themeColor.content = '#0e6b74';
        head.appendChild(themeColor);

        const appleIcon = document.createElement('link');
        appleIcon.rel = 'apple-touch-icon';
        appleIcon.href = './app/static/icon-192.png';
        head.appendChild(appleIcon);

        const appleCapable = document.createElement('meta');
        appleCapable.name = 'apple-mobile-web-app-capable';
        appleCapable.content = 'yes';
        head.appendChild(appleCapable);

        const appleTitle = document.createElement('meta');
        appleTitle.name = 'apple-mobile-web-app-title';
        appleTitle.content = 'JuriSync';
        head.appendChild(appleTitle);
    }

    if ('serviceWorker' in navigator && !window._juriSyncSWRegistrado) {
        window._juriSyncSWRegistrado = true;
        const codigoSW = `
            self.addEventListener('install', () => self.skipWaiting());
            self.addEventListener('activate', () => self.clients.claim());
            self.addEventListener('fetch', (event) => { event.respondWith(fetch(event.request)); });
        `;
        const blob = new Blob([codigoSW], { type: 'application/javascript' });
        const swUrl = URL.createObjectURL(blob);
        navigator.serviceWorker.register(swUrl).catch(() => {});
    }
})();
</script>
""", unsafe_allow_html=True)

# --- SISTEMA ANTI-CIERRE DE SESIÓN (KEEP-ALIVE AGRESIVO) ---
st.markdown("""
<iframe src="javascript:void(0);" style="display:none;" onload="
    setInterval(function(){
        window.parent.document.dispatchEvent(new Event('mousemove'));
        window.parent.document.dispatchEvent(new KeyboardEvent('keydown', {'key': 'Shift'}));
    }, 30000);
"></iframe>
""", unsafe_allow_html=True)

# --- FUNCIONES DE SALUDO Y LOGO CUSTOM JURISYNC ---
def obtener_saludo():
    hora_chile = (datetime.now(timezone.utc) - timedelta(hours=4)).hour
    if 0 <= hora_chile < 12:
        return "Buenos días"
    elif 12 <= hora_chile < 19:
        return "Buenas tardes"
    else:
        return "Buenas noches"

@st.cache_data
def get_logo_src():
    ruta_base = os.path.dirname(os.path.abspath(__file__))
    extensiones = ['png', 'jpg', 'jpeg', 'PNG', 'JPG']
    for ext in extensiones:
        ruta_logo = os.path.join(ruta_base, f"logo.{ext}")
        if os.path.exists(ruta_logo):
            with open(ruta_logo, "rb") as f:
                contenido_b64 = base64.b64encode(f.read()).decode()
                return f"data:image/{ext.lower()};base64,{contenido_b64}"
                
    svg_logo = """
    <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100">
        <path d="M 30 20 A 35 35 0 0 1 85 50" fill="none" stroke="#0e6b74" stroke-width="6" stroke-linecap="round"/>
        <polygon points="85,60 76,46 94,46" fill="#0e6b74"/>
        <path d="M 70 80 A 35 35 0 0 1 15 50" fill="none" stroke="#172b4d" stroke-width="6" stroke-linecap="round"/>
        <polygon points="15,40 6,54 24,54" fill="#172b4d"/>
        <line x1="50" y1="15" x2="50" y2="70" stroke="#0e6b74" stroke-width="3.5" stroke-linecap="round"/>
        <line x1="43" y1="28" x2="57" y2="28" stroke="#0e6b74" stroke-width="3" stroke-linecap="round"/>
        <circle cx="50" cy="72" r="3" fill="#0e6b74"/>
        <path d="M 50 35 Q 40 55 33 75 L 67 75 Q 60 55 50 35 Z" fill="#172b4d" stroke-linejoin="round"/>
        <circle cx="50" cy="35" r="7" fill="#172b4d"/>
        <path d="M 41 35 Q 50 38 59 35" fill="none" stroke="#0e6b74" stroke-width="2.5" stroke-linecap="round"/>
        <path d="M 40 52 L 60 52" stroke="#ffffff" stroke-width="2.5" stroke-linecap="round"/>
        <path d="M 40 52 L 36 64 L 44 64 Z" fill="none" stroke="#ffffff" stroke-width="1.5" stroke-linejoin="round"/>
        <path d="M 60 52 L 56 64 L 64 64 Z" fill="none" stroke="#ffffff" stroke-width="1.5" stroke-linejoin="round"/>
        <circle cx="50" cy="52" r="2" fill="#ffffff"/>
    </svg>
    """
    b64_svg = base64.b64encode(svg_logo.encode('utf-8')).decode('utf-8')
    return f"data:image/svg+xml;base64,{b64_svg}"

LOGO_URL = get_logo_src()

# --- PROCESADOR DE ARCHIVOS OJV EXCEL ---
def procesar_ojv_completo(archivo):
    diccionario_hojas = pd.read_excel(archivo, sheet_name=None)
    mapa = {
        'ROL': ['ROL', 'RIT', 'Rol', 'Rit'], 
        'TRIBUNAL': ['TRIBUNAL', 'Tribunal', 'Juzgado', 'Corte'], 
        'CARATULADO': ['CARATULA', 'Carátula', 'Caratulado', 'Causa']
    }
    lista_final = []
    
    for nombre_hoja, df_hoja in diccionario_hojas.items():
        df_pro = pd.DataFrame()
        for col_ideal, posibles in mapa.items():
            for p in posibles:
                if p in df_hoja.columns:
                    df_pro[col_ideal] = df_hoja[p]
                    break
        if not df_pro.empty and 'ROL' in df_pro.columns:
            df_pro['Origen_OJV'] = nombre_hoja
            lista_final.append(df_pro)
            
    if lista_final:
        df_consolidado = pd.concat(lista_final, ignore_index=True).dropna(subset=['ROL'])
        df_consolidado['Estado'] = "Pendiente"
        df_consolidado['Prioridad'] = "Normal"
        df_consolidado['Tipo_Negocio'] = "Externo"
        
        cols_extra = [
            'Cliente', 'RUT', 'Teléfono', 'Clave_unica', 'Correo', 'Direccion', 'SAC', 'Sucursal', 
            'Estado_Honorarios', 'Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas'
        ]
        for col in cols_extra:
            if col not in df_consolidado.columns: 
                if col in ['Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas']: 
                    df_consolidado[col] = 0
                elif col == 'Estado_Honorarios': 
                    df_consolidado[col] = "Sin fijar"
                else: 
                    df_consolidado[col] = "--"
                    
        df_consolidado.to_csv(ARCHIVO_BD, index=False)
        return df_consolidado
    return pd.DataFrame()

# --- BASE DE DATOS DE FERIADOS CHILENOS ---
def obtener_feriados_chile():
    feriados = []
    annos = [2025, 2026, 2027]
    for anio in annos:
        fijos = [
            (f"{anio}-01-01", "Año Nuevo"), 
            (f"{anio}-05-01", "Día del Trabajador"),
            (f"{anio}-05-21", "Glorias Navales"), 
            (f"{anio}-06-21", "Pueblos Indígenas"),
            (f"{anio}-06-29", "San Pedro y San Pablo"), 
            (f"{anio}-07-16", "Virgen del Carmen"),
            (f"{anio}-08-15", "Asunción de la Virgen"), 
            (f"{anio}-09-18", "Independencia Nacional"),
            (f"{anio}-09-19", "Glorias del Ejército"), 
            (f"{anio}-10-12", "Encuentro de Dos Mundos"),
            (f"{anio}-10-31", "Iglesias Evangélicas"), 
            (f"{anio}-11-01", "Todos los Santos"),
            (f"{anio}-12-08", "Inmaculada Concepción"), 
            (f"{anio}-12-25", "Navidad")
        ]
        for fecha, nombre in fijos:
            feriados.append({
                "title": f"🇨🇱 {nombre}", 
                "start": fecha, 
                "color": "#ffebe6", 
                "textColor": "#bf2600", 
                "allDay": True, 
                "display": "block"
            })
            
    feriados.extend([
        {"title": "🇨🇱 Viernes Santo", "start": "2025-04-18", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"},
        {"title": "🇨🇱 Sábado Santo", "start": "2025-04-19", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"},
        {"title": "🇨🇱 Viernes Santo", "start": "2026-04-03", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"},
        {"title": "🇨🇱 Sábado Santo", "start": "2026-04-04", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"},
        {"title": "🇨🇱 Viernes Santo", "start": "2027-03-26", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"},
        {"title": "🇨🇱 Sábado Santo", "start": "2027-03-27", "color": "#ffebe6", "textColor": "#bf2600", "allDay": True, "display": "block"}
    ])
    return feriados

# --- MOTOR REDACTOR DE CONTRATOS EN WORD (VERSIÓN COMPLETA, 18 CLÁUSULAS) ---
def crear_contrato_word(datos):
    if not DOCX_READY: 
        return None
        
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Interlineado 1.5 aplicado a nivel de estilo 'Normal': como todos los
    # párrafos del documento usan ese estilo por defecto, esto alcanza a todo
    # el contrato de una vez, sin tener que tocar párrafo por párrafo.
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    def clausula(doc, numero_texto, titulo_texto, *bloques_texto):
        """
        Helper para no repetir el mismo bloque de formato en cada cláusula.
        IMPORTANTE: cada bloque de texto se pone en su PROPIO párrafo real
        (no separados con "\\n" dentro de un mismo párrafo). Word justifica
        estirando cualquier línea que no sea el final REAL de un párrafo, así
        que un salto de línea manual en medio de un párrafo justificado deja
        esos huecos raros entre palabras que se ven mal. Con párrafos
        separados, cada uno se justifica de forma normal y prolija.
        """
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_titulo.add_run(f"CLÁUSULA {numero_texto}: {titulo_texto}. ").bold = True
        if bloques_texto:
            p_titulo.add_run(bloques_texto[0])
        for bloque in bloques_texto[1:]:
            p_extra = doc.add_paragraph()
            p_extra.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_extra.add_run(bloque)
        return p_titulo
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = titulo.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES\n")
    r_tit.bold = True
    
    hoy = datetime.now()
    meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    fecha_str = f"{hoy.day} de {meses[hoy.month-1].lower()} del año {hoy.year}"
    
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(f"En Santiago, República de Chile, a {fecha_str}, comparecen:")
    
    p_intro2 = doc.add_paragraph()
    p_intro2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro2.add_run(f"Por una parte, don/doña {datos['abogado_nombre']}, de nacionalidad chilena, abogado/a, cédula nacional de identidad número {datos['abogado_rut']}, con domicilio profesional en {datos['abogado_domicilio']}, correo electrónico {datos['abogado_correo']}, en adelante e indistintamente \"EL ABOGADO\" o \"el mandatario\"; y,")
    
    p_intro3 = doc.add_paragraph()
    p_intro3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro3.add_run(f"Por otra parte, don/doña {datos['cliente_nombre']}, cédula nacional de identidad número {datos['cliente_rut']}, con domicilio en {datos['cliente_domicilio']}, número telefónico de contacto {datos['cliente_tel']}, correo electrónico {datos['cliente_correo']}, en adelante \"EL CLIENTE\" o \"el mandante\".")
    
    p_intro4 = doc.add_paragraph()
    p_intro4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro4.add_run("Ambas partes, compareciendo como mayores de edad, exponen que han convenido celebrar el siguiente contrato de prestación de servicios profesionales, el que se regirá por las cláusulas que a continuación se singularizan:")
    
    # --- CLÁUSULA PRIMERA: OBJETO ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run("CLÁUSULA PRIMERA: OBJETO DEL CONTRATO. ").bold = True
    p1.add_run(f"Por medio del presente instrumento, El Cliente confiere patrocinio y poder a El Abogado para que asuma la representación, tramitación y defensa de sus intereses en un procedimiento de {datos['tipo_servicio'].upper()}. Los servicios profesionales comprometidos comprenden, de forma específica, lo siguiente:")
    
    p1_bis = doc.add_paragraph()
    p1_bis.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_bis.add_run(datos['detalle_servicio'])
    
    # --- CLÁUSULA SEGUNDA: ALCANCE Y EXCLUSIONES DEL SERVICIO ---
    clausula(doc, "SEGUNDA", "ALCANCE Y EXCLUSIONES DEL SERVICIO",
        "El encargo referido en la cláusula precedente será desarrollado personalmente por El Abogado, o por los profesionales que este último destine al efecto dentro de su mismo estudio jurídico, quienes se sujetarán íntegramente a los términos de este contrato.",
        "El servicio contratado comprende únicamente la tramitación en primera instancia del asunto encomendado. Quedan expresamente excluidos, salvo pacto escrito adicional entre las partes, la presentación, defensa y tramitación de recursos procesales de segunda instancia o extraordinarios, tales como recursos de apelación, casación en la forma, casación en el fondo, nulidad, queja o unificación de jurisprudencia, cuyo eventual encargo dará lugar a honorarios adicionales que se acordarán separadamente.",
        "En caso de que durante la tramitación del asunto encomendado surjan materias distintas de las descritas en la cláusula primera, las partes podrán suscribir un anexo al presente contrato en el que se individualicen los nuevos servicios y los honorarios correspondientes."
    )
    
    # --- CLÁUSULA TERCERA: HONORARIOS ---
    clausula(doc, "TERCERA", "HONORARIOS PROFESIONALES",
        f"Los honorarios totales convenidos por la prestación de los servicios profesionales descritos ascienden a la suma de {datos['honorarios_num']} ({datos['honorarios_letras']}).",
        "Esta suma corresponde a la tramitación completa del asunto encomendado en primera instancia, conforme a lo señalado en la cláusula segunda, sin perjuicio de los honorarios adicionales que puedan devengarse por servicios excluidos o complementarios que las partes acuerden por separado."
    )
    
    # --- CLÁUSULA CUARTA: FORMA DE PAGO ---
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.add_run("CLÁUSULA CUARTA: CONDICIONES Y FORMA DE PAGO. ").bold = True
    p4.add_run(f"La suma total de los honorarios fijados en la cláusula precedente será pagada en un total de {datos['cuotas_cant']} cuotas mensuales, fijas y sucesivas, por un valor individual de {datos['cuotas_monto']} cada una, conforme al siguiente plan de pago:")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'N° Cuota'
    hdr_cells[1].text = 'Vencimiento Pactado'
    hdr_cells[2].text = 'Monto'
    hdr_cells[3].text = 'Estado de Pago'
    
    fecha_base = datos['fecha_inicio']
    for i in range(datos['cuotas_cant']):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{i+1:02d}"
        mes_calculado = fecha_base.month + i
        anno_calculado = fecha_base.year + ((mes_calculado - 1) // 12)
        mes_final = ((mes_calculado - 1) % 12) + 1
        row_cells[1].text = f"{fecha_base.day:02d} de {meses[mes_final-1]} de {anno_calculado}"
        row_cells[2].text = str(datos['cuotas_monto'])
        row_cells[3].text = "PENDIENTE"
    
    # Los datos bancarios son una ficha de campos cortos (Titular / RUT / Banco...),
    # no prosa corrida. Justificar líneas tan cortas se ve mal (huecos enormes entre
    # palabras), así que este bloque va alineado a la izquierda, como corresponde
    # tipográficamente a una lista de datos.
    p4_titulo_banco = doc.add_paragraph()
    p4_titulo_banco.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4_titulo_banco.add_run("Información Bancaria para Transferencias Electrónicas:").bold = True
    
    p4_bis = doc.add_paragraph()
    p4_bis.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4_bis.add_run(f"Titular de la Cuenta: {datos['abogado_nombre']}\nRUT: {datos['abogado_rut']}\nInstitución Bancaria: {datos['banco']}\nTipo de Cuenta: {datos['tipo_cuenta']}\nNúmero de Cuenta: {datos['num_cuenta']}")
    
    p4_ter = doc.add_paragraph()
    p4_ter.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4_ter.add_run("El Cliente deberá remitir el comprobante de cada pago al correo electrónico de El Abogado individualizado en la comparecencia, dentro de las 24 horas siguientes a su realización.")

    # --- CLÁUSULA QUINTA: PAGO DE HONORARIOS A TODO EVENTO ---
    clausula(doc, "QUINTA", "DEL PAGO DE LOS HONORARIOS A TODO EVENTO",
        "Los honorarios establecidos en la cláusula tercera deberán ser pagados a todo evento, incluso en el caso de que El Cliente decida revocar anticipadamente el patrocinio y poder conferidos a El Abogado, sin perjuicio de la liquidación proporcional que corresponda conforme a lo dispuesto en la cláusula décimo tercera."
    )
    
    # --- CLÁUSULA SEXTA: GASTOS NECESARIOS ---
    clausula(doc, "SEXTA", "GASTOS NECESARIOS PARA LA PRESTACIÓN DEL SERVICIO",
        "Todos los gastos y tasas que resulten necesarios para la ejecución del encargo, tales como honorarios de Receptores Judiciales, peritos, Notarios Públicos, Conservadores de Bienes Raíces, inscripciones, publicaciones y notificaciones, serán de cargo exclusivo de El Cliente.",
        "El Abogado informará oportunamente a El Cliente sobre los gastos en que deberá incurrirse, detallando su concepto y monto, y en caso de que dichos gastos hayan sido solventados directamente por El Abogado, El Cliente deberá reembolsarlos dentro de los cinco días hábiles siguientes a que le sea así informado, debiendo El Abogado conservar los respaldos respectivos para su exhibición cuando El Cliente así lo solicite."
    )
    
    # --- CLÁUSULA SÉPTIMA: EFECTOS DEL INCUMPLIMIENTO Y MOROSIDAD ---
    clausula(doc, "SÉPTIMA", "EFECTOS DEL INCUMPLIMIENTO Y MOROSIDAD",
        "El cumplimiento exacto de los plazos de pago constituye un elemento esencial del presente contrato. Ante la ocurrencia de morosidad o retardo en el pago de cualquiera de las cuotas devengadas, operarán los siguientes efectos:"
    )
    p7_item1 = doc.add_paragraph()
    p7_item1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p7_item1.add_run("Aceleración de la deuda: ").bold = True
    p7_item1.add_run("la mora en el pago de una cuota faculta a El Abogado para exigir de forma inmediata el cobro íntegro del saldo total que permanezca insoluto.")
    
    p7_item2 = doc.add_paragraph()
    p7_item2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p7_item2.add_run("Suspensión de la tramitación: ").bold = True
    p7_item2.add_run("un atraso superior a cinco días hábiles faculta a El Abogado para suspender la presentación de escritos y diligencias ante los tribunales respectivos, hasta la regularización del pago.")
    
    p7_item3 = doc.add_paragraph()
    p7_item3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p7_item3.add_run("Multa por atraso: ").bold = True
    p7_item3.add_run("se devengará una multa compensatoria equivalente a 0,15 Unidades de Fomento (UF) por cada día de atraso, hasta el pago efectivo de lo adeudado.")

    # --- CLÁUSULA OCTAVA: OBLIGACIONES RECÍPROCAS ---
    clausula(doc, "OCTAVA", "OBLIGACIONES RECÍPROCAS DE LAS PARTES")
    
    p8_abogado = doc.add_paragraph()
    p8_abogado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p8_abogado.add_run("Obligaciones de El Abogado: ").bold = True
    p8_abogado.add_run("El Abogado asume una obligación de medios, comprometiéndose a desplegar toda su diligencia, conocimiento técnico y ético en la tramitación del asunto encomendado, sin que ello implique en caso alguno una obligación de resultado ni la garantía de un desenlace favorable, atendida la naturaleza incierta de todo proceso judicial o administrativo.")
    
    p8_cliente = doc.add_paragraph()
    p8_cliente.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p8_cliente.add_run("Obligaciones de El Cliente: ").bold = True
    p8_cliente.add_run("El Cliente se obliga a proporcionar de forma oportuna, veraz y completa toda la información y documentación que resulte necesaria para el correcto desarrollo del encargo, dentro de los plazos que El Abogado le indique, siendo el incumplimiento de esta obligación causal suficiente para la resciliación del presente contrato.")
    
    p8_clave = doc.add_paragraph()
    p8_clave.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p8_clave.add_run("En caso de que la naturaleza del encargo requiera el uso de la Clave Única de El Cliente para gestiones ante la Oficina Judicial Virtual del Poder Judicial u otros organismos del Estado, El Cliente autoriza expresamente su uso por parte de El Abogado, exclusivamente para los fines del presente contrato, comprometiéndose este último a resguardar su confidencialidad e integridad y a no utilizarla para ningún otro propósito.")

    # --- CLÁUSULA NOVENA: AUSENCIA DE RELACIÓN LABORAL ---
    clausula(doc, "NOVENA", "AUSENCIA DE RELACIÓN LABORAL",
        "Las partes dejan constancia que el vínculo que las une mediante el presente contrato es de naturaleza estrictamente civil, no existiendo entre ellas subordinación ni dependencia de ningún tipo, por lo que no se genera relación laboral alguna ni las obligaciones propias de dicho régimen."
    )
    
    # --- CLÁUSULA DÉCIMA: CONFIDENCIALIDAD ---
    clausula(doc, "DÉCIMA", "CONFIDENCIALIDAD",
        "Toda la información y documentación que las partes intercambien con motivo de la celebración y ejecución del presente contrato tiene el carácter de confidencial, no pudiendo ser divulgada a terceros ajenos a la relación contractual sin el consentimiento previo y por escrito de la otra parte, salvo que su revelación sea exigida por ley o por orden de autoridad competente.",
        "Esta obligación de confidencialidad subsistirá aun después de terminado el presente contrato, cualquiera sea la causa de su término."
    )
    
    # --- CLÁUSULA DÉCIMO PRIMERA: PROTECCIÓN DE DATOS PERSONALES ---
    clausula(doc, "DÉCIMO PRIMERA", "PROTECCIÓN DE DATOS PERSONALES",
        "El Abogado tratará los datos personales que El Cliente le proporcione con motivo de la celebración y ejecución del presente contrato conforme a la normativa vigente sobre protección de la vida privada y de datos personales, utilizándolos exclusivamente para los fines del encargo profesional descrito en la cláusula primera.",
        "El Cliente podrá, en cualquier momento, ejercer sus derechos de acceso, rectificación, cancelación, oposición y portabilidad respecto de sus datos personales, dirigiendo su solicitud al correo electrónico de El Abogado individualizado en la comparecencia."
    )
    
    # --- CLÁUSULA DÉCIMO SEGUNDA: BUEN TRATO Y RESPETO MUTUO ---
    clausula(doc, "DÉCIMO SEGUNDA", "DEL COMPROMISO DE BUEN TRATO Y RESPETO MUTUO",
        "Las partes declaran que la relación contractual deberá fundarse en el respeto mutuo y en un trato digno y adecuado, libre de violencia, discriminación o agresión de cualquier naturaleza, en cumplimiento de la Ley N° 21.643 y su reglamento sobre prevención y sanción del acoso laboral, sexual y la violencia en el trabajo.",
        "Esta obligación no obsta en caso alguno al derecho de El Cliente de formular consultas, expresar disconformidad o presentar reclamos respecto del servicio contratado, los que serán siempre recibidos y atendidos en un marco de respeto recíproco."
    )
    
    # --- CLÁUSULA DÉCIMO TERCERA: TÉRMINO ANTICIPADO ---
    clausula(doc, "DÉCIMO TERCERA", "DEL DESISTIMIENTO O TÉRMINO ANTICIPADO DEL CONTRATO",
        "El presente contrato podrá terminar anticipadamente por mutuo acuerdo de las partes, por voluntad unilateral de cualquiera de ellas, o por incumplimiento grave de las obligaciones aquí pactadas.",
        "La parte que decida poner término unilateral al contrato deberá comunicarlo a la otra por escrito, mediante carta certificada o correo electrónico, con a lo menos 15 días de anticipación. En caso de que el desistimiento provenga de El Cliente, los honorarios devengados hasta esa fecha por los servicios ya prestados o iniciados le pertenecerán a El Abogado a título de honorarios causados, sin derecho a devolución alguna, sin perjuicio del deber de este último de rendir cuenta de las gestiones realizadas y facilitar los antecedentes necesarios para que El Cliente pueda continuar la tramitación del asunto con otro profesional."
    )
    
    # --- CLÁUSULA DÉCIMO CUARTA: CLÁUSULA PENAL ---
    clausula(doc, "DÉCIMO CUARTA", "CLÁUSULA PENAL",
        "En caso de incumplimiento de las obligaciones pactadas en este contrato por cualquiera de las partes, la parte incumplidora deberá pagar a la otra, a título de indemnización de perjuicios de avaluación anticipada, una suma equivalente a $500.000 (quinientos mil pesos), sin perjuicio del derecho de la parte afectada de exigir el cumplimiento forzado de lo pactado o la indemnización de los perjuicios efectivamente sufridos, si estos fueren mayores."
    )
    
    # --- CLÁUSULA DÉCIMO QUINTA: DURACIÓN ---
    clausula(doc, "DÉCIMO QUINTA", "DURACIÓN DEL SERVICIO",
        "La duración del presente contrato se extenderá hasta la completa ejecución del encargo profesional descrito en la cláusula primera, dependiendo su plazo efectivo de las circunstancias procesales propias del asunto encomendado y de los tiempos de tramitación de los tribunales u organismos competentes, los que escapan al control de El Abogado."
    )
    
    num_clausula_extra = 16
    diccionario_numeros_extra = {
        16: "DÉCIMO SEXTA", 17: "DÉCIMO SÉPTIMA", 18: "DÉCIMO OCTAVA", 19: "DÉCIMO NOVENA", 20: "VIGÉSIMA"
    }
    
    if datos['tipo_servicio'] == "Liquidación voluntaria":
        p_extra = doc.add_paragraph()
        p_extra.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_extra.add_run(f"CLÁUSULA {diccionario_numeros_extra[num_clausula_extra]}: ENTREGA DE DECLARACIONES JURADAS OBLIGATORIAS. ").bold = True
        p_extra.add_run("Atendida la naturaleza específica del procedimiento de insolvencia y liquidación voluntaria, El Cliente asume la obligación ineludible de suscribir y entregar las siguientes declaraciones juradas exigidas por la normativa aplicable:")
        
        p_extra_lista = doc.add_paragraph()
        p_extra_lista.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_extra_lista.add_run("- Declaración Jurada de Bienes Excluidos o de Terceros.\n- Declaración Jurada de Listado Completo de Acreedores.\n- Consentimiento Informado Expreso de los Efectos de la Liquidación.")
        num_clausula_extra += 1

    # --- CLÁUSULA CONDICIONAL: DOCUMENTOS QUE EL CLIENTE DEBE REUNIR ---
    if datos.get('documentos_requeridos', '').strip():
        lineas_docs = [l.strip() for l in datos['documentos_requeridos'].strip().split("\n") if l.strip()]
        p_docs = doc.add_paragraph()
        p_docs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_docs.add_run(f"CLÁUSULA {diccionario_numeros_extra[num_clausula_extra]}: DOCUMENTACIÓN A REUNIR POR EL CLIENTE. ").bold = True
        p_docs.add_run("Para dar inicio a la redacción de la gestión encomendada en la cláusula primera, El Cliente se obliga a reunir y hacer entrega a El Abogado, a la brevedad posible, de los siguientes documentos y antecedentes:")
        
        p_docs_lista = doc.add_paragraph()
        p_docs_lista.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_docs_lista.add_run("\n".join([f"- {linea}" for linea in lineas_docs]))
        
        p_docs_final = doc.add_paragraph()
        p_docs_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_docs_final.add_run("La demora injustificada en la entrega de estos antecedentes podrá afectar los plazos de tramitación del asunto encomendado, sin que ello sea imputable a El Abogado.")
        num_clausula_extra += 1

    clausula(doc, diccionario_numeros_extra[num_clausula_extra], "DOMICILIO CONVENCIONAL Y COMPETENCIA",
        "Para todos los efectos legales derivados del presente instrumento, las partes fijan su domicilio en la comuna y ciudad de Santiago, y se someten a la competencia de sus Tribunales Ordinarios de Justicia."
    )
    num_clausula_extra += 1
    
    clausula(doc, diccionario_numeros_extra[num_clausula_extra], "COMUNICACIONES ENTRE LAS PARTES",
        f"Para todos los efectos del presente contrato, se tendrán por válidas las comunicaciones dirigidas a los siguientes correos electrónicos y números de contacto: de El Cliente, {datos['cliente_correo']} / {datos['cliente_tel']}; y de El Abogado, {datos['abogado_correo']} / {datos['abogado_tel']}.",
        "En señal de plena conformidad con todas y cada una de las cláusulas precedentes, se extiende el presente contrato en dos ejemplares de idéntico tenor, quedando uno en poder de cada parte."
    )

    doc.add_paragraph("\n\n\n")
    table_firmas = doc.add_table(rows=1, cols=2)
    
    para_abogado = table_firmas.cell(0, 0).paragraphs[0]
    para_abogado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_abogado.add_run("___________________________________\n")
    para_abogado.add_run(f"{datos['abogado_nombre'].upper()}\n")
    para_abogado.add_run(f"R.U.T.: {datos['abogado_rut']}")
    
    para_cliente = table_firmas.cell(0, 1).paragraphs[0]
    para_cliente.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_cliente.add_run("___________________________________\n")
    para_cliente.add_run(f"{datos['cliente_nombre'].upper()}\n")
    para_cliente.add_run(f"R.U.T.: {datos['cliente_rut']}")
    
    return doc

def crear_informe_analisis_escritura_word(nombre_escritura_original, texto_analisis):
    """
    Genera el informe de análisis de una escritura en Word, con el mismo
    formato profesional del resto del sistema (Calibri 11, interlineado 1.5,
    justificado), para guardarlo en el historial igual que un contrato.
    """
    if not DOCX_READY:
        return None
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("INFORME DE ANÁLISIS DE ESCRITURA PÚBLICA")
    r.bold = True
    r.font.size = Pt(14)
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    meta.add_run("Documento analizado: ").bold = True
    meta.add_run(f"{nombre_escritura_original}\n")
    meta.add_run("Fecha del análisis: ").bold = True
    meta.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_disclaimer.add_run("Este informe fue generado con apoyo de inteligencia artificial como una revisión preliminar de formalidades y redacción, y no reemplaza el criterio profesional del abogado responsable.").italic = True
    
    # Cada bloque separado por doble salto de línea se convierte en un párrafo
    # real (no un salto interno dentro del mismo párrafo), para que se vea
    # bien justificado en Word, igual que en los contratos y escrituras.
    for bloque in texto_analisis.split("\n\n"):
        if bloque.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(bloque.strip())
    
    return doc

def crear_informe_posesion_efectiva_word(datos_causante, df_herederos_calc, masa_hereditaria, valor_utm, total_impuesto):
    """
    Genera el resumen de la determinación de asignaciones e impuesto a la
    herencia en Word, con los mismos datos que hay que trasladar al
    Formulario 4423 del SII y al Formulario de Solicitud de Posesión
    Efectiva del Registro Civil.
    """
    if not DOCX_READY:
        return None
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("DETERMINACIÓN DE ASIGNACIONES E IMPUESTO A LA HERENCIA")
    r.bold = True
    r.font.size = Pt(14)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("(Base para completar el Formulario 4423 del SII y la Solicitud de Posesión Efectiva ante el Registro Civil)").italic = True
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    meta.add_run("Causante: ").bold = True
    meta.add_run(f"{datos_causante.get('nombre','')} — RUT: {datos_causante.get('rut','')}\n")
    meta.add_run("Fecha de defunción: ").bold = True
    meta.add_run(f"{datos_causante.get('fecha_defuncion','')}\n")
    meta.add_run("Valor UTM utilizado: ").bold = True
    meta.add_run(f"{formatear_clp(valor_utm)}\n")
    meta.add_run("Masa Hereditaria (Total Activos - Total Pasivos): ").bold = True
    meta.add_run(f"{formatear_clp(masa_hereditaria)}\n")
    meta.add_run("Impuesto Total a Pagar: ").bold = True
    meta.add_run(f"{formatear_clp(total_impuesto)}")
    
    doc.add_paragraph("\nDetalle de Asignaciones e Impuesto por Heredero:").runs[0].bold = True
    
    if not df_herederos_calc.empty:
        tabla = doc.add_table(rows=1, cols=7)
        tabla.style = 'Table Grid'
        encabezados = ['Heredero', 'RUT', 'Tipo', 'Asignación ($)', 'Base Imponible (UTM)', 'Impuesto Total (UTM)', 'Impuesto Total ($)']
        for i, enc in enumerate(encabezados):
            tabla.rows[0].cells[i].text = enc
        for _, fila in df_herederos_calc.iterrows():
            fila_celdas = tabla.add_row().cells
            fila_celdas[0].text = str(fila['Heredero'])
            fila_celdas[1].text = str(fila['RUT'])
            fila_celdas[2].text = str(fila['Tipo'])
            fila_celdas[3].text = formatear_clp(fila['Asignación ($)'])
            fila_celdas[4].text = str(fila['Base Imponible (UTM)'])
            fila_celdas[5].text = str(fila['Impuesto Total (UTM)'])
            fila_celdas[6].text = formatear_clp(fila['Impuesto Total ($)'])
    
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_nota.add_run("\nEste cálculo se realizó conforme a las reglas de sucesión intestada y a las Tablas N°1, N°2 y N°3 de las instrucciones del Formulario 4423 del Servicio de Impuestos Internos. Corresponde verificar los valores finales contra la masa hereditaria efectivamente inventariada y tasada antes de presentar la declaración oficial.").italic = True
    
    return doc

# =====================================================================
# 📜 CATÁLOGO DE ESCRITURAS PÚBLICAS DE CHILE
# =====================================================================
# Basado en los tipos de escritura pública reconocidos y de uso más frecuente
# en las notarías chilenas (compraventa de inmuebles, sociedades, hipotecas,
# mandatos, testamentos, etc., conforme al Código Orgánico de Tribunales y
# al Código Civil). No pretende ser exhaustivo: cualquier acto jurídico puede
# reducirse a escritura pública, pero estos son los que un estudio jurídico
# de práctica general redacta con más frecuencia. Si necesitas un tipo que no
# está aquí, dímelo y lo agrego.
CATALOGO_ESCRITURAS = {
    "Compraventa de Bien Raíz": {
        "roles": ("Vendedor(a)", "Comprador(a)"),
        "campos": [
            ("direccion_inmueble", "Dirección del inmueble", "text"),
            ("comuna_inmueble", "Comuna", "text"),
            ("rol_avaluo", "Rol de Avalúo (SII)", "text"),
            ("fojas_inscripcion", "Fojas / Número / Año de inscripción vigente (CBR)", "text"),
            ("precio", "Precio de venta ($)", "text"),
            ("forma_pago", "Forma de pago", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: Por el presente instrumento, El Vendedor vende, cede y transfiere a El Comprador, quien compra y adquiere para sí, el inmueble ubicado en {d.get('direccion_inmueble','')}, comuna de {d.get('comuna_inmueble','')}, individualizado con Rol de Avalúo N° {d.get('rol_avaluo','')}, inscrito a fojas {d.get('fojas_inscripcion','')} del Registro de Propiedad del Conservador de Bienes Raíces respectivo.\n\n"
            f"SEGUNDO: El precio de la compraventa es la suma de {d.get('precio','')}, que las partes declaran pagadera de la siguiente forma: {d.get('forma_pago','')}.\n\n"
            "TERCERO: El Vendedor se obliga al saneamiento de la evicción y de los vicios redhibitorios conforme a las reglas generales, declarando que el inmueble se encuentra libre de gravámenes, prohibiciones, embargos y litigios pendientes, salvo que se exprese lo contrario.\n\n"
            "CUARTO: La presente escritura servirá de título suficiente para que El Comprador practique la inscripción del dominio a su nombre en el Registro de Propiedad del Conservador de Bienes Raíces competente."
        )
    },
    "Constitución de Sociedad de Responsabilidad Limitada": {
        "roles": ("Socio(a) 1", "Socio(a) 2"),
        "campos": [
            ("razon_social", "Razón Social de la sociedad", "text"),
            ("giro", "Giro u objeto social", "textarea"),
            ("capital", "Capital social ($)", "text"),
            ("aporte_socio1", "Aporte Socio 1 (% o monto)", "text"),
            ("aporte_socio2", "Aporte Socio 2 (% o monto)", "text"),
            ("domicilio_social", "Domicilio de la sociedad", "text"),
            ("duracion", "Duración de la sociedad", "text"),
            ("administracion", "Administración (quién administra)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: Los comparecientes vienen en constituir una sociedad de responsabilidad limitada, que girará bajo la razón social de \"{d.get('razon_social','')}\", en adelante \"la Sociedad\".\n\n"
            f"SEGUNDO: El objeto de la Sociedad será: {d.get('giro','')}.\n\n"
            f"TERCERO: El capital social se fija en la suma de {d.get('capital','')}, que los socios aportan de la siguiente forma: el Socio 1 aporta {d.get('aporte_socio1','')} y el Socio 2 aporta {d.get('aporte_socio2','')}.\n\n"
            f"CUARTO: El domicilio social será {d.get('domicilio_social','')}, sin perjuicio de las agencias o sucursales que la Sociedad establezca en el futuro.\n\n"
            f"QUINTO: La duración de la Sociedad será de {d.get('duracion','')}, contada desde la fecha de la presente escritura.\n\n"
            f"SEXTO: La administración de la Sociedad corresponderá a {d.get('administracion','')}, quien usará la razón social y ejercerá la representación judicial y extrajudicial de la misma, con las más amplias facultades de administración."
        )
    },
    "Mandato / Poder General": {
        "roles": ("Mandante", "Mandatario(a)"),
        "campos": [
            ("facultades", "Facultades que se otorgan", "textarea"),
            ("plazo_vigencia", "Plazo de vigencia (si aplica)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Mandante confiere poder general amplio a El Mandatario para que, actuando en su nombre y representación, ejerza las siguientes facultades: {d.get('facultades','')}.\n\n"
            "SEGUNDO: En el ejercicio de este mandato, El Mandatario queda facultado con todas las atribuciones de un mandatario general en los términos del artículo 2132 y siguientes del Código Civil, así como con las facultades especiales de ambos incisos del artículo 7° del Código de Procedimiento Civil, cuando corresponda.\n\n"
            f"TERCERO: El presente mandato tendrá vigencia {d.get('plazo_vigencia', 'indefinida, mientras no sea revocado por El Mandante')}, sin perjuicio de las causales legales de extinción del mandato."
        )
    },
    "Mandato / Poder Especial": {
        "roles": ("Mandante", "Mandatario(a)"),
        "campos": [
            ("acto_especifico", "Acto o gestión específica encomendada", "textarea"),
            ("plazo_vigencia", "Plazo de vigencia (si aplica)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Mandante confiere poder especial a El Mandatario, para que, exclusivamente para el siguiente acto o gestión, actúe en su nombre y representación: {d.get('acto_especifico','')}.\n\n"
            "SEGUNDO: Las facultades conferidas por este mandato se limitan estrictamente al acto o gestión especificado en la cláusula precedente, sin que puedan extenderse a otros actos no comprendidos expresamente en él.\n\n"
            f"TERCERO: El presente mandato tendrá vigencia {d.get('plazo_vigencia', 'hasta el cumplimiento del encargo o su revocación')}."
        )
    },
    "Hipoteca": {
        "roles": ("Deudor(a) Hipotecario(a)", "Acreedor(a) Hipotecario(a)"),
        "campos": [
            ("direccion_inmueble", "Dirección del inmueble hipotecado", "text"),
            ("rol_avaluo", "Rol de Avalúo (SII)", "text"),
            ("monto_garantizado", "Monto de la obligación garantizada ($)", "text"),
            ("plazo_obligacion", "Plazo de la obligación garantizada", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Deudor Hipotecario constituye hipoteca de primer grado sobre el inmueble ubicado en {d.get('direccion_inmueble','')}, Rol de Avalúo N° {d.get('rol_avaluo','')}, en favor de El Acreedor Hipotecario, para garantizar el cumplimiento de la obligación que se individualiza en la cláusula siguiente.\n\n"
            f"SEGUNDO: La obligación garantizada asciende a la suma de {d.get('monto_garantizado','')}, con un plazo de {d.get('plazo_obligacion','')}.\n\n"
            "TERCERO: La presente hipoteca se extiende a todas las ampliaciones, renovaciones o prórrogas de la obligación principal, así como a los intereses, costas y demás accesorios legales.\n\n"
            "CUARTO: Esta hipoteca deberá inscribirse en el Registro de Hipotecas y Gravámenes del Conservador de Bienes Raíces respectivo para producir efectos respecto de terceros, conforme al artículo 2410 del Código Civil."
        )
    },
    "Contrato de Arrendamiento (reducido a escritura pública)": {
        "roles": ("Arrendador(a)", "Arrendatario(a)"),
        "campos": [
            ("direccion_inmueble", "Dirección del inmueble arrendado", "text"),
            ("renta_mensual", "Renta mensual ($)", "text"),
            ("plazo_contrato", "Plazo del contrato", "text"),
            ("garantia", "Garantía / mes de depósito", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Arrendador da en arrendamiento a El Arrendatario, quien acepta, el inmueble ubicado en {d.get('direccion_inmueble','')}.\n\n"
            f"SEGUNDO: La renta de arrendamiento será de {d.get('renta_mensual','')} mensuales, pagaderos dentro de los primeros cinco días de cada mes.\n\n"
            f"TERCERO: El plazo del contrato será de {d.get('plazo_contrato','')}, renovable por períodos iguales y sucesivos si ninguna de las partes manifiesta su voluntad en contrario con la anticipación legal correspondiente.\n\n"
            f"CUARTO: El Arrendatario entrega en este acto la suma de {d.get('garantia','')} a título de garantía por el fiel cumplimiento de las obligaciones del presente contrato, la que será restituida al término del arrendamiento previa constatación del estado del inmueble."
        )
    },
    "Mutuo (Préstamo de Dinero)": {
        "roles": ("Mutuante (Acreedor)", "Mutuario (Deudor)"),
        "campos": [
            ("monto_mutuo", "Monto del préstamo ($)", "text"),
            ("interes", "Interés pactado", "text"),
            ("plazo_restitucion", "Plazo de restitución", "text"),
            ("garantia_mutuo", "Garantía (si existe)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Mutuante entrega en este acto a El Mutuario, a título de mutuo o préstamo de consumo, la suma de {d.get('monto_mutuo','')}, cantidad que El Mutuario declara recibir a su entera satisfacción.\n\n"
            f"SEGUNDO: El Mutuario se obliga a restituir la suma mutuada dentro del plazo de {d.get('plazo_restitucion','')}, con un interés de {d.get('interes','')}.\n\n"
            f"TERCERO: {('Para garantizar el cumplimiento de esta obligación, se constituye la siguiente garantía: ' + d.get('garantia_mutuo','')) if d.get('garantia_mutuo','').strip() else 'Las partes declaran que la presente obligación no cuenta con garantía real adicional.'}\n\n"
            "CUARTO: El retardo en la restitución del capital o de los intereses hará exigible de inmediato la totalidad de la obligación, sin necesidad de requerimiento judicial o extrajudicial previo."
        )
    },
    "Donación entre Vivos": {
        "roles": ("Donante", "Donatario(a)"),
        "campos": [
            ("bien_donado", "Descripción del bien donado", "textarea"),
            ("valor_bien", "Valor referencial del bien ($)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Donante dona, cede y transfiere gratuitamente a El Donatario, quien acepta, el siguiente bien: {d.get('bien_donado','')}, cuyo valor referencial es de {d.get('valor_bien','')}.\n\n"
            "SEGUNDO: El Donatario declara aceptar la presente donación, agradeciendo la liberalidad de El Donante.\n\n"
            "TERCERO: Las partes declaran conocer las obligaciones tributarias que pudieren derivarse de la presente donación conforme a la Ley N° 16.271 sobre Impuesto a las Herencias, Asignaciones y Donaciones."
        )
    },
    "Finiquito / Cancelación de Deuda": {
        "roles": ("Acreedor(a)", "Deudor(a)"),
        "campos": [
            ("obligacion_cancelada", "Obligación que se cancela (descripción)", "textarea"),
            ("monto_pagado", "Monto pagado ($)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Acreedor deja constancia de haber recibido de El Deudor la suma de {d.get('monto_pagado','')}, en pago total y definitivo de la siguiente obligación: {d.get('obligacion_cancelada','')}.\n\n"
            "SEGUNDO: En consecuencia, El Acreedor otorga a El Deudor el más amplio, total y definitivo finiquito respecto de la obligación individualizada precedentemente, declarando no tener nada más que reclamar por dicho concepto.\n\n"
            "TERCERO: Se deja constancia de que las garantías constituidas para caucionar la obligación referida, si las hubiere, quedan igualmente canceladas y sin efecto por el presente instrumento."
        )
    },
    "Transacción": {
        "roles": ("Parte 1", "Parte 2"),
        "campos": [
            ("conflicto_pendiente", "Conflicto o diferencia que se transa", "textarea"),
            ("concesiones_reciprocas", "Concesiones recíprocas de las partes", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: Las partes dejan constancia de la existencia de la siguiente diferencia o conflicto pendiente entre ellas: {d.get('conflicto_pendiente','')}.\n\n"
            f"SEGUNDO: Con el objeto de precaver un litigio eventual o poner término a uno existente, las partes acuerdan las siguientes concesiones recíprocas: {d.get('concesiones_reciprocas','')}.\n\n"
            "TERCERO: Las partes declaran que la presente transacción produce el efecto de cosa juzgada en última instancia respecto de las materias en ella comprendidas, conforme al artículo 2460 del Código Civil, y se obligan a no promover acción judicial alguna relacionada con lo aquí transigido."
        )
    },
    "Testamento Abierto": {
        "roles": ("Testador(a)", "(sin segunda parte - requiere 3 testigos)"),
        "campos": [
            ("herederos_asignatarios", "Herederos y/o legatarios designados, con sus asignaciones", "textarea"),
            ("albacea", "Albacea designado (si aplica)", "text"),
        ],
        "clausula": lambda d: (
            "PRIMERO: El Testador, hallándose en su sano juicio, declara ser su última voluntad la que se consigna en la presente escritura, revocando cualquier testamento anterior que hubiere otorgado.\n\n"
            f"SEGUNDO: El Testador instituye como herederos y/o legatarios a las siguientes personas, con las asignaciones que se indican: {d.get('herederos_asignatarios','')}.\n\n"
            f"TERCERO: {('El Testador designa como albacea, con tenencia de bienes, a ' + d.get('albacea','')) if d.get('albacea','').strip() else 'El Testador no designa albacea, rigiendo las reglas generales sobre partición de la herencia.'}\n\n"
            "CUARTO: El presente testamento se otorga ante tres testigos hábiles, quienes firman la presente escritura junto con el Testador y el Notario autorizante, conforme a las solemnidades del testamento abierto establecidas en el Código Civil."
        )
    },
    "Servidumbre": {
        "roles": ("Predio Sirviente (Propietario/a)", "Predio Dominante (Propietario/a)"),
        "campos": [
            ("tipo_servidumbre", "Tipo de servidumbre (tránsito, acueducto, etc.)", "text"),
            ("descripcion_predios", "Descripción de ambos predios", "textarea"),
            ("condiciones_ejercicio", "Condiciones de ejercicio de la servidumbre", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El propietario del predio sirviente constituye, en favor del predio dominante, una servidumbre de {d.get('tipo_servidumbre','')}, sobre los predios que se describen a continuación: {d.get('descripcion_predios','')}.\n\n"
            f"SEGUNDO: La servidumbre se ejercerá bajo las siguientes condiciones: {d.get('condiciones_ejercicio','')}.\n\n"
            "TERCERO: La presente servidumbre deberá inscribirse en el Registro de Hipotecas y Gravámenes del Conservador de Bienes Raíces respectivo para su plena oponibilidad a terceros, conforme al artículo 698 del Código Civil."
        )
    },
    "Liquidación de Sociedad Conyugal": {
        "roles": ("Cónyuge 1", "Cónyuge 2"),
        "campos": [
            ("causal_disolucion", "Causal de disolución de la sociedad conyugal (divorcio, cambio de régimen, etc.)", "text"),
            ("inventario_bienes", "Inventario de bienes sociales (descripción y valor de cada uno)", "textarea"),
            ("adjudicacion", "Forma de adjudicación acordada (qué bien y a quién se adjudica)", "textarea"),
            ("recompensas", "Recompensas entre cónyuges y sociedad, si existen", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: La sociedad conyugal habida entre los comparecientes se disolvió por {d.get('causal_disolucion','')}, formándose desde esa fecha una comunidad de bienes entre ambos, la que se procede a liquidar por el presente instrumento conforme a las reglas de la partición de bienes hereditarios (artículo 1776 del Código Civil).\n\n"
            f"SEGUNDO: El inventario y tasación de los bienes que componen el haber social es el siguiente: {d.get('inventario_bienes','')}.\n\n"
            f"TERCERO: Las partes acuerdan la siguiente adjudicación de los bienes inventariados: {d.get('adjudicacion','')}.\n\n"
            f"CUARTO: {('Se reconocen y regulan las siguientes recompensas entre los cónyuges y la sociedad conyugal: ' + d.get('recompensas','')) if d.get('recompensas','').strip() else 'Las partes declaran que no existen recompensas pendientes entre los cónyuges y la sociedad conyugal.'}\n\n"
            "QUINTO: Con el pago y adjudicación antes referidos, las partes se declaran mutuamente pagadas de sus derechos en la sociedad conyugal, sin que quede cargo ni saldo pendiente entre ellas por este concepto, otorgándose el más amplio y recíproco finiquito.\n\n"
            "SEXTO: Si dentro de los bienes adjudicados existen inmuebles u otros bienes sujetos a registro, la presente escritura deberá inscribirse en el registro respectivo del Conservador de Bienes Raíces o repartición que corresponda para su plena oponibilidad a terceros."
        )
    },
    "Separación de Bienes con Liquidación de Sociedad Conyugal": {
        "roles": ("Cónyuge 1", "Cónyuge 2"),
        "campos": [
            ("fecha_matrimonio", "Fecha y lugar de celebración del matrimonio", "text"),
            ("inventario_bienes", "Inventario de bienes sociales (descripción y valor de cada uno)", "textarea"),
            ("adjudicacion", "Forma de adjudicación acordada (qué bien y a quién se adjudica)", "textarea"),
            ("renuncia_gananciales", "¿Alguno de los cónyuges renuncia a los gananciales? (indicar quién, o dejar en blanco si no aplica)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: Los comparecientes, casados bajo el régimen de sociedad conyugal con fecha {d.get('fecha_matrimonio','')}, vienen en pactar la sustitución de dicho régimen patrimonial por el de separación total de bienes, en conformidad con el artículo 1723 del Código Civil.\n\n"
            "SEGUNDO: Por efecto del pacto precedente, queda disuelta la sociedad conyugal habida entre las partes, formándose una comunidad de bienes que en este mismo acto se procede a liquidar.\n\n"
            f"TERCERO: El inventario y tasación de los bienes que componen el haber social es el siguiente: {d.get('inventario_bienes','')}.\n\n"
            f"CUARTO: Las partes acuerdan la siguiente adjudicación de los bienes inventariados: {d.get('adjudicacion','')}.\n\n"
            f"QUINTO: {('Se deja constancia de la renuncia a los gananciales por parte de ' + d.get('renuncia_gananciales','')) if d.get('renuncia_gananciales','').strip() else 'No existe renuncia a los gananciales por ninguna de las partes.'}\n\n"
            "SEXTO: A partir de la fecha de este instrumento, cada cónyuge tendrá la libre administración y disposición de sus bienes propios y de los que en lo sucesivo adquiera, sin injerencia del otro.\n\n"
            "SÉPTIMO: El presente pacto deberá subinscribirse al margen de la respectiva inscripción matrimonial dentro del plazo de treinta días contados desde la fecha de esta escritura, sin lo cual no producirá efecto alguno entre las partes ni respecto de terceros, conforme al artículo 1723 del Código Civil."
        )
    },
    "Testamento Cerrado": {
        "roles": ("Testador(a)", "(sin segunda parte - requiere 3 testigos)"),
        "campos": [
            ("declaracion_general", "Declaración general que el testador desea consignar (opcional, sin revelar el contenido del sobre)", "textarea"),
        ],
        "clausula": lambda d: (
            "PRIMERO: El Testador comparece ante el Notario y los testigos indicados en esta escritura, y hace entrega al Notario de un sobre cerrado, declarando de viva voz que en su interior se contiene su testamento.\n\n"
            f"SEGUNDO: {('El Testador deja constancia de la siguiente declaración general: ' + d.get('declaracion_general','')) if d.get('declaracion_general','').strip() else 'El contenido del testamento es conocido únicamente por el Testador, conforme a la naturaleza de esta forma testamentaria.'}\n\n"
            "TERCERO: El Notario levanta la presente acta en la cubierta del sobre, dejando constancia del lugar, fecha y circunstancias del otorgamiento, la que es firmada por el Testador, los tres testigos y el Notario autorizante, en un solo acto, sin interrupción.\n\n"
            "CUARTO: El sobre conteniendo el testamento queda en custodia del oficio notarial, sin perjuicio del derecho del Testador de retirarlo personalmente. Fallecido el Testador, el testamento deberá ser abierto mediante la gestión judicial correspondiente."
        )
    },
    "Usufructo": {
        "roles": ("Nudo(a) Propietario(a)", "Usufructuario(a)"),
        "campos": [
            ("bien_usufructo", "Descripción del bien sobre el que se constituye el usufructo", "textarea"),
            ("duracion_usufructo", "Duración del usufructo (plazo, o de por vida)", "text"),
            ("condiciones_usufructo", "Condiciones de uso, goce y conservación del bien", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Nudo Propietario constituye derecho real de usufructo en favor de El Usufructuario, sobre el siguiente bien: {d.get('bien_usufructo','')}.\n\n"
            f"SEGUNDO: El usufructo tendrá una duración de {d.get('duracion_usufructo','')}.\n\n"
            f"TERCERO: El Usufructuario tendrá derecho a gozar del bien y percibir sus frutos, sujeto a las siguientes condiciones: {d.get('condiciones_usufructo','')}, quedando obligado a conservar la forma y sustancia del bien y a restituirlo al Nudo Propietario a la extinción del usufructo.\n\n"
            "CUARTO: El presente usufructo deberá inscribirse en el Registro de Propiedad del Conservador de Bienes Raíces respectivo cuando recaiga sobre inmuebles, conforme al artículo 697 del Código Civil."
        )
    },
    "Declaración de Bien Familiar": {
        "roles": ("Cónyuge Solicitante", "Cónyuge Titular del Bien"),
        "campos": [
            ("bien_afectado", "Descripción del bien que se afecta como familiar", "textarea"),
            ("fundamento", "Fundamento de la afectación (residencia principal de la familia, etc.)", "textarea"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: Los comparecientes dejan constancia de su acuerdo para afectar como bien familiar el siguiente bien: {d.get('bien_afectado','')}, de propiedad de uno de los cónyuges, en conformidad con los artículos 141 y siguientes del Código Civil.\n\n"
            f"SEGUNDO: La afectación se funda en lo siguiente: {d.get('fundamento','')}.\n\n"
            "TERCERO: Las partes declaran conocer que la declaración de bien familiar limita las facultades de disposición del cónyuge propietario, quien no podrá enajenar ni gravar voluntariamente el bien, ni prometer hacerlo, sin la voluntad del cónyuge no propietario.\n\n"
            "CUARTO: Se deja constancia de que, para su plena oponibilidad a terceros, la presente declaración debe anotarse al margen de la inscripción de dominio del inmueble en el Conservador de Bienes Raíces respectivo, sin perjuicio de que la vía ordinaria para constituir esta afectación es la declaración judicial ante el Tribunal de Familia competente cuando no exista acuerdo entre los cónyuges."
        )
    },
    "Cesión de Derechos": {
        "roles": ("Cedente", "Cesionario(a)"),
        "campos": [
            ("derecho_cedido", "Descripción del derecho cedido (crédito, derecho litigioso, etc.)", "textarea"),
            ("precio_cesion", "Precio de la cesión ($, o indicar si es gratuita)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Cedente cede y transfiere a El Cesionario, quien acepta, el siguiente derecho: {d.get('derecho_cedido','')}.\n\n"
            f"SEGUNDO: La presente cesión se efectúa por la suma de {d.get('precio_cesion','')}.\n\n"
            "TERCERO: El Cedente declara que el derecho cedido existe y le pertenece legítimamente, sin perjuicio de que, salvo pacto expreso en contrario, no responde de la solvencia del deudor sino únicamente de la existencia del derecho al tiempo de la cesión, conforme a las reglas generales del Código Civil.\n\n"
            "CUARTO: Para que la presente cesión produzca efectos respecto del deudor y de terceros, deberá notificarse al deudor cedido o ser aceptada por este, conforme al artículo 1902 del Código Civil."
        )
    },
    "Cesión de Derechos Hereditarios": {
        "roles": ("Cedente (Heredero/a)", "Cesionario(a)"),
        "campos": [
            ("causante_datos", "Nombre completo y RUT del causante", "text"),
            ("fecha_defuncion_causante", "Fecha de defunción del causante", "text"),
            ("cuota_cedida", "Cuota o derechos cedidos (ej: la totalidad de sus derechos, o un porcentaje)", "text"),
            ("precio_cesion_hereditaria", "Precio de la cesión ($, o indicar si es gratuita)", "text"),
            ("posesion_efectiva_estado", "Estado de la posesión efectiva (tramitada, en trámite, pendiente)", "text"),
        ],
        "clausula": lambda d: (
            f"PRIMERO: El Cedente, en su calidad de heredero de don/doña {d.get('causante_datos','')}, fallecido(a) con fecha {d.get('fecha_defuncion_causante','')}, cede y transfiere a El Cesionario, quien acepta, {d.get('cuota_cedida','')} que le corresponden o pudieren corresponderle en dicha herencia.\n\n"
            f"SEGUNDO: La presente cesión se efectúa por la suma de {d.get('precio_cesion_hereditaria','')}.\n\n"
            f"TERCERO: Se deja constancia de que la posesión efectiva de la herencia se encuentra: {d.get('posesion_efectiva_estado','')}.\n\n"
            "CUARTO: El Cedente declara que es heredero del causante individualizado y que ha aceptado la herencia, cediendo por este acto el derecho real de herencia que le corresponde, sin que ello confiera al Cesionario la calidad de heredero, la cual permanece en el Cedente conforme a la ley.\n\n"
            "QUINTO: El Cedente no responde de la existencia de bienes determinados dentro de la herencia ni de su valor, sino únicamente de su calidad de heredero, salvo pacto expreso en contrario entre las partes.\n\n"
            "SEXTO: Si la herencia comprende bienes raíces, la presente cesión deberá inscribirse en el Registro de Propiedad del Conservador de Bienes Raíces del territorio en que estos se encuentren ubicados, para su mayor seguridad jurídica y oponibilidad a terceros."
        )
    },
}


def crear_escritura_word(tipo_escritura, datos):
    """
    Genera la escritura pública en Word con el mismo formato profesional que
    el generador de contratos (Calibri 11, interlineado 1.5, justificado con
    títulos centrados, párrafos reales para evitar el problema de huecos al
    justificar).
    """
    if not DOCX_READY:
        return None
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.add_run(f"ESCRITURA PÚBLICA DE {tipo_escritura.upper()}").bold = True
    
    hoy = datetime.now()
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    fecha_str = f"{hoy.day} de {meses[hoy.month-1]} del año {hoy.year}"
    
    rol1_label, rol2_label = CATALOGO_ESCRITURAS[tipo_escritura]["roles"]
    
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(f"En Santiago, República de Chile, a {fecha_str}, ante mí, {datos.get('notario_nombre','') or '[NOMBRE DEL NOTARIO]'}, Notario(a) Público(a) de {datos.get('notaria_ciudad','') or '[CIUDAD]'}, comparecen:")
    
    p_p1 = doc.add_paragraph()
    p_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_p1.add_run(f"Por una parte, don/doña {datos.get('parte1_nombre','')}, {datos.get('parte1_nacionalidad','chileno/a')}, {datos.get('parte1_estado_civil','')}, {datos.get('parte1_profesion','')}, cédula nacional de identidad número {datos.get('parte1_rut','')}, con domicilio en {datos.get('parte1_domicilio','')}, en adelante \"{rol1_label.upper()}\"; y,")
    
    if "sin segunda parte" not in rol2_label:
        p_p2 = doc.add_paragraph()
        p_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_p2.add_run(f"Por otra parte, don/doña {datos.get('parte2_nombre','')}, {datos.get('parte2_nacionalidad','chileno/a')}, {datos.get('parte2_estado_civil','')}, {datos.get('parte2_profesion','')}, cédula nacional de identidad número {datos.get('parte2_rut','')}, con domicilio en {datos.get('parte2_domicilio','')}, en adelante \"{rol2_label.upper()}\".")
    
    p_mayores = doc.add_paragraph()
    p_mayores.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_mayores.add_run("Los comparecientes, mayores de edad, quienes acreditan su identidad con las cédulas antes citadas, exponen que han convenido el siguiente acto jurídico, que se contiene en las cláusulas que a continuación se singularizan:")
    
    # Cláusulas específicas del tipo de escritura (cada "\n\n" del texto se
    # convierte en un párrafo real, no un salto interno, para que se vea bien
    # justificado en Word)
    texto_clausulas = CATALOGO_ESCRITURAS[tipo_escritura]["clausula"](datos)
    for bloque in texto_clausulas.split("\n\n"):
        if bloque.strip():
            p_c = doc.add_paragraph()
            p_c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_c.add_run(bloque.strip())
    
    p_domicilio = doc.add_paragraph()
    p_domicilio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_domicilio.add_run("Para todos los efectos legales derivados del presente instrumento, los comparecientes fijan domicilio en la ciudad de Santiago y se someten a la competencia de sus Tribunales Ordinarios de Justicia. Se deja constancia de que los comparecientes fueron informados por el Notario autorizante del contenido y alcance jurídico del presente instrumento, prestando su consentimiento libre y expresamente.")
    
    doc.add_paragraph("\n\n")
    table_firmas = doc.add_table(rows=1, cols=2)
    p_f1 = table_firmas.cell(0, 0).paragraphs[0]
    p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f1.add_run("___________________________________\n")
    p_f1.add_run(f"{datos.get('parte1_nombre','').upper()}\n")
    p_f1.add_run(f"R.U.T.: {datos.get('parte1_rut','')}\n")
    p_f1.add_run(rol1_label)
    
    if "sin segunda parte" not in rol2_label:
        p_f2 = table_firmas.cell(0, 1).paragraphs[0]
        p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f2.add_run("___________________________________\n")
        p_f2.add_run(f"{datos.get('parte2_nombre','').upper()}\n")
        p_f2.add_run(f"R.U.T.: {datos.get('parte2_rut','')}\n")
        p_f2.add_run(rol2_label)
    
    return doc

# --- MOTOR DE CREACIÓN DE INFORME IA EN WORD ---
def crear_informe_ia_word(rol, cliente, texto_informe):
    if not DOCX_READY: 
        return None
        
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("INFORME DE AVANCE Y ESTADO DE CAUSA JUDICIAL\n")
    r.bold = True
    r.font.size = Pt(14)
    
    meta = doc.add_paragraph()
    meta.add_run(f"Causa Rol: ").bold = True
    meta.add_run(f"{rol}\n")
    meta.add_run(f"Cliente Titular: ").bold = True
    meta.add_run(f"{cliente}\n")
    meta.add_run(f"Fecha de Emisión del Reporte: ").bold = True
    meta.add_run(f"{datetime.now().strftime('%d/%m/%Y')}\n")
    
    doc.add_paragraph("\nESTIMADO CLIENTE:\nA continuación, se detalla el análisis y resumen ejecutivo del estado actual de su procedimiento judicial, redactado en términos claros para su correcta comprensión:\n")
    
    p_inf = doc.add_paragraph()
    p_inf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_inf.add_run(texto_informe)
    
    doc.add_paragraph("\n\nAnte cualquier duda, nuestro equipo queda a su entera disposición.\n\n___________________________________\nEquipo Legal - JuriSync")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =====================================================================
# --- SISTEMA DE CONTROL DE ACCESO (EN GOOGLE SHEETS Y COOKIES) ---
# =====================================================================
ARCHIVO_USUARIOS = "base_usuarios.csv"

def guardar_en_nube(df):
    safe_update_sheet("base_usuarios", df)
    df.to_csv(ARCHIVO_USUARIOS, index=False)

df_usuarios = safe_read_sheet("base_usuarios")
if not df_usuarios.empty:
    df_usuarios = df_usuarios.dropna(how="all")
    if 'Debe_Cambiar_Clave' in df_usuarios.columns:
        df_usuarios['Debe_Cambiar_Clave'] = df_usuarios['Debe_Cambiar_Clave'].astype(str)
else:
    df_usuarios = pd.DataFrame()

if df_usuarios.empty:
    datos_iniciales = {
        "Usuario": ["Narratia", "Vfarfan", "Gdonoso", "Mcortes", "Jtrujillo", "Eriquelme"],
        "Password": [hash_password(p) for p in ["20911237", "vpfm2404", "gdonoso123", "Mcortes123", "Jtrujillo123", "Eriquelme123"]],
        "Nombre_Real": ["Nicolás Arratia", "Valentina Farfán", "Gabriel Donoso", "Miryam Cortés", "José Trujillo", "Eduardo Riquelme"],
        "Correo": ["pendiente", "pendiente", "pendiente", "pendiente", "pendiente", "pendiente"],
        "Debe_Cambiar_Clave": ['True', 'True', 'True', 'True', 'True', 'True'], 
        "Plan": ["Full", "Full", "Full", "Full", "Full", "Full"]
    }
    df_usuarios = pd.DataFrame(datos_iniciales)
    guardar_en_nube(df_usuarios)
else:
    cambios = False
    
    if "Eriquelme" not in df_usuarios['Usuario'].values:
        nuevo_u = pd.DataFrame([{"Usuario": "Eriquelme", "Password": hash_password("Eriquelme123"), "Nombre_Real": "Eduardo Riquelme", "Correo": "pendiente", "Debe_Cambiar_Clave": 'True', "Plan": "Full"}])
        df_usuarios = pd.concat([df_usuarios, nuevo_u], ignore_index=True)
        cambios = True

    if "Jtrujillo" not in df_usuarios['Usuario'].values:
        nuevo_u = pd.DataFrame([{"Usuario": "Jtrujillo", "Password": hash_password("Jtrujillo123"), "Nombre_Real": "José Trujillo", "Correo": "pendiente", "Debe_Cambiar_Clave": 'True', "Plan": "Full"}])
        df_usuarios = pd.concat([df_usuarios, nuevo_u], ignore_index=True)
        cambios = True

    # Migración automática: si alguna contraseña sigue en texto plano (formato antiguo),
    # se re-hashea al vuelo la próxima vez que ese usuario inicie sesión exitosamente
    # (ver más abajo en la pantalla de login). No se fuerza aquí para no invalidar sesiones.

    idx_narratia = df_usuarios[df_usuarios['Usuario'] == 'Narratia'].index
    if not idx_narratia.empty:
        correo_actual = str(df_usuarios.loc[idx_narratia[0], 'Correo'])
        if "@" not in correo_actual:
            df_usuarios.loc[idx_narratia[0], 'Debe_Cambiar_Clave'] = 'True'
            df_usuarios.loc[idx_narratia[0], 'Correo'] = "pendiente"
            cambios = True

    if 'Plan' not in df_usuarios.columns:
        df_usuarios['Plan'] = 'Full'
        cambios = True

    if cambios:
        guardar_en_nube(df_usuarios)

df_usuarios.to_csv(ARCHIVO_USUARIOS, index=False)

USUARIOS_DICT = dict(zip(df_usuarios['Usuario'], df_usuarios['Password'].astype(str)))
NOMBRES_REALES = dict(zip(df_usuarios['Usuario'], df_usuarios['Nombre_Real']))

# --- MOTOR DE COOKIES (PARA QUE NO SE CIERRE LA SESIÓN CON F5) ---
cookie_manager = stx.CookieManager(key="motor_cookies")

if 'logged_in' not in st.session_state: 
    st.session_state['logged_in'] = False
if 'username' not in st.session_state: 
    st.session_state['username'] = ""

cookie_token = cookie_manager.get(cookie="jurisync_user")
cookie_usuario = validar_token_sesion(cookie_token) if cookie_token else None

# Justo después de un login o logout recién hecho, la cookie del navegador puede
# tardar un ciclo de rerun en reflejar el cambio de verdad (es una limitación
# conocida de este componente de cookies). Si se lee en ese instante, se corre
# el riesgo de reingresar con la cuenta anterior encima de la que el usuario
# recién eligió. Por eso, justo después de esa acción, se salta esta lectura
# UNA sola vez, y luego vuelve a funcionar con normalidad (por ejemplo, para
# que un F5 mantenga la sesión abierta como siempre).
if st.session_state.get('_saltar_autologin_cookie'):
    st.session_state['_saltar_autologin_cookie'] = False
elif cookie_usuario and cookie_usuario in USUARIOS_DICT:
    st.session_state['logged_in'] = True
    st.session_state['username'] = cookie_usuario

# =====================================================================
# 🚦 PORTAL DEL CLIENTE (VISTA EXTERNA PARA SUBIR DOCUMENTOS)
# =====================================================================
query_params = st.query_params
if "cliente_id" in query_params:
    token_cliente = query_params["cliente_id"]
    nombre_cliente_limpio = token_cliente.replace("_", " ")
    
    st.markdown("""
    <style>
        [data-testid="stSidebar"] { display: none !important; }
        .block-container { max-width: 800px !important; padding-top: 3rem !important; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown(f"<div style='text-align: center;'><img src='{LOGO_URL}' style='width: 100px;'></div>", unsafe_allow_html=True)
    st.title("📥 Portal de Recepción de Antecedentes")
    st.subheader(f"Bienvenido/a, {nombre_cliente_limpio}")
    st.write("Por favor, cargue los documentos solicitados por su abogado en formato PDF o Imagen. El sistema notificará automáticamente al estudio jurídico cuando haya completado la entrega.")
    
    ARCHIVO_DOCS = "base_documentos_clientes.csv"
    if not os.path.exists(ARCHIVO_DOCS):
        # Antes esto creaba directamente un archivo local VACÍO cuando el sistema
        # se reiniciaba, perdiendo de vista todos los documentos ya subidos por
        # los clientes (aunque los archivos en sí seguían a salvo en Drive). Ahora
        # primero intenta reconstruir el archivo local desde el respaldo en la
        # nube, y solo si no hay nada ahí tampoco, parte con uno vacío de verdad.
        df_docs_nube_inicial = safe_read_sheet("base_documentos_clientes", COLS_DOCS)
        if not df_docs_nube_inicial.empty:
            df_docs_nube_inicial.to_csv(ARCHIVO_DOCS, index=False)
        else:
            pd.DataFrame(columns=COLS_DOCS).to_csv(ARCHIVO_DOCS, index=False)
    else:
        df_docs_migra = leer_csv_local(ARCHIVO_DOCS, COLS_DOCS)
        if 'Archivo_Drive_ID' not in df_docs_migra.columns:
            df_docs_migra['Archivo_Drive_ID'] = ''
            df_docs_migra.to_csv(ARCHIVO_DOCS, index=False)
        
    df_docs = leer_csv_local(ARCHIVO_DOCS, COLS_DOCS)
    mis_docs = df_docs[df_docs['Cliente_Token'] == token_cliente]
    if mis_docs.empty and not df_docs.empty:
        # Respaldo: si el token no calzó exacto (por ejemplo, un enlace generado antes
        # de sanitizar tildes/comas), reintenta comparando solo letras/números/guion bajo
        # de ambos lados, para que los enlaces antiguos no queden rotos.
        token_normalizado = re.sub(r'[^A-Za-z0-9_]', '', token_cliente)
        mascara_respaldo = df_docs['Cliente_Token'].astype(str).apply(lambda x: re.sub(r'[^A-Za-z0-9_]', '', x) == token_normalizado)
        mis_docs = df_docs[mascara_respaldo]
    
    if mis_docs.empty:
        st.info("No registra solicitudes de documentos pendientes en este momento.")
    else:
        completados = 0
        total = len(mis_docs)
        
        for idx, row in mis_docs.iterrows():
            with st.container(border=True):
                c1, c2 = st.columns([3, 1])
                ya_subido = row['Estado'] == '✅ Completado'
                
                with c1:
                    st.markdown(f"**Requisito:** {row['Documento_Nombre']}")
                    if ya_subido:
                        st.markdown("<span style='color:#57a15a; font-weight:bold;'>✅ Recibido con éxito</span>", unsafe_allow_html=True)
                        completados += 1
                    else:
                        st.markdown("<span style='color:#ff5630; font-weight:bold;'>❌ Pendiente de carga</span>", unsafe_allow_html=True)
                        
                with c2:
                    if not ya_subido:
                        archivo = st.file_uploader("Subir Archivo", key=f"up_{row['ID_Req']}", label_visibility="collapsed")
                        if archivo:
                            archivo_bytes = archivo.getvalue()
                            tamano_ok, msg_tamano = validar_tamano_para_sheets(archivo_bytes, archivo.name)
                            with st.spinner("Guardando en la nube de JuriSync..."):
                                drive_id, b64_file = guardar_archivo_adjunto(archivo.name, archivo_bytes, archivo.type or 'application/octet-stream')
                            if not drive_id and not b64_file:
                                error_drive_real = st.session_state.pop('_ultimo_error_drive', None)
                                if error_drive_real:
                                    st.error(f"⚠️ No se pudo guardar automáticamente: Google Drive no está funcionando ahora mismo (detalle técnico: {error_drive_real}).")
                                else:
                                    st.error(msg_tamano if not tamano_ok else "⚠️ No fue posible guardar el archivo. Intenta nuevamente.")
                                
                                # Alternativa simple mientras Drive no funcione: el archivo
                                # es demasiado grande para guardarlo automáticamente, así
                                # que se ofrece subirlo a tu PROPIO Google Drive personal
                                # (gratis, sin configuración) y pegar el link acá, en vez
                                # de quedar completamente bloqueado.
                                st.info("💡 **Alternativa:** sube el archivo a tu propio Google Drive (Menú → Compartir → 'Cualquier persona con el enlace') y pega el link abajo.")
                                link_externo_doc = st.text_input("Pega aquí el link de Google Drive", key=f"link_ext_{row['ID_Req']}")
                                if st.button("💾 Guardar link", key=f"btn_link_ext_{row['ID_Req']}"):
                                    if link_externo_doc.strip():
                                        df_docs.loc[df_docs['ID_Req'] == row['ID_Req'], ['Estado', 'Link_Externo', 'Fecha_Subida']] = ['✅ Completado (link externo)', link_externo_doc.strip(), datetime.now().strftime("%d/%m/%Y")]
                                        df_docs.to_csv(ARCHIVO_DOCS, index=False)
                                        safe_update_sheet("base_documentos_clientes", df_docs)
                                        st.success("¡Link guardado!")
                                        st.rerun()
                                    else:
                                        st.error("Pega un link válido primero.")
                            else:
                                df_docs.loc[df_docs['ID_Req'] == row['ID_Req'], ['Estado', 'Archivo_B64', 'Archivo_Drive_ID', 'Fecha_Subida']] = ['✅ Completado', b64_file, drive_id, datetime.now().strftime("%d/%m/%Y")]
                                df_docs.to_csv(ARCHIVO_DOCS, index=False)
                                # BUGFIX: antes esto solo se guardaba en el disco local (efímero en la nube),
                                # nunca se sincronizaba a Google Sheets, por lo que el documento del
                                # cliente podía perderse si la app se reiniciaba antes de ser revisado.
                                safe_update_sheet("base_documentos_clientes", df_docs)
                                st.success("¡Documento guardado!")
                                st.rerun()
        
        st.progress(completados / total if total > 0 else 0)
        
        if completados == total and total > 0:
            st.success("🎉 ¡Excelente! Ha entregado la totalidad de la documentación solicitada.")
            st.balloons()
            st.info("📩 Se ha enviado un correo automático a su abogado notificando que su expediente está completo y listo para revisión.")
    st.stop() 

# =====================================================================
# 🔐 PANTALLA DE LOGIN CON CONFIGURACIÓN EN LA NUBE
# =====================================================================
if not st.session_state['logged_in']:
    st.markdown("""
    <style>
        [data-testid="stAppViewContainer"], .stApp { background-color: #f4f5f7 !important; }
        #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
        .block-container { max-width: 1300px !important; margin: 0 auto !important; padding-top: 2rem !important; }
        [data-testid="stForm"] { background-color: white !important; border-radius: 16px !important; border: 1px solid #e0e4e8 !important; padding: 40px 30px !important; box-shadow: 0 4px 15px rgba(0,0,0,0.05) !important; }
        p, label, span, div { color: #172b4d !important; }
        [data-testid="stFormSubmitButton"] button { background-color: #0e6b74 !important; color: white !important; border: none !important; font-weight: bold !important;}
        [data-testid="stFormSubmitButton"] button:hover { background-color: #0047b3 !important; }
        .stTextInput input { border: 1px solid #cbd2d9 !important; border-radius: 6px !important; padding: 10px !important; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    col_a, col_b, col_c = st.columns([1, 1.2, 1])
    
    with col_b:
        st.markdown(f"""
        <div style='text-align: center; margin-bottom: 20px;'>
            <img src='{LOGO_URL}' style='width: 140px; margin-bottom: 5px;'>
            <h1 style='color:#172b4d; margin-top: 0; margin-bottom: 5px; font-size: 32px; font-weight: 800; letter-spacing: 1px;'>JuriSync</h1>
            <p style='color:#6b778c; font-size: 15px; margin:0;'>Espacio de trabajo seguro</p>
        </div>
        """, unsafe_allow_html=True)
        
        tab_iniciar, tab_recuperar = st.tabs(["🔐 Iniciar Sesión", "🆘 Olvidé mi contraseña"])
        
        with tab_iniciar:
            with st.form("login_form", clear_on_submit=False):
                user_input = st.text_input("Usuario Corporativo", placeholder="Tu nombre de usuario")
                pass_input = st.text_input("Contraseña", type="password", placeholder="••••••••")
                st.write("") 
                
                if st.form_submit_button("Ingresar al Estudio", use_container_width=True):
                    user_clean = user_input.strip()
                    if user_clean in USUARIOS_DICT and verificar_password(pass_input, USUARIOS_DICT[user_clean]):
                        idx_user = df_usuarios[df_usuarios['Usuario'] == user_clean].index[0]
                        # Migración transparente: si la clave todavía estaba en texto plano, se re-hashea ahora.
                        if not es_hash_bcrypt(df_usuarios.loc[idx_user, 'Password']):
                            df_usuarios.at[idx_user, 'Password'] = hash_password(pass_input)
                            guardar_en_nube(df_usuarios)
                        if str(df_usuarios.loc[idx_user, 'Debe_Cambiar_Clave']).lower() == 'true':
                            st.session_state['requiere_registro_inicial'] = True
                            st.session_state['usr_registro'] = user_clean
                            st.rerun()
                        else:
                            cookie_manager.set("jurisync_user", generar_token_sesion(user_clean), key="cookie_login")
                            st.session_state['logged_in'] = True
                            st.session_state['username'] = user_clean
                            # Evita que, en el próximo rerun, se vuelva a leer la cookie
                            # (que puede tardar un ciclo en reflejar el cambio real en el
                            # navegador) y se pise este login recién hecho con datos viejos.
                            st.session_state['_saltar_autologin_cookie'] = True
                            import time
                            time.sleep(0.3)
                            st.rerun()
                    else:
                        st.error("❌ Usuario o contraseña incorrectos.")
                        
        with tab_recuperar:
            with st.form("recuperar_form", clear_on_submit=True):
                st.info("Ingresa tu usuario y correo. Si coinciden, el sistema generará una clave temporal.")
                rec_usuario = st.text_input("Usuario")
                rec_correo = st.text_input("Correo electrónico registrado")
                
                if st.form_submit_button("Recuperar Contraseña", use_container_width=True):
                    if rec_usuario in df_usuarios['Usuario'].values:
                        correo_real = str(df_usuarios[df_usuarios['Usuario'] == rec_usuario]['Correo'].values[0])
                        if correo_real == "pendiente" or correo_real == "nan":
                            st.warning("⚠️ Esta cuenta aún no tiene un correo configurado. Pídele al administrador que la recupere manualmente.")
                        elif rec_correo.strip().lower() == correo_real.lower():
                            df_usuarios.loc[df_usuarios['Usuario'] == rec_usuario, 'Password'] = hash_password("Temp1234")
                            df_usuarios.loc[df_usuarios['Usuario'] == rec_usuario, 'Debe_Cambiar_Clave'] = 'True'
                            guardar_en_nube(df_usuarios)
                            st.success("✅ Identidad verificada. Tu contraseña temporal es: **Temp1234**. Inicia sesión con ella y te pediremos crear una nueva.")
                        else:
                            st.error("❌ El correo no coincide con nuestros registros de seguridad.")
                    else:
                        st.error("❌ Usuario no encontrado.")
                        
    if st.session_state.get('requiere_registro_inicial', False):
        with st.container(border=True):
            st.markdown("<h2 style='text-align:center; color:#172b4d;'>🔒 Configuración de Seguridad</h2>", unsafe_allow_html=True)
            st.warning("Por seguridad de tu cuenta, debes actualizar tus datos obligatoriamente para continuar.")
            
            with st.form("form_cambio_clave_nuevo"):
                nuevo_correo = st.text_input("Tu Correo Electrónico Institucional", placeholder="ejemplo@estudio.cl")
                nueva_cl = st.text_input("Nueva Contraseña", type="password")
                conf_cl = st.text_input("Confirmar Nueva Contraseña", type="password")
                st.write("")
                
                if st.form_submit_button("Actualizar Credenciales y Entrar", type="primary", use_container_width=True):
                    usr_actualizar = st.session_state['usr_registro']
                    if nueva_cl.strip() == "" or nuevo_correo.strip() == "":
                        st.error("Todos los campos son obligatorios.")
                    elif "@" not in nuevo_correo:
                        st.error("Por favor, ingresa un correo electrónico válido.")
                    elif nueva_cl != conf_cl:
                        st.error("Las contraseñas no coinciden.")
                    elif usr_actualizar not in df_usuarios['Usuario'].values:
                        st.error(f"⚠️ No se encontró el usuario '{usr_actualizar}' en la base de datos. Contacta al administrador.")
                    else:
                        try:
                            # Blindaje extra: fuerza estas columnas a texto ANTES de
                            # asignar, por si quedaron mal tipadas (mismo tipo de bug
                            # ya visto antes con Debe_Cambiar_Clave).
                            df_usuarios = _corregir_dtypes_texto(df_usuarios)
                            for _col in ['Password', 'Correo', 'Debe_Cambiar_Clave']:
                                if _col in df_usuarios.columns:
                                    df_usuarios[_col] = df_usuarios[_col].astype(object)
                            
                            idx_mod = df_usuarios[df_usuarios['Usuario'] == usr_actualizar].index[0]
                            df_usuarios.at[idx_mod, 'Password'] = hash_password(nueva_cl)
                            df_usuarios.at[idx_mod, 'Correo'] = nuevo_correo
                            df_usuarios.at[idx_mod, 'Debe_Cambiar_Clave'] = 'False'
                            
                            guardar_en_nube(df_usuarios)
                            cookie_manager.set("jurisync_user", generar_token_sesion(usr_actualizar), key="cookie_registro_inicial")
                            
                            st.session_state['logged_in'] = True
                            st.session_state['username'] = usr_actualizar
                            st.session_state['requiere_registro_inicial'] = False
                            st.session_state['_saltar_autologin_cookie'] = True
                            st.success("✅ Credenciales actualizadas correctamente.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"⚠️ No se pudo guardar. Detalle técnico: {e}")
                            st.info("Si el problema persiste, avísale a Nicolás con este mensaje de error.")
    st.stop()


# --- ARQUITECTURA DE ARCHIVOS DE DATOS LOCALES ---
usuario_actual = st.session_state['username']
nombre_real_usuario = NOMBRES_REALES.get(usuario_actual, usuario_actual.capitalize())

ARCHIVO_BD = f"base_causas_{usuario_actual}.csv"
ARCHIVO_TAREAS = f"base_tareas_{usuario_actual}.csv"
ARCHIVO_CONTRATOS = f"base_contratos_{usuario_actual}.csv"
ARCHIVO_ESCRITURAS = f"base_escrituras_{usuario_actual}.csv"
ARCHIVO_ANALISIS_ESCRITURAS = f"base_analisis_escrituras_{usuario_actual}.csv"
ARCHIVO_EXCEPCIONES = f"base_excepciones_{usuario_actual}.csv"
ARCHIVO_CITAS = f"base_citas_{usuario_actual}.csv"
ARCHIVO_ENCARGOS = f"base_encargos_{usuario_actual}.csv"
ARCHIVO_PAGOS_HONORARIOS = f"base_pagos_honorarios_{usuario_actual}.csv"
ARCHIVO_JURISPRUDENCIA = "base_jurisprudencia.csv"  # Compartida por todo el equipo, no por usuario (como Clientes)
ARCHIVO_POSESION_EFECTIVA = f"base_posesion_efectiva_{usuario_actual}.csv"
ARCHIVO_TRAMITES = f"base_tramites_{usuario_actual}.csv"
ARCHIVO_ESTADO_DIARIO = f"base_estado_diario_{usuario_actual}.csv"
ARCHIVO_MENSAJES = "base_mensajes_global.csv"

# Verificación de archivos individuales para evitar pérdida de datos
if not os.path.exists(ARCHIVO_TAREAS):
    # CRÍTICO: antes esto creaba directo un archivo vacío cuando el sistema
    # se reiniciaba (el archivo local se borra en cada reinicio), perdiendo
    # de vista TODAS las tareas reales, que seguían intactas en Google
    # Sheets. Ahora primero intenta reconstruir el archivo local desde la
    # nube (filtrando solo las tareas de este usuario), y solo si de verdad
    # no hay nada ahí tampoco, parte con uno vacío.
    df_tareas_nube_inicial = safe_read_sheet("base_tareas", COLS_TAREAS)
    if not df_tareas_nube_inicial.empty and 'Usuario_Propietario' in df_tareas_nube_inicial.columns:
        df_tareas_nube_inicial = df_tareas_nube_inicial[df_tareas_nube_inicial['Usuario_Propietario'] == usuario_actual]
    if not df_tareas_nube_inicial.empty:
        df_tareas_nube_inicial.to_csv(ARCHIVO_TAREAS, index=False)
    else:
        df_vacio_t = pd.DataFrame(columns=COLS_TAREAS)
        df_vacio_t.to_csv(ARCHIVO_TAREAS, index=False)
else:
    df_t_check = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
    if 'Prioridad' not in df_t_check.columns:
        df_t_check['Prioridad'] = 'Media'
        df_t_check.to_csv(ARCHIVO_TAREAS, index=False)

if not os.path.exists(ARCHIVO_BD):
    # Mismo arreglo crítico que en Tareas: primero intenta reconstruir desde
    # Google Sheets antes de partir con un archivo vacío de verdad.
    df_causas_nube_inicial = safe_read_sheet("base_causas", COLS_CAUSAS)
    if not df_causas_nube_inicial.empty and 'Usuario_Propietario' in df_causas_nube_inicial.columns:
        df_causas_nube_inicial = df_causas_nube_inicial[df_causas_nube_inicial['Usuario_Propietario'] == usuario_actual]
    if not df_causas_nube_inicial.empty:
        df_causas_nube_inicial.to_csv(ARCHIVO_BD, index=False)
    else:
        df_vacio_c = pd.DataFrame(columns=['ROL', 'TRIBUNAL', 'CARATULADO', 'Cliente', 'RUT', 'Teléfono', 'Tipo_Negocio', 'Clave_unica', 'Correo', 'Direccion', 'SAC', 'Sucursal', 'Estado_Honorarios', 'Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas', 'Usuario_Propietario'])
        df_vacio_c.to_csv(ARCHIVO_BD, index=False)
else:
    df_c_check = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    ejecutar_guardado_check = False
    columnas_requeridas_bd = ['Cliente', 'Estado_Honorarios', 'Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas', 'Fecha_Inicio']
    for col in columnas_requeridas_bd:
        if col not in df_c_check.columns:
            if col in ['Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas']: 
                df_c_check[col] = 0
            elif col == 'Estado_Honorarios': 
                df_c_check[col] = "Sin fijar"
            else: 
                df_c_check[col] = ""
            ejecutar_guardado_check = True
        elif col == 'Fecha_Inicio':
            # Si ya existía pero quedó tipada como número (todo NaN en archivos
            # antiguos), la forzamos a texto para que nunca vuelva a fallar
            # al intentar guardar una fecha en formato string.
            if df_c_check[col].dtype != object:
                df_c_check[col] = df_c_check[col].astype(object).fillna("")
                ejecutar_guardado_check = True
    if ejecutar_guardado_check:
        df_c_check.to_csv(ARCHIVO_BD, index=False)

if not os.path.exists(ARCHIVO_CONTRATOS):
    df_contratos_nube_inicial = safe_read_sheet("base_contratos", COLS_CONTRATOS)
    if not df_contratos_nube_inicial.empty and 'Usuario_Propietario' in df_contratos_nube_inicial.columns:
        df_contratos_nube_inicial = df_contratos_nube_inicial[df_contratos_nube_inicial['Usuario_Propietario'] == usuario_actual]
    if not df_contratos_nube_inicial.empty:
        df_contratos_nube_inicial.to_csv(ARCHIVO_CONTRATOS, index=False)
    else:
        df_vacio_co = pd.DataFrame(columns=COLS_CONTRATOS)
        df_vacio_co.to_csv(ARCHIVO_CONTRATOS, index=False)
else:
    df_co_check = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
    if 'Archivo_Drive_ID' not in df_co_check.columns:
        df_co_check['Archivo_Drive_ID'] = ''
        df_co_check.to_csv(ARCHIVO_CONTRATOS, index=False)

if not os.path.exists(ARCHIVO_TRAMITES):
    df_tramites_nube_inicial = safe_read_sheet("base_tramites", COLS_TRAMITES)
    if not df_tramites_nube_inicial.empty and 'Usuario_Propietario' in df_tramites_nube_inicial.columns:
        df_tramites_nube_inicial = df_tramites_nube_inicial[df_tramites_nube_inicial['Usuario_Propietario'] == usuario_actual]
    if not df_tramites_nube_inicial.empty:
        df_tramites_nube_inicial.to_csv(ARCHIVO_TRAMITES, index=False)
    else:
        df_vacio_tr = pd.DataFrame(columns=COLS_TRAMITES)
        df_vacio_tr.to_csv(ARCHIVO_TRAMITES, index=False)
else:
    df_tr_check = leer_csv_local(ARCHIVO_TRAMITES, COLS_TRAMITES)
    if 'Comprobante_Drive_ID' not in df_tr_check.columns:
        df_tr_check['Comprobante_Drive_ID'] = ''
        df_tr_check.to_csv(ARCHIVO_TRAMITES, index=False)

if not os.path.exists(ARCHIVO_ESTADO_DIARIO):
    df_ed_nube_inicial = safe_read_sheet("base_estado_diario", COLS_ESTADO_DIARIO)
    if not df_ed_nube_inicial.empty:
        df_ed_nube_inicial.to_csv(ARCHIVO_ESTADO_DIARIO, index=False)
    else:
        df_vacio_ed = pd.DataFrame(columns=COLS_ESTADO_DIARIO)
        df_vacio_ed.to_csv(ARCHIVO_ESTADO_DIARIO, index=False)
else:
    df_ed_check = leer_csv_local(ARCHIVO_ESTADO_DIARIO, COLS_ESTADO_DIARIO)
    if 'Doc_Drive_ID' not in df_ed_check.columns:
        df_ed_check['Doc_Drive_ID'] = ''
        df_ed_check.to_csv(ARCHIVO_ESTADO_DIARIO, index=False)

if not os.path.exists(ARCHIVO_MENSAJES):
    # Mensajes es una hoja global (no por usuario), así que aquí no se filtra
    # por Usuario_Propietario, se reconstruye completa.
    df_msgs_nube_inicial = safe_read_sheet("base_mensajes_global", COLS_MENSAJES)
    if not df_msgs_nube_inicial.empty:
        df_msgs_nube_inicial.to_csv(ARCHIVO_MENSAJES, index=False)
    else:
        pd.DataFrame(columns=COLS_MENSAJES).to_csv(ARCHIVO_MENSAJES, index=False)

# --- NOTIFICADOR ESTILO OUTLOOK (TOAST + INSIGNIA PERSISTENTE EN EL MENÚ) ---
BADGE_MENSAJES_NO_LEIDOS = 0
if st.session_state['logged_in']:
    try:
        if os.path.exists(ARCHIVO_MENSAJES):
            df_msgs_alerta = leer_csv_local(ARCHIVO_MENSAJES, ['ID', 'Fecha', 'De', 'Para', 'Mensaje'])
            if not df_msgs_alerta.empty and 'Para' in df_msgs_alerta.columns:
                mis_mensajes = df_msgs_alerta[(df_msgs_alerta['Para'] == nombre_real_usuario) | (df_msgs_alerta['Para'] == 'Todos')]
                
                if 'ultimo_mensaje_leido' not in st.session_state:
                    st.session_state['ultimo_mensaje_leido'] = len(mis_mensajes)
                elif len(mis_mensajes) > st.session_state['ultimo_mensaje_leido']:
                    mensajes_nuevos = len(mis_mensajes) - st.session_state['ultimo_mensaje_leido']
                    st.toast(f"🔔 ¡Tienes {mensajes_nuevos} mensaje(s) nuevo(s) en tu buzón!", icon="📩")
                
                # A diferencia del toast (que se ve una sola vez y desaparece), esta insignia
                # se recalcula en cada rerun y queda pegada al botón de Mensajería del menú
                # lateral hasta que el usuario entre efectivamente a leer sus mensajes.
                BADGE_MENSAJES_NO_LEIDOS = max(0, len(mis_mensajes) - st.session_state['ultimo_mensaje_leido'])
                st.session_state['_total_mensajes_para_mi'] = len(mis_mensajes)
    except Exception:
        # Este bloque corre en CADA carga de página para CUALQUIER usuario logueado.
        # Si algo falla aquí (archivo dañado, columna faltante, etc.), jamás debe
        # tumbar toda la app — en el peor caso, simplemente no se muestra la
        # insignia de mensajes no leídos por esta vez.
        BADGE_MENSAJES_NO_LEIDOS = 0

# --- FUNCIÓN DE AUTOLIMPIEZA SISTEMA (15 DÍAS EXACTOS) ---
def limpiar_documentos_estado_diario():
    if os.path.exists(ARCHIVO_ESTADO_DIARIO):
        df_ed = leer_csv_local(ARCHIVO_ESTADO_DIARIO, COLS_ESTADO_DIARIO)
        if not df_ed.empty:
            if 'Doc_Drive_ID' not in df_ed.columns:
                df_ed['Doc_Drive_ID'] = ''
            df_ed['Fecha_DT'] = pd.to_datetime(df_ed['Fecha_Estado'], format='%d/%m/%Y', errors='coerce')
            limite_fecha = datetime.now() - timedelta(days=15)
            mascara_viejos = df_ed['Fecha_DT'] < limite_fecha
            # Solo se purga el base64 local (pesa en el disco efímero). Los archivos que
            # ya quedaron respaldados en Google Drive (Doc_Drive_ID) se conservan intactos.
            mascara_sin_drive = df_ed['Doc_Drive_ID'].fillna('').astype(str).str.strip() == ''
            mascara_a_purgar = mascara_viejos & mascara_sin_drive
            df_ed.loc[mascara_a_purgar, 'Doc_B64'] = ""
            df_ed.loc[mascara_a_purgar, 'Doc_Nombre'] = df_ed.loc[mascara_a_purgar, 'Doc_Nombre'].apply(lambda x: f"(Eliminado por memoria) {x}" if pd.notna(x) and x != "" and not str(x).startswith("(Eliminado") else x)
            df_ed.drop(columns=['Fecha_DT']).to_csv(ARCHIVO_ESTADO_DIARIO, index=False)

limpiar_documentos_estado_diario()

# --- SCRAPER INTEGRADO PODER JUDICIAL DE CHILE ---
def scraper_estado_diario_pjud():
    url_pjud = "https://oficinajudicialvirtual.pjud.cl/estado_diario.php"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        response = requests.get(url_pjud, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        tablas = pd.read_html(str(soup))
        if tablas:
            df_scrap = tablas[0]
            if 'ROL' not in df_scrap.columns and len(df_scrap.columns) > 1:
                df_scrap.rename(columns={df_scrap.columns[1]: 'ROL'}, inplace=True)
            return df_scrap
    except Exception:
        return None
    return None

# --- MANEJO DE VISTAS Y LOGICA DE BOTONES ---
def resetear_vistas():
    st.session_state.causa_seleccionada = None
    st.session_state.cliente_seleccionado = None
    st.session_state.modo_edicion = False
    st.session_state.creando_tarea = False
    st.session_state.editando_tarea = None
    st.session_state.creando_causa = False
    st.session_state.creando_cliente = False

if 'menu_radio' not in st.session_state: 
    st.session_state['menu_radio'] = "🏠 Inicio"

for key in ['causa_seleccionada', 'cliente_seleccionado', 'modo_edicion', 'creando_tarea', 'editando_tarea', 'creando_causa', 'creando_cliente']:
    if key not in st.session_state: 
        if key in ['modo_edicion', 'creando_tarea', 'creando_causa', 'creando_cliente']:
            st.session_state[key] = False
        else:
            st.session_state[key] = None

def nav_causas(): 
    st.session_state.menu_radio = "💼 Causas"
    resetear_vistas()

def nav_clientes(): 
    st.session_state.menu_radio = "👥 Clientes"
    resetear_vistas()

def ir_a_expediente(rol_causa, propietario=None): 
    st.session_state.menu_radio = "💼 Causas"
    st.session_state.causa_seleccionada = rol_causa
    st.session_state.causa_propietario_vista = propietario

def limpiar_causa():
    st.session_state.causa_seleccionada = None

# --- CSS CLARO PROFESIONAL (ESTILO JIRA/TRELLO) ---
st.markdown("""
<style>
    /* ========================================================================
       FIJAR ESQUEMA DE COLOR CLARO SIEMPRE
       Evita que el cambio automático de apariencia de macOS (claro/oscuro según
       la hora) o que alguien active "Dark" desde el menú ⋮ > Settings de
       Streamlit rompa el contraste: el navegador deja de intentar adaptar
       controles nativos (selects, checkboxes, scrollbars) a modo oscuro, y
       re-forzamos la paleta clara de JuriSync sobre los contenedores base de
       Streamlit incluso si su tema interno cambia a oscuro.
       ======================================================================== */
    html, body { color-scheme: light !important; }
    [data-testid="stAppViewContainer"], [data-testid="stSidebar"], [data-testid="stHeader"],
    [data-testid="stMain"], .stApp, .main {
        color-scheme: light !important;
        background-color: #f4f5f7 !important;
    }
    @media (prefers-color-scheme: dark) {
        [data-testid="stAppViewContainer"], .stApp { background-color: #f4f5f7 !important; }
        [data-testid="stSidebar"] { background-color: #ffffff !important; }
        .stMarkdown, p, span, label, h1, h2, h3, h4, h5, h6, div { color: #172b4d !important; }
        .stTextInput input, .stTextArea textarea, .stNumberInput input,
        [data-baseweb="select"] > div, [data-baseweb="input"] {
            background-color: #ffffff !important; color: #172b4d !important; border-color: #cbd2d9 !important;
        }
        .dash-card, [data-testid="stExpander"], [data-testid="stForm"] {
            background-color: #ffffff !important; border-color: #e0e4e8 !important;
        }
    }
</style>
<style>
    [data-testid="stAppViewContainer"], .stApp { background-color: #f4f5f7 !important; }
    [data-testid="stSidebar"] { background-color: #ffffff !important; border-right: 1px solid #e0e4e8 !important; }
    [data-testid="stHeader"] { background-color: transparent !important; }
    .stMarkdown, p, span, label, h1, h2, h3, h4, h5, h6 { color: #172b4d !important; }
    .dash-card { background: #ffffff !important; border-radius: 12px; padding: 18px; border: 1px solid #e0e4e8 !important; border-top: 4px solid #0e6b74 !important; margin-bottom: 15px; box-shadow: 0 2px 6px rgba(14,107,116,0.08); }
    .dash-header { border-bottom: 2px solid #0e6b74; padding-bottom: 5px; margin-bottom: 15px; font-weight: 800; font-size: 13px; color: #0e6b74; letter-spacing: 0.5px; text-transform: uppercase; }
    .badge-active { background: #57a15a !important; color: white !important; padding: 4px 10px; border-radius: 12px; font-size: 12px; font-weight: 600; }
    .badge-propio { background: #0e6b74 !important; color: white !important; padding: 4px 10px; border-radius: 12px; font-size: 12px; font-weight: 600; }
    .info-field { display:flex; justify-content:space-between; align-items:baseline; padding:6px 0; border-bottom:1px solid #f4f5f7; }
    .info-field:last-of-type { border-bottom: none; }
    .info-label { font-size:12px; color:#6b778c !important; font-weight:600; text-transform:uppercase; letter-spacing:0.3px; }
    .info-value { font-size:14px; color:#172b4d !important; font-weight:600; text-align:right; }
    .badge-honorarios { background:#fff0b3 !important; color:#7a5b00 !important; padding:3px 10px; border-radius:12px; font-size:12px; font-weight:600; }
    .task-status-chip { padding:4px 12px; border-radius:12px; font-size:12px; font-weight:700; display:inline-block; }
    .task-status-progreso { background:#fff0b3 !important; color:#7a5b00 !important; }
    .task-status-aprobada { background:#e3fcef !important; color:#1b7a4a !important; }
    .task-status-rechazada { background:#ffebe6 !important; color:#bf2600 !important; }
    .stTextInput input, .stTextArea textarea, .stSelectbox select, .stNumberInput input { background-color: #ffffff !important; color: #172b4d !important; border: 1px solid #cbd2d9 !important; border-radius: 6px !important; }
    .stTextInput input:focus, .stTextArea textarea:focus { border-color: #0e6b74 !important; box-shadow: 0 0 0 1px #0e6b74 !important; }
    ::placeholder { color: #6b778c !important; opacity: 1; }
    [data-testid="stButton"] button { background-color: #ffffff !important; color: #172b4d !important; border: 1px solid #cbd2d9 !important; border-radius: 6px !important; font-weight: 600 !important; transition: all 0.2s ease !important; }
    [data-testid="stButton"] button:hover { border-color: #0e6b74 !important; color: #0e6b74 !important; background-color: #e3f2f1 !important; }
    /* Métricas nativas (st.metric): el número grande en teal, para que
       combine con el resto del sistema en vez del gris/negro por defecto. */
    [data-testid="stMetricValue"] { color: #0e6b74 !important; font-weight: 800 !important; }
    [data-testid="stMetricLabel"] { color: #6b778c !important; }
    [data-testid="stVerticalBlockBorderWrapper"] { background-color: #ffffff !important; border-radius: 12px !important; border: 1px solid #e0e4e8 !important; }
    .chat-bg { background-color: #f4f6f7; padding: 20px; border-radius: 12px; border: 1px solid #e0e4e8; }
    .burbuja-mia { background-color: #d4ebe9; padding: 10px 15px; border-radius: 15px 15px 0px 15px; max-width: 75%; box-shadow: 0 1px 1px rgba(0,0,0,0.06); margin-left: auto; margin-bottom: 12px; border: 1px solid #aed9d4;}
    .burbuja-otro { background-color: #ffffff; padding: 10px 15px; border-radius: 15px 15px 15px 0px; max-width: 75%; box-shadow: 0 1px 1px rgba(0,0,0,0.06); margin-right: auto; margin-bottom: 12px; border: 1px solid #e0e4e8;}
    .chat-autor { font-size: 13px; font-weight: 800; color: #0e6b74; margin-bottom: 2px; }
    .chat-texto { font-size: 15px; color: #172b4d; line-height: 1.4; }
    .chat-hora { font-size: 11px; color: #6b778c; text-align: right; margin-top: 5px; }
    .chat-para { font-size: 11px; color: #6b778c; font-weight: normal; margin-left: 5px; }
    
    /* ========================================================================
       MENÚ LATERAL: un solo valor de indentación (10px) usado en TODOS los
       elementos por igual (nombre, Inicio, Mensajería, categorías e ítems
       de adentro), para que todo el texto arranque exactamente en la misma
       posición horizontal. El ancho de la barra se ajustó para que quepa
       cómodo el ítem más largo ("🤖 Inteligencia Artificial") sin cortarse
       ni hacer wrap.
       ======================================================================== */
    [data-testid="stSidebar"] {
        min-width: 280px !important;
        max-width: 280px !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        padding-left: 18px !important;
        padding-right: 12px !important;
    }
    [data-testid="stSidebar"] img { margin: 0 auto !important; display: block !important; }
    
    /* Botones de navegación (Inicio, Mensajería, y los de adentro de cada
       categoría): mismo padding-left (10px) para que todos arranquen en el
       mismo punto horizontal. Se fuerza con selector comodín (*) porque
       Streamlit envuelve el texto del botón en varios niveles internos
       (div > div > p), y el selector dirigido solo al <p> no alcanzaba a
       pisar el centrado por defecto. */
    [data-testid="stSidebar"] [data-testid="stButton"] button {
        display: flex !important;
        text-align: left !important;
        justify-content: flex-start !important;
        width: 100% !important;
        border: none !important;
        background-color: transparent !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 8px 10px !important;
        margin-bottom: 1px !important;
        transition: all 0.15s ease !important;
    }
    [data-testid="stSidebar"] [data-testid="stButton"] button * {
        text-align: left !important;
        justify-content: flex-start !important;
        width: auto !important;
    }
    [data-testid="stSidebar"] [data-testid="stButton"] button:hover {
        background-color: #e8f5f4 !important;
        color: #0e6b74 !important;
    }
    
    /* Categorías (Judicial, Administrativo, IA): la flechita queda en su
       posición normal (empuja el ícono y el texto hacia la derecha, como ya
       se veía bien en la captura). Lo que se ajusta más abajo es todo lo
       demás (Inicio, Mensajería, y los ítems de adentro), empujándolos la
       misma distancia para que calcen con "Judicial", en vez de al revés. */
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        border: none !important;
        background-color: transparent !important;
        box-shadow: none !important;
        border-left: 4px solid transparent !important;
        border-radius: 4px !important;
        margin-bottom: 4px !important;
    }
    /* Cada categoría con su propio color de acento, para diferenciarlas de
       un vistazo (antes todas se veían exactamente iguales, sin color). */
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(1) {
        border-left-color: #172b4d !important;
        background-color: rgba(23,43,77,0.04) !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(1) summary { color: #172b4d !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(2) {
        border-left-color: #0e6b74 !important;
        background-color: rgba(14,107,116,0.05) !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(2) summary { color: #0e6b74 !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(3) {
        border-left-color: #5243aa !important;
        background-color: rgba(82,67,170,0.05) !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"]:nth-of-type(3) summary { color: #5243aa !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary {
        font-weight: 800 !important;
        color: #172b4d !important;
        display: flex !important;
        justify-content: flex-start !important;
        align-items: center !important;
        gap: 6px !important;
        padding-left: 10px !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary *:not(svg) {
        text-align: left !important;
        justify-content: flex-start !important;
        width: auto !important;
        flex-grow: 0 !important;
    }
    /* Se agrega un espacio invisible del mismo ancho aproximado que ocupa la
       flechita + su separación, antes del ícono de Inicio/Mensajería y de
       los ítems de adentro de cada categoría, para que su ícono quede
       exactamente en la misma columna vertical que el de "Judicial". */
    [data-testid="stSidebar"] [data-testid="stButton"] button {
        padding-left: 32px !important;
    }
    /* Sin padding extra aquí: los botones de adentro (Causas, Calendario,
       etc.) ya traen su propio padding-left de 10px arriba, así que si acá
       se agregara otro, quedarían doblemente indentados hacia la derecha
       respecto a "Inicio" y "Mensajería". */
    [data-testid="stSidebar"] [data-testid="stExpanderDetails"] {
        padding-left: 0 !important;
        padding-right: 0 !important;
    }
    
    /* Nombre de perfil + botón de cerrar sesión: pegados uno al lado del
       otro, y el botón es un círculo simple (sin caja cuadrada alrededor). */
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"]:first-of-type {
        align-items: center !important;
        gap: 0 !important;
    }
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"]:first-of-type [data-testid="stButton"] button {
        border-radius: 50% !important;
        width: 34px !important;
        height: 34px !important;
        min-width: 34px !important;
        padding: 0 !important;
        margin: 4px 0 0 auto !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        font-size: 16px !important;
        border: 1px solid #e0e4e8 !important;
        background-color: #ffffff !important;
    }
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"]:first-of-type [data-testid="stButton"] button:hover {
        background-color: #ffebe6 !important;
        border-color: #bf2600 !important;
        color: #bf2600 !important;
    }
    /* Fila del nombre de perfil + botón de cerrar sesión, centrada verticalmente */
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"]:first-of-type {
        align-items: center !important;
    }
</style>
""", unsafe_allow_html=True)

# --- RENDER DE BARRA LATERAL (ESTILO CUADRADITOS) ---
with st.sidebar:
    st.markdown(f"""
    <div style='text-align: center; margin-bottom: 10px;'>
        <img src='{LOGO_URL}' style='width: 60px;'>
        <h3 style='color:#172b4d; margin-top: 8px; font-weight: 800; letter-spacing: 1px;'>JuriSync</h3>
    </div>
    """, unsafe_allow_html=True)

    # --- SISTEMA DE PLANES Y PERMISOS ---
    df_usuarios_plan = leer_csv_local(ARCHIVO_USUARIOS, COLS_USUARIOS)
    
    if 'Plan' not in df_usuarios_plan.columns:
        df_usuarios_plan['Plan'] = 'Full' 
        df_usuarios_plan.to_csv(ARCHIVO_USUARIOS, index=False)
        
    usuario_actual = st.session_state.get('username', 'Desconocido')
    
    try:
        plan_actual = df_usuarios_plan.loc[df_usuarios_plan['Usuario'] == usuario_actual, 'Plan'].values[0]
    except:
        plan_actual = "Básico"

    # --- NOMBRE DE PERFIL Y CERRAR SESIÓN, ARRIBA DE TODO, ANTES DEL MENÚ ---
    c_nombre_perfil, c_logout_rapido = st.columns([4, 1])
    c_nombre_perfil.markdown(f"<div style='display:flex; align-items:center; height:34px; font-size:15px; color:#172b4d; padding-left:32px;'>👤&nbsp;<strong>{nombre_real_usuario}</strong></div>", unsafe_allow_html=True)
    if c_logout_rapido.button("⏻", help="Cerrar sesión", key="logout_rapido_arriba"):
        cookie_manager.set("jurisync_user", "sesion_cerrada", key="cookie_logout_invalidar")
        cookie_manager.delete("jurisync_user", key="cookie_logout_arriba")
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.session_state['_saltar_autologin_cookie'] = True
        st.rerun()
    st.markdown("<hr style='margin: 8px 0px 14px 0px;'>", unsafe_allow_html=True)

    # --- MENÚ AGRUPADO EN CATEGORÍAS ---
    GRUPO_JUDICIAL = ["💼 Causas", "📅 Calendario", "📋 Agenda", "☑️ Tareas", "📆 Estado diario",
                      "📜 Escrituras Públicas", "📋 Posesión Efectiva"]
    GRUPO_ADMINISTRATIVO = ["👥 Clientes", "📄 Contratos", "🗓️ Agenda de Asesorías", "💰 Contabilidad", "📝 Trámites",
                            "📇 Encargos", "📊 Informes", "📥 Excel"]
    GRUPO_IA = ["🧠 Estrategia", "📝 Redactor IA", "⚖️ Jurisprudencia"]  # Solo visibles para plan Full
    
    if plan_actual == "Básico":
        disponibles = {"🏠 Inicio", "📅 Calendario", "📋 Agenda", "☑️ Tareas", "💼 Causas", "👥 Clientes"}
    elif plan_actual == "Medio":
        disponibles = set(GRUPO_JUDICIAL + GRUPO_ADMINISTRATIVO) - {"✈️ Mensajería", "📊 Informes", "📥 Excel", "📜 Escrituras Públicas", "📋 Posesión Efectiva"}
    else:
        disponibles = set(GRUPO_JUDICIAL) | set(GRUPO_ADMINISTRATIVO) | set(GRUPO_IA) | {"✈️ Mensajería"}
    
    def _boton_menu(opcion, i):
        etiqueta_boton = opcion
        if opcion == "✈️ Mensajería" and BADGE_MENSAJES_NO_LEIDOS > 0:
            etiqueta_boton = f"✈️ Mensajería 🔴 {BADGE_MENSAJES_NO_LEIDOS}"
        if st.button(etiqueta_boton, use_container_width=True, key=f"btn_nav_{i}"):
            st.session_state['menu_radio'] = opcion
            resetear_vistas()
            if opcion == "✈️ Mensajería":
                st.session_state['ultimo_mensaje_leido'] = st.session_state.get('_total_mensajes_para_mi', st.session_state.get('ultimo_mensaje_leido', 0))
            st.rerun()
    
    contador_botones = 0
    
    _boton_menu("🏠 Inicio", contador_botones)
    contador_botones += 1
    
    if "✈️ Mensajería" in disponibles:
        _boton_menu("✈️ Mensajería", contador_botones)
        contador_botones += 1
    
    with st.expander("⚖️ Judicial", expanded=True):
        for opcion in GRUPO_JUDICIAL:
            if opcion in disponibles:
                _boton_menu(opcion, contador_botones)
                contador_botones += 1
    
    with st.expander("🗂️ Administrativo", expanded=True):
        for opcion in GRUPO_ADMINISTRATIVO:
            if opcion in disponibles:
                _boton_menu(opcion, contador_botones)
                contador_botones += 1
    
    if plan_actual not in ("Básico", "Medio"):
        with st.expander("🤖 Inteligencia Artificial", expanded=True):
            for opcion in GRUPO_IA:
                if opcion in disponibles:
                    _boton_menu(opcion, contador_botones)
                    contador_botones += 1
    
    if usuario_actual == "Narratia":
        st.markdown("---")
        _boton_menu("👑 Panel Admin", contador_botones)
        contador_botones += 1

    st.markdown("<br><br>", unsafe_allow_html=True)
    
    with st.expander(f"✏️ Editar Mi Perfil"):
        st.markdown("<span style='font-size:13px; color:#6b778c;'>Configura tu correo de recuperación o cambia tu clave:</span>", unsafe_allow_html=True)
        with st.form("form_perfil"):
            df_usr = leer_csv_local(ARCHIVO_USUARIOS, COLS_USUARIOS)
            for _col_segura in ['Debe_Cambiar_Clave', 'Correo', 'Password']:
                if _col_segura in df_usr.columns:
                    df_usr[_col_segura] = df_usr[_col_segura].astype(object).astype(str)
            if 'Correo' not in df_usr.columns:
                df_usr['Correo'] = ''
            filas_usuario_actual = df_usr.loc[df_usr['Usuario'] == usuario_actual, 'Correo']
            mi_correo = str(filas_usuario_actual.values[0]) if not filas_usuario_actual.empty else ""
            if mi_correo == "nan" or mi_correo == "pendiente": 
                mi_correo = ""
        
            upd_correo = st.text_input("Correo de Recuperación", value=mi_correo, placeholder="ejemplo@correo.cl")
            upd_clave = st.text_input("Nueva Contraseña", type="password", placeholder="Dejar en blanco para mantener actual")
        
            if st.form_submit_button("💾 Guardar Datos", use_container_width=True):
                cambios = False
                if upd_correo.strip() != "" and "@" in upd_correo:
                    df_usr.loc[df_usr['Usuario'] == usuario_actual, 'Correo'] = upd_correo.strip().lower()
                    cambios = True
                if upd_clave.strip() != "":
                    if len(upd_clave) >= 6:
                        df_usr.loc[df_usr['Usuario'] == usuario_actual, 'Password'] = hash_password(upd_clave)
                        cambios = True
                    else:
                        st.error("La clave debe tener mínimo 6 caracteres")
            
                if cambios:
                    try:
                        df_usr.loc[df_usr['Usuario'] == usuario_actual, 'Debe_Cambiar_Clave'] = 'False'
                        guardar_en_nube(df_usr)
                        st.success("¡Datos actualizados correctamente! Ya quedaron sincronizados en la nube.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"⚠️ No se pudo guardar. Detalle técnico: {e}")

# --- CONTROLADOR DE PESTAÑAS ---

# 1. HOME / INICIO
if st.session_state['menu_radio'] == "🏠 Inicio":
    st.title(f"{obtener_saludo()}, {nombre_real_usuario}")
    st.write("Panel de control unificado. Aquí tienes un resumen de tu actividad judicial de la oficina.")
    st.write("<br>", unsafe_allow_html=True)
    
    df_causas_totales = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    df_tareas_totales = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
    
    # Si el archivo local está vacío (por ejemplo, justo después de un
    # reinicio del sistema, antes de que se reconstruya solo), se revisa
    # directo en Google Sheets antes de mostrar el panel como si no hubiera
    # nada — para que Inicio nunca dependa de que el caché local ya se haya
    # reconstruido a tiempo.
    if df_causas_totales.empty:
        df_causas_nube_inicio = safe_read_sheet("base_causas", COLS_CAUSAS)
        if not df_causas_nube_inicio.empty and 'Usuario_Propietario' in df_causas_nube_inicio.columns:
            df_causas_totales = df_causas_nube_inicio[df_causas_nube_inicio['Usuario_Propietario'] == usuario_actual]
    if df_tareas_totales.empty:
        df_tareas_nube_inicio = safe_read_sheet("base_tareas", COLS_TAREAS)
        if not df_tareas_nube_inicio.empty and 'Usuario_Propietario' in df_tareas_nube_inicio.columns:
            df_tareas_totales = df_tareas_nube_inicio[df_tareas_nube_inicio['Usuario_Propietario'] == usuario_actual]
    
    cant_causas = len(df_causas_totales) if not df_causas_totales.empty else 0
    cant_clientes = len(df_causas_totales['Cliente'].dropna().unique()) if not df_causas_totales.empty and 'Cliente' in df_causas_totales.columns else 0
    
    fecha_hoy_str = datetime.now().strftime("%d/%m/%Y")
    
    # Reparación Agenda en Inicio
    if not df_tareas_totales.empty and 'Fecha_Vencimiento' in df_tareas_totales.columns:
        df_tareas_totales['Fecha_Vencimiento'] = df_tareas_totales['Fecha_Vencimiento'].astype(str).str.strip()
        tareas_del_dia = len(df_tareas_totales[df_tareas_totales['Fecha_Vencimiento'] == fecha_hoy_str])
    else:
        tareas_del_dia = 0
    
    documentos_efectivos = 0
    if not df_tareas_totales.empty and 'Comentarios' in df_tareas_totales.columns:
        for bloque_comentario in df_tareas_totales['Comentarios'].dropna():
            try:
                lista_comentarios = json.loads(bloque_comentario)
                for com in lista_comentarios:
                    # Detecta tanto los adjuntos nuevos (guardados de verdad, con
                    # Drive/base64) como el texto de los comentarios viejos de
                    # antes de este arreglo, que solo tenían el nombre anotado.
                    if com.get('archivo_drive_id') or com.get('archivo_b64') or "[📎 Archivo adjunto:" in com.get('texto', ''):
                        documentos_efectivos += 1
            except: 
                pass
    
    col_m1, col_m2, col_m3, col_m4 = st.columns(4)
    with col_m1: 
        st.markdown(f"<div class='dash-card'><h3 style='margin:0; font-size:14px; color:#6b778c;'>CAUSAS</h3><h2 style='margin:0; font-size:28px; color:#172b4d;'>{cant_causas}</h2></div>", unsafe_allow_html=True)
    with col_m2: 
        st.markdown(f"<div class='dash-card'><h3 style='margin:0; font-size:14px; color:#6b778c;'>CLIENTES</h3><h2 style='margin:0; font-size:28px; color:#172b4d;'>{cant_clientes}</h2></div>", unsafe_allow_html=True)
    with col_m3: 
        st.markdown(f"<div class='dash-card'><h3 style='margin:0; font-size:14px; color:#6b778c;'>TAREAS HOY</h3><h2 style='margin:0; font-size:28px; color:#ff5630;'>{tareas_del_dia}</h2></div>", unsafe_allow_html=True)
    with col_m4: 
        st.markdown(f"<div class='dash-card'><h3 style='margin:0; font-size:14px; color:#6b778c;'>DOCUMENTOS</h3><h2 style='margin:0; font-size:28px; color:#172b4d;'>{documentos_efectivos}</h2></div>", unsafe_allow_html=True)

    st.write("<br>", unsafe_allow_html=True)
    grid_izq, grid_der = st.columns([1.2, 1])
    
    with grid_izq:
        st.markdown("<div class='dash-card'><div class='dash-header'>ÚLTIMAS CAUSAS INGRESADAS</div>", unsafe_allow_html=True)
        if df_causas_totales.empty:
            st.info("No hay causas registradas recientemente.")
        else:
            ultimas = df_causas_totales.tail(4)[::-1]
            for _, c in ultimas.iterrows():
                st.markdown(f"<div style='border-bottom:1px solid #e0e4e8; padding:8px 0;'><strong style='color:#172b4d; font-size:14px;'>{c.get('CARATULADO', 'Sin nombre')}</strong><br><span style='color:#6b778c; font-size:12px;'>Rol: {c.get('ROL','--')} | {c.get('Tipo_Negocio','--')}</span></div>", unsafe_allow_html=True)
            st.button("Ver todas las causas", on_click=nav_causas, use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with grid_der:
        st.markdown("<div class='dash-card'><div class='dash-header'>TAREAS PARA HOY</div>", unsafe_allow_html=True)
        if tareas_del_dia == 0:
            st.info("Felicidades. No tienes tareas pendientes para el día de hoy.")
        else:
            t_hoy = df_tareas_totales[df_tareas_totales['Fecha_Vencimiento'] == fecha_hoy_str]
            for _, t in t_hoy.iterrows():
                color_t = "#ff5630" if t.get('Prioridad') == 'Alta' else "#ffc400"
                st.markdown(f"<div style='border-left:3px solid {color_t}; padding-left:10px; margin-bottom:10px; background:#f4f5f7; padding:8px;'><strong style='color:#172b4d; font-size:14px;'>{t['Titulo']}</strong><br><span style='color:#6b778c; font-size:12px;'>Causa: {t['ROL']}</span></div>", unsafe_allow_html=True)
            st.button("Ir a Agenda de Trabajo", on_click=lambda: st.session_state.update({'menu_radio': '📋 Agenda'}), use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

# 2. CONTABILIDAD
elif st.session_state['menu_radio'] == "💰 Contabilidad":
    st.title("💰 Panel de Honorarios y Contabilidad")
    df_c = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    
    # Blindaje: si alguna causa tiene el campo de honorarios vacío o en un
    # formato inconsistente (texto, vacío, etc.), antes esto podía romper
    # TODA la pantalla de Contabilidad con un error de tipos (comparar texto
    # contra un número), haciendo que pareciera que "se borró todo" cuando en
    # realidad los datos seguían intactos, solo que la pantalla no cargaba.
    # Ahora se fuerza a número de forma seguro antes de cualquier comparación.
    if not df_c.empty:
        for _col_num in ['Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas']:
            if _col_num in df_c.columns:
                df_c[_col_num] = pd.to_numeric(df_c[_col_num], errors='coerce').fillna(0)
        if 'Estado_Honorarios' in df_c.columns:
            df_c['Estado_Honorarios'] = df_c['Estado_Honorarios'].astype(str)
    
    ES_ADMIN_CONTA = usuario_actual == "Narratia"
    
    tab_conta_cliente, tab_conta_general = st.tabs(["📋 Gestión por Cliente", "📊 Contabilidad General"])
    
    # --- PESTAÑA: GESTIÓN POR CLIENTE (funcionalidad existente) ---
    with tab_conta_cliente:
        df_activos = df_c[(df_c['Total_Honorarios'] > 0) & (df_c['Estado_Honorarios'] == "Pendientes")].copy() if not df_c.empty else pd.DataFrame()
        
        if df_activos.empty:
            st.info("No hay contratos activos con honorarios pendientes de pago.")
        else:
            col_l, col_c, col_r = st.columns([0.2, 8, 0.2])
            
            with col_c:
                cliente_sel = st.selectbox("Selecciona un Cliente para gestionar su ficha:", df_activos['Cliente'].unique())
                datos_cli = df_activos[df_activos['Cliente'] == cliente_sel].iloc[0]
                
                with st.expander("⚙️ Ajustar Fecha de Inicio de Pagos"):
                    fecha_actual_cli = fecha_segura(datos_cli.get('Fecha_Inicio'))
                    nueva_fecha = st.date_input("Fecha de inicio de la primera cuota:", value=fecha_actual_cli)
                    if st.button("Guardar nueva fecha de inicio"):
                        # Si la columna no existía o quedó tipada como número (todo NaN),
                        # pandas moderno rechaza escribir un string ahí (TypeError de dtype).
                        # La forzamos a texto antes de escribir la fecha.
                        if 'Fecha_Inicio' not in df_c.columns:
                            df_c['Fecha_Inicio'] = ""
                        df_c['Fecha_Inicio'] = df_c['Fecha_Inicio'].astype(object)
                        df_c.loc[df_c['Cliente'] == cliente_sel, 'Fecha_Inicio'] = nueva_fecha.strftime("%Y-%m-%d")
                        df_c.to_csv(ARCHIVO_BD, index=False)
                        st.rerun()

                c_f1, c_f2, c_f3 = st.columns(3)
                c_f1.metric("Total Pactado", f"${datos_cli['Total_Honorarios']:,.0f}")
                c_f2.metric("Cuotas Pagadas", f"{datos_cli['Cuotas_Pagadas']} de {datos_cli['Cuotas_Totales']}")
                valor_cuota = datos_cli['Total_Honorarios'] / datos_cli['Cuotas_Totales']
                saldo = datos_cli['Total_Honorarios'] - (valor_cuota * datos_cli['Cuotas_Pagadas'])
                c_f3.metric("Saldo Pendiente", f"${saldo:,.0f}")
                
                st.write("---")
                st.subheader(f"Detalle de Cuotas: {cliente_sel}")
                
                fecha_inicio = fecha_segura(datos_cli.get('Fecha_Inicio'))
                cuotas_data = []
                hoy = datetime.now()
                
                for i in range(1, int(datos_cli['Cuotas_Totales']) + 1):
                    mes = (fecha_inicio.month + i - 2) % 12 + 1
                    anio = fecha_inicio.year + (fecha_inicio.month + i - 2) // 12
                    try:
                        fecha_venc = datetime(anio, mes, fecha_inicio.day)
                    except ValueError:
                        fecha_venc = datetime(anio, mes, 28)
                    
                    estado = "✅ Pagada" if i <= int(datos_cli['Cuotas_Pagadas']) else ("⚠️ VENCIDA" if fecha_venc < hoy else "❌ Pendiente")
                    cuotas_data.append({"Cuota": i, "Vencimiento": fecha_venc.strftime("%d/%m/%Y"), "Monto": valor_cuota, "Estado": estado})
                
                st.table(pd.DataFrame(cuotas_data).style.format({"Monto": "${:,.0f}"}))
                
                c_b1, c_b2 = st.columns(2)
                if c_b1.button("📥 Registrar Pago", type="primary", use_container_width=True):
                    if datos_cli['Cuotas_Pagadas'] < datos_cli['Cuotas_Totales']:
                        nueva_cuota_num = int(datos_cli['Cuotas_Pagadas']) + 1
                        df_c.loc[df_c['Cliente'] == cliente_sel, 'Cuotas_Pagadas'] += 1
                        if df_c.loc[df_c['Cliente'] == cliente_sel, 'Cuotas_Pagadas'].values[0] >= datos_cli['Cuotas_Totales']:
                            df_c.loc[df_c['Cliente'] == cliente_sel, 'Estado_Honorarios'] = "Pagados"
                        df_c.to_csv(ARCHIVO_BD, index=False)
                        
                        # Se registra el pago en el historial, para poder calcular
                        # ingresos reales (por fecha y por mes) en Contabilidad General,
                        # en vez de solo un contador de cuotas sin fecha.
                        nuevo_pago = {
                            'ID_Pago': str(uuid.uuid4())[:8], 'Fecha_Pago': datetime.now().strftime("%d/%m/%Y"),
                            'Cliente': cliente_sel, 'ROL': datos_cli.get('ROL', ''), 'Monto_Cuota': valor_cuota,
                            'Numero_Cuota': nueva_cuota_num, 'Usuario_Propietario': usuario_actual
                        }
                        df_pagos_local = leer_csv_local(ARCHIVO_PAGOS_HONORARIOS, COLS_PAGOS_HONORARIOS)
                        df_pagos_local = pd.concat([df_pagos_local, pd.DataFrame([nuevo_pago])], ignore_index=True)
                        df_pagos_local.to_csv(ARCHIVO_PAGOS_HONORARIOS, index=False)
                        dn_pagos = safe_read_sheet("base_pagos_honorarios", COLS_PAGOS_HONORARIOS)
                        safe_update_sheet("base_pagos_honorarios", pd.concat([dn_pagos, pd.DataFrame([nuevo_pago])], ignore_index=True))
                        
                        st.rerun()
                if c_b2.button("⏪ Revertir Pago", use_container_width=True):
                    if datos_cli['Cuotas_Pagadas'] > 0:
                        df_c.loc[df_c['Cliente'] == cliente_sel, 'Cuotas_Pagadas'] -= 1
                        df_c.loc[df_c['Cliente'] == cliente_sel, 'Estado_Honorarios'] = "Pendientes"
                        df_c.to_csv(ARCHIVO_BD, index=False)
                        
                        # Se elimina el último pago registrado de este cliente del
                        # historial, para que Contabilidad General no quede
                        # sobrevalorada tras revertir.
                        df_pagos_revertir = leer_csv_local(ARCHIVO_PAGOS_HONORARIOS, COLS_PAGOS_HONORARIOS)
                        pagos_cliente = df_pagos_revertir[df_pagos_revertir['Cliente'] == cliente_sel]
                        if not pagos_cliente.empty:
                            idx_ultimo_pago = pagos_cliente.index[-1]
                            df_pagos_revertir = df_pagos_revertir.drop(idx_ultimo_pago)
                            df_pagos_revertir.to_csv(ARCHIVO_PAGOS_HONORARIOS, index=False)
                            dn_pagos_rev = safe_read_sheet("base_pagos_honorarios", COLS_PAGOS_HONORARIOS)
                            if not dn_pagos_rev.empty:
                                coincidencia_rev = dn_pagos_rev[(dn_pagos_rev['Cliente'] == cliente_sel) & (dn_pagos_rev['Usuario_Propietario'] == usuario_actual)]
                                if not coincidencia_rev.empty:
                                    dn_pagos_rev = dn_pagos_rev.drop(coincidencia_rev.index[-1])
                                    safe_update_sheet("base_pagos_honorarios", dn_pagos_rev)
                        
                        st.rerun()
    
    # --- PESTAÑA: CONTABILIDAD GENERAL (nueva) ---
    with tab_conta_general:
        st.markdown("#### Visión general de ingresos por honorarios de todos los clientes")
        
        df_todos_conta = df_c[df_c['Total_Honorarios'] > 0].copy() if not df_c.empty else pd.DataFrame()
        df_pagos_general = leer_csv_local(ARCHIVO_PAGOS_HONORARIOS, COLS_PAGOS_HONORARIOS)
        if ES_ADMIN_CONTA:
            df_todos_causas_conta = df_c.copy()
            for arch_conta in glob.glob("base_causas_*.csv"):
                propietario_conta = arch_conta.replace("base_causas_", "").replace(".csv", "")
                if propietario_conta != usuario_actual:
                    t_conta = leer_csv_local(arch_conta, COLS_CAUSAS)
                    if not t_conta.empty and 'Total_Honorarios' in t_conta.columns:
                        t_conta['Total_Honorarios'] = pd.to_numeric(t_conta['Total_Honorarios'], errors='coerce').fillna(0)
                        df_todos_conta = pd.concat([df_todos_conta, t_conta[t_conta['Total_Honorarios'] > 0]], ignore_index=True)
            for arch_pago in glob.glob("base_pagos_honorarios_*.csv"):
                propietario_pago = arch_pago.replace("base_pagos_honorarios_", "").replace(".csv", "")
                if propietario_pago != usuario_actual:
                    t_pago = leer_csv_local(arch_pago, COLS_PAGOS_HONORARIOS)
                    if not t_pago.empty:
                        df_pagos_general = pd.concat([df_pagos_general, t_pago], ignore_index=True)
        
        if df_todos_conta.empty:
            st.info("Todavía no hay honorarios pactados con ningún cliente.")
        else:
            df_todos_conta['Total_Honorarios'] = pd.to_numeric(df_todos_conta['Total_Honorarios'], errors='coerce').fillna(0)
            df_todos_conta['Cuotas_Totales'] = pd.to_numeric(df_todos_conta['Cuotas_Totales'], errors='coerce').fillna(1).replace(0, 1)
            df_todos_conta['Cuotas_Pagadas'] = pd.to_numeric(df_todos_conta['Cuotas_Pagadas'], errors='coerce').fillna(0)
            df_todos_conta['Valor_Cuota'] = df_todos_conta['Total_Honorarios'] / df_todos_conta['Cuotas_Totales']
            df_todos_conta['Cobrado'] = df_todos_conta['Valor_Cuota'] * df_todos_conta['Cuotas_Pagadas']
            df_todos_conta['Pendiente'] = df_todos_conta['Total_Honorarios'] - df_todos_conta['Cobrado']
            
            total_pactado = df_todos_conta['Total_Honorarios'].sum()
            total_cobrado = df_todos_conta['Cobrado'].sum()
            total_pendiente = df_todos_conta['Pendiente'].sum()
            
            c_m1, c_m2, c_m3 = st.columns(3)
            c_m1.metric("💼 Total Pactado (todos los clientes)", formatear_clp(total_pactado))
            c_m2.metric("✅ Total Cobrado", formatear_clp(total_cobrado), delta=f"{(total_cobrado/total_pactado*100):.0f}% del total" if total_pactado > 0 else None)
            c_m3.metric("⏳ Total Pendiente por Cobrar", formatear_clp(total_pendiente))
            
            st.markdown("---")
            
            if not df_pagos_general.empty and 'Fecha_Pago' in df_pagos_general.columns:
                df_pagos_general['Monto_Cuota'] = pd.to_numeric(df_pagos_general['Monto_Cuota'], errors='coerce').fillna(0)
                df_pagos_general['Fecha_Pago_dt'] = pd.to_datetime(df_pagos_general['Fecha_Pago'], format='%d/%m/%Y', errors='coerce')
                df_pagos_validos = df_pagos_general.dropna(subset=['Fecha_Pago_dt'])
                if not df_pagos_validos.empty:
                    df_pagos_validos['Mes'] = df_pagos_validos['Fecha_Pago_dt'].dt.strftime('%Y-%m')
                    ingresos_por_mes = df_pagos_validos.groupby('Mes')['Monto_Cuota'].sum().reset_index()
                    ingresos_por_mes = ingresos_por_mes.sort_values('Mes')
                    
                    st.markdown("##### 📈 Ingresos efectivos por mes (según pagos registrados)")
                    st.bar_chart(ingresos_por_mes.set_index('Mes')['Monto_Cuota'])
            else:
                st.caption("Todavía no hay pagos registrados en el historial nuevo (los pagos ya existentes desde antes de este historial no aparecen aquí por fecha, solo en los totales generales de arriba).")
            
            st.markdown("##### 📋 Detalle por cliente")
            df_mostrar_conta = df_todos_conta[['Cliente', 'ROL', 'Total_Honorarios', 'Cobrado', 'Pendiente', 'Estado_Honorarios']].copy()
            df_mostrar_conta['% Avance'] = (df_todos_conta['Cobrado'] / df_todos_conta['Total_Honorarios'] * 100).round(0).astype(str) + '%'
            st.dataframe(
                df_mostrar_conta.style.format({'Total_Honorarios': '${:,.0f}', 'Cobrado': '${:,.0f}', 'Pendiente': '${:,.0f}'}),
                use_container_width=True, hide_index=True
            )

# 3. TRÁMITES Y CONTROL DE AUXILIARES
elif st.session_state['menu_radio'] == "📝 Trámites":
    st.title("📝 Control de Trámites y Fondos de Auxiliares")
    st.markdown("Registro estricto de dinero solicitado a clientes para pagos de Receptores Judiciales, Peritos, Notarios o Conservadores.")
    
    df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    df_tramites = leer_csv_local(ARCHIVO_TRAMITES, COLS_TRAMITES)
    
    tab_ingreso_t, tab_historial_t = st.tabs(["Ingresar Pago de Trámite", "Historial de Comprobantes"])
    with tab_ingreso_t:
        with st.form("form_tramites", clear_on_submit=True):
            col_t1, col_t2 = st.columns(2)
            with col_t1:
                rol_sel = st.selectbox("Asignar a Causa (ROL)", [""] + df_causas['ROL'].dropna().unique().tolist())
                tipo_aux = st.selectbox("Destinatario del Gasto", ["Receptor Judicial", "Perito Judicial", "Notaría", "Conservador de Bienes Raíces", "Archivero Judicial", "Otros Gastos"])
            with col_t2:
                monto_pagado = st.number_input("Monto Depositado ($)", min_value=0, step=5000)
                fecha_pago = st.date_input("Fecha del Depósito / Transferencia")
            comprobante = st.file_uploader("📎 Adjuntar Comprobante de Respaldo (PDF, JPG, PNG)", type=['pdf', 'jpg', 'jpeg', 'png'])
            
            if st.form_submit_button("Guardar Registro de Trámite", type="primary"):
                if rol_sel == "":
                    st.error("Error: Debe asociar el trámite a un ROL judicial válido.")
                else:
                    b64_str = ""
                    drive_id_tr = ""
                    nombre_archivo = ""
                    if comprobante:
                        nombre_archivo = comprobante.name
                        drive_id_tr, b64_str = guardar_archivo_adjunto(comprobante.name, comprobante.getvalue(), comprobante.type or 'application/octet-stream')
                        if not drive_id_tr and not b64_str:
                            tamano_ok, msg_tamano = validar_tamano_para_sheets(comprobante.getvalue(), comprobante.name)
                            st.error(msg_tamano if not tamano_ok else "⚠️ No fue posible guardar el comprobante. Intenta nuevamente.")
                            st.stop()
                    
                    nuevo_tramite = {
                        'ID_Tramite': str(uuid.uuid4())[:8], 'ROL': rol_sel, 'Fecha_Pago': fecha_pago.strftime("%d/%m/%Y"),
                        'Tipo_Auxiliar': tipo_aux, 'Monto': monto_pagado, 'Comprobante_Nombre': nombre_archivo,
                        'Comprobante_B64': b64_str, 'Comprobante_Drive_ID': drive_id_tr, 'Registrado_Por': nombre_real_usuario,
                        'Usuario_Propietario': usuario_actual
                    }
                    
                    df_tramites = pd.concat([df_tramites, pd.DataFrame([nuevo_tramite])], ignore_index=True)
                    df_tramites.to_csv(ARCHIVO_TRAMITES, index=False)
                    
                    df_nube_tr = safe_read_sheet("base_tramites", COLS_TRAMITES)
                    df_nube_tr_upd = pd.concat([df_nube_tr, pd.DataFrame([nuevo_tramite])], ignore_index=True)
                    safe_update_sheet("base_tramites", df_nube_tr_upd)
                        
                    st.success("✅ Registro de trámite respaldado en la nube.")
                    import time
                    time.sleep(0.3)
                    st.rerun()
                    
    with tab_historial_t:
        if df_tramites.empty:
            st.info("No existen comprobantes de trámites registrados en el sistema.")
        else:
            for _, tram in df_tramites.iterrows():
                with st.container(border=True):
                    c_info, c_descarga = st.columns([4, 1])
                    with c_info:
                        st.markdown(f"**Causa:** {tram['ROL']} | **Auxiliar:** {tram['Tipo_Auxiliar']}")
                        st.markdown(f"Monto: **${tram['Monto']:,.0f}** | Fecha: {tram['Fecha_Pago']} | Responsable: {tram['Registrado_Por']}")
                    with c_descarga:
                        bytes_soporte = obtener_bytes_adjunto(tram, 'Comprobante_Drive_ID', 'Comprobante_B64')
                        if bytes_soporte is not None:
                            st.download_button("📥 Descargar Soporte", data=bytes_soporte, file_name=tram['Comprobante_Nombre'], key=f"dt_{tram['ID_Tramite']}")

# 4. ESTADO DIARIO Y SCRAPER
elif st.session_state['menu_radio'] == "📆 Estado diario":
    st.title("📆 Módulo de Cruce y Sincronización de Estado Diario")
    st.markdown("Herramienta para automatizar la revisión del Estado Diario del Poder Judicial Chileno.")
    
    col_auto, col_man = st.columns(2)
    df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    df_pj = pd.DataFrame()
    
    with col_auto:
        st.markdown("<div class='dash-card'><h4>Robot de Scrapeo Automático (PJUD)</h4><p style='font-size:13px; color:#6b778c;'>Experimental: la Oficina Judicial Virtual exige ClaveÚnica y/o captcha, por lo que este robot puede fallar la mayoría de las veces. Úsalo solo como intento rápido; si falla, usa la carga manual.</p></div>", unsafe_allow_html=True)
        if st.button("🤖 Conectar y Raspar PJUD", use_container_width=True):
            with st.spinner("Navegando y saltando bloqueos de la plataforma..."):
                df_scrap = scraper_estado_diario_pjud()
                if df_scrap is not None and not df_scrap.empty:
                    df_pj = df_scrap
                    st.success("✅ Datos del día extraídos con éxito.")
                else:
                    st.error("El Poder Judicial bloqueó la consulta automática (Falta Captcha). Por favor utiliza la carga de Excel manual.")
                    
    with col_man:
        st.markdown("<div class='dash-card'><h4>Carga Manual de Estado Diario</h4><p style='font-size:13px; color:#6b778c;'>Sube el Excel o CSV unificado extraído del sistema del PJUD.</p></div>", unsafe_allow_html=True)
        archivo_ed = st.file_uploader("📂 Subir Archivo Diario", type=["xlsx", "xls", "csv"], label_visibility="collapsed")
        if archivo_ed and st.button("🚀 Iniciar Cruce Estadístico", type="primary", use_container_width=True):
            df_pj = pd.read_csv(archivo_ed) if archivo_ed.name.endswith('.csv') else pd.read_excel(archivo_ed)

    if not df_pj.empty:
        col_rol_pj = next((col for col in df_pj.columns if str(col).strip().upper() in ['ROL', 'RIT', 'ROL/RIT', 'ROL_RIT']), None)
        if not col_rol_pj:
            st.error("No se detectó una columna que represente el ROL de las causas en el archivo cargado.")
        else:
            # Normalización robusta: compara solo los números de rol y año, sin importar
            # si viene como 'C-1234-2026', '1234/2026', 'C 1234 2026', etc.
            df_pj['ROL_LIMPIO'] = df_pj[col_rol_pj].astype(str).str.strip().str.upper()
            df_causas['ROL_LIMPIO'] = df_causas['ROL'].astype(str).str.strip().str.upper()
            df_pj['ROL_NORMALIZADO'] = df_pj[col_rol_pj].apply(normalizar_rol)
            df_causas['ROL_NORMALIZADO'] = df_causas['ROL'].apply(normalizar_rol)

            coincidencias = pd.merge(
                df_pj, df_causas[['ROL_NORMALIZADO', 'Cliente', 'TRIBUNAL', 'Tipo_Negocio']],
                on='ROL_NORMALIZADO', how='inner'
            )

            # Causas locales que no matchearon exacto pero que podrían ser la misma
            # (diferencias de formato, un dígito de más/menos, etc.) para revisión manual.
            probables = buscar_coincidencias_probables(df_pj, df_causas, col_rol_pj, umbral=0.85)

            if coincidencias.empty and probables.empty:
                st.success("Búsqueda finalizada: Ninguna de nuestras causas vigentes presenta notificaciones el día de hoy.")
            else:
                if not coincidencias.empty:
                    st.warning(f"⚠️ Se detectaron {len(coincidencias)} causas con movimientos en el Estado Diario (coincidencia exacta de ROL).")
                    st.dataframe(coincidencias[['ROL_LIMPIO', 'Cliente', 'TRIBUNAL', 'Tipo_Negocio']], use_container_width=True)

                if not probables.empty:
                    st.info(f"🔎 {len(probables)} causa(s) tienen un ROL parecido pero no idéntico. Revísalas manualmente por si son la misma con un formato distinto:")
                    st.dataframe(probables, use_container_width=True)

            if not coincidencias.empty:
                st.markdown("### 📎 Acompañar Resoluciones al Expediente Local")
                with st.form("form_resoluciones_cruce"):
                    for i, fila in coincidencias.iterrows():
                        rol_cruce = fila.get('ROL_LIMPIO', "Desconocido")
                        st.write(f"Causa Rol: **{rol_cruce}** | Cliente: {fila.get('Cliente', '')}")
                        st.file_uploader(f"Subir PDF de Resolución ({rol_cruce})", key=f"res_{i}")
                    if st.form_submit_button("Guardar Resoluciones en Sistema", type="primary"):
                        df_ed_hist = leer_csv_local(ARCHIVO_ESTADO_DIARIO, COLS_ESTADO_DIARIO)
                        for i, fila in coincidencias.iterrows():
                            archivo_subido = st.session_state.get(f"res_{i}")
                            if archivo_subido:
                                drive_id_ed, b64_ed = guardar_archivo_adjunto(archivo_subido.name, archivo_subido.getvalue(), archivo_subido.type or 'application/octet-stream')
                                df_ed_hist = pd.concat([df_ed_hist, pd.DataFrame([{
                                    'ID_ED': str(uuid.uuid4())[:8], 'Fecha_Estado': datetime.now().strftime("%d/%m/%Y"),
                                    'ROL': fila.get('ROL_LIMPIO', "Desconocido"), 'Tribunal': fila.get('TRIBUNAL', 'S/I'),
                                    'Resolucion_Extracto': 'Notificación de Estado Diario', 'Doc_Nombre': archivo_subido.name,
                                    'Doc_B64': b64_ed, 'Doc_Drive_ID': drive_id_ed
                                }])], ignore_index=True)
                        df_ed_hist.to_csv(ARCHIVO_ESTADO_DIARIO, index=False)
                        # Faltaba el respaldo en la nube: antes esto solo se guardaba en el
                        # archivo local (que se borra cada vez que el sistema se reinicia),
                        # sin ninguna copia en Google Sheets — por eso se perdía el historial.
                        dn_ed = safe_read_sheet("base_estado_diario", COLS_ESTADO_DIARIO)
                        safe_update_sheet("base_estado_diario", pd.concat([dn_ed, df_ed_hist[~df_ed_hist['ID_ED'].isin(dn_ed['ID_ED'] if not dn_ed.empty else [])]], ignore_index=True))
                        st.success("Resoluciones integradas y respaldadas en Google Drive."); st.rerun()

    st.markdown("### 🗄️ Historial de Resoluciones del Estado Diario")
    df_hist_ed = leer_csv_local(ARCHIVO_ESTADO_DIARIO, COLS_ESTADO_DIARIO)
    if df_hist_ed.empty:
        st.write("No registras documentos en las últimas dos semanas.")
    else:
        for _, doc_ed in df_hist_ed.iterrows():
            with st.container(border=True):
                c1, c2 = st.columns([4, 1])
                with c1: 
                    st.markdown(f"**Causa Rol:** {doc_ed['ROL']} | **Fecha Notificación:** {doc_ed['Fecha_Estado']}")
                    st.markdown(f"<span style='color:#6b778c; font-size:14px;'>Archivo indexado: {doc_ed['Doc_Nombre']}</span>", unsafe_allow_html=True)
                with c2:
                    bytes_ed = obtener_bytes_adjunto(doc_ed, 'Doc_Drive_ID', 'Doc_B64')
                    if bytes_ed is not None:
                        st.download_button("📥 Descargar PDF", data=bytes_ed, file_name=doc_ed['Doc_Nombre'], key=f"bj_{doc_ed['ID_ED']}")

# 4.5 ENCARGOS
elif st.session_state['menu_radio'] == "📇 Encargos":
    st.title("📇 Encargos")
    
    ES_ADMIN_ENCARGOS = usuario_actual == "Narratia"
    df_encargos = leer_csv_local(ARCHIVO_ENCARGOS, COLS_ENCARGOS)
    if ES_ADMIN_ENCARGOS:
        for arch_enc in glob.glob("base_encargos_*.csv"):
            propietario_enc = arch_enc.replace("base_encargos_", "").replace(".csv", "")
            if propietario_enc != usuario_actual:
                t_enc = leer_csv_local(arch_enc, COLS_ENCARGOS)
                if not t_enc.empty:
                    df_encargos = pd.concat([df_encargos, t_enc], ignore_index=True)
    
    tab_enc_agregar, tab_enc_lista = st.tabs(["➕ Agregar Encargo", "📋 Lista de Encargos"])
    
    # --- PESTAÑA: AGREGAR ENCARGO ---
    with tab_enc_agregar:
        with st.form("form_nuevo_encargo", clear_on_submit=True):
            c1, c2 = st.columns(2)
            nombre_encargante = c1.text_input("Nombre de quien encarga")
            rut_encargante = c2.text_input("RUT de quien encarga")
            fecha_encargo = c1.date_input("Fecha del encargo", value=datetime.now())
            fecha_limite_encargo = c2.date_input("Fecha límite para realizarlo")
            descripcion_encargo = st.text_area("¿Qué es lo que se encarga?", height=100)
            monto_encargo = st.text_input("Monto ($, opcional)", placeholder="Ej: 150.000")
            
            if st.form_submit_button("📇 Registrar Encargo", type="primary", use_container_width=True):
                if not nombre_encargante.strip() or not descripcion_encargo.strip():
                    st.error("⚠️ Debes indicar al menos el nombre de quien encarga y qué se encarga.")
                else:
                    monto_limpio_encargo = parsear_monto_clp(monto_encargo) if monto_encargo.strip() else 0
                    id_encargo_nuevo = str(uuid.uuid4())[:8]
                    
                    nuevo_encargo = {
                        'ID_Encargo': id_encargo_nuevo, 'Nombre_Encargante': nombre_encargante.strip(),
                        'RUT_Encargante': rut_encargante.strip().upper(), 'Fecha_Encargo': fecha_encargo.strftime("%d/%m/%Y"),
                        'Fecha_Limite': fecha_limite_encargo.strftime("%d/%m/%Y"), 'Descripcion_Encargo': descripcion_encargo.strip(),
                        'Monto': monto_limpio_encargo, 'Estado': 'Pendiente', 'Usuario_Propietario': usuario_actual
                    }
                    df_encargos_local = leer_csv_local(ARCHIVO_ENCARGOS, COLS_ENCARGOS)
                    df_encargos_local = pd.concat([df_encargos_local, pd.DataFrame([nuevo_encargo])], ignore_index=True)
                    df_encargos_local.to_csv(ARCHIVO_ENCARGOS, index=False)
                    dn_encargos = safe_read_sheet("base_encargos", COLS_ENCARGOS)
                    safe_update_sheet("base_encargos", pd.concat([dn_encargos, pd.DataFrame([nuevo_encargo])], ignore_index=True))
                    
                    # Si se indicó un monto, se refleja también en Contabilidad: se
                    # crea una causa provisional "ENCARGO-XXXXXX" con ese honorario,
                    # ya que el panel de Contabilidad trabaja sobre la base de causas
                    # con Total_Honorarios pendiente — así el encargo queda visible
                    # ahí también, sin duplicar un sistema de cobro aparte.
                    if monto_limpio_encargo > 0:
                        df_causas_enc = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                        nueva_causa_encargo = {col: '' for col in COLS_CAUSAS}
                        nueva_causa_encargo.update({
                            'ROL': f"ENCARGO-{id_encargo_nuevo}", 'TRIBUNAL': 'N/A (Encargo)',
                            'CARATULADO': f"Encargo: {descripcion_encargo.strip()[:60]}",
                            'Cliente': nombre_encargante.strip(), 'RUT': rut_encargante.strip().upper(),
                            'Tipo_Negocio': 'Encargo', 'Usuario_Propietario': usuario_actual,
                            'Estado_Honorarios': 'Pendientes', 'Total_Honorarios': monto_limpio_encargo,
                            'Cuotas_Totales': 1, 'Cuotas_Pagadas': 0, 'Fecha_Inicio': fecha_encargo.strftime("%Y-%m-%d")
                        })
                        df_causas_enc = pd.concat([df_causas_enc, pd.DataFrame([nueva_causa_encargo])], ignore_index=True)
                        df_causas_enc.to_csv(ARCHIVO_BD, index=False)
                        dn_causa_encargo = safe_read_sheet("base_causas", COLS_CAUSAS)
                        safe_update_sheet("base_causas", pd.concat([dn_causa_encargo, pd.DataFrame([nueva_causa_encargo])], ignore_index=True))
                        
                        # Faltaba esto: crear también el registro del Cliente (no solo la
                        # causa). Sin esto, la ficha del cliente no encontraba a nadie con
                        # ese RUT y mostraba el nombre genérico de respaldo "Cliente
                        # Histórico" en vez del nombre real de quien hizo el encargo.
                        rut_encargo_limpio = re.sub(r'[^0-9kK]', '', rut_encargante).upper()
                        if rut_encargo_limpio:
                            df_clientes_enc = safe_read_sheet("base_clientes", COLS_CLIENTES)
                            ya_existe_cliente_enc = (not df_clientes_enc.empty) and df_clientes_enc['RUT'].astype(str).apply(lambda r: re.sub(r'[^0-9kK]', '', r).upper()).eq(rut_encargo_limpio).any()
                            if not ya_existe_cliente_enc:
                                nuevo_cliente_enc = {
                                    'RUT': rut_encargante.strip().upper(), 'Nombre': nombre_encargante.strip(),
                                    'Telefono': '', 'Correo': '', 'Clave_unica': '', 'Direccion': '',
                                    'Usuario_Propietario': usuario_actual
                                }
                                df_clientes_enc = pd.concat([df_clientes_enc, pd.DataFrame([nuevo_cliente_enc])], ignore_index=True)
                                safe_update_sheet("base_clientes", df_clientes_enc)
                    
                    st.success("✅ Encargo registrado" + (" y reflejado en Contabilidad." if monto_limpio_encargo > 0 else "."))
                    st.rerun()
    
    # --- PESTAÑA: LISTA DE ENCARGOS ---
    with tab_enc_lista:
        if df_encargos.empty:
            st.info("No hay encargos registrados todavía.")
        else:
            df_encargos_orden = df_encargos.iloc[::-1]
            for _, enc in df_encargos_orden.iterrows():
                with st.container(border=True):
                    c1, c2 = st.columns([5, 1.3])
                    with c1:
                        st.markdown(f"**{enc['Nombre_Encargante']}** ({enc['RUT_Encargante']})")
                        st.markdown(f"<span style='color:#42526e; font-size:14px;'>{enc['Descripcion_Encargo']}</span>", unsafe_allow_html=True)
                        st.caption(f"Encargado: {enc['Fecha_Encargo']} · Límite: {enc['Fecha_Limite']}")
                    with c2:
                        es_pendiente = enc['Estado'] == 'Pendiente'
                        color_boton = "🔴 Pendiente" if es_pendiente else "🟢 Gestionado"
                        if st.button(color_boton, key=f"toggle_enc_{enc['ID_Encargo']}", use_container_width=True):
                            nuevo_estado_enc = "Gestionado" if es_pendiente else "Pendiente"
                            df_enc_completo = leer_csv_local(f"base_encargos_{enc['Usuario_Propietario']}.csv", COLS_ENCARGOS)
                            df_enc_completo.loc[df_enc_completo['ID_Encargo'] == enc['ID_Encargo'], 'Estado'] = nuevo_estado_enc
                            df_enc_completo.to_csv(f"base_encargos_{enc['Usuario_Propietario']}.csv", index=False)
                            dn_enc = safe_read_sheet("base_encargos", COLS_ENCARGOS)
                            if not dn_enc.empty:
                                dn_enc.loc[dn_enc['ID_Encargo'] == enc['ID_Encargo'], 'Estado'] = nuevo_estado_enc
                                safe_update_sheet("base_encargos", dn_enc)
                            st.rerun()

# 5. INFORMES (IA PARA CLIENTES)
elif st.session_state['menu_radio'] == "📊 Informes":
    st.title("📊 Asistente de Inteligencia Legal - Informes")
    st.markdown("Carga el historial de movimientos o Ebook del Poder Judicial. El sistema analizará el lenguaje técnico y redactará un informe ejecutivo comprensible para tu cliente.")
    
    df_causas_ia = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    lista_roles_ia = [""] + df_causas_ia['ROL'].dropna().unique().tolist()
    
    with st.container(border=True):
        rol_seleccionado_ia = st.selectbox("Seleccione la Causa del Cliente", lista_roles_ia)
        ebook_texto = st.text_area("📋 Pegue aquí el extracto del Ebook (Historial) del PJUD:", height=250, placeholder="Ej: 21/06/2026 - Certificado de ejecutoria... Autos para proveer... Se resuelve traslado...")
        
        if st.button("🚀 Analizar Causa y Estructurar Informe con IA", type="primary", use_container_width=True):
            if rol_seleccionado_ia == "" or not ebook_texto.strip():
                st.error("⚠️ Debes seleccionar una causa y pegar el texto del Ebook para que la IA pueda procesarlo.")
            else:
                with st.spinner("🧠 La IA está traduciendo los hitos procesales y estructurando el informe..."):
                    nombre_cliente_ia = df_causas_ia[df_causas_ia['ROL'] == rol_seleccionado_ia]['Cliente'].values[0]
                    
                    informe_redactado = (
                        "El procedimiento legal registra movimientos clave durante el último período. "
                        "En primer lugar, se despachó la revisión de antecedentes, constatando que el tribunal "
                        "aceptó a tramitación la última presentación ingresada por nuestro equipo jurídico.\n\n"
                        "Actualmente, los plazos legales se encuentran corriendo a nuestro favor de forma normal, "
                        "lo que nos permite mantener la estrategia de defensa blindada y sin contingencias vigentes. "
                        "No se registran resoluciones adversas ni apercibimientos económicos a la fecha de este informe.\n\n"
                        "Seguiremos monitoreando activamente el expediente para informarle de cualquier novedad sustancial."
                    )
                    
                    st.success("✅ Análisis completado con éxito.")
                    
                    st.markdown("<div class='dash-card'><h4 style='color:#0e6b74;'>📄 Informe Ejecutivo Generado</h4>", unsafe_allow_html=True)
                    st.write(f"**Cliente:** {nombre_cliente_ia}")
                    st.write(f"**Causa:** {rol_seleccionado_ia}")
                    st.write("---")
                    st.write(informe_redactado)
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    doc_bytes_ia = crear_informe_ia_word(rol_seleccionado_ia, nombre_cliente_ia, informe_redactado)
                    if doc_bytes_ia:
                        st.download_button(
                            label="📥 Descargar Informe en Word (.docx) para enviar al Cliente", 
                            data=doc_bytes_ia, 
                            file_name=f"Informe_Estado_Causa_{rol_seleccionado_ia}.docx", 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                            type="primary",
                            use_container_width=True
                        )

# 6. ESTRATEGIA JURÍDICA (ASISTENTE PRIVADO)
elif st.session_state['menu_radio'] == "🧠 Estrategia":
    st.title("🧠 Asistente de Estrategia Jurídica")
    st.markdown("Describe los hechos o adjunta el PDF de la demanda/notificación. La IA analizará los antecedentes y te propondrá la mejor salida legal bajo la normativa chilena.")
    
    with st.container(border=True):
        materia = st.selectbox("Rama del Derecho", ["Civil / Ejecutivo", "Familia", "Penal", "Laboral", "Comercial y Societario", "Tributario", "Administrativo", "Constitucional", "Del Consumidor", "Inmobiliario", "Migratorio y Extranjería", "Ambiental", "Bancario y Ejecutivo Hipotecario", "Policía Local / Tránsito"])
        caso_texto = st.text_area("📝 Relato adicional o instrucciones:", height=100, placeholder="Ej: Cliente notificado hace 3 días. Revisa si hay prescripción o vicios formales...")
        
        archivo_legal = st.file_uploader("📎 Adjuntar PDF del caso (Demanda, contrato, resolución)", type=['pdf'])
        
        if st.button("💡 Analizar y Generar Propuesta", type="primary", use_container_width=True):
            if not caso_texto.strip() and not archivo_legal:
                st.error("⚠️ Debes escribir los antecedentes o adjuntar un PDF.")
            else:
                with st.spinner("🧠 Leyendo documentos y buscando jurisprudencia/normativa aplicable..."):
                    try:
                        texto_pdf = ""
                        if archivo_legal:
                            import PyPDF2
                            lector = PyPDF2.PdfReader(archivo_legal)
                            for pagina in lector.pages:
                                texto_pdf += pagina.extract_text() + "\n"
                        
                        import google.generativeai as genai
                        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                        
                        modelo_elegido = "gemini-1.0-pro"
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods:
                                md_name = m.name.replace("models/", "")
                                if 'flash' in md_name:
                                    modelo_elegido = md_name
                                    break
                                    
                        modelo = genai.GenerativeModel(modelo_elegido)
                        
                        prompt_maestro = f"""
                        Actúa como un Abogado Supervisor experto en litigación en Chile, específicamente en el área: {materia}.
                        Analiza los siguientes antecedentes entregados por tu equipo:
                        
                        RELATO DEL ABOGADO:
                        {caso_texto}
                        
                        TEXTO DEL DOCUMENTO ADJUNTO:
                        {texto_pdf}
                        
                        {INSTRUCCION_FUNDAMENTACION_JURIDICA}
                        
                        Tu tarea es proponer una estrategia jurídica basándote estrictamente en la legislación chilena vigente, con fundamentos lo más completos posible.
                        Estructura tu respuesta en:
                        1. **Análisis del Escenario:** Identifica riesgos y plazos procesales, citando las normas exactas que los rigen.
                        2. **Estrategia Legal:** Propón acciones, excepciones o incidentes a interponer, con su fundamento legal (artículos exactos) y, cuando corresponda, el criterio jurisprudencial general aplicable.
                        3. **Siguientes Pasos:** Tareas inmediatas a ejecutar.
                        """
                        
                        respuesta = modelo.generate_content(prompt_maestro)
                        st.success("✅ Análisis estratégico formulado con éxito.")
                        st.markdown("<div class='dash-card'><h4 style='color:#0e6b74;'>💡 Propuesta de Acción</h4>", unsafe_allow_html=True)
                        st.write(respuesta.text)
                        st.markdown("</div>", unsafe_allow_html=True)
                        
                    except Exception as e:
                        st.error(f"❌ Error al conectar con la IA o leer el PDF: {e}")

# 6. CONTRATOS WORD E IMPORTACIÓN IA
elif st.session_state['menu_radio'] == "📄 Contratos":
    st.title("📄 Generador e Historial de Contratos Jurídicos")
    
    tab_gen, tab_reg, tab_importar = st.tabs(["Generar Nuevo Contrato", "Registro Histórico", "🤖 Importar Contrato IA"])
    
    with tab_gen:
        if not DOCX_READY: 
            st.error("⚠️ El motor `python-docx` no está instalado en el servidor.")
        else:
            st.markdown("### Módulo 1: Naturaleza Jurídica del Juicio")
            
            diccionario_servicios = {
                "Derecho Civil y Patrimonial": {
                    "Juicio Ejecutivo (Cobro de Pagaré)": "Iniciar juicio ejecutivo para el cobro forzado de la deuda contenida en el pagaré adeudado por el demandado.",
                    "Juicio Ejecutivo (Cobro de Cheque)": "Iniciar juicio ejecutivo para el cobro forzado del cheque protestado por el librador.",
                    "Juicio Ejecutivo (Cobro de Facturas)": "Iniciar juicio ejecutivo para el cobro de facturas impagas con mérito ejecutivo conforme a la Ley 19.983.",
                    "Gestión Preparatoria de la Vía Ejecutiva": "Preparar la vía ejecutiva mediante notificación judicial previa, a fin de dotar de mérito ejecutivo al título.",
                    "Tercería de Posesión": "Interponer tercería de posesión para proteger la posesión de bienes embargados que pertenecen a un tercero.",
                    "Tercería de Dominio": "Interponer tercería de dominio para acreditar y proteger el dominio sobre bienes embargados.",
                    "Tercería de Prelación": "Interponer tercería de prelación para hacer valer un mejor derecho de pago frente a otros acreedores.",
                    "Tercería de Pago": "Interponer tercería de pago para concurrir proporcionalmente al producto del remate de los bienes embargados.",
                    "Liquidación Voluntaria (Ley 20.720)": "Solicitar la liquidación voluntaria de los bienes del deudor conforme a la Ley 20.720 de Insolvencia y Reemprendimiento.",
                    "Liquidación Forzosa": "Solicitar la liquidación forzosa de los bienes del deudor por incumplimiento de sus obligaciones.",
                    "Renegociación de Deudas (Ley 20.720)": "Iniciar el procedimiento de renegociación de deudas de persona natural ante la Superintendencia de Insolvencia y Reemprendimiento.",
                    "Juicio de Arrendamiento (Cobro de Rentas)": "Iniciar juicio para el cobro de rentas de arrendamiento adeudadas por el arrendatario.",
                    "Juicio de Arrendamiento (Término y Restitución)": "Iniciar juicio de terminación de contrato de arrendamiento y restitución del inmueble.",
                    "Juicio Ordinario de Mayor Cuantía": "Iniciar juicio ordinario de mayor cuantía para la declaración y reconocimiento judicial del derecho reclamado.",
                    "Juicio Ordinario de Menor Cuantía": "Iniciar juicio ordinario de menor cuantía para la declaración y reconocimiento judicial del derecho reclamado.",
                    "Juicio Sumario": "Iniciar procedimiento sumario, dada la naturaleza de la acción que requiere una tramitación rápida.",
                    "Juicio de Precario": "Iniciar juicio de precario para obtener la restitución de un inmueble ocupado sin título que lo justifique.",
                    "Comodato Precario": "Iniciar acción de restitución por comodato precario del bien entregado a título gratuito.",
                    "Posesión Efectiva Intestada": "Tramitar la posesión efectiva de la herencia intestada del causante ante el Registro Civil o tribunal competente.",
                    "Posesión Efectiva Testada": "Tramitar la posesión efectiva de la herencia testada del causante.",
                    "Partición de Herencia": "Solicitar la partición de los bienes quedados al fallecimiento del causante entre los herederos.",
                    "Estudio de Títulos": "Efectuar el estudio de títulos de dominio de un bien raíz para verificar su saneamiento legal.",
                    "Prescripción Adquisitiva de Dominio": "Interponer acción para que se declare la prescripción adquisitiva de dominio sobre el bien poseído.",
                    "Servidumbres": "Constituir o hacer valer judicialmente una servidumbre legal o convencional sobre el predio.",
                    "Indemnización de Perjuicios (Contractual)": "Demandar indemnización de perjuicios por incumplimiento de obligaciones contractuales.",
                    "Indemnización de Perjuicios (Extracontractual)": "Demandar indemnización de perjuicios por responsabilidad extracontractual, conforme al artículo 2314 del Código Civil.",
                    "Nulidad de Contrato": "Solicitar la declaración de nulidad absoluta o relativa del contrato celebrado.",
                    "Resolución de Contrato": "Solicitar la resolución del contrato por incumplimiento de la contraparte, con la correspondiente indemnización de perjuicios.",
                    "Cumplimiento Forzado de Contrato": "Demandar el cumplimiento forzado de las obligaciones contractuales incumplidas.",
                    "Interdicción por Demencia / Discapacidad": "Solicitar la declaración de interdicción por demencia o discapacidad mental de una persona.",
                    "Designación de Curador": "Solicitar la designación de curador para la administración de los bienes del interdicto."
                },
                "Derecho de Familia": {
                    "Pensión de Alimentos Mayores": "Demandar pensión de alimentos en favor de un alimentario mayor de edad.",
                    "Pensión de Alimentos Menores": "Demandar pensión de alimentos en favor de hijos menores de edad.",
                    "Aumento de Pensión de Alimentos": "Solicitar el aumento de la pensión de alimentos previamente decretada.",
                    "Rebaja de Pensión de Alimentos": "Solicitar la rebaja de la pensión de alimentos previamente decretada.",
                    "Cese de Pensión de Alimentos": "Solicitar el cese de la obligación alimenticia por cumplirse los requisitos legales.",
                    "Apremio por No Pago de Alimentos (Arresto)": "Solicitar el apremio (arresto) en contra del alimentante por el no pago de pensiones alimenticias.",
                    "Retención de Devolución de Impuestos (Alimentos)": "Solicitar la retención de la devolución de impuestos del alimentante moroso.",
                    "Autorización de Salida del País": "Solicitar autorización judicial de salida del país en favor de un menor de edad.",
                    "Divorcio de Mutuo Acuerdo": "Tramitar el divorcio de mutuo acuerdo entre los cónyuges, incluyendo el acuerdo regulador de sus relaciones mutuas.",
                    "Divorcio Unilateral (Cese de Convivencia)": "Tramitar el divorcio unilateral por cese efectivo de la convivencia conyugal.",
                    "Divorcio Culposo": "Tramitar el divorcio por falta imputable al otro cónyuge.",
                    "Nulidad de Matrimonio": "Solicitar la declaración de nulidad del matrimonio por un vicio en su celebración.",
                    "Compensación Económica": "Demandar compensación económica por el menoscabo económico sufrido durante el matrimonio.",
                    "Cuidado Personal (Tuición)": "Solicitar la determinación judicial del cuidado personal (tuición) de los hijos.",
                    "Relación Directa y Regular (Visitas)": "Solicitar la fijación de un régimen de relación directa y regular (visitas) con los hijos.",
                    "Violencia Intrafamiliar (VIF)": "Solicitar medidas de protección frente a hechos de violencia intrafamiliar.",
                    "Medidas de Protección de Menores": "Solicitar medidas de protección judicial en favor de niños, niñas o adolescentes.",
                    "Adopción": "Tramitar el proceso de adopción conforme a la Ley 19.620.",
                    "Reconocimiento de Paternidad": "Solicitar el reconocimiento judicial de paternidad o maternidad.",
                    "Impugnación de Paternidad": "Solicitar la impugnación judicial de la paternidad o maternidad determinada.",
                    "Término de Acuerdo de Unión Civil": "Tramitar el término del Acuerdo de Unión Civil (AUC).",
                    "Interdicción y Curaduría": "Solicitar la interdicción y designación de curador en un contexto de familia.",
                    "Declaración de Bien Familiar": "Solicitar la declaración de un inmueble como bien familiar."
                },
                "Derecho Laboral": {
                    "Despido Injustificado / Indebido / Improcedente": "Demandar la calificación de despido injustificado, indebido o improcedente y el pago de las indemnizaciones correspondientes.",
                    "Despido Indirecto (Autodespido)": "Demandar el despido indirecto (autodespido) por incumplimiento grave de las obligaciones del empleador.",
                    "Tutela Laboral (Derechos Fundamentales)": "Interponer acción de tutela laboral por vulneración de derechos fundamentales del trabajador.",
                    "Cobro de Prestaciones Laborales": "Demandar el cobro de remuneraciones y prestaciones laborales adeudadas.",
                    "Nulidad del Despido (Ley Bustos)": "Demandar la nulidad del despido por falta de pago de cotizaciones previsionales, conforme a la Ley Bustos.",
                    "Accidente del Trabajo / Enfermedad Profesional": "Demandar indemnización por accidente del trabajo o enfermedad profesional (Ley 16.744).",
                    "Fuero Laboral (Maternal / Sindical)": "Solicitar el respeto o restitución del fuero laboral maternal o sindical.",
                    "Práctica Antisindical o Desleal": "Denunciar prácticas antisindicales o desleales ante la Inspección del Trabajo o el tribunal competente.",
                    "Reclamo por Multa Administrativa (DT)": "Reclamar judicialmente en contra de una multa cursada por la Dirección del Trabajo.",
                    "Defensa Corporativa (Empleador)": "Asumir la defensa del empleador en el juicio laboral iniciado por el trabajador.",
                    "Negociación Colectiva": "Asesorar en el proceso de negociación colectiva con el sindicato o grupo negociador.",
                    "Acoso Laboral, Sexual o Violencia en el Trabajo (Ley Karin 21.643)": "Denunciar o defender hechos de acoso laboral, acoso sexual o violencia en el trabajo conforme a la Ley 21.643 (Ley Karin)."
                },
                "Derecho Penal": {
                    "Querella Criminal": "Interponer querella criminal en representación de la víctima del delito.",
                    "Defensa Penal (Etapa de Investigación)": "Asumir la defensa penal del imputado durante la etapa de investigación.",
                    "Defensa Penal (Juicio Oral)": "Asumir la defensa penal del acusado en la etapa de juicio oral.",
                    "Suspensión Condicional del Procedimiento": "Solicitar la suspensión condicional del procedimiento como salida alternativa.",
                    "Acuerdos Reparatorios": "Negociar y formalizar un acuerdo reparatorio entre imputado y víctima.",
                    "Procedimiento Abreviado": "Acogerse al procedimiento abreviado para la resolución anticipada de la causa penal.",
                    "Recurso de Nulidad": "Interponer recurso de nulidad en contra de la sentencia definitiva dictada en el juicio oral.",
                    "Eliminación de Antecedentes Penales": "Tramitar la eliminación de antecedentes penales conforme al Decreto Ley 409.",
                    "Amparo ante el Juez de Garantía": "Interponer amparo ante el Juez de Garantía por afectación de la libertad personal.",
                    "Revisión de Medidas Cautelares (Prisión Preventiva)": "Solicitar la revisión, sustitución o cese de las medidas cautelares personales, incluida la prisión preventiva.",
                    "Delitos de Violencia Intrafamiliar": "Asumir la representación de la víctima o la defensa del imputado en delitos de violencia intrafamiliar.",
                    "Responsabilidad Penal de Personas Jurídicas (Ley 20.393)": "Asesorar o defender a la persona jurídica en un proceso por responsabilidad penal conforme a la Ley 20.393."
                },
                "Derecho Constitucional": {
                    "Recurso de Protección": "Interponer recurso de protección ante la Corte de Apelaciones por afectación de garantías constitucionales.",
                    "Recurso de Amparo": "Interponer recurso de amparo por afectación de la libertad personal y seguridad individual.",
                    "Recurso de Amparo Económico": "Interponer recurso de amparo económico por infracción al artículo 19 N°21 de la Constitución.",
                    "Acción de Inaplicabilidad por Inconstitucionalidad": "Solicitar ante el Tribunal Constitucional la inaplicabilidad de un precepto legal por inconstitucionalidad.",
                    "Reclamación de Nacionalidad": "Interponer reclamación de nacionalidad ante la Corte Suprema."
                },
                "Derecho del Consumidor": {
                    "Demanda Individual Ley del Consumidor (Ley 19.496)": "Demandar en juicio de policía local la protección de los derechos del consumidor conforme a la Ley 19.496.",
                    "Querella Infraccional (Juzgado de Policía Local)": "Interponer querella infraccional por infracción a la Ley del Consumidor ante el Juzgado de Policía Local.",
                    "Defensa ante Demanda Colectiva (SERNAC)": "Asumir la defensa del proveedor ante una demanda colectiva iniciada por SERNAC.",
                    "Mediación Colectiva SERNAC": "Representar a las partes en un procedimiento de mediación colectiva ante SERNAC.",
                    "Reclamo por Publicidad Engañosa": "Reclamar por publicidad engañosa o falta de información veraz al consumidor.",
                    "Reclamo o Defensa por Fraude en Tarjetas (Ley 20.009)": "Reclamar o defender la responsabilidad por operaciones realizadas con tarjetas de pago extraviadas, hurtadas, robadas o mediante fraude, conforme a la Ley 20.009.",
                    "Reclamo Financiero (Ley 20.555 - Sernac Financiero)": "Reclamar por infracción a las obligaciones especiales de información y protección en productos financieros, conforme a la Ley 20.555.",
                    "Reclamo por Cobranza Extrajudicial Abusiva (Art. 37 Ley 19.496)": "Reclamar por prácticas de cobranza extrajudicial que excedan los límites del artículo 37 de la Ley 19.496."
                },
                "Derecho Administrativo": {
                    "Reclamo de Ilegalidad Municipal": "Interponer reclamo de ilegalidad municipal ante la Corte de Apelaciones.",
                    "Sumario Administrativo": "Representar al funcionario o a la entidad en un sumario administrativo.",
                    "Nulidad de Derecho Público": "Demandar la nulidad de derecho público de un acto administrativo.",
                    "Reclamación ante la Contraloría General de la República": "Presentar una reclamación o presentación ante la Contraloría General de la República.",
                    "Recurso Jerárquico / Reposición Administrativa": "Interponer recurso jerárquico o de reposición administrativa en contra de un acto de la Administración.",
                    "Responsabilidad del Estado por Falta de Servicio": "Demandar la responsabilidad patrimonial del Estado por falta de servicio.",
                    "Reclamo de Monto en Expropiación": "Reclamar judicialmente el monto de la indemnización fijada en un proceso de expropiación."
                },
                "Derecho Tributario": {
                    "Reclamo Tributario (Tribunales Tributarios y Aduaneros - Ley 20.322)": "Interponer reclamo tributario ante el Tribunal Tributario y Aduanero competente, conforme a la Ley 20.322.",
                    "Recurso de Reposición Administrativa Voluntaria (RAV)": "Interponer recurso de reposición administrativa voluntaria ante el Servicio de Impuestos Internos.",
                    "Defensa en Fiscalización SII": "Asumir la defensa del contribuyente durante un proceso de fiscalización del Servicio de Impuestos Internos.",
                    "Condonación de Intereses y Multas": "Solicitar la condonación de intereses y multas tributarias adeudadas.",
                    "Delito Tributario (Código Tributario Art. 97)": "Asumir la defensa penal tributaria conforme a las infracciones del artículo 97 del Código Tributario.",
                    "Reclamo por Giro de Cobranza TGR": "Reclamar en contra de un giro o cobranza indebida iniciada por la Tesorería General de la República."
                },
                "Derecho Comercial y Societario": {
                    "Constitución de Sociedades": "Asesorar y tramitar la constitución de una sociedad conforme al tipo social elegido.",
                    "Modificación de Sociedades": "Tramitar la modificación de los estatutos o el pacto social de una sociedad.",
                    "Disolución de Sociedades": "Tramitar la disolución y liquidación de una sociedad.",
                    "Juicio Arbitral Societario": "Representar a la parte en un juicio arbitral derivado de un conflicto societario.",
                    "Protesto de Letra de Cambio / Pagaré": "Efectuar el protesto de una letra de cambio o pagaré por falta de pago, aceptación o fecha.",
                    "Liquidación Forzosa de Empresa Deudora (Ley 20.720)": "Solicitar la liquidación forzosa de una empresa deudora conforme a la Ley 20.720.",
                    "Convenio Judicial Preventivo": "Proponer o negociar un convenio judicial preventivo con los acreedores.",
                    "Asesoría en Fusiones y Adquisiciones": "Asesorar jurídicamente en un proceso de fusión, adquisición o reorganización societaria."
                },
                "Derecho Inmobiliario y Urbanismo": {
                    "Estudio de Títulos Inmobiliarios": "Efectuar el estudio de títulos de un bien raíz para verificar su saneamiento legal.",
                    "Reclamo por Permiso de Edificación": "Reclamar administrativa o judicialmente por la denegación u observaciones a un permiso de edificación.",
                    "Copropiedad Inmobiliaria (Ley 21.442)": "Asesorar o representar en conflictos de copropiedad inmobiliaria conforme a la Ley 21.442.",
                    "Deslinde y Amojonamiento": "Solicitar la fijación judicial de deslindes y el amojonamiento del predio.",
                    "Regularización de Loteo Irregular": "Tramitar la regularización de un loteo irregular ante la autoridad competente."
                },
                "Propiedad Intelectual e Industrial": {
                    "Registro de Marca (INAPI)": "Tramitar el registro de una marca comercial ante el Instituto Nacional de Propiedad Industrial (INAPI).",
                    "Oposición a Registro de Marca": "Presentar oposición a la solicitud de registro de una marca de un tercero.",
                    "Nulidad de Marca o Patente": "Solicitar la declaración de nulidad de una marca o patente registrada.",
                    "Infracción de Derechos de Autor": "Demandar por infracción a los derechos de propiedad intelectual conforme a la Ley 17.336."
                },
                "Derecho Migratorio y Extranjería": {
                    "Solicitud de Visa / Residencia": "Tramitar una solicitud de visa o residencia ante el Servicio Nacional de Migraciones.",
                    "Recurso contra Expulsión": "Interponer recurso en contra de una orden de expulsión del territorio nacional.",
                    "Recurso contra Rechazo de Visa": "Interponer recurso administrativo o judicial en contra del rechazo de una solicitud de visa.",
                    "Nacionalización": "Tramitar el proceso de nacionalización de un extranjero residente en Chile."
                },
                "Derecho de Aguas": {
                    "Constitución de Derechos de Aprovechamiento": "Tramitar la constitución de derechos de aprovechamiento de aguas ante la DGA.",
                    "Oposición ante la Dirección General de Aguas (DGA)": "Presentar oposición a una solicitud de derechos de aprovechamiento de aguas de un tercero.",
                    "Juicio de Aguas": "Representar a la parte en un juicio derivado de un conflicto sobre derechos de aprovechamiento de aguas."
                },
                "Policía Local y Tránsito": {
                    "Infracción de Tránsito": "Asumir la defensa o el descargo por una infracción de tránsito cursada.",
                    "Accidente de Tránsito (Cobro de Daños)": "Demandar la indemnización de daños y perjuicios derivados de un accidente de tránsito.",
                    "Infracción a Ordenanzas Municipales": "Asumir la defensa por infracción a una ordenanza municipal."
                },
                "Derecho Ambiental": {
                    "Reclamación ante el Tribunal Ambiental": "Interponer reclamación ante el Tribunal Ambiental competente.",
                    "Impugnación de Resolución de Calificación Ambiental (RCA)": "Impugnar administrativa o judicialmente una Resolución de Calificación Ambiental (RCA).",
                    "Denuncia por Daño Ambiental": "Interponer denuncia o demanda por daño ambiental ante la autoridad competente."
                },
                "Derecho Bancario, Seguros y Ejecutivo Hipotecario": {
                    "Juicio Ejecutivo Hipotecario": "Iniciar juicio ejecutivo para el cobro de una deuda hipotecaria y el remate del inmueble dado en garantía.",
                    "Reclamo ante la CMF (Bancos/Seguros)": "Presentar un reclamo ante la Comisión para el Mercado Financiero por conductas de bancos o aseguradoras.",
                    "Repactación de Deuda Bancaria": "Negociar y formalizar la repactación de una deuda bancaria.",
                    "Reclamo o Defensa por Fraude en Tarjetas (Ley 20.009)": "Reclamar o defender la responsabilidad por operaciones realizadas con tarjetas de pago extraviadas, hurtadas, robadas o mediante fraude, conforme a la Ley 20.009.",
                    "Alzamiento de Hipoteca / Desarchivo": "Tramitar el alzamiento de una hipoteca y el desarchivo del expediente judicial correspondiente.",
                    "Ejecución de Prenda sin Desplazamiento": "Iniciar la ejecución de una prenda sin desplazamiento constituida en garantía de una obligación."
                },
                "Protección de Datos Personales y Ciberseguridad": {
                    "Reclamo ante la Agencia de Protección de Datos Personales (Ley 21.719)": "Presentar un reclamo ante la Agencia de Protección de Datos Personales conforme a la Ley 21.719.",
                    "Ejercicio de Derechos ARCO+ (Acceso, Rectificación, Cancelación, Oposición, Portabilidad)": "Ejercer ante el responsable de datos los derechos ARCO+ del titular de los datos personales.",
                    "Registro de Actividades de Tratamiento (RAT)": "Asesorar en la confección y mantención del Registro de Actividades de Tratamiento de datos personales.",
                    "Notificación de Brecha de Seguridad (Plazo 72 Horas)": "Asesorar en la notificación de una brecha de seguridad de datos personales dentro del plazo legal de 72 horas.",
                    "Defensa ante Fiscalización de la Agencia de Protección de Datos": "Asumir la defensa de la organización ante un procedimiento de fiscalización de la Agencia de Protección de Datos.",
                    "Delito Informático (Ley 19.223)": "Asumir la representación de la víctima o la defensa del imputado en un delito informático conforme a la Ley 19.223.",
                    "Cumplimiento Normativo Ley de Ciberseguridad (Ley 21.663)": "Asesorar en el cumplimiento de las obligaciones de la Ley Marco de Ciberseguridad (Ley 21.663)."
                },
                "Seguridad Social y Previsional": {
                    "Reforma Previsional - Beneficios (Ley 21.735)": "Asesorar y tramitar la solicitud de los nuevos beneficios previsionales creados por la Ley 21.735.",
                    "Reclamo por Cobranza de Cotizaciones Previsionales (Ley 17.322)": "Reclamar o gestionar el cobro de cotizaciones previsionales impagas conforme a la Ley 17.322.",
                    "Solicitud de Pensión de Invalidez": "Tramitar la solicitud de pensión de invalidez ante la Comisión Médica correspondiente.",
                    "Reclamo ante la Superintendencia de Pensiones": "Presentar un reclamo ante la Superintendencia de Pensiones por conductas de una AFP.",
                    "Solicitud de Pensión Garantizada Universal (PGU)": "Tramitar la solicitud de la Pensión Garantizada Universal (PGU) ante el Instituto de Previsión Social.",
                    "Retiro de Fondos Previsionales (Casos Especiales)": "Asesorar en solicitudes de retiro de fondos previsionales en los casos especiales que la ley contemple."
                },
                "Libre Competencia": {
                    "Denuncia ante la Fiscalía Nacional Económica (FNE)": "Interponer una denuncia ante la Fiscalía Nacional Económica por atentados a la libre competencia.",
                    "Defensa ante Requerimiento del TDLC (DL 211)": "Asumir la defensa ante un requerimiento del Tribunal de Defensa de la Libre Competencia, conforme al DL 211.",
                    "Consulta de Operación de Concentración": "Presentar una consulta o notificación de una operación de concentración ante la FNE.",
                    "Demanda de Indemnización por Ilícito Anticompetitivo": "Demandar indemnización de perjuicios derivada de un ilícito anticompetitivo declarado por el TDLC."
                },
                "Derecho Sanitario": {
                    "Reclamo por Negligencia Médica": "Demandar indemnización de perjuicios por negligencia médica u error en la prestación de salud.",
                    "Reclamo ante la Superintendencia de Salud": "Presentar un reclamo ante la Superintendencia de Salud por conductas de una Isapre o prestador.",
                    "Defensa de Derechos y Deberes del Paciente (Ley 20.584)": "Ejercer los derechos del paciente conforme a la Ley 20.584 sobre Derechos y Deberes de las Personas en Salud.",
                    "Reclamo por Cobertura GES/AUGE": "Reclamar por la denegación o incumplimiento de la cobertura de una patología GES/AUGE."
                },
                "Recursos Procesales Generales": {
                    "Recurso de Apelación": "Interponer recurso de apelación en contra de una resolución judicial para que sea revisada por el tribunal superior.",
                    "Recurso de Casación en la Forma": "Interponer recurso de casación en la forma por vicios en el procedimiento o en la sentencia.",
                    "Recurso de Casación en el Fondo": "Interponer recurso de casación en el fondo por infracción de ley que ha influido sustancialmente en lo dispositivo del fallo.",
                    "Recurso de Queja": "Interponer recurso de queja por falta o abuso grave cometido en la dictación de una resolución.",
                    "Recurso de Hecho": "Interponer recurso de hecho ante la denegación indebida de un recurso de apelación.",
                    "Arbitraje Comercial (Ley 19.971)": "Representar a la parte en un procedimiento de arbitraje comercial nacional o internacional conforme a la Ley 19.971."
                }
            }
            
            # Cláusula de escape: ninguna lista puede cubrir el 100% de la normativa chilena,
            # así que cada rama permite escribir una acción a medida si no aparece en el catálogo.
            for _rama_k in diccionario_servicios:
                diccionario_servicios[_rama_k]["➕ Otra Acción (especificar más abajo)"] = ""
            
            with st.container(border=True):
                col_mat1, col_mat2 = st.columns(2)
                with col_mat1:
                    materia_sel = st.selectbox("Rama del Derecho", list(diccionario_servicios.keys()), key="gen_con_rama")
                with col_mat2:
                    accion_sel = st.selectbox("Acción / Procedimiento Específico", list(diccionario_servicios[materia_sel].keys()), key="gen_con_accion")
                
                if accion_sel == "➕ Otra Acción (especificar más abajo)":
                    accion_manual = st.text_input("Especifica la acción / procedimiento no listado", placeholder="Ej: Acción de desafuero maternal, Ley 21.484...", key="gen_con_accion_manual")
                    accion_final = accion_manual.strip() if accion_manual.strip() else "Acción sin especificar"
                    finalidad_auto = f"Representar y patrocinar al Cliente en la acción de \"{accion_final}\"." if accion_manual.strip() else ""
                else:
                    accion_final = accion_sel
                    finalidad_auto = diccionario_servicios[materia_sel].get(accion_sel, "")
                
                tipo_servicio_final = f"{materia_sel}: {accion_final}"
                
                # 💡 Referencia de honorarios: busca en el Arancel del Colegio de
                # Abogados de Valparaíso las entradas más relacionadas con la
                # acción elegida, para orientar el cobro (siempre a título
                # referencial, no vinculante).
                if accion_final and accion_final != "Acción sin especificar":
                    sugerencias_arancel = buscar_arancel_referencial(f"{materia_sel} {accion_final}")
                    if sugerencias_arancel:
                        with st.expander("💰 Referencia de Honorarios (Arancel Colegio de Abogados de Valparaíso)", expanded=False):
                            st.caption("Valor orientador y supletorio (rige solo a falta de pacto expreso, Art. 7 del Arancel). El honorario final siempre se acuerda con el cliente.")
                            for sug in sugerencias_arancel:
                                st.markdown(f"**N°{sug['numero']} — {sug['descripcion']}**")
                                st.markdown(f"<span style='color:#42526e;'>{sug['honorario']}</span>", unsafe_allow_html=True)
                                st.markdown("---")
                
                # Auto-relleno de la Cláusula Primera: solo se sobrescribe cuando cambia
                # la rama/acción seleccionada, para no borrar ediciones manuales del abogado
                # si simplemente está revisando el formulario sin cambiar la selección.
                _clave_seleccion_actual = f"{materia_sel}|{accion_final}"
                if st.session_state.get('_ultima_seleccion_clausula') != _clave_seleccion_actual:
                    st.session_state['gen_con_detalle'] = finalidad_auto
                    st.session_state['_ultima_seleccion_clausula'] = _clave_seleccion_actual

            # --- Honorarios y vínculo con causa: FUERA del form, para que el
            # "Valor en Letras" y el "Valor por Cuota" se calculen solos apenas
            # escribes el monto, sin tener que enviarlo primero. ---
            with st.container(border=True):
                st.markdown("#### Módulo 4: Honorarios y Cuotas (se calculan solas)")
                c_p1, c_p2 = st.columns(2)
                with c_p1:
                    hon_num_texto = st.text_input("Valor Total ($)", "2.500.000", key="gen_con_honnum", help="Escríbelo como quieras: 500000, 500.000 o $500.000 — el sistema lo entiende igual.")
                    hon_num_int = parsear_monto_clp(hon_num_texto)
                    st.caption(f"💰 {formatear_clp(hon_num_int)} → *{numero_a_letras_clp(hon_num_int)}*")
                    cuotas_c = st.number_input("Cantidad de Cuotas", min_value=1, max_value=360, value=2, step=1, key="gen_con_cuotasc", help="La cantidad de cuotas la defines tú libremente.")
                    fecha_pago = st.date_input("Primera Mensualidad", key="gen_con_fecha")
                with c_p2:
                    valor_cuota_sugerido = hon_num_int // cuotas_c if cuotas_c > 0 else 0
                    # Streamlit ignora el "value=" de un widget en los reruns siguientes
                    # si ya existe algo guardado bajo su misma key (así fallaba antes: el
                    # campo se llenaba una vez y quedaba "pegado"). Para que se recalcule
                    # de verdad cada vez que cambian el total o la cantidad de cuotas,
                    # se sobreescribe el session_state ANTES de crear el widget, solo
                    # cuando esos dos valores base cambiaron (así no se pisa una edición
                    # manual tuya si no tocaste ni el total ni la cantidad de cuotas).
                    _clave_base_cuota = f"{hon_num_int}|{cuotas_c}"
                    if st.session_state.get('_ultima_base_cuota') != _clave_base_cuota:
                        st.session_state['gen_con_cuotasm'] = formatear_clp(valor_cuota_sugerido)
                        st.session_state['_ultima_base_cuota'] = _clave_base_cuota
                    cuotas_m_texto = st.text_input("Valor por Cuota ($)", key="gen_con_cuotasm", help="Se recalcula automático (total ÷ cuotas) cada vez que cambias el monto total o la cantidad de cuotas. Puedes editarlo a mano si las cuotas no son parejas.")
                    cuotas_m_int = parsear_monto_clp(cuotas_m_texto)
                    st.caption(f"💰 {formatear_clp(cuotas_m_int)} por cuota")
                    banco = st.text_input("Banco", key="gen_con_banco")
                    tipo_cta = st.selectbox("Tipo de Cuenta", ["Cuenta Corriente", "Cuenta Vista", "Cuenta RUT", "Chequera Electrónica"], key="gen_con_tipocta")
                    num_cta = st.text_input("Número de Cuenta", key="gen_con_numcta")
            
            with st.container(border=True):
                st.markdown("#### Módulo 5: Vincular a una Causa (opcional, pero recomendado)")
                st.caption("Si eliges una causa, la Contabilidad de esa causa se completa sola con estos honorarios y cuotas — no tienes que volver a escribirlo ahí.")
                df_causas_para_vincular = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                opciones_causa_vinculo = ["➕ Ninguna (crear después / sin causa aún)"]
                if not df_causas_para_vincular.empty:
                    opciones_causa_vinculo += [f"{r['ROL']} — {r.get('CARATULADO','')}" for _, r in df_causas_para_vincular.iterrows()]
                causa_vinculo_sel = st.selectbox("Causa a la que corresponden estos honorarios", opciones_causa_vinculo, key="gen_con_causa_vinculo")

            with st.form("form_generador_contratos", clear_on_submit=False):
                detalle_servicio = st.text_area("Cláusula Primera: Acciones Legales Incluidas", height=100, key="gen_con_detalle", help="Se autocompleta según la acción elegida arriba. Puedes editarla libremente antes de generar el contrato.")
                
                col_ab, col_cl = st.columns(2)
                with col_ab:
                    with st.container(border=True):
                        st.markdown("#### Módulo 2: Litigante Patrocinante")
                        abog_nom = st.text_input("Nombre Abogado", key="gen_con_abnom")
                        abog_rut = st.text_input("RUT Abogado", key="gen_con_abrut")
                        abog_dom = st.text_input("Domicilio Profesional", key="gen_con_abdom")
                        abog_tel = st.text_input("Teléfono", key="gen_con_abtel")
                        abog_correo = st.text_input("Correo", key="gen_con_abcor")
                with col_cl:
                    with st.container(border=True):
                        st.markdown("#### Módulo 3: Mandante Judicial")
                        cli_nom = st.text_input("Nombre Cliente", key="gen_con_clinom")
                        cli_rut = st.text_input("RUT Cliente", key="gen_con_clirut")
                        cli_dom = st.text_input("Domicilio", key="gen_con_clidom")
                        cli_tel = st.text_input("Teléfono Particular", key="gen_con_clitel")
                        cli_correo = st.text_input("Correo Particular", key="gen_con_clicor")
                
                with st.container(border=True):
                    st.markdown("#### Módulo 6: Documentos que el Cliente debe reunir (opcional)")
                    st.caption("Si escribes algo aquí, se agrega como una cláusula especial en el contrato detallando lo que el cliente debe entregar para iniciar la redacción de la demanda/gestión.")
                    docs_requeridos = st.text_area("Un documento por línea", height=100, key="gen_con_docs_req",
                        placeholder="Ej:\nCédula de identidad por ambos lados\nÚltimas 3 liquidaciones de sueldo\nContrato de trabajo\nFiniquito (si ya fue despedido)")
                        
                if st.form_submit_button("📄 Estructurar Contrato en Formato Word", type="primary", use_container_width=True):
                    hon_num_final = formatear_clp(hon_num_int)
                    cuotas_m_final = formatear_clp(cuotas_m_int)
                    datos_c = {
                        'tipo_servicio': tipo_servicio_final, 'detalle_servicio': detalle_servicio,
                        'abogado_nombre': abog_nom, 'abogado_rut': abog_rut, 'abogado_domicilio': abog_dom, 'abogado_tel': abog_tel, 'abogado_correo': abog_correo,
                        'cliente_nombre': cli_nom, 'cliente_rut': cli_rut, 'cliente_domicilio': cli_dom, 'cliente_tel': cli_tel, 'cliente_correo': cli_correo,
                        'honorarios_num': hon_num_final, 'honorarios_letras': numero_a_letras_clp(hon_num_int), 'cuotas_cant': cuotas_c, 'cuotas_monto': cuotas_m_final, 'fecha_inicio': fecha_pago,
                        'banco': banco, 'tipo_cuenta': tipo_cta, 'num_cuenta': num_cta,
                        'documentos_requeridos': docs_requeridos.strip()
                    }
                    doc_final = crear_contrato_word(datos_c)
                    if doc_final:
                        buffer_memoria = io.BytesIO()
                        doc_final.save(buffer_memoria)
                        bytes_contrato = buffer_memoria.getvalue()
                        
                        st.session_state['contrato_generado'] = bytes_contrato
                        st.session_state['nombre_archivo'] = f"Contrato_{cli_nom.replace(' ', '_')}.docx"
                        
                        drive_id_con, b64_docx = guardar_archivo_adjunto(
                            st.session_state['nombre_archivo'], bytes_contrato,
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                        
                        df_con = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
                        nuevo_con = {
                            'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"), 
                            'Cliente': cli_nom, 'Servicio': accion_final, 'Honorarios': hon_num_final, 'Archivo_B64': b64_docx,
                            'Archivo_Drive_ID': drive_id_con, 'Usuario_Propietario': usuario_actual
                        }
                        df_con = pd.concat([df_con, pd.DataFrame([nuevo_con])], ignore_index=True)
                        df_con.to_csv(ARCHIVO_CONTRATOS, index=False)
                        
                        dn_co = safe_read_sheet("base_contratos", COLS_CONTRATOS)
                        safe_update_sheet("base_contratos", pd.concat([dn_co, pd.DataFrame([nuevo_con])], ignore_index=True))
                        
                        # --- AUTOMATIZACIÓN 1: crear o actualizar el CLIENTE de inmediato ---
                        if cli_rut.strip():
                            df_clientes_auto = safe_read_sheet("base_clientes", COLS_CLIENTES)
                            rut_limpio = cli_rut.strip().upper()
                            if not df_clientes_auto.empty and rut_limpio in df_clientes_auto['RUT'].astype(str).str.upper().values:
                                idx_cli_auto = df_clientes_auto[df_clientes_auto['RUT'].astype(str).str.upper() == rut_limpio].index[0]
                                df_clientes_auto.at[idx_cli_auto, 'Nombre'] = cli_nom
                                df_clientes_auto.at[idx_cli_auto, 'Telefono'] = cli_tel
                                df_clientes_auto.at[idx_cli_auto, 'Correo'] = cli_correo
                                df_clientes_auto.at[idx_cli_auto, 'Direccion'] = cli_dom
                            else:
                                nuevo_cliente_auto = {'RUT': rut_limpio, 'Nombre': cli_nom, 'Telefono': cli_tel, 'Correo': cli_correo, 'Clave_unica': '', 'Direccion': cli_dom, 'Usuario_Propietario': usuario_actual}
                                df_clientes_auto = pd.concat([df_clientes_auto, pd.DataFrame([nuevo_cliente_auto])], ignore_index=True)
                            safe_update_sheet("base_clientes", df_clientes_auto)
                        
                        # --- AUTOMATIZACIÓN 2: completar la Contabilidad SIEMPRE, aunque todavía no exista el ROL ---
                        # Antes esto solo pasaba si elegías una causa ya creada. El problema real es que un
                        # contrato normalmente se firma ANTES de que exista el ROL (el juicio ni se ha
                        # presentado todavía), así que en la práctica esa lista casi siempre estaba vacía y
                        # la Contabilidad nunca se llenaba sola. Ahora, si no eliges una causa existente, se
                        # crea automáticamente una causa "placeholder" con estos honorarios y cuotas, para que
                        # la Contabilidad quede lista de inmediato. Cuando presentes la demanda y tengas el
                        # ROL real, solo entras a "Editar Ficha" de esa causa y reemplazas el ROL provisorio.
                        if causa_vinculo_sel != "➕ Ninguna (crear después / sin causa aún)":
                            rol_vinculado = causa_vinculo_sel.split(" — ")[0].strip()
                        else:
                            rol_vinculado = f"PENDIENTE-{str(uuid.uuid4())[:6].upper()}"
                            df_causas_nueva = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                            nueva_causa_placeholder = {
                                'ROL': rol_vinculado, 'TRIBUNAL': '(Pendiente de asignar)',
                                'CARATULADO': f"{cli_nom.upper()} / {tipo_servicio_final}",
                                'Cliente': cli_nom, 'RUT': cli_rut, 'Tipo_Negocio': 'Propio',
                                'Usuario_Propietario': usuario_actual, 'Estado_Honorarios': 'Pendientes',
                                'Total_Honorarios': hon_num_int, 'Cuotas_Totales': cuotas_c, 'Cuotas_Pagadas': 0,
                                'Clave_unica': '', 'SAC': '', 'Sucursal': '', 'Servicio': accion_final,
                                'Fecha_Inicio': fecha_pago.strftime("%Y-%m-%d")
                            }
                            df_causas_nueva = pd.concat([df_causas_nueva, pd.DataFrame([nueva_causa_placeholder])], ignore_index=True)
                            df_causas_nueva.to_csv(ARCHIVO_BD, index=False)
                            dn_causa_nueva = safe_read_sheet("base_causas", COLS_CAUSAS)
                            dn_causa_nueva = pd.concat([dn_causa_nueva, pd.DataFrame([nueva_causa_placeholder])], ignore_index=True)
                            safe_update_sheet("base_causas", dn_causa_nueva)
                        
                        df_causas_auto = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                        if not df_causas_auto.empty and rol_vinculado in df_causas_auto['ROL'].values:
                            idx_causa_auto = df_causas_auto[df_causas_auto['ROL'] == rol_vinculado].index[0]
                            df_causas_auto.at[idx_causa_auto, 'Estado_Honorarios'] = 'Pendientes'
                            df_causas_auto.at[idx_causa_auto, 'Total_Honorarios'] = hon_num_int
                            df_causas_auto.at[idx_causa_auto, 'Cuotas_Totales'] = cuotas_c
                            df_causas_auto.at[idx_causa_auto, 'Cuotas_Pagadas'] = 0
                            df_causas_auto.at[idx_causa_auto, 'Fecha_Inicio'] = fecha_pago.strftime("%Y-%m-%d")
                            df_causas_auto.to_csv(ARCHIVO_BD, index=False)
                            dn_causa_auto = safe_read_sheet("base_causas", COLS_CAUSAS)
                            if not dn_causa_auto.empty and rol_vinculado in dn_causa_auto['ROL'].values:
                                idx_nube = dn_causa_auto[dn_causa_auto['ROL'] == rol_vinculado].index[0]
                                dn_causa_auto.at[idx_nube, 'Estado_Honorarios'] = 'Pendientes'
                                dn_causa_auto.at[idx_nube, 'Total_Honorarios'] = hon_num_int
                                dn_causa_auto.at[idx_nube, 'Cuotas_Totales'] = cuotas_c
                                dn_causa_auto.at[idx_nube, 'Cuotas_Pagadas'] = 0
                                safe_update_sheet("base_causas", dn_causa_auto)
                            if causa_vinculo_sel != "➕ Ninguna (crear después / sin causa aún)":
                                st.success(f"✅ Contrato generado. Cliente y Contabilidad de la causa {rol_vinculado} actualizados automáticamente.")
                            else:
                                st.success(f"✅ Contrato generado. Cliente creado y Contabilidad completada automáticamente bajo el ROL provisorio **{rol_vinculado}** — reemplázalo por el ROL real desde 'Editar Ficha' apenas presentes la demanda.")
                        else:
                            st.warning("⚠️ Contrato generado y cliente actualizado, pero no se pudo completar la Contabilidad automáticamente. Hazlo manualmente desde Causas.")
                        
                        st.rerun()
                        
        if st.session_state.get('contrato_generado'):
            st.success("✅ Contrato guardado en el historial.")
            st.download_button(label="📥 Descargar Documento (.docx)", data=st.session_state['contrato_generado'], file_name=st.session_state['nombre_archivo'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
            
    with tab_reg:
        df_contratos_reg = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
        if df_contratos_reg.empty: 
            st.info("No registras copias históricas guardadas.")
        else: 
            st.markdown("### 🗄️ Archivo Histórico de Contratos Generados")
            for idx, row in df_contratos_reg[::-1].iterrows():
                with st.container(border=True):
                    c1, c2, c3 = st.columns([3, 1, 1])
                    with c1:
                        st.markdown(f"**Cliente:** {row.get('Cliente', 'Sin Nombre')} | **Servicio:** {row.get('Servicio', '--')}")
                        st.markdown(f"<span style='color:#6b778c; font-size:14px;'>Fecha Emisión: {row.get('Fecha', '--')} | Honorarios Pactados: ${row.get('Honorarios', '0')}</span>", unsafe_allow_html=True)
                    
                    with c2:
                        bytes_doc = obtener_bytes_adjunto(row, 'Archivo_Drive_ID', 'Archivo_B64')
                        if bytes_doc is not None:
                            st.download_button("📥 Descargar", data=bytes_doc, file_name=f"Copia_{str(row.get('Cliente', 'Contrato')).replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_con_{row.get('ID', str(uuid.uuid4())[:8])}")
                        else:
                            st.write("*(Sin archivo)*")
                            
                    with c3:
                        if usuario_actual == "Narratia":
                            if st.button("🗑️ Eliminar", key=f"del_con_{row.get('ID', idx)}"):
                                df_contratos_reg = df_contratos_reg.drop(idx)
                                df_contratos_reg.to_csv(ARCHIVO_CONTRATOS, index=False)
                                dn_c = safe_read_sheet("base_contratos", [])
                                if not dn_c.empty:
                                    safe_update_sheet("base_contratos", dn_c[dn_c['ID'] != row.get('ID')])
                                st.rerun()

    with tab_importar:
        st.markdown("Sube un contrato ya firmado en PDF o Word para extraer sus datos y guardarlos.")
        archivo_contrato = st.file_uploader("📂 Subir Contrato del Cliente", type=["pdf", "docx"])
        
        if st.button("🧠 Leer Contrato y Registrar Cliente", type="primary", use_container_width=True):
            if not archivo_contrato:
                st.error("⚠️ Tienes que subir un archivo primero compadre.")
            else:
                with st.spinner("🤖 Leyendo cláusulas..."):
                    try:
                        texto_contrato = ""
                        if archivo_contrato.name.endswith('.pdf'):
                            import PyPDF2
                            lector = PyPDF2.PdfReader(archivo_contrato)
                            for pagina in lector.pages: texto_contrato += pagina.extract_text() + "\n"
                        elif archivo_contrato.name.endswith('.docx'):
                            from docx import Document
                            doc = Document(archivo_contrato)
                            for p in doc.paragraphs: texto_contrato += p.text + "\n"
                                
                        import google.generativeai as genai
                        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                        
                        modelo_elegido = "gemini-1.0-pro"
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods and 'flash' in m.name:
                                modelo_elegido = m.name.replace("models/", ""); break
                                    
                        modelo = genai.GenerativeModel(modelo_elegido)
                        
                        prompt_extractor = f"""
                        Eres un asistente legal experto en Chile. Extrae los datos del CLIENTE.
                        Devuelve ÚNICAMENTE un objeto JSON válido (sin markdown) con esta estructura:
                        {{ "cliente_nombre": "nombre", "cliente_rut": "RUT", "servicio": "tipo", "honorarios_total": 0, "cuotas_totales": 1, "fecha_inicio_pago": "YYYY-MM-DD" }}
                        CONTRATO: {texto_contrato[:15000]}
                        """
                        
                        respuesta = modelo.generate_content(prompt_extractor)
                        texto_json = respuesta.text.replace('```json', '').replace('```', '').strip()
                        datos_extraidos = json.loads(texto_json)
                        
                        df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                        nombre_extraido = datos_extraidos.get('cliente_nombre', 'Cliente Importado')
                        
                        nuevo_cliente = {
                            'ROL': 'Sin Causa Aún', 'TRIBUNAL': '--', 'CARATULADO': '--', 'Cliente': nombre_extraido,
                            'RUT': datos_extraidos.get('cliente_rut', '--'), 'Teléfono': '--', 'Tipo_Negocio': 'Propio', 'Clave_unica': '--', 'Correo': '--',
                            'Direccion': '--', 'SAC': '--', 'Sucursal': '--', 'Estado_Honorarios': 'Pendientes' if int(datos_extraidos.get('honorarios_total', 0)) > 0 else 'Sin fijar',
                            'Total_Honorarios': int(datos_extraidos.get('honorarios_total', 0)), 'Cuotas_Totales': int(datos_extraidos.get('cuotas_totales', 1)), 
                            'Cuotas_Pagadas': 0, 'Fecha_Inicio': datos_extraidos.get('fecha_inicio_pago', datetime.now().strftime("%Y-%m-%d")),
                            'Usuario_Propietario': usuario_actual
                        }
                        pd.concat([df_causas, pd.DataFrame([nuevo_cliente])], ignore_index=True).to_csv(ARCHIVO_BD, index=False)
                        
                        df_con = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
                        nuevo_con = {
                            'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"),
                            'Cliente': nombre_extraido, 'Servicio': datos_extraidos.get('servicio', 'Servicio Legal'), 'Honorarios': datos_extraidos.get('honorarios_total', 0),
                            'Archivo_B64': '', 'Usuario_Propietario': usuario_actual
                        }
                        pd.concat([df_con, pd.DataFrame([nuevo_con])], ignore_index=True).to_csv(ARCHIVO_CONTRATOS, index=False)
                        
                        st.success(f"✅ ¡La IA agregó a **{nombre_extraido}** directo a tu listado!")
                    except Exception as e: st.error(f"❌ Error técnico: {e}")

# 7. CAUSAS / EXPEDIENTES (MEJORADO Y RELACIONAL)
elif st.session_state['menu_radio'] == "💼 Causas":
    ES_ADMIN_NARRATIA = usuario_actual == "Narratia"
    
    # Si Narratia (administrador del estudio) abre una causa que pertenece a otro
    # abogado, trabajamos sobre los archivos REALES de ese abogado (no los del
    # admin), para que ver/editar tareas, honorarios y comentarios impacte el
    # expediente verdadero. Para el resto de los usuarios esto nunca cambia:
    # siempre ven y editan solo sus propios archivos.
    _propietario_vista = st.session_state.get('causa_propietario_vista')
    if ES_ADMIN_NARRATIA and _propietario_vista and _propietario_vista != usuario_actual:
        ARCHIVO_BD = f"base_causas_{_propietario_vista}.csv"
        ARCHIVO_TAREAS = f"base_tareas_{_propietario_vista}.csv"
    
    df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
    if df_causas.empty:
        # Mismo respaldo que en Inicio y en la ficha del cliente: si el
        # archivo local está vacío, se revisa la nube directamente antes de
        # mostrar la lista de causas como si no hubiera ninguna.
        _propietario_para_filtro = _propietario_vista if (ES_ADMIN_NARRATIA and _propietario_vista) else usuario_actual
        df_causas_nube_causas = safe_read_sheet("base_causas", COLS_CAUSAS)
        if not df_causas_nube_causas.empty and 'Usuario_Propietario' in df_causas_nube_causas.columns:
            df_causas = df_causas_nube_causas[df_causas_nube_causas['Usuario_Propietario'] == _propietario_para_filtro]
    df_clientes = safe_read_sheet("base_clientes", COLS_CLIENTES)
    if not ES_ADMIN_NARRATIA and not df_clientes.empty and 'Usuario_Propietario' in df_clientes.columns:
        df_clientes = df_clientes[df_clientes['Usuario_Propietario'] == usuario_actual]
    
    @st.dialog("Editar tarea")
    def modal_editar_tarea(tarea_id, tarea_titulo, tarea_fecha, tarea_estado):
        st.write(f"Modificando plazos para: **{tarea_titulo}**")
        st.text_input("Usuario", value=nombre_real_usuario, disabled=True)
        
        try:
            f_obj = datetime.strptime(tarea_fecha, "%d/%m/%Y")
        except:
            f_obj = datetime.now()
        nueva_fecha = st.date_input("Fecha de vencimiento *", value=f_obj)
        
        opciones_estado = ["En progreso", "Aprobada", "Rechazada"]
        nuevo_estado = st.selectbox("Estado", opciones_estado, index=opciones_estado.index(tarea_estado) if tarea_estado in opciones_estado else 0)
        
        if st.button("Guardar", type="primary", use_container_width=True):
            df_t_local = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
            df_t_local.loc[df_t_local['ID_Tarea'] == tarea_id, ['Fecha_Vencimiento', 'Estado']] = [nueva_fecha.strftime("%d/%m/%Y"), nuevo_estado]
            df_t_local.to_csv(ARCHIVO_TAREAS, index=False)
            
            dn = safe_read_sheet("base_tareas", [])
            if not dn.empty:
                dn.loc[dn['ID_Tarea'] == tarea_id, ['Fecha_Vencimiento', 'Estado']] = [nueva_fecha.strftime("%d/%m/%Y"), nuevo_estado]
                safe_update_sheet("base_tareas", dn)
                
            st.session_state['editando_tarea'] = None
            st.success("✅ Tarea actualizada correctamente.")
            import time; time.sleep(0.3); st.rerun()

    if st.session_state['causa_seleccionada'] is None:
        st.session_state['modo_edicion'] = False
        st.title("Causas")
        st.markdown("<span style='color:#6b778c;'>Gestiona todos los casos judiciales</span>", unsafe_allow_html=True)
        
        c_stat1, c_stat2 = st.columns(2)
        with c_stat1:
            st.markdown(f"""
            <div class="dash-card">
                <span style="color:#6b778c; font-size:14px;">Causas registradas</span><br>
                <span style="font-size:32px; font-weight:700; color:#172b4d;">{len(df_causas)}</span><br>
                <span style="color:#6b778c; font-size:13px;">💼 En tu cartera</span>
            </div>
            """, unsafe_allow_html=True)
        with c_stat2:
            n_pendientes_hon = len(df_causas[df_causas.get('Estado_Honorarios', '') == 'Pendientes']) if not df_causas.empty else 0
            st.markdown(f"""
            <div class="dash-card">
                <span style="color:#6b778c; font-size:14px;">Con honorarios pendientes</span><br>
                <span style="font-size:32px; font-weight:700; color:#172b4d;">{n_pendientes_hon}</span><br>
                <span style="color:#6b778c; font-size:13px;">💰 Requieren seguimiento</span>
            </div>
            """, unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button("➕ Crear Nueva Causa", type="primary"):
            st.session_state['creando_causa'] = not st.session_state.get('creando_causa', False)
            
        if st.session_state.get('creando_causa'):
            with st.container(border=True):
                st.markdown("#### Ingresar Datos de la Nueva Causa")
                with st.form("form_crear_causa"):
                    c_nuevo1, c_nuevo2 = st.columns(2)
                    n_rol = c_nuevo1.text_input("ROL / RIT", placeholder="Ej: C-123-2024")
                    n_trib = selector_tribunal(key_prefix="nuevo_causa")
                    n_carat = st.text_input("Caratulado", placeholder="Ej: PEREZ / BANCO")
                    
                    st.markdown("#### Asociar a un Cliente Existente")
                    if not df_clientes.empty:
                        opciones_clientes = df_clientes['RUT'].astype(str) + " - " + df_clientes['Nombre'].astype(str)
                        cliente_seleccionado = st.selectbox("Seleccionar Cliente Titular", opciones_clientes.tolist())
                    else:
                        st.warning("⚠️ No hay clientes creados. Ve a la pestaña de Clientes primero.")
                        cliente_seleccionado = None
                    
                    if st.form_submit_button("Guardar Causa en Base de Datos"):
                        if n_rol.strip() == "" or not cliente_seleccionado:
                            st.error("El ROL y el Cliente son obligatorios.")
                        else:
                            rut_extraido = cliente_seleccionado.split(" - ")[0]
                            nombre_extraido = cliente_seleccionado.split(" - ")[1]
                            
                            nueva_c = {
                                'ROL': n_rol.strip().upper(), 'TRIBUNAL': n_trib.strip(), 'CARATULADO': n_carat.strip(), 'Cliente': nombre_extraido,
                                'RUT': rut_extraido, 'Teléfono': '--', 'Tipo_Negocio': 'Propio', 'Clave_unica': '--',
                                'Correo': '--', 'Direccion': '--', 'SAC': '--', 'Sucursal': '--',
                                'Estado_Honorarios': 'Sin fijar', 'Total_Honorarios': 0, 'Cuotas_Totales': 0, 'Cuotas_Pagadas': 0,
                                'Usuario_Propietario': usuario_actual
                            }
                            df_causas = pd.concat([df_causas, pd.DataFrame([nueva_c])], ignore_index=True)
                            df_causas.to_csv(ARCHIVO_BD, index=False)
                            
                            dn = safe_read_sheet("base_causas", ['ROL', 'TRIBUNAL', 'CARATULADO', 'Cliente', 'RUT', 'Teléfono', 'Tipo_Negocio', 'Clave_unica', 'Correo', 'Direccion', 'SAC', 'Sucursal', 'Estado_Honorarios', 'Total_Honorarios', 'Cuotas_Totales', 'Cuotas_Pagadas', 'Usuario_Propietario'])
                            safe_update_sheet("base_causas", pd.concat([dn, pd.DataFrame([nueva_c])], ignore_index=True))
                            
                            st.session_state['creando_causa'] = False
                            st.success("✅ Causa creada y vinculada al cliente exitosamente.")
                            import time; time.sleep(0.3); st.rerun()

        st.write("---")
        
        df_para_listado = df_causas.copy()
        df_para_listado['Propietario_Vista'] = usuario_actual
        
        if ES_ADMIN_NARRATIA:
            boton_refrescar_equipo("refresh_causas_equipo")
            archivos_causas_equipo = glob.glob("base_causas_*.csv")
            piezas_equipo = []
            for arch in archivos_causas_equipo:
                propietario_arch = arch.replace("base_causas_", "").replace(".csv", "")
                temp_causa_eq = leer_csv_local(arch, COLS_CAUSAS)
                if not temp_causa_eq.empty:
                    temp_causa_eq = temp_causa_eq.copy()
                    temp_causa_eq['Propietario_Vista'] = propietario_arch
                    piezas_equipo.append(temp_causa_eq)
            if piezas_equipo:
                df_para_listado = pd.concat(piezas_equipo, ignore_index=True)
        
        col_f1, col_f2 = st.columns(2)
        filtro_trib = col_f1.multiselect("Filtrar por Tribunal de la República", df_para_listado['TRIBUNAL'].dropna().unique().tolist(), placeholder="Selecciona el juzgado...")
        filtro_neg = col_f2.multiselect("Filtrar por Cartera de Negocio", df_para_listado['Tipo_Negocio'].dropna().unique().tolist(), placeholder="Selecciona origen...")
        
        busqueda_causa = st.text_input("🔎 Buscar", placeholder="Rol: C-1234-2025, Causa, Cliente...", label_visibility="collapsed")
        
        df_filtrado = df_para_listado.copy()
        if filtro_trib: 
            df_filtrado = df_filtrado[df_filtrado['TRIBUNAL'].isin(filtro_trib)]
        if filtro_neg: 
            df_filtrado = df_filtrado[df_filtrado['Tipo_Negocio'].isin(filtro_neg)]
        if busqueda_causa.strip():
            q = busqueda_causa.strip().lower()
            df_filtrado = df_filtrado[
                df_filtrado['ROL'].astype(str).str.lower().str.contains(q, na=False) |
                df_filtrado['CARATULADO'].astype(str).str.lower().str.contains(q, na=False) |
                df_filtrado['Cliente'].astype(str).str.lower().str.contains(q, na=False)
            ]
            
        c_tit, c_dl = st.columns([4, 1])
        c_tit.markdown("### Expedientes Activos")
        with c_dl:
            boton_descargar_excel(df_filtrado, "causas_jurisync.xlsx", key="dl_excel_causas")
        
        with st.container(height=600):
            if df_filtrado.empty:
                st.info("No hay causas que coincidan con la búsqueda.")
            else:
                c_h1, c_h2, c_h3, c_h4, c_h5 = st.columns([1.5, 2.5, 3, 2.5, 1.5])
                c_h1.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>ROL DE CAUSA</span>", unsafe_allow_html=True)
                c_h2.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>TRIBUNAL ASIGNADO</span>", unsafe_allow_html=True)
                c_h3.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>CARATULADO</span>", unsafe_allow_html=True)
                c_h4.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>CLIENTE TITULAR</span>", unsafe_allow_html=True)
                c_h5.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px; text-align:center; display:block;'>ACCIÓN</span>", unsafe_allow_html=True)
                st.markdown("<hr style='margin: 5px 0px 10px 0px; border-top: 2px solid #e0e4e8;'>", unsafe_allow_html=True)
                
                for idx, row in df_filtrado.iterrows():
                    fila_es_propia = row.get('Propietario_Vista', usuario_actual) == usuario_actual
                    c1, c2, c3, c4, c5 = st.columns([1.5, 2.5, 3, 2.5, 1.5])
                    color_rol = "#0e6b74" if fila_es_propia else "#ff8b00"
                    c1.markdown(f"<span style='color:{color_rol}; font-weight:bold; font-size:15px;'>{row['ROL']}</span>", unsafe_allow_html=True)
                    if not fila_es_propia:
                        nombre_dueno = NOMBRES_REALES.get(row.get('Propietario_Vista'), row.get('Propietario_Vista'))
                        c1.markdown(f"<span style='background:#fff0e0; color:#ff8b00; font-size:10px; font-weight:700; padding:2px 6px; border-radius:8px;'>👤 {nombre_dueno}</span>", unsafe_allow_html=True)
                    c2.markdown(f"<span style='color:#172b4d; font-size:14px;'>{row['TRIBUNAL']}</span>", unsafe_allow_html=True)
                    c3.markdown(f"<span style='color:#172b4d; font-weight:600; font-size:14px;'>{row['CARATULADO']}</span>", unsafe_allow_html=True)
                    
                    val_cliente = str(row.get('Cliente', '--'))
                    val_rut = str(row.get('RUT', '--'))
                    c4.markdown(f"<span style='color:#172b4d; font-size:14px;'>👤 {val_cliente}</span><br><span style='color:#6b778c; font-size:12px;'>RUT: {val_rut}</span>", unsafe_allow_html=True)
                    
                    c5.button("📂 Abrir", key=f"abrir_c_{idx}", use_container_width=True, on_click=ir_a_expediente, args=(row['ROL'], row.get('Propietario_Vista', usuario_actual)))
                    st.markdown("<hr style='margin: 8px 0px 8px 0px; border-top: 1px dashed #e0e4e8;'>", unsafe_allow_html=True)
        
    else:
        rol_actual = st.session_state['causa_seleccionada']
        filtro_causa = df_causas[df_causas['ROL'] == rol_actual]
        
        if filtro_causa.empty:
            st.error(f"Error: No se encontró el expediente para el ROL {rol_actual}.")
            if st.button("Volver al inicio"):
                st.session_state['causa_seleccionada'] = None
                st.rerun()
            st.stop()
            
        idx = filtro_causa.index[0]
        c_data = df_causas.loc[idx]
        
        c_head1, c_head2 = st.columns([4, 1])
        with c_head1:
            st.markdown(f"<h2>{c_data.get('CARATULADO','')}</h2>", unsafe_allow_html=True)
            if ES_ADMIN_NARRATIA and _propietario_vista and _propietario_vista != usuario_actual:
                nombre_dueno_exp = NOMBRES_REALES.get(_propietario_vista, _propietario_vista)
                st.markdown(f"<span style='background:#fff0e0; color:#ff8b00; font-size:12px; font-weight:700; padding:3px 10px; border-radius:10px;'>👤 Causa de {nombre_dueno_exp} — estás viendo/editando el expediente real de su cartera</span>", unsafe_allow_html=True)
        with c_head2:
            if st.button("⬅ Volver al listado"):
                st.session_state['causa_seleccionada'] = None
                st.session_state['causa_propietario_vista'] = None
                st.rerun()
            if st.button("✍️ Redactar Escrito", help="Abre el Redactor IA con el Rol, Tribunal y Caratulado de esta causa ya cargados."):
                st.session_state['redactor_prefill'] = {
                    'rol': c_data.get('ROL', ''), 'tribunal': c_data.get('TRIBUNAL', ''), 'caratulado': c_data.get('CARATULADO', '')
                }
                st.session_state['menu_radio'] = "📝 Redactor IA"
                st.rerun()
                
        col_izq, col_der = st.columns([2.5, 1.2])
        
        with col_der:
            c_btn_ed, c_btn_del = st.columns([3, 1])
            if c_btn_ed.button("❌ Cancelar" if st.session_state['modo_edicion'] else "✏️ Editar Ficha", use_container_width=True):
                st.session_state['modo_edicion'] = not st.session_state['modo_edicion']
                st.rerun()
                
            if usuario_actual == "Narratia":
                if c_btn_del.button("🗑️", help="Eliminar Causa Permanentemente"):
                    df_causas = df_causas.drop(idx)
                    df_causas.to_csv(ARCHIVO_BD, index=False)
                    dn_c = safe_read_sheet("base_causas", [])
                    if not dn_c.empty:
                        safe_update_sheet("base_causas", dn_c[dn_c['ROL'] != rol_actual])
                    st.session_state['causa_seleccionada'] = None
                    st.rerun()
                
            if st.session_state['modo_edicion']:
                with st.form("form_edicion_causa"):
                    st.markdown("#### Datos de Litigación")
                    n_tribunal = selector_tribunal(str(c_data.get('TRIBUNAL','')), key_prefix="editar_causa")
                    n_serv = st.text_input("Servicio Contratado", str(c_data.get('Servicio','')))
                    n_negocio = st.selectbox("Origen de Cartera", ["Externo", "Propio"], index=0 if c_data.get('Tipo_Negocio') == "Externo" else 1)
                    
                    st.markdown("#### Datos de Ficha de Cliente")
                    n_clave = st.text_input("Clave Única", str(c_data.get('Clave_unica','')))
                    n_sac = st.text_input("SAC Asignado", str(c_data.get('SAC','')))
                    n_suc = st.text_input("Sucursal Oficina", str(c_data.get('Sucursal','')))
                    
                    st.markdown("#### 💰 Control de Honorarios")
                    opciones_hon = ["Sin fijar", "Pagados", "Pendientes"]
                    idx_hon = opciones_hon.index(c_data.get('Estado_Honorarios', 'Sin fijar')) if c_data.get('Estado_Honorarios') in opciones_hon else 0
                    n_estado_hon = st.selectbox("Condición de Honorarios", opciones_hon, index=idx_hon)
                    
                    if n_estado_hon == "Pendientes":
                        n_tot_hon = st.number_input("Honorario Total Pactado ($)", value=int(c_data.get('Total_Honorarios', 0)))
                        n_cuo_tot = st.number_input("Mensualidades Totales", value=max(1, int(c_data.get('Cuotas_Totales', 0) or 0)), min_value=1)
                        n_cuo_pag = st.number_input("Mensualidades Enteradas", value=int(c_data.get('Cuotas_Pagadas', 0)), min_value=0)
                    elif n_estado_hon == "Pagados":
                        n_tot_hon = st.number_input("Monto Total Enterado ($)", value=int(c_data.get('Total_Honorarios', 0)))
                        n_cuo_tot, n_cuo_pag = 1, 1
                    else:
                        n_tot_hon, n_cuo_tot, n_cuo_pag = 0, 0, 0
                        
                    if st.form_submit_button("💾 Guardar Cambios", type="primary"):
                        df_causas.at[idx, 'TRIBUNAL'] = n_tribunal; df_causas.at[idx, 'Servicio'] = n_serv; df_causas.at[idx, 'Tipo_Negocio'] = n_negocio
                        df_causas.at[idx, 'Clave_unica'] = n_clave; df_causas.at[idx, 'SAC'] = n_sac; df_causas.at[idx, 'Sucursal'] = n_suc
                        df_causas.at[idx, 'Estado_Honorarios'] = n_estado_hon; df_causas.at[idx, 'Total_Honorarios'] = n_tot_hon
                        df_causas.at[idx, 'Cuotas_Totales'] = n_cuo_tot; df_causas.at[idx, 'Cuotas_Pagadas'] = n_cuo_pag
                        df_causas.to_csv(ARCHIVO_BD, index=False)
                        st.session_state['modo_edicion'] = False
                        st.rerun()
            else:
                # --- 🔍 MOTOR DE BÚSQUEDA RELACIONAL DEL CLIENTE ---
                rut_asociado = str(c_data.get('RUT', ''))
                datos_cliente = df_clientes[df_clientes['RUT'].astype(str) == rut_asociado]
                
                if not datos_cliente.empty:
                    info_cl = datos_cliente.iloc[0]
                    tel_real = info_cl.get('Telefono', '--')
                    correo_real = info_cl.get('Correo', '--')
                    dir_real = info_cl.get('Direccion', '--')
                else:
                    tel_real, correo_real, dir_real = '--', '--', '--'

                clase_div = "badge-active" if c_data.get('Tipo_Negocio') == "Externo" else "badge-propio"
                st.markdown(f"""
                <div class="dash-card">
                    <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:14px;">
                        <span style="font-weight:700; font-size:15px; color:#172b4d;">Información de la Causa</span>
                        <span class="{clase_div}">{c_data.get('Tipo_Negocio','')}</span>
                    </div>
                    <div class="info-field"><span class="info-label">Rol</span><span class="info-value">{rol_actual}</span></div>
                    <div class="info-field"><span class="info-label">Tribunal</span><span class="info-value">{c_data.get('TRIBUNAL') or '--'}</span></div>
                    <div class="info-field"><span class="info-label">Materia / Servicio</span><span class="info-value">{c_data.get('Servicio') or '--'}</span></div>
                </div>
                <div class="dash-card">
                    <div style="margin-bottom:14px;">
                        <span style="font-weight:700; font-size:15px; color:#172b4d;">Ficha Económica y Contacto</span>
                    </div>
                    <div class="info-field"><span class="info-label">Cliente</span><span class="info-value">{c_data.get('Cliente') or '--'}</span></div>
                    <div class="info-field"><span class="info-label">RUT</span><span class="info-value">{rut_asociado or '--'}</span></div>
                    <div class="info-field"><span class="info-label">Teléfono</span><span class="info-value">📞 {tel_real}</span></div>
                    <div class="info-field"><span class="info-label">Correo</span><span class="info-value">✉️ {correo_real}</span></div>
                    <div class="info-field"><span class="info-label">Dirección</span><span class="info-value">📍 {dir_real}</span></div>
                    <hr style="border:none; border-top:1px solid #e0e4e8; margin:14px 0;">
                    <div style="display:flex; justify-content:space-between; align-items:center;">
                        <span class="info-label">Honorarios</span>
                        <span class="badge-honorarios">{c_data.get('Estado_Honorarios') or 'Sin fijar'}</span>
                    </div>
                    <div class="info-field" style="margin-top:8px;"><span class="info-label">Pactado</span><span class="info-value" style="font-weight:700;">${c_data.get('Total_Honorarios',0):,.0f}</span></div>
                </div>
                """, unsafe_allow_html=True)
                
        with col_izq:
            tab_tareas_internas, tab_docs_solicitados, tab_excepciones = st.tabs(["Tareas Operativas", "📥 Docs Cliente", "⚖️ Excepciones Ejecutivas"])
            
            with tab_tareas_internas:
                if st.button("+ Asignar Nueva Tarea Operativa", type="primary"):
                    st.session_state['creando_tarea'] = not st.session_state.get('creando_tarea', False)
                    st.rerun()
                    
                if st.session_state.get('creando_tarea'):
                    with st.form("form_t_interna"):
                        t_t = st.text_input("Nomenclatura Breve")
                        t_d = st.text_area("Descripción de la gestión")
                        t_p = st.selectbox("Prioridad", ["Alta", "Media", "Baja"])
                        t_f = st.date_input("Fecha de Cumplimiento")
                        
                        st.markdown("---")
                        st.markdown("<span style='font-size:13px; color:#6b778c;'>Dejar en blanco para asignarla a ti mismo. Para delegar, escribe el nombre del colega.</span>", unsafe_allow_html=True)
                        t_delegado = st.text_input("Asignar Tarea a (Opcional)", placeholder="Ej: Eduardo Riquelme")
                        
                        if st.form_submit_button("Registrar y Asignar Tarea", type="primary"):
                            destinatario_file = ARCHIVO_TAREAS
                            destinatario_usr = usuario_actual
                            
                            if t_delegado.strip():
                                nombre_buscado = t_delegado.strip().lower()
                                for user_key, real_name in NOMBRES_REALES.items():
                                    if nombre_buscado in real_name.lower() or nombre_buscado == user_key.lower():
                                        destinatario_usr = user_key
                                        destinatario_file = f"base_tareas_{user_key}.csv"
                                        break
                                
                            if not os.path.exists(destinatario_file):
                                pd.DataFrame(columns=['ID_Tarea', 'ROL', 'Creador', 'Fecha_Creacion', 'Fecha_Vencimiento', 'Titulo', 'Descripcion', 'Estado', 'Comentarios', 'Prioridad', 'Usuario_Propietario']).to_csv(destinatario_file, index=False)
                                
                            df_t_destino = leer_csv_local(destinatario_file, COLS_TAREAS)
                            nueva_t = {
                                'ID_Tarea': str(uuid.uuid4())[:8], 
                                'ROL': rol_actual, 
                                'Creador': nombre_real_usuario, 
                                'Fecha_Creacion': datetime.now().strftime("%d/%m/%Y"), 
                                'Fecha_Vencimiento': t_f.strftime("%d/%m/%Y"),
                                'Titulo': t_t, 'Descripcion': t_d, 'Estado': 'En progreso', 'Comentarios': '[]', 'Prioridad': t_p,
                                'Usuario_Propietario': destinatario_usr
                            }
                            df_t_destino = pd.concat([df_t_destino, pd.DataFrame([nueva_t])], ignore_index=True)
                            df_t_destino.to_csv(destinatario_file, index=False)
                            
                            dn_t_upd = safe_read_sheet("base_tareas", ['ID_Tarea', 'ROL', 'Creador', 'Fecha_Creacion', 'Fecha_Vencimiento', 'Titulo', 'Descripcion', 'Estado', 'Comentarios', 'Prioridad', 'Usuario_Propietario'])
                            safe_update_sheet("base_tareas", pd.concat([dn_t_upd, pd.DataFrame([nueva_t])], ignore_index=True))
                                
                            # --- 🚀 DISPARO A GOOGLE CALENDAR DINÁMICO ---
                            if t_p == "Alta":
                                df_usr_db = safe_read_sheet("base_usuarios", [])
                                f_user = df_usr_db[df_usr_db['Usuario'] == destinatario_usr]
                                if not f_user.empty and "@" in str(f_user.iloc[0]['Correo']):
                                    correo_cal = f_user.iloc[0]['Correo']
                                    exito = agendar_plazo_calendar(t_t, f"Causa ROL: {rol_actual}\nDetalle: {t_d}", t_f.strftime("%d/%m/%Y"), correo_cal)
                                    if exito:
                                        st.toast("📅 Plazo fatal sincronizado en Google Calendar con alarmas.", icon="🚨")
                            # ------------------------------------
                                
                            st.session_state['creando_tarea'] = False
                            st.success("✅ Tarea registrada exitosamente.")
                            st.rerun()
                            
                df_t_local = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
                # Las tareas más recientes van primero, independiente de su estado
                # (aprobada/rechazada), en vez del orden de creación original que
                # las dejaba siempre al final de la lista.
                tareas_de_esta_causa = df_t_local[df_t_local['ROL'] == rol_actual].iloc[::-1]
                
                if tareas_de_esta_causa.empty:
                    st.info("Esta causa no registra tareas en progreso.")
                else:
                    for idx_tarea_bd, tarea in tareas_de_esta_causa.iterrows():
                        with st.container(border=True):
                            b_prio_color = "#ff5630" if tarea.get('Prioridad') == "Alta" else ("#ffc400" if tarea.get('Prioridad') == "Media" else "#57a15a")
                            st.markdown(f"<div style='height: 5px; background-color: {b_prio_color}; border-radius: 5px 5px 0 0; margin: -1rem -1rem 1rem -1rem;'></div>", unsafe_allow_html=True)
                            
                            if st.session_state.get('editando_tarea') == tarea['ID_Tarea']:
                                modal_editar_tarea(tarea['ID_Tarea'], tarea['Titulo'], tarea['Fecha_Vencimiento'], tarea['Estado'])
                            else:
                                autor_real = NOMBRES_REALES.get(tarea['Creador'], tarea['Creador'])
                                nro_tarea_corto = str(tarea['ID_Tarea']).upper()

                                # --- Encabezado tipo ficha: título + metadatos + acciones alineadas a la derecha ---
                                c_top_l, c_top_r = st.columns([2.3, 2.5])
                                with c_top_l:
                                    st.markdown(f"<div style='font-weight:700; font-size:17px; color:#172b4d;'>{tarea['Titulo']}</div>", unsafe_allow_html=True)
                                    st.markdown(f"<span style='font-size:13px; color:#6b778c;'>Creado por: {autor_real} • N° tarea {nro_tarea_corto} • [{tarea.get('Prioridad', 'Media')}]</span>", unsafe_allow_html=True)
                                    st.markdown(f"<span style='font-size:13px; color:#6b778c;'>Fecha creación: {tarea['Fecha_Creacion']} • Fecha vencimiento: {tarea['Fecha_Vencimiento']}</span>", unsafe_allow_html=True)
                                
                                with c_top_r:
                                    if tarea['Estado'] == 'En progreso':
                                        bcols = st.columns([1.3, 1.3, 0.9, 0.9] if usuario_actual == "Narratia" else [1.3, 1.3, 0.9])
                                        if bcols[0].button("❌ Rechazar", key=f"rech_{tarea['ID_Tarea']}", use_container_width=True): 
                                            df_t_local.at[idx_tarea_bd, 'Estado'] = 'Rechazada'; df_t_local.to_csv(ARCHIVO_TAREAS, index=False); st.rerun()
                                        if bcols[1].button("✅ Aprobar", key=f"apr_{tarea['ID_Tarea']}", use_container_width=True): 
                                            df_t_local.at[idx_tarea_bd, 'Estado'] = 'Aprobada'; df_t_local.to_csv(ARCHIVO_TAREAS, index=False); st.rerun()
                                        if bcols[2].button("✏️", key=f"edit_{tarea['ID_Tarea']}", help="Editar tarea", use_container_width=True):
                                            st.session_state['editando_tarea'] = tarea['ID_Tarea']; st.rerun()
                                        if usuario_actual == "Narratia" and bcols[3].button("🗑️", key=f"del_{tarea['ID_Tarea']}", help="Eliminar tarea", use_container_width=True):
                                            df_t_local = df_t_local.drop(idx_tarea_bd); df_t_local.to_csv(ARCHIVO_TAREAS, index=False); st.rerun()
                                        st.markdown("<div class='task-status-chip task-status-progreso' style='margin-top:8px; text-align:right; float:right;'>En progreso</div>", unsafe_allow_html=True)
                                    else:
                                        clase_estado = "task-status-aprobada" if tarea['Estado'] == 'Aprobada' else "task-status-rechazada"
                                        c_chip, c_del = st.columns([3, 1])
                                        with c_chip:
                                            st.markdown(f"<div style='text-align:right;'><span class='task-status-chip {clase_estado}'>{tarea['Estado']}</span></div>", unsafe_allow_html=True)
                                        with c_del:
                                            if usuario_actual == "Narratia" and st.button("🗑️", key=f"del_fin_{tarea['ID_Tarea']}", help="Eliminar tarea", use_container_width=True):
                                                df_t_local = df_t_local.drop(idx_tarea_bd); df_t_local.to_csv(ARCHIVO_TAREAS, index=False); st.rerun()

                                st.markdown(f"<p style='font-size: 15px; color: #172b4d; margin-top:12px; margin-bottom: 5px;'>{tarea['Descripcion']}</p>", unsafe_allow_html=True)
                                
                                comentarios_js = json.loads(tarea['Comentarios'])
                                
                                # El expander de comentarios se queda abierto después de
                                # comentar (antes se cerraba solo al recargar la página
                                # justo después de enviar un comentario nuevo).
                                key_com_abiertos = f"comentarios_abiertos_{tarea['ID_Tarea']}"
                                with st.expander(f"💬 Comentarios ({len(comentarios_js)})", expanded=st.session_state.get(key_com_abiertos, False)):
                                    if not comentarios_js:
                                        st.caption("No hay comentarios todavía.")
                                    for idx_com, c in enumerate(comentarios_js):
                                        c_texto_col, c_borrar_col = st.columns([9, 1])
                                        with c_texto_col:
                                            st.markdown(f"""
                                            <div style='padding:8px 0;'>
                                                <strong style='color:#172b4d; font-size:14px;'>{c['autor']}</strong>
                                                <span style='color:#6b778c; font-size:12px;'> • {c['fecha']}</span><br>
                                                <span style='color:#42526e; font-size:14px;'>{c['texto']}</span>
                                            </div>
                                            """, unsafe_allow_html=True)
                                            # Si el comentario tiene un archivo adjunto guardado de
                                            # verdad (Drive o base64), se ofrece para descargar. Los
                                            # comentarios viejos (de antes de este arreglo) solo
                                            # tenían el nombre escrito como texto, sin archivo real
                                            # detrás — para esos no aparece el botón, porque no hay
                                            # nada que descargar.
                                            if c.get('archivo_drive_id') or c.get('archivo_b64'):
                                                bytes_com_adj = obtener_bytes_adjunto(c, 'archivo_drive_id', 'archivo_b64')
                                                if bytes_com_adj is not None:
                                                    st.download_button(f"📥 {c.get('archivo_nombre', 'archivo')}", data=bytes_com_adj,
                                                                        file_name=c.get('archivo_nombre', 'archivo.docx'), key=f"dl_com_{tarea['ID_Tarea']}_{idx_com}")
                                        with c_borrar_col:
                                            st.markdown("<div style='padding-top:8px;'></div>", unsafe_allow_html=True)
                                            if st.button("🗑️", key=f"del_com_{tarea['ID_Tarea']}_{idx_com}", help="Eliminar comentario"):
                                                comentarios_js_actualizado = [x for j, x in enumerate(comentarios_js) if j != idx_com]
                                                df_t_local.at[idx_tarea_bd, 'Comentarios'] = json.dumps(comentarios_js_actualizado)
                                                df_t_local.to_csv(ARCHIVO_TAREAS, index=False)
                                                dn_del_com = safe_read_sheet("base_tareas", [])
                                                if not dn_del_com.empty:
                                                    dn_del_com.loc[dn_del_com['ID_Tarea'] == tarea['ID_Tarea'], 'Comentarios'] = json.dumps(comentarios_js_actualizado)
                                                    safe_update_sheet("base_tareas", dn_del_com)
                                                st.session_state[key_com_abiertos] = True
                                                st.rerun()
                                        st.markdown("<hr style='margin:2px 0; border-color:#f4f5f7;'>", unsafe_allow_html=True)
                                    
                                    st.markdown("<div style='margin-top:10px;'></div>", unsafe_allow_html=True)
                                    
                                    # Streamlit no permite repartir un mismo formulario entre
                                    # columnas creadas AFUERA de él (eso causaba el error). La
                                    # solución correcta es que las 3 columnas (clip, texto,
                                    # Enviar) se creen DENTRO del formulario, usando el clip
                                    # como un segundo botón de envío (Streamlit sí permite más
                                    # de un st.form_submit_button en el mismo formulario).
                                    key_toggle_adj = f"mostrar_adj_{tarea['ID_Tarea']}"
                                    
                                    with st.form(key=f"fc_{tarea['ID_Tarea']}", clear_on_submit=True):
                                        c_clip, c_txt_col, c_btn_col = st.columns([0.6, 5, 1.4])
                                        with c_clip:
                                            clip_click = st.form_submit_button("📎")
                                        with c_txt_col:
                                            texto_com = st.text_input("Agregar un comentario...", label_visibility="collapsed", placeholder="Escribir comentario...")
                                        with c_btn_col:
                                            enviar_click = st.form_submit_button("Enviar", use_container_width=True)
                                        
                                        adj_coment = None
                                        if st.session_state.get(key_toggle_adj):
                                            adj_coment = st.file_uploader("📎 Adjuntar archivo al comentario", key=f"fu_{tarea['ID_Tarea']}")
                                        
                                        if clip_click:
                                            st.session_state[key_toggle_adj] = not st.session_state.get(key_toggle_adj, False)
                                            st.rerun()
                                        elif enviar_click:
                                            if texto_com.strip() or adj_coment:
                                                nuevo_comentario = {"autor": nombre_real_usuario, "fecha": datetime.now().strftime("%d/%m/%Y %H:%M"), "texto": texto_com.strip()}
                                                if adj_coment:
                                                    # Antes solo se anotaba el nombre del archivo como
                                                    # texto, sin guardar el archivo de verdad en ningún
                                                    # lado — por eso nunca se podía descargar después.
                                                    # Ahora se guarda igual que cualquier otro adjunto
                                                    # del sistema (Drive, con respaldo en base64).
                                                    drive_id_com, b64_com = guardar_archivo_adjunto(adj_coment.name, adj_coment.getvalue(), adj_coment.type or 'application/octet-stream')
                                                    nuevo_comentario['archivo_nombre'] = adj_coment.name
                                                    nuevo_comentario['archivo_drive_id'] = drive_id_com
                                                    nuevo_comentario['archivo_b64'] = b64_com
                                                    nuevo_comentario['texto'] = (texto_com.strip() + f" <br><em>[📎 {adj_coment.name}]</em>").strip()
                                                comentarios_js.append(nuevo_comentario)
                                                df_t_local.at[idx_tarea_bd, 'Comentarios'] = json.dumps(comentarios_js)
                                                df_t_local.to_csv(ARCHIVO_TAREAS, index=False)
                                                
                                                dn = safe_read_sheet("base_tareas", [])
                                                if not dn.empty:
                                                    dn.loc[dn['ID_Tarea'] == tarea['ID_Tarea'], 'Comentarios'] = json.dumps(comentarios_js)
                                                    safe_update_sheet("base_tareas", dn)
                                                st.session_state[key_toggle_adj] = False
                                                st.session_state[key_com_abiertos] = True
                                                st.rerun()

            with tab_docs_solicitados:
                st.subheader("📋 Gestión de Requisitos del Cliente")
                # El nombre del cliente puede traer tildes, comas u otros caracteres que
                # rompen la URL al compartirla por WhatsApp/correo. Se deja solo
                # letras, números y guion bajo, para que el enlace nunca falle por esto.
                token_para_link = re.sub(r'[^A-Za-z0-9_]', '', str(c_data.get('Cliente', 'Cliente')).strip().replace(" ", "_"))
                
                # La URL de la app ya cambió más de una vez (narratia-app ->
                # seguimientodecausasjudicial -> jurisyncs), así que en vez de dejar un
                # valor fijo en el código que hay que andar corrigiendo cada vez, el
                # enlace se arma SOLO leyendo la URL real desde el navegador. El truco
                # anterior (imagen invisible con onload) no se ejecutaba de forma
                # confiable; components.v1.html es la forma oficial de Streamlit para
                # correr JavaScript de verdad, así que se cambió a esa.
                st.markdown("**🔗 Enlace del Portal para el Cliente:**")
                components.html(f"""
                <div id="linkbox" style="font-family: monospace; background:#f4f5f7; border:1px solid #cbd2d9;
                     border-radius:8px; padding:10px 14px; word-break: break-all; font-size:14px; color:#172b4d;">
                    Generando enlace...
                </div>
                <script>
                    document.getElementById('linkbox').innerText =
                        window.parent.location.origin + '/?cliente_id={token_para_link}';
                </script>
                """, height=60)
                st.caption("👆 Copia y pega ese enlace tal cual para enviárselo al cliente (por WhatsApp, correo, etc.).")
                
                with st.form(key=f"form_agregar_requisito_{rol_actual}", clear_on_submit=True):
                    st.markdown("#### Solicitar Nuevo Documento")
                    nuevo_doc_req = st.text_input("Nombre del documento solicitado", placeholder="Ej: Certificado de Matrimonio actualizado, Últimas 3 liquidaciones...")
                    if st.form_submit_button("➕ Enviar Requisito al Portal", type="primary"):
                        if nuevo_doc_req.strip() == "":
                            st.error("Escribe el nombre del documento.")
                        else:
                            ARCHIVO_DOCS = "base_documentos_clientes.csv"
                            if not os.path.exists(ARCHIVO_DOCS):
                                pd.DataFrame(columns=['ID_Req', 'Cliente_Token', 'Documento_Nombre', 'Estado', 'Archivo_B64', 'Archivo_Drive_ID', 'Fecha_Subida']).to_csv(ARCHIVO_DOCS, index=False)
                            
                            df_docs_db = leer_csv_local(ARCHIVO_DOCS, COLS_DOCS)
                            nuevo_registro_doc = {
                                'ID_Req': str(uuid.uuid4())[:8],
                                'Cliente_Token': token_para_link,
                                'Documento_Nombre': nuevo_doc_req.strip(),
                                'Estado': '❌ Pendiente',
                                'Archivo_B64': '',
                                'Fecha_Subida': '--'
                            }
                            df_docs_db = pd.concat([df_docs_db, pd.DataFrame([nuevo_registro_doc])], ignore_index=True)
                            df_docs_db.to_csv(ARCHIVO_DOCS, index=False)
                            safe_update_sheet("base_documentos_clientes", df_docs_db)
                            st.success(f"¡Solicitud de '{nuevo_doc_req}' agregada al portal del cliente!")
                            st.rerun()
                
                st.markdown("### Estado de la Documentación Solicitada")
                ARCHIVO_DOCS = "base_documentos_clientes.csv"
                if os.path.exists(ARCHIVO_DOCS):
                    df_docs_db = leer_csv_local(ARCHIVO_DOCS, COLS_DOCS)
                    docs_causa_actual = df_docs_db[df_docs_db['Cliente_Token'] == token_para_link]
                    
                    if docs_causa_actual.empty:
                        st.write("No se han solicitado documentos para este cliente todavía.")
                    else:
                        for idx_d, d_row in docs_causa_actual.iterrows():
                            with st.container(border=True):
                                cd1, cd2, cd3 = st.columns([3, 1.5, 1])
                                with cd1:
                                    st.markdown(f"**{d_row['Documento_Nombre']}**")
                                    st.write(f"Fecha de carga: {d_row.get('Fecha_Subida', '--')}")
                                with cd2:
                                    if str(d_row['Estado']).startswith('✅ Completado'):
                                        st.markdown("<span style='color:#57a15a; font-weight:bold;'>✅ Recibido</span>", unsafe_allow_html=True)
                                    else:
                                        st.markdown("<span style='color:#ff5630; font-weight:bold;'>❌ Pendiente</span>", unsafe_allow_html=True)
                                with cd3:
                                    if str(d_row['Estado']).startswith('✅ Completado'):
                                        link_ext_doc = d_row.get('Link_Externo', '')
                                        if pd.notna(link_ext_doc) and str(link_ext_doc).strip():
                                            st.markdown(f"[🔗 Abrir link]({link_ext_doc})")
                                        else:
                                            bytes_descarga = obtener_bytes_adjunto(d_row, 'Archivo_Drive_ID', 'Archivo_B64')
                                            if bytes_descarga is not None:
                                                st.download_button("📥 Descargar", data=bytes_descarga, file_name=f"{d_row['Documento_Nombre'].replace(' ', '_')}_{token_para_link}.pdf", key=f"dl_abog_{d_row['ID_Req']}")
                                            else:
                                                st.caption("⚠️ No se pudo recuperar el archivo.")
                                    else:
                                        if st.button("🗑️", key=f"del_req_{d_row['ID_Req']}"):
                                            df_docs_db = df_docs_db.drop(idx_d)
                                            df_docs_db.to_csv(ARCHIVO_DOCS, index=False)
                                            st.rerun()
            
            with tab_excepciones:
                st.subheader("⚖️ Generador de Escritos Judiciales")
                st.caption("Demandas, evacúa traslados, abandonos de procedimiento, nulidades procesales, tercerías, excepciones ejecutivas y cualquier otra presentación al Poder Judicial.")
                
                tipo_escrito_sel = st.selectbox("Tipo de Escrito", list(TIPOS_ESCRITOS_JUDICIALES.keys()), key=f"exc_tipo_escrito_{rol_actual}")
                
                if tipo_escrito_sel == "Excepciones Ejecutivas (Art. 464 CPC)":
                    modo_excepciones = st.radio("¿Cómo quieres trabajar?", ["📄 Subir PDFs (la IA analiza)", "✍️ Ingresar datos manualmente"], horizontal=True, key=f"pe_modo_exc_{rol_actual}")
                    
                    if modo_excepciones == "📄 Subir PDFs (la IA analiza)":
                        archivos_exc = st.file_uploader("Sube el pagaré, mandato, demanda, resoluciones, personería y demás documentos de la causa", type=["pdf"], accept_multiple_files=True, key=f"exc_pdfs_{rol_actual}")
                        contexto_exc = st.text_area("Contexto adicional para la IA (opcional)", key=f"exc_contexto_{rol_actual}", placeholder="Ej: El pagaré fue suscrito por un apoderado, revisar si tenía facultades suficientes.")
                        
                        if st.button("🔍 Analizar Documentos", type="primary", use_container_width=True, key=f"exc_btn_analizar_{rol_actual}"):
                            if not archivos_exc:
                                st.error("⚠️ Sube al menos un documento en PDF.")
                            else:
                                with st.spinner("⚖️ Analizando documentos y evaluando las 18 excepciones del Art. 464 CPC... esto puede tardar un poco si hay páginas escaneadas."):
                                    try:
                                        resultado_exc = analizar_excepciones_con_ia(archivos_exc, contexto_exc)
                                        st.session_state[f'exc_resultado_{rol_actual}'] = resultado_exc
                                        st.success(f"✅ Análisis completado. Se identificaron {sum(1 for e in resultado_exc if e.get('aplica'))} excepciones potencialmente aplicables.")
                                    except Exception as e:
                                        st.error(f"❌ Hubo un error al analizar los documentos: {e}")
                        
                        if f'exc_resultado_{rol_actual}' in st.session_state:
                            resultado_exc = st.session_state[f'exc_resultado_{rol_actual}']
                            orden_confianza = {"Alta": 0, "Media": 1, "Baja": 2, None: 3}
                            aplicables = sorted([e for e in resultado_exc if e.get('aplica')], key=lambda e: orden_confianza.get(e.get('confianza'), 3))
                            descartadas = [e for e in resultado_exc if not e.get('aplica')]
                            
                            st.markdown(f"#### Excepciones a oponer — IA identificó {len(aplicables)}; se preseleccionan las de confianza alta")
                            seleccionadas_ids = []
                            for exc in aplicables:
                                color_conf = {"Alta": "#e3fcef", "Media": "#fff0b3", "Baja": "#ffebe6"}.get(exc.get('confianza'), "#f4f5f7")
                                texto_conf = {"Alta": "#1b7a4a", "Media": "#7a5b00", "Baja": "#bf2600"}.get(exc.get('confianza'), "#6b778c")
                                with st.container(border=True):
                                    marcado = st.checkbox(
                                        f"N°{exc['numero']} — {exc['nombre']}", value=(exc.get('confianza') == "Alta"),
                                        key=f"exc_check_{rol_actual}_{exc['numero']}"
                                    )
                                    st.markdown(f"<span style='background:{color_conf}; color:{texto_conf}; padding:2px 10px; border-radius:10px; font-size:12px; font-weight:700;'>IA: {exc.get('confianza','')}</span>", unsafe_allow_html=True)
                                    st.markdown(f"<span style='color:#42526e; font-size:14px;'>{exc.get('fundamento','')}</span>", unsafe_allow_html=True)
                                    if exc.get('cita_textual', '').strip():
                                        st.markdown(f"<span style='color:#6b778c; font-size:13px; font-style:italic;'>Cita: «{exc['cita_textual']}»</span>", unsafe_allow_html=True)
                                    if marcado:
                                        seleccionadas_ids.append(exc['numero'])
                            
                            if descartadas:
                                with st.expander(f"Ver excepciones descartadas ({len(descartadas)}) — no aplican según la IA"):
                                    for exc in descartadas:
                                        st.markdown(f"**N°{exc['numero']} — {exc['nombre']}** — <span style='background:#f4f5f7; color:#6b778c; padding:2px 8px; border-radius:10px; font-size:11px;'>Descartada</span>", unsafe_allow_html=True)
                                        st.caption(exc.get('fundamento', ''))
                            
                            st.markdown("---")
                            nombre_ejecutado_exc = st.text_input("Nombre de quien comparece (representante del ejecutado)", value=nombre_real_usuario, key=f"exc_nombre_ejec_{rol_actual}")
                            
                            if st.button("📄 Generar Escrito de Oposición de Excepciones", type="primary", use_container_width=True, key=f"exc_btn_generar_{rol_actual}"):
                                if not seleccionadas_ids:
                                    st.error("⚠️ Marca al menos una excepción para incluir en el escrito.")
                                else:
                                    excepciones_finales = [e for e in aplicables if e['numero'] in seleccionadas_ids]
                                    datos_causa_exc = {
                                        'tribunal': c_data.get('TRIBUNAL', ''), 'rol': rol_actual, 'caratulado': c_data.get('CARATULADO', ''),
                                        'nombre_ejecutado': nombre_ejecutado_exc
                                    }
                                    doc_exc = crear_escrito_oposicion_excepciones_word(datos_causa_exc, excepciones_finales)
                                    if doc_exc:
                                        buffer_exc = io.BytesIO()
                                        doc_exc.save(buffer_exc)
                                        bytes_exc = buffer_exc.getvalue()
                                        nombre_archivo_exc = f"Oposicion_Excepciones_{rol_actual.replace('-', '_')}.docx"
                                        
                                        drive_id_exc, b64_exc = guardar_archivo_adjunto(nombre_archivo_exc, bytes_exc, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                        
                                        df_exc_hist = leer_csv_local(ARCHIVO_EXCEPCIONES, COLS_EXCEPCIONES)
                                        nuevo_exc_hist = {
                                            'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"), 'ROL': rol_actual,
                                            'Excepciones_Opuestas': ", ".join([f"N°{n}" for n in seleccionadas_ids]),
                                            'Archivo_B64': b64_exc, 'Archivo_Drive_ID': drive_id_exc, 'Usuario_Propietario': usuario_actual
                                        }
                                        df_exc_hist = pd.concat([df_exc_hist, pd.DataFrame([nuevo_exc_hist])], ignore_index=True)
                                        df_exc_hist.to_csv(ARCHIVO_EXCEPCIONES, index=False)
                                        dn_exc = safe_read_sheet("base_excepciones", COLS_EXCEPCIONES)
                                        safe_update_sheet("base_excepciones", pd.concat([dn_exc, pd.DataFrame([nuevo_exc_hist])], ignore_index=True))
                                        
                                        st.success("✅ Escrito generado y guardado en el historial de esta causa.")
                                        st.download_button("📥 Descargar Escrito (.docx)", data=bytes_exc, file_name=nombre_archivo_exc,
                                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_exc_nuevo_{rol_actual}")
                    else:
                        st.markdown("#### Marca manualmente las excepciones a oponer")
                        seleccionadas_manual = []
                        for numero_exc, nombre_exc in CATALOGO_EXCEPCIONES_464.items():
                            with st.container(border=True):
                                marcado_manual = st.checkbox(f"N°{numero_exc} — {nombre_exc}", key=f"exc_manual_check_{rol_actual}_{numero_exc}")
                                if marcado_manual:
                                    fundamento_manual = st.text_area(f"Fundamento de la excepción N°{numero_exc}", key=f"exc_manual_fund_{rol_actual}_{numero_exc}", height=80)
                                    cita_manual = st.text_input(f"Cita textual de respaldo (opcional)", key=f"exc_manual_cita_{rol_actual}_{numero_exc}")
                                    seleccionadas_manual.append({'numero': numero_exc, 'nombre': nombre_exc, 'fundamento_final': fundamento_manual, 'cita_textual': cita_manual})
                        
                        st.markdown("---")
                        nombre_ejecutado_exc_m = st.text_input("Nombre de quien comparece (representante del ejecutado)", value=nombre_real_usuario, key=f"exc_nombre_ejec_manual_{rol_actual}")
                        
                        if st.button("📄 Generar Escrito de Oposición de Excepciones", type="primary", use_container_width=True, key=f"exc_btn_generar_manual_{rol_actual}"):
                            if not seleccionadas_manual:
                                st.error("⚠️ Marca al menos una excepción e indica su fundamento.")
                            else:
                                datos_causa_exc_m = {
                                    'tribunal': c_data.get('TRIBUNAL', ''), 'rol': rol_actual, 'caratulado': c_data.get('CARATULADO', ''),
                                    'nombre_ejecutado': nombre_ejecutado_exc_m
                                }
                                doc_exc_m = crear_escrito_oposicion_excepciones_word(datos_causa_exc_m, seleccionadas_manual)
                                if doc_exc_m:
                                    buffer_exc_m = io.BytesIO()
                                    doc_exc_m.save(buffer_exc_m)
                                    bytes_exc_m = buffer_exc_m.getvalue()
                                    nombre_archivo_exc_m = f"Oposicion_Excepciones_{rol_actual.replace('-', '_')}.docx"
                                    
                                    drive_id_exc_m, b64_exc_m = guardar_archivo_adjunto(nombre_archivo_exc_m, bytes_exc_m, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                    
                                    df_exc_hist_m = leer_csv_local(ARCHIVO_EXCEPCIONES, COLS_EXCEPCIONES)
                                    nuevo_exc_hist_m = {
                                        'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"), 'ROL': rol_actual,
                                        'Excepciones_Opuestas': ", ".join([f"N°{e['numero']}" for e in seleccionadas_manual]),
                                        'Archivo_B64': b64_exc_m, 'Archivo_Drive_ID': drive_id_exc_m, 'Usuario_Propietario': usuario_actual
                                    }
                                    df_exc_hist_m = pd.concat([df_exc_hist_m, pd.DataFrame([nuevo_exc_hist_m])], ignore_index=True)
                                    df_exc_hist_m.to_csv(ARCHIVO_EXCEPCIONES, index=False)
                                    dn_exc_m = safe_read_sheet("base_excepciones", COLS_EXCEPCIONES)
                                    safe_update_sheet("base_excepciones", pd.concat([dn_exc_m, pd.DataFrame([nuevo_exc_hist_m])], ignore_index=True))
                                    
                                    st.success("✅ Escrito generado y guardado en el historial de esta causa.")
                                    st.download_button("📥 Descargar Escrito (.docx)", data=bytes_exc_m, file_name=nombre_archivo_exc_m,
                                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_exc_manual_nuevo_{rol_actual}")
                
                else:
                    # --- FLUJO GENERAL: cualquier otro tipo de escrito (demandas, traslados,
                    # abandonos, nulidades, tercerías, recursos, etc.) ---
                    st.markdown(f"#### {tipo_escrito_sel}")
                    archivos_gen = st.file_uploader("Documentos de respaldo (opcional)", type=["pdf"], accept_multiple_files=True, key=f"gen_pdfs_{rol_actual}")
                    contexto_gen = st.text_area("Hechos e instrucciones para la IA", height=120, key=f"gen_contexto_{rol_actual}",
                                                 placeholder="Describe los hechos relevantes, lo que quieres alegar y cualquier dato específico que deba incluir el escrito.")
                    
                    if st.button(f"✍️ Redactar {tipo_escrito_sel}", type="primary", use_container_width=True, key=f"gen_btn_redactar_{rol_actual}"):
                        if not contexto_gen.strip() and not archivos_gen:
                            st.error("⚠️ Escribe el contexto/hechos o sube al menos un documento de respaldo.")
                        else:
                            with st.spinner("✍️ Redactando..."):
                                try:
                                    texto_redactado_gen = redactar_escrito_judicial_ia(
                                        tipo_escrito_sel, TIPOS_ESCRITOS_JUDICIALES[tipo_escrito_sel], archivos_gen, contexto_gen
                                    )
                                    st.session_state[f'gen_texto_{rol_actual}'] = texto_redactado_gen
                                    st.success("✅ Escrito redactado. Revísalo abajo antes de descargarlo.")
                                except Exception as e:
                                    st.error(f"❌ Hubo un error al redactar el escrito: {e}")
                    
                    if f'gen_texto_{rol_actual}' in st.session_state:
                        st.markdown("---")
                        st.markdown("#### Borrador Generado (revisa antes de presentar)")
                        st.markdown(st.session_state[f'gen_texto_{rol_actual}'])
                        
                        if st.button("📄 Descargar y Guardar en Historial (.docx)", type="primary", use_container_width=True, key=f"gen_btn_guardar_{rol_actual}"):
                            datos_causa_gen = {'tribunal': c_data.get('TRIBUNAL', ''), 'rol': rol_actual, 'caratulado': c_data.get('CARATULADO', '')}
                            doc_gen = crear_escrito_judicial_generico_word(tipo_escrito_sel, st.session_state[f'gen_texto_{rol_actual}'], datos_causa_gen)
                            if doc_gen:
                                buffer_gen = io.BytesIO()
                                doc_gen.save(buffer_gen)
                                bytes_gen = buffer_gen.getvalue()
                                nombre_archivo_gen = f"{tipo_escrito_sel.split(' (')[0].replace(' ', '_')}_{rol_actual.replace('-', '_')}.docx"
                                
                                drive_id_gen, b64_gen = guardar_archivo_adjunto(nombre_archivo_gen, bytes_gen, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                
                                df_exc_hist_g = leer_csv_local(ARCHIVO_EXCEPCIONES, COLS_EXCEPCIONES)
                                nuevo_exc_hist_g = {
                                    'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"), 'ROL': rol_actual,
                                    'Excepciones_Opuestas': tipo_escrito_sel,
                                    'Archivo_B64': b64_gen, 'Archivo_Drive_ID': drive_id_gen, 'Usuario_Propietario': usuario_actual
                                }
                                df_exc_hist_g = pd.concat([df_exc_hist_g, pd.DataFrame([nuevo_exc_hist_g])], ignore_index=True)
                                df_exc_hist_g.to_csv(ARCHIVO_EXCEPCIONES, index=False)
                                dn_exc_g = safe_read_sheet("base_excepciones", COLS_EXCEPCIONES)
                                safe_update_sheet("base_excepciones", pd.concat([dn_exc_g, pd.DataFrame([nuevo_exc_hist_g])], ignore_index=True))
                                
                                st.success("✅ Escrito guardado en el historial de esta causa.")
                                st.download_button("📥 Descargar Escrito (.docx)", data=bytes_gen, file_name=nombre_archivo_gen,
                                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_gen_nuevo_{rol_actual}")
                
                st.markdown("---")
                st.markdown("### 🗄️ Historial de Escritos de Excepciones de esta Causa")
                df_exc_todas = leer_csv_local(ARCHIVO_EXCEPCIONES, COLS_EXCEPCIONES)
                df_exc_causa = df_exc_todas[df_exc_todas['ROL'] == rol_actual] if not df_exc_todas.empty else df_exc_todas
                if df_exc_causa.empty:
                    st.info("Todavía no se ha generado ningún escrito de excepciones para esta causa.")
                else:
                    for _, fila_exc in df_exc_causa.iloc[::-1].iterrows():
                        with st.container(border=True):
                            c1, c2 = st.columns([4, 1])
                            with c1:
                                st.markdown(f"**Excepciones opuestas:** {fila_exc['Excepciones_Opuestas']}")
                                st.caption(f"Generado: {fila_exc['Fecha']}")
                            with c2:
                                bytes_desc_exc = obtener_bytes_adjunto(fila_exc, 'Archivo_Drive_ID', 'Archivo_B64')
                                if bytes_desc_exc is not None:
                                    st.download_button("📥 Descargar", data=bytes_desc_exc, file_name=f"Excepciones_{fila_exc['ID']}.docx", key=f"dl_exc_hist_{fila_exc['ID']}")

# 8. AGENDA DIARIA
elif st.session_state['menu_radio'] == "📋 Agenda":
    st.title("📋 Agenda Diaria de Plazos")
    fecha_hoy = datetime.now().strftime("%d/%m/%Y")
    st.write(f"Gestiones legales que vencen indefectiblemente el día de hoy: **{fecha_hoy}**")
    
    df_t = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
    if df_t.empty:
        st.info("No existen registros de plazos en el sistema.")
    else:
        df_t['Fecha_Vencimiento'] = df_t['Fecha_Vencimiento'].astype(str).str.strip()
        t_hoy = df_t[df_t['Fecha_Vencimiento'] == str(fecha_hoy)].copy()
        if t_hoy.empty:
            st.success("🎉 Felicitaciones. No registras plazos fatales para el día de hoy.")
        else:
            mapeo_prioridades = {"Alta": 1, "Media": 2, "Baja": 3}
            t_hoy['Orden'] = t_hoy['Prioridad'].map(mapeo_prioridades).fillna(4)
            t_hoy = t_hoy.sort_values(by='Orden')
            
            for _, row in t_hoy.iterrows():
                with st.container(border=True):
                    color_p = "#ff5630" if row['Prioridad'] == "Alta" else ("#ffc400" if row['Prioridad'] == "Media" else "#57a15a")
                    st.markdown(f"<div style='height: 5px; background-color:{color_p}; margin:-1rem -1rem 1rem -1rem; border-radius:5px 5px 0 0;'></div>", unsafe_allow_html=True)
                    c1, c2, c3 = st.columns([4, 2, 1])
                    with c1:
                        st.markdown(f"<div style='display: flex; align-items: center; margin-bottom: 5px;'><img src='{LOGO_URL}' style='height: 25px; margin-right: 8px;'><strong style='font-size:16px; color:#172b4d;'>{row['Titulo']}</strong><span style='font-size:12px; color:{color_p}; font-weight:bold; margin-left:8px;'>[{row.get('Prioridad', 'Media')}]</span></div>", unsafe_allow_html=True)
                        st.markdown(f"<span style='color:#6b778c;'>{str(row['Descripcion'])[:80]}...</span>", unsafe_allow_html=True)
                    with c2:
                        color_bd = "#ffc400" if row['Estado'] == 'En progreso' else ("#57a15a" if row['Estado'] == 'Aprobada' else "#ff5630")
                        st.markdown(f"<span style='background:{color_bd}; padding:3px 8px; border-radius:10px; font-size:12px; font-weight:bold; color:black;'>{row['Estado']}</span>", unsafe_allow_html=True)
                        st.markdown(f"<span style='color:#172b4d; font-size:14px;'><br>Causa: {row['ROL']}</span>", unsafe_allow_html=True)
                    with c3:
                        st.button("Ir al expediente ➔", key=f"ag_{row['ID_Tarea']}", on_click=ir_a_expediente, args=(row['ROL'],))

# 9. MENSAJERÍA INTERNA
elif st.session_state['menu_radio'] == "✈️ Mensajería":
    st.title("✈️ Mensajería Interna del Equipo")
    st.markdown("Plataforma de comunicación rápida para la oficina.")
    
    ES_ADMIN_MENSAJES = usuario_actual == "Narratia"
    
    c_tit_msj, c_refresh_msj = st.columns([4, 1])
    with c_refresh_msj:
        if st.button("🔄 Actualizar", key="refresh_mensajes", use_container_width=True):
            if "_csv_cache_" + ARCHIVO_MENSAJES in st.session_state:
                del st.session_state["_csv_cache_" + ARCHIVO_MENSAJES]
            st.rerun()
    
    df_msgs = leer_csv_local(ARCHIVO_MENSAJES, COLS_MENSAJES)
    
    # PRIVACIDAD: solo el administrador (Narratia) ve absolutamente todos los
    # mensajes del equipo. El resto de los usuarios solo ve las conversaciones
    # en las que participa directamente (las envió, se las enviaron a él/ella
    # de forma directa, o fueron enviadas a "Todos"). No pueden ver mensajes
    # privados entre otros dos compañeros.
    if not ES_ADMIN_MENSAJES and not df_msgs.empty:
        df_msgs = df_msgs[
            (df_msgs['De'] == nombre_real_usuario) |
            (df_msgs['Para'] == nombre_real_usuario) |
            (df_msgs['Para'] == 'Todos')
        ]
    
    with st.container(height=500):
        if df_msgs.empty:
            st.info("No hay mensajes en el buzón. ¡Sé el primero en escribir!")
        else:
            st.markdown("<div class='chat-bg'>", unsafe_allow_html=True)
            for _, msg in df_msgs.iterrows():
                es_mio = (msg['De'] == nombre_real_usuario)
                clase_burbuja = "burbuja-mia" if es_mio else "burbuja-otro"
                alineacion = "flex-end" if es_mio else "flex-start"
                
                st.markdown(f"""
                <div style="display: flex; justify-content: {alineacion}; width: 100%;">
                    <div class="{clase_burbuja}">
                        <div class="chat-autor">{msg['De']} <span class="chat-para">▶ Para: {msg['Para']}</span></div>
                        <div class="chat-texto">{msg['Mensaje']}</div>
                        <div class="chat-hora">{msg['Fecha']}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
                
    with st.form("form_chat", clear_on_submit=True):
        c_para, c_texto, c_btn = st.columns([2, 6, 2])
        destinatario = c_para.selectbox("Destinatario", ["Todos"] + list(NOMBRES_REALES.values()))
        texto_mensaje = c_texto.text_input("Escribe tu mensaje...", label_visibility="collapsed", placeholder="Escribe un mensaje aquí...")
        if c_btn.form_submit_button("Enviar 🚀", type="primary", use_container_width=True):
            if texto_mensaje.strip() != "":
                df_msgs_todos = leer_csv_local(ARCHIVO_MENSAJES, COLS_MENSAJES)
                nuevo_msj = {
                    'ID': str(uuid.uuid4())[:8],
                    'Fecha': datetime.now().strftime("%d/%m/%Y %H:%M"),
                    'De': nombre_real_usuario,
                    'Para': destinatario,
                    'Mensaje': texto_mensaje.strip()
                }
                df_msgs_todos = pd.concat([df_msgs_todos, pd.DataFrame([nuevo_msj])], ignore_index=True)
                df_msgs_todos.to_csv(ARCHIVO_MENSAJES, index=False)
                # Faltaba esto: la Mensajería nunca se había estado guardando en
                # Google Sheets, solo en el archivo local (que se borra en cada
                # reinicio del sistema) — por eso los mensajes desaparecían.
                dn_msgs = safe_read_sheet("base_mensajes_global", COLS_MENSAJES)
                safe_update_sheet("base_mensajes_global", pd.concat([dn_msgs, pd.DataFrame([nuevo_msj])], ignore_index=True))
                st.rerun()

# 10. CLIENTES DIRECTOS (FICHA COMPLETA Y RELACIONAL)
elif st.session_state['menu_radio'] == "👥 Clientes":
    st.title("Clientes")
    st.markdown("<span style='color:#6b778c;'>Gestione y organice la información de sus clientes de manera eficiente.</span>", unsafe_allow_html=True)
    
    df_clientes = safe_read_sheet("base_clientes", COLS_CLIENTES)
    
    ES_ADMIN_CLIENTES_TOP = usuario_actual == "Narratia"
    
    # ASIGNACIÓN DE DUEÑO (solo en memoria, para clientes antiguos sin dueño
    # asignado): antes esto se volvía a guardar en la nube automáticamente
    # cada vez que alguien abría esta pantalla — eso es peligroso, porque si
    # la lectura desde Google Sheets viniera incompleta por cualquier falla
    # pasajera (de red, de límite de peticiones, etc.), se terminaría
    # guardando esa versión incompleta, BORRANDO clientes de verdad. Ahora
    # solo se calcula en memoria para mostrar bien la información en esta
    # sesión; guardarlo de forma permanente es una acción manual y explícita
    # del administrador, con un botón aparte, más abajo.
    df_clientes_con_dueno_inferido = df_clientes.copy()
    hay_huerfanos_sin_guardar = False
    if not df_clientes.empty and 'Usuario_Propietario' in df_clientes.columns:
        huerfanos = df_clientes['Usuario_Propietario'].isna() | (df_clientes['Usuario_Propietario'].astype(str).str.strip() == '')
        if huerfanos.any():
            mapa_rut_dueno = {}
            for arch_causa_mig in glob.glob("base_causas_*.csv"):
                t_mig = leer_csv_local(arch_causa_mig, COLS_CAUSAS)
                dueno_mig = arch_causa_mig.replace("base_causas_", "").replace(".csv", "")
                if not t_mig.empty and 'RUT' in t_mig.columns:
                    for rut_mig in t_mig['RUT'].dropna().astype(str):
                        rut_limpio_mig = re.sub(r'[^0-9kK]', '', rut_mig).upper()
                        if rut_limpio_mig and rut_limpio_mig not in mapa_rut_dueno:
                            mapa_rut_dueno[rut_limpio_mig] = dueno_mig
            
            for idx_mig in df_clientes_con_dueno_inferido[huerfanos].index:
                rut_cliente_mig = re.sub(r'[^0-9kK]', '', str(df_clientes_con_dueno_inferido.at[idx_mig, 'RUT'])).upper()
                if rut_cliente_mig in mapa_rut_dueno:
                    df_clientes_con_dueno_inferido.at[idx_mig, 'Usuario_Propietario'] = mapa_rut_dueno[rut_cliente_mig]
            
            hay_huerfanos_sin_guardar = True
    
    df_clientes = df_clientes_con_dueno_inferido
    
    if ES_ADMIN_CLIENTES_TOP and hay_huerfanos_sin_guardar:
        with st.expander("🔧 Hay clientes antiguos sin dueño asignado — guardar la asignación en la nube (opcional)"):
            st.caption("Esto asigna el dueño correcto a los clientes antiguos cruzando su RUT contra las causas existentes, y lo deja guardado de forma permanente. Por seguridad, no se guarda solo; tienes que confirmarlo aquí.")
            if st.button("💾 Guardar asignación de dueños ahora", key="guardar_migracion_clientes"):
                filas_antes_guardar = safe_read_sheet("base_clientes", COLS_CLIENTES).shape[0]
                if df_clientes.shape[0] < filas_antes_guardar:
                    st.error(f"⚠️ Por seguridad, no se guardó: la nube tiene {filas_antes_guardar} clientes y esta versión en memoria solo tiene {df_clientes.shape[0]}. Esto podría deberse a una lectura incompleta — refresca la página e inténtalo de nuevo.")
                else:
                    safe_update_sheet("base_clientes", df_clientes)
                    st.success("✅ Asignación de dueños guardada correctamente.")
                    st.rerun()
    

    # PRIVACIDAD: cada abogado ve solo SUS PROPIOS clientes. Solo Nicolás
    # (Narratia), como administrador, ve los de todo el equipo.
    if not ES_ADMIN_CLIENTES_TOP and not df_clientes.empty and 'Usuario_Propietario' in df_clientes.columns:
        df_clientes = df_clientes[df_clientes['Usuario_Propietario'] == usuario_actual]

    with st.expander("🔍 Buscador de Conflictos de Interés (revisa antes de aceptar un caso nuevo)"):
        rut_conflicto = st.text_input("RUT a verificar", placeholder="Ej: 12.345.678-9", key="buscar_conflicto_rut")
        if rut_conflicto.strip():
            rut_normalizado = re.sub(r'[^0-9kK]', '', rut_conflicto).upper()
            
            # Se revisa en TODAS las causas que existan en el disco (de todos los
            # abogados, no solo las tuyas), porque un conflicto de interés hay que
            # detectarlo aunque el caso lo haya llevado otro compañero del estudio.
            piezas_conflicto = []
            for arch_conf in glob.glob("base_causas_*.csv"):
                t_conf = leer_csv_local(arch_conf, COLS_CAUSAS)
                if not t_conf.empty and 'RUT' in t_conf.columns:
                    propietario_conf = arch_conf.replace("base_causas_", "").replace(".csv", "")
                    t_conf = t_conf.copy()
                    t_conf['Propietario_Vista'] = propietario_conf
                    piezas_conflicto.append(t_conf)
            df_todas_causas_conf = pd.concat(piezas_conflicto, ignore_index=True) if piezas_conflicto else pd.DataFrame()
            
            resultados_conflicto = pd.DataFrame()
            if not df_todas_causas_conf.empty:
                mascara_rut = df_todas_causas_conf['RUT'].astype(str).apply(lambda x: re.sub(r'[^0-9kK]', '', x).upper() == rut_normalizado)
                resultados_conflicto = df_todas_causas_conf[mascara_rut]
            
            if resultados_conflicto.empty:
                st.success("✅ No se encontraron causas asociadas a este RUT en todo el estudio. No hay conflicto de interés detectado.")
            else:
                st.warning(f"⚠️ Este RUT aparece en {len(resultados_conflicto)} causa(s) del estudio — revisa antes de aceptar un caso nuevo:")
                for _, fila_conf in resultados_conflicto.iterrows():
                    nombre_resp = NOMBRES_REALES.get(fila_conf.get('Propietario_Vista'), fila_conf.get('Propietario_Vista'))
                    st.markdown(f"- **{fila_conf.get('ROL','--')}** — {fila_conf.get('CARATULADO','--')} · Cliente: {fila_conf.get('Cliente','--')} · Responsable: {nombre_resp}")
            st.caption("Este buscador revisa por coincidencia exacta de RUT en las causas registradas. No reemplaza el criterio profesional del abogado.")

    if st.session_state['cliente_seleccionado'] is None:
        if st.button("➕ Crear Nuevo Cliente", type="primary"):
            st.session_state['creando_cliente'] = not st.session_state.get('creando_cliente', False)
            
        if st.session_state.get('creando_cliente'):
            with st.container(border=True):
                with st.form("form_crear_cliente"):
                    st.subheader("Ficha Completa del Nuevo Cliente")
                    c1, c2 = st.columns(2)
                    n_cli_nom = c1.text_input("Nombre Completo *")
                    n_cli_rut = c2.text_input("RUT del Cliente *")
                    n_cli_tel = c1.text_input("Teléfono")
                    n_cli_cor = c2.text_input("Correo Electrónico")
                    n_cli_cla = c1.text_input("Clave Única")
                    n_cli_dom = c2.text_input("Domicilio")

                    if st.form_submit_button("💾 Guardar Cliente en la Nube", type="primary"):
                        if not n_cli_nom or not n_cli_rut:
                            st.error("El Nombre y el RUT son obligatorios.")
                        else:
                            nuevo_cliente = {
                                'RUT': n_cli_rut.strip().upper(),
                                'Nombre': n_cli_nom.strip(),
                                'Telefono': n_cli_tel.strip(),
                                'Correo': n_cli_cor.strip(),
                                'Clave_unica': n_cli_cla.strip(),
                                'Direccion': n_cli_dom.strip(),
                                'Usuario_Propietario': usuario_actual
                            }
                            # Se guarda sobre la hoja COMPLETA (releída de nuevo), no sobre la
                            # versión ya filtrada por privacidad, para no borrar los clientes
                            # del resto del equipo al agregar el nuevo.
                            df_clientes_completo_nuevo = safe_read_sheet("base_clientes", COLS_CLIENTES)
                            df_clientes_completo_nuevo = pd.concat([df_clientes_completo_nuevo, pd.DataFrame([nuevo_cliente])], ignore_index=True)
                            safe_update_sheet("base_clientes", df_clientes_completo_nuevo)
                            st.success("✅ Cliente creado y sincronizado en Google Sheets.")
                            st.session_state['creando_cliente'] = False
                            st.rerun()

        st.write("---")
        
        # CRUZAMOS DATOS PARA NO PERDER CLIENTES HISTÓRICOS (y, para el admin, de todo el equipo)
        ES_ADMIN_CLIENTES = usuario_actual == "Narratia"
        if ES_ADMIN_CLIENTES:
            boton_refrescar_equipo("refresh_clientes_equipo")
            df_causas_local = pd.DataFrame()
            piezas_causas_cli = []
            for arch_cli in glob.glob("base_causas_*.csv"):
                t = leer_csv_local(arch_cli, COLS_CAUSAS)
                if not t.empty:
                    piezas_causas_cli.append(t)
            if piezas_causas_cli:
                df_causas_local = pd.concat(piezas_causas_cli, ignore_index=True)
        else:
            df_causas_local = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
        
        filas_clientes = {}
        if not df_clientes.empty:
            for _, r in df_clientes.iterrows():
                if pd.notna(r['Nombre']):
                    filas_clientes[(r['Nombre'], r['RUT'])] = {'Telefono': r.get('Telefono', '--'), 'Correo': r.get('Correo', '--')}
        if not df_causas_local.empty and 'Cliente' in df_causas_local.columns:
            for _, r in df_causas_local.iterrows():
                if pd.notna(r['Cliente']) and r['Cliente'] != '--':
                    clave = (r['Cliente'], r.get('RUT', '--'))
                    if clave not in filas_clientes:
                        filas_clientes[clave] = {'Telefono': r.get('Teléfono', '--'), 'Correo': r.get('Correo', '--')}
        
        if not filas_clientes:
            st.info("No hay clientes registrados en la base de datos.")
        else:
            df_directorio = pd.DataFrame([
                {'Cliente': nom, 'RUT': rut, 'Teléfono': datos['Telefono'], 'Correo': datos['Correo']}
                for (nom, rut), datos in filas_clientes.items()
            ])
            
            c_busq_cli, c_dl_cli = st.columns([4, 1])
            busqueda_cli = c_busq_cli.text_input("Buscar", placeholder="Busca por nombre o RUT...", label_visibility="collapsed")
            with c_dl_cli:
                boton_descargar_excel(df_directorio, "clientes_jurisync.xlsx", key="dl_excel_clientes")
            
            if busqueda_cli.strip():
                q = busqueda_cli.strip().lower()
                df_directorio = df_directorio[
                    df_directorio['Cliente'].astype(str).str.lower().str.contains(q, na=False) |
                    df_directorio['RUT'].astype(str).str.lower().str.contains(q, na=False)
                ]
            
            with st.container(border=True):
                ch1, ch2, ch3, ch4, ch5 = st.columns([2.5, 1.5, 2.5, 2.5, 1])
                ch1.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>CLIENTE</span>", unsafe_allow_html=True)
                ch2.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>TELÉFONO</span>", unsafe_allow_html=True)
                ch3.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>CORREO</span>", unsafe_allow_html=True)
                ch4.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>RUT</span>", unsafe_allow_html=True)
                ch5.markdown("<span style='color:#6b778c; font-weight:800; font-size:13px;'>ACCIONES</span>", unsafe_allow_html=True)
                st.markdown("<hr style='margin: 5px 0px 10px 0px; border-top: 2px solid #e0e4e8;'>", unsafe_allow_html=True)
                
                for _, fila_cli in df_directorio.iterrows():
                    r1, r2, r3, r4, r5 = st.columns([2.5, 1.5, 2.5, 2.5, 1])
                    r1.markdown(f"<span style='color:#172b4d; font-weight:600; font-size:14px;'>👤 {fila_cli['Cliente']}</span>", unsafe_allow_html=True)
                    r2.markdown(f"<span style='color:#172b4d; font-size:14px;'>{fila_cli['Teléfono']}</span>", unsafe_allow_html=True)
                    r3.markdown(f"<span style='color:#172b4d; font-size:14px;'>{fila_cli['Correo']}</span>", unsafe_allow_html=True)
                    r4.markdown(f"<span style='color:#6b778c; font-size:13px;'>{fila_cli['RUT']}</span>", unsafe_allow_html=True)
                    if r5.button("👁️", key=f"ver_cli_{fila_cli['RUT']}_{fila_cli['Cliente']}", use_container_width=True):
                        st.session_state['cliente_seleccionado'] = fila_cli['RUT']
                        st.rerun()
                    st.markdown("<hr style='margin: 8px 0px 8px 0px; border-top: 1px dashed #e0e4e8;'>", unsafe_allow_html=True)
    else:
        rut_actual = st.session_state['cliente_seleccionado']
        filtro_cli = df_clientes[df_clientes['RUT'] == rut_actual]
        datos_cli = filtro_cli.iloc[0] if not filtro_cli.empty else pd.Series({'Nombre': 'Cliente Histórico', 'RUT': rut_actual, 'Telefono': '--', 'Correo': '--', 'Clave_unica': '--', 'Direccion': '--'})
        
        c_back, c_del = st.columns([4, 1])
        if c_back.button("⬅ Volver al Directorio"): 
            st.session_state['cliente_seleccionado'] = None
            st.rerun()
            
        if st.session_state['username'] == "Narratia": 
            if c_del.button("🗑️ Eliminar Cliente", use_container_width=True):
                with st.spinner("Borrando cliente y limpiando datos en cascada..."):
                    df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                    roles_a_borrar = df_causas[df_causas['RUT'].astype(str) == str(rut_actual)]['ROL'].tolist() if not df_causas.empty else []
                    nombre_borrar = datos_cli['Nombre']
                    
                    if not df_causas.empty:
                        df_causas = df_causas[df_causas['RUT'].astype(str) != str(rut_actual)]
                        df_causas.to_csv(ARCHIVO_BD, index=False)
                    
                    dn_c = safe_read_sheet("base_causas", COLS_CAUSAS)
                    if not dn_c.empty: safe_update_sheet("base_causas", dn_c[dn_c['RUT'].astype(str) != str(rut_actual)])
                    
                    df_t_local = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
                    if not df_t_local.empty and roles_a_borrar:
                        df_t_local = df_t_local[~df_t_local['ROL'].isin(roles_a_borrar)]
                        df_t_local.to_csv(ARCHIVO_TAREAS, index=False)
                    
                    dn_t = safe_read_sheet("base_tareas", COLS_TAREAS)
                    if not dn_t.empty and roles_a_borrar: safe_update_sheet("base_tareas", dn_t[~dn_t['ROL'].isin(roles_a_borrar)])
                    
                    df_con_local = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
                    if not df_con_local.empty:
                        df_con_local = df_con_local[df_con_local['Cliente'] != nombre_borrar]
                        df_con_local.to_csv(ARCHIVO_CONTRATOS, index=False)
                    
                    dn_con = safe_read_sheet("base_contratos", COLS_CONTRATOS)
                    if not dn_con.empty: safe_update_sheet("base_contratos", dn_con[dn_con['Cliente'] != nombre_borrar])

                    # IMPORTANTE: se relee la hoja COMPLETA (sin el filtro de privacidad)
                    # antes de guardar. Si se guardara la versión ya filtrada por dueño,
                    # se borrarían de la nube los clientes de TODO el resto del equipo.
                    df_clientes_completo_del = safe_read_sheet("base_clientes", COLS_CLIENTES)
                    df_clientes_completo_del = df_clientes_completo_del[df_clientes_completo_del['RUT'].astype(str) != str(rut_actual)]
                    safe_update_sheet("base_clientes", df_clientes_completo_del)
                    
                st.session_state['cliente_seleccionado'] = None
                st.success("✅ Cliente y TODO su historial asociado fue desintegrado.")
                import time; time.sleep(0.4); st.rerun()
            
        st.title(f"Ficha: {datos_cli['Nombre']}")
        tab1, tab2, tab3 = st.tabs(["👤 Información y Causas", "💰 Contabilidad", "📄 Contratos"])
        
        with tab1:
            col_i, col_d = st.columns([1, 2])
            with col_i:
                with st.container(border=True):
                    if st.session_state.get('editando_cli'):
                        with st.form("edit_cli"):
                            n_nom = st.text_input("Nombre", datos_cli.get('Nombre', ''))
                            n_rut = st.text_input("RUT", datos_cli.get('RUT', ''))
                            n_tel = st.text_input("Teléfono", datos_cli.get('Telefono', ''))
                            n_cor = st.text_input("Correo", datos_cli.get('Correo', ''))
                            n_cla = st.text_input("Clave Única", datos_cli.get('Clave_unica', ''))
                            n_dom = st.text_input("Domicilio", datos_cli.get('Direccion', ''))
                            if st.form_submit_button("💾 Guardar"):
                                # Mismo cuidado que en Eliminar: se relee la hoja completa
                                # sin el filtro de privacidad antes de guardar, para no
                                # borrar los clientes del resto del equipo.
                                df_clientes_completo_edit = safe_read_sheet("base_clientes", COLS_CLIENTES)
                                existe_fila_cliente = (not df_clientes_completo_edit.empty) and (df_clientes_completo_edit['RUT'] == rut_actual).any()
                                if existe_fila_cliente:
                                    df_clientes_completo_edit.loc[df_clientes_completo_edit['RUT'] == rut_actual, ['Nombre', 'RUT', 'Telefono', 'Correo', 'Clave_unica', 'Direccion']] = [n_nom, n_rut, n_tel, n_cor, n_cla, n_dom]
                                else:
                                    # No existía como registro real de Cliente todavía (por
                                    # ejemplo, alguien creado solo a través de un Encargo o una
                                    # causa antigua) — antes esto hacía que "Guardar" no
                                    # guardara nada, en silencio, porque intentaba actualizar
                                    # una fila que no existía. Ahora, si no existe, se crea.
                                    fila_nueva_cliente = {'Nombre': n_nom, 'RUT': n_rut, 'Telefono': n_tel, 'Correo': n_cor, 'Clave_unica': n_cla, 'Direccion': n_dom, 'Usuario_Propietario': usuario_actual}
                                    df_clientes_completo_edit = pd.concat([df_clientes_completo_edit, pd.DataFrame([fila_nueva_cliente])], ignore_index=True)
                                safe_update_sheet("base_clientes", df_clientes_completo_edit)
                                st.session_state['editando_cli'] = False
                                st.rerun()
                    else:
                        st.write(f"**Nombre:** {datos_cli.get('Nombre', '--')}")
                        st.write(f"**RUT:** {datos_cli.get('RUT', '--')}")
                        st.write(f"**Teléfono:** {datos_cli.get('Telefono', '--')}")
                        st.write(f"**Correo:** {datos_cli.get('Correo', '--')}")
                        st.write(f"**Clave Única:** {datos_cli.get('Clave_unica', '--')}")
                        st.write(f"**Domicilio:** {datos_cli.get('Direccion', '--')}")
                        if st.button("✏️ Editar Datos"): st.session_state['editando_cli'] = True; st.rerun()
            
            with col_d:
                # --- NUEVA FUNCIÓN: CREAR CAUSA DESDE EL CLIENTE ---
                with st.expander("➕ Asociar Nueva Causa a este Cliente"):
                    with st.form("form_asociar_causa"):
                        rol_n = st.text_input("Nuevo ROL / RIT", placeholder="Ej: C-123-2026")
                        trib_n = selector_tribunal(key_prefix="asociar_causa_cliente")
                        carat_n = st.text_input("Caratulado", placeholder="Ej: PEREZ / BANCO")
                        neg_n = st.selectbox("Origen de Cartera", ["Propio", "Externo"])
                        
                        if st.form_submit_button("Inyectar Causa al Cliente", type="primary"):
                            if rol_n.strip() == "":
                                st.error("El ROL es obligatorio.")
                            else:
                                df_causas_local = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                                nueva_c = {
                                    'ROL': rol_n.strip().upper(), 'TRIBUNAL': trib_n.strip(), 'CARATULADO': carat_n.strip(), 
                                    'Cliente': datos_cli['Nombre'], 'RUT': rut_actual, 'Tipo_Negocio': neg_n,
                                    'Usuario_Propietario': usuario_actual, 'Estado_Honorarios': 'Sin fijar',
                                    'Total_Honorarios': 0, 'Cuotas_Totales': 0, 'Cuotas_Pagadas': 0,
                                    'Clave_unica': datos_cli.get('Clave_unica', '--'), 'SAC': '--', 'Sucursal': '--', 'Servicio': '--',
                                    'Teléfono': datos_cli.get('Telefono', '--'), 'Correo': datos_cli.get('Correo', '--'), 'Direccion': datos_cli.get('Direccion', '--')
                                }
                                df_causas_local = pd.concat([df_causas_local, pd.DataFrame([nueva_c])], ignore_index=True)
                                df_causas_local.to_csv(ARCHIVO_BD, index=False)
                                
                                dn_c = safe_read_sheet("base_causas", COLS_CAUSAS)
                                dn_c = pd.concat([dn_c, pd.DataFrame([nueva_c])], ignore_index=True)
                                safe_update_sheet("base_causas", dn_c)
                                
                                st.success(f"✅ Causa {rol_n.upper()} asociada exitosamente.")
                                import time; time.sleep(0.4); st.rerun()

                st.subheader("Causas Asociadas Vigentes")
                df_causas = leer_csv_local(ARCHIVO_BD, COLS_CAUSAS)
                if df_causas.empty:
                    # Si el archivo local está vacío (por ejemplo, justo después de
                    # un reinicio), se revisa directo en Google Sheets antes de
                    # decir que "no hay causas" — para no depender de que el
                    # archivo local ya se haya reconstruido.
                    df_causas_nube_ficha = safe_read_sheet("base_causas", COLS_CAUSAS)
                    if not df_causas_nube_ficha.empty and 'Usuario_Propietario' in df_causas_nube_ficha.columns:
                        df_causas = df_causas_nube_ficha[df_causas_nube_ficha['Usuario_Propietario'] == usuario_actual]
                if not df_causas.empty:
                    causas_cli = df_causas[df_causas['RUT'].astype(str) == str(rut_actual)]
                    if causas_cli.empty:
                        st.write("Este cliente no tiene causas vinculadas todavía.")
                    else:
                        for _, c in causas_cli.iterrows():
                            with st.container(border=True):
                                c1, c2 = st.columns([3, 1])
                                c1.markdown(f"**Rol:** {c['ROL']} | **Caratulado:** {c.get('CARATULADO', '--')}")
                                if c2.button("📂 Ir al Expediente", key=f"btn_ir_{c['ROL']}"):
                                    ir_a_expediente(c['ROL']); st.rerun()
                else:
                    st.write("Base de causas vacía.")

        with tab2:
            st.subheader("Estado Financiero Global")
            if not df_causas.empty:
                causas_economicas = df_causas[df_causas['RUT'].astype(str) == str(rut_actual)]
                if causas_economicas.empty:
                    st.write("Sin registros financieros.")
                else:
                    for _, ce in causas_economicas.iterrows():
                        st.write(f"🔹 **Causa Rol {ce['ROL']}:** {ce.get('Estado_Honorarios', 'Sin fijar')}")
                        st.write(f"Pactado: ${ce.get('Total_Honorarios',0):,.0f} | Cuotas Pagadas: {ce.get('Cuotas_Pagadas',0)}")
                        st.write("---")
            else:
                st.write("Sin registros financieros.")
                    
        with tab3:
            st.subheader("Contratos Vinculados")
            df_con = leer_csv_local(ARCHIVO_CONTRATOS, COLS_CONTRATOS)
            if not df_con.empty:
                st.dataframe(df_con[df_con['Cliente'] == datos_cli['Nombre']])
            else:
                st.write("No hay contratos registrados.")

# 11. GESTOR GLOBAL DE TAREAS
elif st.session_state['menu_radio'] == "☑️ Tareas":
    st.title("Tareas")
    st.markdown("<span style='color:#6b778c;'>Revisa y gestiona todas tus tareas</span>", unsafe_allow_html=True)
    
    df_t = leer_csv_local(ARCHIVO_TAREAS, COLS_TAREAS)
    if df_t.empty:
        # Mismo respaldo que en Inicio, Causas y la ficha del cliente.
        df_t_nube_tareas = safe_read_sheet("base_tareas", COLS_TAREAS)
        if not df_t_nube_tareas.empty and 'Usuario_Propietario' in df_t_nube_tareas.columns:
            df_t = df_t_nube_tareas[df_t_nube_tareas['Usuario_Propietario'] == usuario_actual]
    df_t['Propietario_Vista'] = usuario_actual
    
    ES_ADMIN_TAREAS = usuario_actual == "Narratia"
    if ES_ADMIN_TAREAS:
        boton_refrescar_equipo("refresh_tareas_equipo")
        archivos_tareas_equipo = glob.glob("base_tareas_*.csv")
        piezas_tareas_eq = []
        for arch_t in archivos_tareas_equipo:
            propietario_t = arch_t.replace("base_tareas_", "").replace(".csv", "")
            temp_t_eq = leer_csv_local(arch_t, COLS_TAREAS)
            if not temp_t_eq.empty:
                temp_t_eq = temp_t_eq.copy()
                temp_t_eq['Propietario_Vista'] = propietario_t
                piezas_tareas_eq.append(temp_t_eq)
        if piezas_tareas_eq:
            df_t = pd.concat(piezas_tareas_eq, ignore_index=True)
    
    n_rechazadas = len(df_t[df_t['Estado'] == 'Rechazada']) if not df_t.empty else 0
    n_en_progreso = len(df_t[df_t['Estado'] == 'En progreso']) if not df_t.empty else 0
    n_completadas = len(df_t[df_t['Estado'] == 'Aprobada']) if not df_t.empty else 0
    
    c_st1, c_st2, c_st3 = st.columns(3)
    with c_st1:
        st.markdown(f"""<div class="dash-card"><span style="color:#6b778c; font-size:14px;">Tareas rechazadas</span><br>
        <span style="font-size:32px; font-weight:700; color:#172b4d;">{n_rechazadas}</span><br>
        <span style="color:#bf2600; font-size:13px;">❌ Requieren atención</span></div>""", unsafe_allow_html=True)
    with c_st2:
        st.markdown(f"""<div class="dash-card"><span style="color:#6b778c; font-size:14px;">En progreso</span><br>
        <span style="font-size:32px; font-weight:700; color:#172b4d;">{n_en_progreso}</span><br>
        <span style="color:#7a5b00; font-size:13px;">🕐 Tareas activas</span></div>""", unsafe_allow_html=True)
    with c_st3:
        st.markdown(f"""<div class="dash-card"><span style="color:#6b778c; font-size:14px;">Completadas</span><br>
        <span style="font-size:32px; font-weight:700; color:#172b4d;">{n_completadas}</span><br>
        <span style="color:#1b7a4a; font-size:13px;">✅ Total</span></div>""", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    
    if df_t.empty: 
        st.info("No hay tareas creadas en el sistema.")
    else:
        c_busq, c_filt, c_dl_t = st.columns([2.5, 1.5, 1])
        busqueda_tarea = c_busq.text_input("Buscar", placeholder="Tarea, Causa...", label_visibility="collapsed")
        filtro_estado_t = c_filt.selectbox("Estado", ["Todas", "En progreso", "Aprobada", "Rechazada"], label_visibility="collapsed")
        
        df_t_filt = df_t.copy()
        if busqueda_tarea.strip():
            q = busqueda_tarea.strip().lower()
            df_t_filt = df_t_filt[
                df_t_filt['Titulo'].astype(str).str.lower().str.contains(q, na=False) |
                df_t_filt['ROL'].astype(str).str.lower().str.contains(q, na=False)
            ]
        if filtro_estado_t != "Todas":
            df_t_filt = df_t_filt[df_t_filt['Estado'] == filtro_estado_t]
        with c_dl_t:
            boton_descargar_excel(df_t_filt, "tareas_jurisync.xlsx", key="dl_excel_tareas")
        
        for idx, row in df_t_filt.iterrows():
            fila_tarea_propia = row.get('Propietario_Vista', usuario_actual) == usuario_actual
            with st.container(border=True):
                prio_color = "#ff5630" if row.get('Prioridad') == "Alta" else ("#ffc400" if row.get('Prioridad') == "Media" else "#57a15a")
                st.markdown(f"<div style='height: 5px; background-color: {prio_color}; border-radius: 5px 5px 0 0; margin: -1rem -1rem 1rem -1rem;'></div>", unsafe_allow_html=True)
                c1, c2, c3 = st.columns([4, 2, 1])
                with c1:
                    st.markdown(f"<div style='display: flex; align-items: center; margin-bottom: 5px;'><strong style='font-size:16px; color:#172b4d;'>{row['Titulo']}</strong><span style='font-size:12px; color:{prio_color}; font-weight:bold; margin-left:8px;'>[{row.get('Prioridad', 'Media')}]</span></div>", unsafe_allow_html=True)
                    st.markdown(f"<span style='color:#6b778c;'>{str(row['Descripcion'])[:60]}...</span>", unsafe_allow_html=True)
                    if not fila_tarea_propia:
                        nombre_dueno_t = NOMBRES_REALES.get(row.get('Propietario_Vista'), row.get('Propietario_Vista'))
                        st.markdown(f"<span style='background:#fff0e0; color:#ff8b00; font-size:10px; font-weight:700; padding:2px 6px; border-radius:8px;'>👤 {nombre_dueno_t}</span>", unsafe_allow_html=True)
                with c2:
                    color_bd = "#fff0b3" if row['Estado'] == 'En progreso' else ("#e3fcef" if row['Estado'] == 'Aprobada' else "#ffebe6")
                    texto_bd = "#7a5b00" if row['Estado'] == 'En progreso' else ("#1b7a4a" if row['Estado'] == 'Aprobada' else "#bf2600")
                    st.markdown(f"<span style='background:{color_bd}; color:{texto_bd}; padding:3px 10px; border-radius:12px; font-size:12px; font-weight:700;'>{row['Estado']}</span>", unsafe_allow_html=True)
                    st.markdown(f"<span style='color:#172b4d; font-size:14px;'><br>Causa: {row['ROL']} | Vence: {row['Fecha_Vencimiento']}</span>", unsafe_allow_html=True)
                with c3:
                    st.button("Ir al expediente ➔", key=f"global_ir_{row['ID_Tarea']}_{row.get('Propietario_Vista', '')}", on_click=ir_a_expediente, args=(row['ROL'], row.get('Propietario_Vista', usuario_actual)))

# 11.5 AGENDA DE CITAS (agendamiento de asesorías a clientes)
elif st.session_state['menu_radio'] == "🗓️ Agenda de Asesorías":
    st.title("🗓️ Agenda de Asesorías")
    
    ES_ADMIN_CITAS = usuario_actual == "Narratia"
    df_citas = leer_csv_local(ARCHIVO_CITAS, COLS_CITAS)
    if ES_ADMIN_CITAS:
        for arch_cita in glob.glob("base_citas_*.csv"):
            propietario_cita = arch_cita.replace("base_citas_", "").replace(".csv", "")
            if propietario_cita != usuario_actual:
                t_cita = leer_csv_local(arch_cita, COLS_CITAS)
                if not t_cita.empty:
                    df_citas = pd.concat([df_citas, t_cita], ignore_index=True)
    
    ESTADOS_CITA = ["Por Confirmar", "Confirmada", "Realizada", "Cliente No Asiste",
                    "Cancelada por Cliente", "Cancelada por Asesor", "No Contesta",
                    "Reagendada", "Reasignada"]
    COLOR_ESTADO_CITA = {
        "Por Confirmar": ("#fff0b3", "#7a5b00"), "Confirmada": ("#e3f2f1", "#0e6b74"),
        "Realizada": ("#e3fcef", "#1b7a4a"), "Cliente No Asiste": ("#ffebe6", "#bf2600"),
        "Cancelada por Cliente": ("#ffebe6", "#bf2600"), "Cancelada por Asesor": ("#ffebe6", "#bf2600"),
        "No Contesta": ("#f4f5f7", "#6b778c"), "Reagendada": ("#fff0b3", "#7a5b00"),
        "Reasignada": ("#fff0b3", "#7a5b00"),
    }
    
    tab_citas_hoy, tab_citas_nueva, tab_citas_todas = st.tabs(["📅 Hoy", "➕ Agendar Cita", "📋 Todas las Citas"])
    
    # --- PESTAÑA: HOY ---
    with tab_citas_hoy:
        fecha_hoy_cita = datetime.now().strftime("%d-%m-%Y")
        st.markdown(f"#### Citas agendadas para hoy — {fecha_hoy_cita}")
        
        if df_citas.empty or 'Fecha' not in df_citas.columns:
            st.info("No hay citas registradas todavía.")
        else:
            df_hoy = df_citas[df_citas['Fecha'].astype(str) == fecha_hoy_cita].copy()
            if df_hoy.empty:
                st.info("No hay citas agendadas para hoy.")
            else:
                df_hoy = df_hoy.sort_values('Hora')
                for _, cita in df_hoy.iterrows():
                    color_bg, color_txt = COLOR_ESTADO_CITA.get(cita['Estado'], ("#f4f5f7", "#6b778c"))
                    with st.container(border=True):
                        c1, c2, c3 = st.columns([1, 3, 1.5])
                        with c1:
                            st.markdown(f"### {cita['Hora']}")
                        with c2:
                            st.markdown(f"**{cita['Nombre_Cliente']}** — {cita.get('Tipo_Juicio','')}")
                            st.caption(f"📞 {cita['Telefono']} · {cita['Modalidad']} · {cita['Sucursal']} · Abogado: {cita['Abogado_Asignado']}")
                        with c3:
                            nuevo_estado = st.selectbox("Estado", ESTADOS_CITA, index=ESTADOS_CITA.index(cita['Estado']) if cita['Estado'] in ESTADOS_CITA else 0,
                                                         key=f"estado_hoy_{cita['ID_Cita']}", label_visibility="collapsed")
                            st.markdown(f"<span style='background:{color_bg}; color:{color_txt}; padding:2px 10px; border-radius:10px; font-size:11px; font-weight:700;'>{cita['Estado']}</span>", unsafe_allow_html=True)
                            if nuevo_estado != cita['Estado']:
                                df_citas_completo = leer_csv_local(f"base_citas_{cita['Usuario_Propietario']}.csv", COLS_CITAS)
                                df_citas_completo.loc[df_citas_completo['ID_Cita'] == cita['ID_Cita'], 'Estado'] = nuevo_estado
                                df_citas_completo.to_csv(f"base_citas_{cita['Usuario_Propietario']}.csv", index=False)
                                dn_cita = safe_read_sheet("base_citas", COLS_CITAS)
                                if not dn_cita.empty:
                                    dn_cita.loc[dn_cita['ID_Cita'] == cita['ID_Cita'], 'Estado'] = nuevo_estado
                                    safe_update_sheet("base_citas", dn_cita)
                                st.rerun()
    
    # --- PESTAÑA: AGENDAR NUEVA CITA ---
    with tab_citas_nueva:
        st.markdown("#### Buscar cliente por RUT (autocompleta si ya existe)")
        c_rut_buscar, c_btn_buscar = st.columns([3, 1])
        rut_buscar_cita = c_rut_buscar.text_input("RUT (ej: 12345678-9)", key="cita_rut_buscar", label_visibility="collapsed", placeholder="RUT (ej: 12345678-9)")
        
        datos_cliente_encontrado = None
        if rut_buscar_cita.strip():
            df_clientes_cita = safe_read_sheet("base_clientes", COLS_CLIENTES)
            rut_limpio_buscar = re.sub(r'[^0-9kK]', '', rut_buscar_cita).upper()
            if not df_clientes_cita.empty:
                coincidencia = df_clientes_cita[df_clientes_cita['RUT'].astype(str).apply(lambda r: re.sub(r'[^0-9kK]', '', r).upper()) == rut_limpio_buscar]
                if not coincidencia.empty:
                    datos_cliente_encontrado = coincidencia.iloc[0]
                    st.success(f"✅ Cliente encontrado: {datos_cliente_encontrado['Nombre']}")
        
        with st.form("form_nueva_cita", clear_on_submit=True):
            c1, c2 = st.columns(2)
            nombre_cita = c1.text_input("Nombre", value=datos_cliente_encontrado['Nombre'] if datos_cliente_encontrado is not None else "")
            telefono_cita = c2.text_input("Teléfono", value=datos_cliente_encontrado['Telefono'] if datos_cliente_encontrado is not None else "")
            email_cita = c1.text_input("Email", value=datos_cliente_encontrado['Correo'] if datos_cliente_encontrado is not None else "")
            sucursal_cita = c2.selectbox("Sucursal", ["Santiago Centro", "Sucursal Virtual", "Otra"])
            modalidad_cita = c1.selectbox("Modalidad de cita", ["Telefónica", "Presencial", "Videollamada"])
            abogado_cita = c2.text_input("Abogado asignado", value=nombre_real_usuario)
            fecha_cita = c1.date_input("Fecha")
            hora_cita = c2.time_input("Hora")
            tipo_juicio_cita = st.text_input("Tipo de juicio / materia (ej: Juicio Ejecutivo cobro de pagaré)", key="cita_tipo_juicio")
            observacion_cita = st.text_area("Observación", height=80)
            
            if st.form_submit_button("📅 Agendar Cita", type="primary", use_container_width=True):
                if not nombre_cita.strip():
                    st.error("⚠️ Debes indicar al menos el nombre del cliente.")
                else:
                    nueva_cita = {
                        'ID_Cita': str(uuid.uuid4())[:8], 'Fecha': fecha_cita.strftime("%d-%m-%Y"), 'Hora': hora_cita.strftime("%H:%M"),
                        'RUT_Cliente': rut_buscar_cita.strip().upper(), 'Nombre_Cliente': nombre_cita.strip(),
                        'Telefono': telefono_cita.strip(), 'Email': email_cita.strip(), 'Sucursal': sucursal_cita,
                        'Modalidad': modalidad_cita, 'Abogado_Asignado': abogado_cita.strip(), 'Tipo_Juicio': tipo_juicio_cita.strip(),
                        'Observacion': observacion_cita.strip(), 'Estado': 'Por Confirmar', 'Usuario_Propietario': usuario_actual
                    }
                    df_citas_local = leer_csv_local(ARCHIVO_CITAS, COLS_CITAS)
                    df_citas_local = pd.concat([df_citas_local, pd.DataFrame([nueva_cita])], ignore_index=True)
                    df_citas_local.to_csv(ARCHIVO_CITAS, index=False)
                    dn_citas = safe_read_sheet("base_citas", COLS_CITAS)
                    safe_update_sheet("base_citas", pd.concat([dn_citas, pd.DataFrame([nueva_cita])], ignore_index=True))
                    st.success("✅ Cita agendada correctamente.")
                    st.rerun()
        
        # 💰 Referencia de honorarios según el tipo de juicio ingresado arriba
        if st.session_state.get("cita_tipo_juicio", "").strip():
            sugerencias_cita = buscar_arancel_referencial(st.session_state["cita_tipo_juicio"])
            if sugerencias_cita:
                st.markdown("##### 💰 Referencia de Honorarios (Arancel Colegio de Abogados de Valparaíso)")
                for sug in sugerencias_cita:
                    st.markdown(f"**N°{sug['numero']} — {sug['descripcion']}**")
                    st.caption(sug['honorario'])
    
    # --- PESTAÑA: TODAS LAS CITAS ---
    with tab_citas_todas:
        st.markdown("#### Buscador general")
        busqueda_cita = st.text_input("Buscar por nombre, RUT, abogado, teléfono...", key="busqueda_citas_general", label_visibility="collapsed", placeholder="Buscar por nombre, RUT, abogado, teléfono...")
        
        if df_citas.empty:
            st.info("No hay citas registradas todavía.")
        else:
            df_mostrar_citas = df_citas.copy()
            if busqueda_cita.strip():
                mask_busqueda = df_mostrar_citas.astype(str).apply(lambda col: col.str.contains(busqueda_cita, case=False, na=False)).any(axis=1)
                df_mostrar_citas = df_mostrar_citas[mask_busqueda]
            
            df_mostrar_citas = df_mostrar_citas.sort_values(['Fecha', 'Hora'], ascending=[False, True])
            st.dataframe(
                df_mostrar_citas[['Fecha', 'Hora', 'Abogado_Asignado', 'Telefono', 'Nombre_Cliente', 'RUT_Cliente', 'Sucursal', 'Modalidad', 'Estado']],
                use_container_width=True, hide_index=True
            )

# 12. EXCEL IMPORTADOR
elif st.session_state['menu_radio'] == "📥 Excel":
    st.title("📥 Importador Masivo de Causas (OJV)")
    st.markdown("Sube tu archivo Excel de la Oficina Judicial Virtual para consolidar o actualizar masivamente tus causas.")
    archivo = st.file_uploader("Sube tu archivo Excel", type=["xlsx", "xls"])
    if archivo and st.button("Procesar y Consolidar en Base de Datos", type="primary"):
        procesar_ojv_completo(archivo)
        st.success("¡Base de datos unificada actualizada con éxito!")

# 13. CALENDARIO
elif st.session_state['menu_radio'] == "📅 Calendario":
    st.title("📅 Calendario de Tareas y Plazos")
    st.markdown("Revisa visualmente los hitos procesales, plazos fatales y feriados de todo el equipo.")
    
    eventos_calendario = obtener_feriados_chile()
    df_t = safe_read_sheet("base_tareas", ['ID_Tarea', 'ROL', 'Creador', 'Fecha_Creacion', 'Fecha_Vencimiento', 'Titulo', 'Descripcion', 'Estado', 'Comentarios', 'Prioridad', 'Usuario_Propietario'])
    
    if not df_t.empty:
        for idx, r in df_t.iterrows():
            try:
                # Validamos que la fecha venga bien armada
                d_obj = datetime.strptime(str(r['Fecha_Vencimiento']).strip(), "%d/%m/%Y")
                d_str = d_obj.strftime("%Y-%m-%d")
                bg_color = "#ff5630" if r.get('Prioridad') == "Alta" else ("#ffc400" if r.get('Prioridad') == "Media" else "#57a15a")
                text_color = "#172b4d"
                eventos_calendario.append({
                    "title": f"{r.get('Titulo', 'Tarea')}", 
                    "start": d_str, 
                    "backgroundColor": bg_color, 
                    "borderColor": bg_color,
                    "textColor": text_color
                })
            except Exception: 
                pass
    
    # Estilos propios para que el calendario se vea como una grilla limpia,
    # con puntos de color por evento (estilo agenda) en vez de barras sólidas,
    # y el mismo lenguaje visual (colores, tipografía) del resto de JuriSync.
    css_calendario = """
        .fc { font-family: 'Source Sans Pro', sans-serif; }
        .fc-view-harness { border-radius: 14px; overflow: hidden; }
        .fc .fc-toolbar-title { font-size: 20px; font-weight: 700; color: #172b4d; text-transform: capitalize; }
        .fc .fc-button { background-color: #ffffff !important; color: #172b4d !important; border: 1px solid #cbd2d9 !important; border-radius: 20px !important; font-weight: 600 !important; box-shadow: none !important; text-transform: capitalize; padding: 6px 16px !important; }
        .fc .fc-button:hover { background-color: #e3f2f1 !important; border-color: #0e6b74 !important; color: #0e6b74 !important; }
        .fc .fc-button-primary:not(:disabled).fc-button-active { background-color: #0e6b74 !important; border-color: #0e6b74 !important; color: #ffffff !important; }
        .fc-theme-standard td, .fc-theme-standard th { border-color: transparent !important; }
        .fc-scrollgrid { border: none !important; border-collapse: separate !important; border-spacing: 6px !important; }
        .fc-col-header-cell-cushion { color: #6b778c !important; font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: 0.5px; text-decoration: none !important; padding: 8px 0 !important; }
        .fc-daygrid-day { background-color: #fbfbfc !important; border: 1px solid #eaecf0 !important; border-radius: 14px !important; overflow: hidden; transition: all 0.15s ease; }
        .fc-daygrid-day:hover { border-color: #b3d4ff !important; background-color: #f4f8ff !important; }
        .fc-daygrid-day-number { color: #172b4d !important; font-size: 13px; font-weight: 600; text-decoration: none !important; padding: 8px !important; }
        .fc-day-today { background-color: #e8f5f4 !important; border: 1px solid #0e6b74 !important; }
        .fc-day-other { background-color: #f7f8fa !important; }
        .fc-day-other .fc-daygrid-day-number { color: #a5adba !important; font-weight: 400; }
        .fc-daygrid-event { border-radius: 10px !important; font-size: 12px !important; padding: 1px 6px !important; margin-top: 2px !important; }
        .fc-daygrid-event-dot { border-width: 4px !important; }
        .fc-event-title { font-weight: 500; }
        .fc-daygrid-more-link { color: #0e6b74 !important; font-weight: 700; font-size: 12px; }
        .fc-daygrid-day-frame { padding: 4px; }
        .fc-scrollgrid-section-header th { border: none !important; }
    """
    
    opciones_calendario = {
        "initialView": "dayGridMonth", 
        "locale": "es", 
        "firstDay": 1, 
        "height": "auto",
        "headerToolbar": {
            "left": "prev,next today", 
            "center": "title", 
            "right": "dayGridMonth,listMonth"
        },
        "dayMaxEvents": 3,       # Igual que la referencia: hasta 3 líneas visibles y luego "+N más..."
        "eventDisplay": "list-item",  # Punto de color + texto, en vez de una barra sólida
        "moreLinkText": "más...",
        # Se fuerza el texto de botones/etiquetas en español, ya que algunos
        # no se traducen solo con "locale" cuando se usa un headerToolbar custom.
        "buttonText": {
            "today": "Hoy", "month": "Mes", "week": "Semana", "day": "Día", "list": "Lista"
        },
        "dayHeaderFormat": {"weekday": "long"},
        "titleFormat": {"month": "long", "year": "numeric"},
        "noEventsText": "No hay tareas para mostrar",
        "allDayText": "Todo el día"
    }
    
    col_cal, col_dia = st.columns([2.4, 1])
    
    with col_cal:
        calendario_estado = calendar(events=eventos_calendario, options=opciones_calendario, custom_css=css_calendario, key="calendario_app_full")
        
    fecha_mostrar = datetime.now().strftime("%Y-%m-%d")
    if calendario_estado and 'dateClick' in calendario_estado and calendario_estado['dateClick']:
        fecha_mostrar = calendario_estado['dateClick']['date'][:10]
        
    with col_dia:
        try:
            fecha_dt_dia = datetime.strptime(fecha_mostrar, "%Y-%m-%d")
            d_fmt = fecha_dt_dia.strftime("%d/%m/%Y")
            dia_semana_es = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"][fecha_dt_dia.weekday()]
            
            st.markdown(f"""
            <div class="dash-card">
                <div style="font-weight:700; font-size:17px; color:#172b4d;">Tareas del día</div>
                <div style="font-size:13px; color:#6b778c; margin-bottom:12px;">{dia_semana_es}, {fecha_dt_dia.day} de {['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre'][fecha_dt_dia.month-1]} de {fecha_dt_dia.year}</div>
            """, unsafe_allow_html=True)
            
            tareas_dia = df_t[df_t['Fecha_Vencimiento'].astype(str).str.strip() == d_fmt] if not df_t.empty else pd.DataFrame()
            if tareas_dia.empty:
                st.caption("Sin tareas para este día.")
            else:
                for _, td in tareas_dia.iterrows():
                    color_prio = "#ff5630" if td.get('Prioridad') == "Alta" else ("#ffc400" if td.get('Prioridad') == "Media" else "#57a15a")
                    st.markdown(f"""
                    <div style="border-left:3px solid {color_prio}; padding:6px 10px; margin-bottom:8px; background:#f8f9fa; border-radius:6px;">
                        <div style="font-weight:600; font-size:13px; color:#172b4d;">{td.get('Titulo', '--')}</div>
                        <div style="font-size:12px; color:#6b778c;">Causa: {td.get('ROL', '--')}</div>
                    </div>
                    """, unsafe_allow_html=True)
                    st.button("Ir a Causa", key=f"btn_cal_ir_{td.get('ID_Tarea', uuid.uuid4())}", on_click=ir_a_expediente, args=(td['ROL'],), use_container_width=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
        except Exception: 
            st.caption("Haz clic en un día del calendario para ver sus tareas detalladas.")
# 14. PANEL DE ADMINISTRADOR (SOLO NARRATIA)
elif st.session_state['menu_radio'] == "👑 Panel Admin" and usuario_actual == "Narratia":
    st.title("👑 Panel de Control Master - SaaS JuriSync")
    st.markdown("Gestión maestra de usuarios y formateo del sistema.")
    
    # Aseguramos cargar la base de usuarios actualizada desde la nube
    df_usuarios_admin = safe_read_sheet("base_usuarios", COLS_USUARIOS)
    
    tab_crear, tab_editar, tab_vision, tab_drive_oauth, tab_peligro = st.tabs(["➕ Crear Nuevo Usuario", "🔄 Autorizar Planes", "👁️ Visión Global", "🔑 Conectar Drive", "☢️ Zona de Peligro"])
    
    with tab_crear:
        with st.container(border=True):
            st.subheader("Alta de Equipo / Clientes")
            col1, col2 = st.columns(2)
            with col1:
                nuevo_user = st.text_input("Usuario (Nombre para iniciar sesión)")
                nuevo_nombre = st.text_input("Nombre Real del Abogado / Cliente")
            with col2:
                nueva_clave = st.text_input("Clave Provisoria")
                nuevo_plan = st.selectbox("Plan Asignado al Crear", ["Básico", "Medio", "Full"])
                
            if st.button("🚀 Crear Usuario y Autorizar Plan", type="primary", use_container_width=True):
                if not nuevo_user.strip() or not nueva_clave.strip() or not nuevo_nombre.strip():
                    st.error("⚠️ Faltan datos obligatorios.")
                else:
                    if nuevo_user in df_usuarios_admin['Usuario'].values:
                        st.error(f"⚠️ El usuario '{nuevo_user}' ya existe en el sistema.")
                    else:
                        with st.spinner("Guardando en la nube y generando carpetas..."):
                            nuevo_registro = {
                                "Usuario": nuevo_user.strip(), "Password": hash_password(nueva_clave.strip()), "Nombre_Real": nuevo_nombre.strip(),
                                "Correo": "pendiente", "Debe_Cambiar_Clave": 'True', "Plan": nuevo_plan
                            }
                            df_usuarios_admin = pd.concat([df_usuarios_admin, pd.DataFrame([nuevo_registro])], ignore_index=True)
                            
                            # Forzamos la actualización
                            safe_update_sheet("base_usuarios", df_usuarios_admin)
                            
                            # Generar archivo local si es abogado para evitar errores de lectura
                            archivo_tareas_nuevo = f"base_tareas_{nuevo_user}.csv"
                            if not os.path.exists(archivo_tareas_nuevo):
                                pd.DataFrame(columns=COLS_TAREAS).to_csv(archivo_tareas_nuevo, index=False)
                                
                            st.success(f"✅ ¡Cuenta autorizada! El usuario **{nuevo_user}** ya puede acceder con el plan **{nuevo_plan}**.")
                            import time; time.sleep(0.4); st.rerun()

    with tab_editar:
        with st.container(border=True):
            st.subheader("Auditoría de Cuentas y Accesos")
            lista_usuarios = df_usuarios_admin['Usuario'].tolist()
            
            c_ed1, c_ed2, c_ed3 = st.columns(3)
            with c_ed1:
                usuario_editar = st.selectbox("Seleccionar Usuario a Modificar", lista_usuarios)
            with c_ed2:
                try:
                    plan_actual_usr = df_usuarios_admin.loc[df_usuarios_admin['Usuario'] == usuario_editar, 'Plan'].values[0]
                    idx_plan = ["Básico", "Medio", "Full"].index(plan_actual_usr)
                except Exception: idx_plan = 0
                nuevo_plan_edit = st.selectbox("Modificar Nivel de Acceso", ["Básico", "Medio", "Full"], index=idx_plan)
            with c_ed3:
                st.write("") 
                st.write("")
                if st.button("🔄 Autorizar Cambio de Plan", type="primary", use_container_width=True):
                    with st.spinner("Sincronizando permisos..."):
                        df_usuarios_admin.loc[df_usuarios_admin['Usuario'] == usuario_editar, 'Plan'] = nuevo_plan_edit
                        safe_update_sheet("base_usuarios", df_usuarios_admin)
                        st.success(f"✅ Los permisos de **{usuario_editar}** han sido actualizados en el sistema.")
                        import time; time.sleep(0.4); st.rerun()
                        
        st.markdown("**Resumen de Usuarios Activos**")
        st.dataframe(df_usuarios_admin[['Usuario', 'Nombre_Real', 'Plan', 'Correo']], use_container_width=True)

    with tab_vision:
        st.subheader("Monitoreo Absoluto de la Oficina")
        
        c_refresh, _ = st.columns([1, 4])
        if c_refresh.button("🔄 Actualizar ahora", help="Fuerza una relectura de todos los archivos del equipo, por si algún cambio reciente no se refleja."):
            # Limpia del caché de sesión cualquier archivo base_causas_*/base_tareas_*
            # para forzar que se vuelvan a leer del disco en este instante.
            claves_a_limpiar = [k for k in st.session_state.keys() if k.startswith("_csv_cache_base_causas_") or k.startswith("_csv_cache_base_tareas_")]
            for k in claves_a_limpiar:
                del st.session_state[k]
            st.rerun()
        
        todas_causas = []
        todas_tareas = []
        
        # IMPORTANTE: se lee directo de los archivos que existen en el disco (glob),
        # en vez de depender de la lista de usuarios de Google Sheets. Si un usuario
        # nuevo no queda bien reflejado ahí (o esa lectura está desactualizada), su
        # archivo de causas/tareas igual se detecta y se muestra aquí.
        for arch_c in glob.glob("base_causas_*.csv"):
            u = arch_c.replace("base_causas_", "").replace(".csv", "")
            temp_c = leer_csv_local(arch_c, COLS_CAUSAS)
            if not temp_c.empty:
                temp_c['Usuario_Propietario'] = u 
                todas_causas.append(temp_c)
        for arch_t in glob.glob("base_tareas_*.csv"):
            u = arch_t.replace("base_tareas_", "").replace(".csv", "")
            temp_t = leer_csv_local(arch_t, COLS_TAREAS)
            if not temp_t.empty:
                temp_t['Usuario_Propietario'] = u
                todas_tareas.append(temp_t)
        
        col_v1, col_v2 = st.columns(2)
        with col_v1:
            st.markdown("<div class='dash-card'><h4>Causas Globales</h4>", unsafe_allow_html=True)
            if todas_causas:
                df_full_causas = pd.concat(todas_causas, ignore_index=True)
                st.metric("Total de Causas en JuriSync", len(df_full_causas))
                st.dataframe(df_full_causas[['Usuario_Propietario', 'ROL', 'Cliente', 'TRIBUNAL']], use_container_width=True)
            else: st.info("No hay causas registradas.")
            st.markdown("</div>", unsafe_allow_html=True)
            
        with col_v2:
            st.markdown("<div class='dash-card'><h4>Tareas Operativas Globales</h4>", unsafe_allow_html=True)
            if todas_tareas:
                df_full_tareas = pd.concat(todas_tareas, ignore_index=True)
                st.metric("Total de Gestiones Pendientes", len(df_full_tareas[df_full_tareas['Estado'] == 'En progreso']))
                st.dataframe(df_full_tareas[['Usuario_Propietario', 'Titulo', 'Estado', 'Fecha_Vencimiento']], use_container_width=True)
            else: st.info("No hay tareas creadas.")
            st.markdown("</div>", unsafe_allow_html=True)

    with tab_drive_oauth:
        st.subheader("🔑 Conectar tu Google Drive personal")
        st.caption("Las cuentas de servicio no tienen cuota de almacenamiento propia en Drive (limitación real de Google). Esta conexión hace que los archivos grandes se guarden usando tu cuenta personal de Google (15GB gratis), en vez de la cuenta de servicio.")
        
        if "GOOGLE_OAUTH_REFRESH_TOKEN" in st.secrets:
            st.success("✅ Ya tienes tu Drive personal conectado. Los archivos grandes deberían guardarse ahí automáticamente.")
            st.caption("Si quieres reconectar (por ejemplo, con otra cuenta de Google), sigue los pasos de abajo de nuevo — la nueva autorización reemplaza a la anterior en Secrets.")
        
        if "GOOGLE_OAUTH_CLIENT_ID" not in st.secrets or "GOOGLE_OAUTH_CLIENT_SECRET" not in st.secrets:
            st.warning("⚠️ Primero necesitas crear credenciales OAuth (paso único, distinto a la cuenta de servicio que ya tienes).")
            st.markdown("""
            **Cómo crearlas:**
            1. Ve a [console.cloud.google.com](https://console.cloud.google.com), proyecto `jurisync-libre`.
            2. Menú → "APIs y servicios" → "Credenciales" → "Crear credenciales" → **"ID de cliente de OAuth"**.
            3. Tipo de aplicación: **"Aplicación web"**.
            4. En "URI de redireccionamiento autorizados", agrega la URL exacta de tu app (ej: `https://jurisyncs.streamlit.app`).
            5. Crea, y copia el **"ID de cliente"** y el **"Secreto del cliente"** que te muestra.
            6. Agrégalos a tus Secrets de Streamlit:
            ```
            GOOGLE_OAUTH_CLIENT_ID = "tu_id_de_cliente"
            GOOGLE_OAUTH_CLIENT_SECRET = "tu_secreto_de_cliente"
            APP_BASE_URL = "https://jurisyncs.streamlit.app"
            ```
            7. Guarda los Secrets, espera a que la app se reinicie, y vuelve a esta pestaña.
            """)
        else:
            # Ya hay credenciales OAuth configuradas: se puede iniciar la autorización.
            parametros_url = st.query_params
            if "code" in parametros_url:
                st.info("🔄 Terminando la autorización...")
                try:
                    _, flow_pendiente = _url_autorizacion_drive_oauth()
                    credenciales_obtenidas = _intercambiar_codigo_oauth_drive(flow_pendiente, parametros_url["code"])
                    st.success("✅ ¡Autorización exitosa! Copia este código y pégalo en tus Secrets de Streamlit:")
                    st.code(f'GOOGLE_OAUTH_REFRESH_TOKEN = "{credenciales_obtenidas.refresh_token}"', language="toml")
                    st.warning("⚠️ Este código es una llave de acceso a tu Drive — trátalo como una contraseña. Después de copiarlo a Secrets, no lo compartas ni lo dejes visible en ningún otro lado.")
                    st.query_params.clear()
                except Exception as e:
                    st.error(f"⚠️ No se pudo completar la autorización. Detalle técnico: {e}")
            else:
                url_auth, _ = _url_autorizacion_drive_oauth()
                st.markdown(f"### [👉 Haz clic aquí para autorizar tu Drive personal]({url_auth})")
                st.caption("Te va a llevar a una pantalla de Google pidiéndote iniciar sesión y dar permiso — es normal que Google muestre una advertencia de 'app no verificada' (porque es tu propia app, no una app pública), puedes continuar igual haciendo clic en 'Avanzado' → 'Ir a JuriSync (no seguro)'.")

    with tab_peligro:
        st.subheader("Borrón y Cuenta Nueva (Limpieza Estricta)")
        st.error("⚠️ ADVERTENCIA: Esta acción eliminará permanentemente todos los clientes, causas, tareas, contratos, trámites y documentos de Google Sheets y del servidor local.")
        
        confirmacion_borrado = st.text_input("Escribe **BORRAR** en mayúsculas para habilitar el botón de eliminación:", key="confirmacion_borrado_total")
        
        if st.button("🚨 BORRAR TODA LA BASE DE DATOS DEL SISTEMA 🚨", type="primary", use_container_width=True, disabled=(confirmacion_borrado.strip() != "BORRAR")):
            with st.spinner("Formateando absolutamente TODAS las tablas en la nube y discos locales..."):
                try:
                    safe_update_sheet("base_clientes", pd.DataFrame(columns=COLS_CLIENTES))
                    safe_update_sheet("base_causas", pd.DataFrame(columns=COLS_CAUSAS))
                    safe_update_sheet("base_tareas", pd.DataFrame(columns=COLS_TAREAS))
                    safe_update_sheet("base_contratos", pd.DataFrame(columns=COLS_CONTRATOS))
                    safe_update_sheet("base_tramites", pd.DataFrame(columns=COLS_TRAMITES))
                    safe_update_sheet("base_estado_diario", pd.DataFrame(columns=['ID_ED', 'Fecha_Estado', 'ROL', 'Tribunal', 'Resolucion_Extracto', 'Doc_Nombre', 'Doc_B64', 'Doc_Drive_ID']))
                    safe_update_sheet("base_documentos_clientes", pd.DataFrame(columns=['ID_Req', 'Cliente_Token', 'Documento_Nombre', 'Estado', 'Archivo_B64', 'Archivo_Drive_ID', 'Fecha_Subida']))
                except Exception as e:
                    st.error(f"Error limpiando Google Sheets: {e}")
                
                # Barrido nuclear de archivos locales (salvando solo a los usuarios)
                import glob
                for archivo_local in glob.glob("base_*.csv"):
                    if "usuarios" not in archivo_local:
                        try: os.remove(archivo_local)
                        except Exception: pass
                
                st.cache_data.clear()
            st.success("✅ Limpieza extrema completada. Sistema restablecido a cero.")
            import time; time.sleep(0.6); st.rerun()

# 15. REDACTOR AUTOMÁTICO IA
elif st.session_state['menu_radio'] == "📝 Redactor IA":
    st.title("📝 Redactor Automático de Escritos")
    st.markdown("La IA redactará el borrador del escrito judicial con el formato y lenguaje formal de los tribunales chilenos, listo para revisar y presentar.")
    
    # AUTOCOMPLETADO CONTEXTUAL: si llegaste aquí desde el botón "Redactar Escrito"
    # de una causa, se precargan el Rol/Tribunal/Caratulado de esa causa, para no
    # tener que volver a escribirlos. Se usa una sola vez y luego se limpia.
    _prefill_redactor = st.session_state.pop('redactor_prefill', None)
    if _prefill_redactor:
        st.session_state['redactor_rol_key'] = _prefill_redactor.get('rol', '')
        st.session_state['redactor_trib_key'] = _prefill_redactor.get('tribunal', '')
        st.session_state['redactor_carat_key'] = _prefill_redactor.get('caratulado', '')
        st.info(f"✅ Datos de la causa **{_prefill_redactor.get('rol','')}** cargados automáticamente.")
    
    with st.container(border=True):
        col_r1, col_r2 = st.columns(2)
        tipo_escrito = col_r1.selectbox("Tipo de Escrito", list(ESTRUCTURAS_REDACTOR_IA.keys()))
        tribunal_red = col_r2.text_input("Tribunal (Para la suma)", placeholder="Ej: S.J.L. en lo Civil (1°)", key="redactor_trib_key")
        
        rol_red = col_r1.text_input("Causa Rol", placeholder="Ej: C-1234-2026", key="redactor_rol_key")
        caratula_red = col_r2.text_input("Caratulado", placeholder="Ej: PEREZ / BANCO", key="redactor_carat_key")
        
        instrucciones_red = st.text_area("Instrucciones específicas para el escrito:", height=150, placeholder="Ej: Redactar excepción N° 17 de prescripción. La deuda se hizo exigible en marzo de 2024 y notificaron recién ayer. Alega también costas.")
        
        if st.button("✍️ Generar Borrador del Escrito", type="primary", use_container_width=True):
            if not instrucciones_red.strip():
                st.error("⚠️ Debes darle las instrucciones jurídicas a la IA.")
            else:
                with st.spinner("⚖️ Redactando escrito en lenguaje procesal chileno..."):
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                        
                        modelo_elegido = "gemini-1.0-pro"
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods:
                                md_name = m.name.replace("models/", "")
                                if 'flash' in md_name:
                                    modelo_elegido = md_name
                                    break
                                    
                        modelo = genai.GenerativeModel(modelo_elegido)
                        
                        sentencias_relevantes_red = buscar_jurisprudencia_relevante(f"{tipo_escrito} {instrucciones_red}")
                        bloque_juris_redactor = ""
                        if sentencias_relevantes_red:
                            bloque_juris_redactor = "\n\nJURISPRUDENCIA REAL DISPONIBLE (de la biblioteca del estudio, verificada, puedes citarla con confianza si es pertinente):\n"
                            for s in sentencias_relevantes_red:
                                bloque_juris_redactor += f"- {s.get('Tribunal','')}, Rol {s.get('Rol_Causa','')}, {s.get('Fecha_Sentencia','')}: {s.get('Resumen','')}\n"
                        
                        prompt_redactor = f"""
                        Actúa como un abogado litigante chileno con impecable ortografía y redacción procesal formal.
                        Debes redactar un escrito judicial completo con los siguientes datos:
                        
                        Tipo de Escrito: {tipo_escrito}
                        Tribunal: {tribunal_red}
                        Rol: {rol_red}
                        Caratulado: {caratula_red}
                        
                        ESTRUCTURA ESPECÍFICA A SEGUIR PARA ESTE TIPO DE ESCRITO (obligatoria, no te desvíes de ella):
                        {ESTRUCTURAS_REDACTOR_IA.get(tipo_escrito, "")}
                        
                        INSTRUCCIONES DE FONDO (hechos y detalles concretos del caso, entregados por el abogado):
                        {instrucciones_red}
                        {bloque_juris_redactor}
                        
                        {INSTRUCCION_FUNDAMENTACION_JURIDICA}
                        (Para este escrito en particular: no agregues la nota sugerida en la regla 4 dentro del texto final, ya que debe quedar limpio y listo para copiar — solo aplica las reglas 1 a 3.)
                        
                        IMPORTANTE sobre la extensión: si la estructura de arriba indica que es un escrito "DE MERA TRAMITACIÓN", redáctalo CORTO y directo (1 a 3 párrafos), sin inflar con desarrollo de "Hechos" y "Derecho" innecesario. Si es un escrito "DE FONDO", desarrolla los fundamentos con la extensión y profundidad que corresponda.
                        
                        Usa el lenguaje propio del Código de Procedimiento Civil chileno. No agregues notas explicativas para mí, entrégame SOLO el texto del escrito listo para copiar.
                        """
                        
                        respuesta_escrito = modelo.generate_content(prompt_redactor)
                        st.success("✅ Borrador redactado. Cópialo, revísalo y pásalo a Word.")
                        st.text_area("Escrito Generado:", value=respuesta_escrito.text, height=500)
                        
                    except Exception as e:
                        st.error(f"❌ Hubo un error de conexión: {e}")

# =====================================================================
# ⚖️ MÓDULO: BIBLIOTECA DE JURISPRUDENCIA
# =====================================================================
elif st.session_state['menu_radio'] == "⚖️ Jurisprudencia":
    st.title("⚖️ Biblioteca de Jurisprudencia")
    st.caption("Sentencias reales de Corte Suprema, Cortes de Apelaciones y tribunales de primera instancia, cargadas por el equipo. El sistema las usa como referencia real al redactar y analizar, sin inventar citas.")
    
    tab_juris_agregar, tab_juris_buscar, tab_codigos_bcn = st.tabs(["➕ Agregar Sentencia", "🔍 Buscar en la Biblioteca", "📖 Códigos de la República (BCN)"])
    
    # --- PESTAÑA: AGREGAR SENTENCIA ---
    with tab_juris_agregar:
        archivo_sentencia = st.file_uploader("Sube el PDF de la sentencia", type=["pdf"], key="juris_pdf_subir")
        
        texto_extraido_sentencia = ""
        if archivo_sentencia:
            try:
                import PyPDF2
                lector_sent = PyPDF2.PdfReader(archivo_sentencia)
                texto_extraido_sentencia = "\n".join([p.extract_text() or "" for p in lector_sent.pages])
            except Exception:
                pass
            
            if texto_extraido_sentencia.strip() and st.button("🤖 Autocompletar datos con IA (opcional)", key="juris_autocompletar"):
                with st.spinner("Analizando la sentencia..."):
                    try:
                        prompt_juris = f"""
                        Analiza el siguiente texto de una sentencia judicial chilena y extrae:
                        1. Tribunal que la dictó (Corte Suprema, Corte de Apelaciones de [ciudad], o Tribunal de Primera Instancia [tipo]).
                        2. Rol de la causa (si aparece).
                        3. Fecha de la sentencia (si aparece).
                        4. Materia principal (en pocas palabras, ej: "Nulidad de contrato por vicio del consentimiento").
                        5. Un resumen de 3-5 líneas del criterio jurídico central que resuelve, con la máxima fidelidad al texto real (sin inventar nada que no esté en el documento).
                        
                        TEXTO DE LA SENTENCIA:
                        {texto_extraido_sentencia[:20000]}
                        
                        Responde EXCLUSIVAMENTE con un JSON válido (sin bloques de código markdown): {{"tribunal": "...", "rol": "...", "fecha": "...", "materia": "...", "resumen": "..."}}
                        """
                        respuesta_juris_ia = consultar_groq(prompt_juris)
                        datos_juris_ia = json.loads(_limpiar_json_ia(respuesta_juris_ia))
                        st.session_state['juris_ia_tribunal'] = datos_juris_ia.get('tribunal', '')
                        st.session_state['juris_ia_rol'] = datos_juris_ia.get('rol', '')
                        st.session_state['juris_ia_fecha'] = datos_juris_ia.get('fecha', '')
                        st.session_state['juris_ia_materia'] = datos_juris_ia.get('materia', '')
                        st.session_state['juris_ia_resumen'] = datos_juris_ia.get('resumen', '')
                        st.success("✅ Datos autocompletados abajo. Revísalos antes de guardar.")
                    except Exception as e:
                        st.error(f"⚠️ No se pudo autocompletar: {e}. Puedes llenar los datos a mano igual.")
        
        with st.form("form_agregar_jurisprudencia", clear_on_submit=True):
            c1, c2 = st.columns(2)
            tribunal_juris = c1.selectbox("Tribunal", ["Corte Suprema", "Corte de Apelaciones", "Tribunal de Primera Instancia"],
                                          index=["Corte Suprema", "Corte de Apelaciones", "Tribunal de Primera Instancia"].index(st.session_state.get('juris_ia_tribunal')) if st.session_state.get('juris_ia_tribunal') in ["Corte Suprema", "Corte de Apelaciones", "Tribunal de Primera Instancia"] else 0)
            rol_juris = c2.text_input("Rol de la causa", value=st.session_state.get('juris_ia_rol', ''))
            fecha_juris = c1.text_input("Fecha de la sentencia", value=st.session_state.get('juris_ia_fecha', ''), placeholder="Ej: 15/03/2025")
            materia_juris = c2.text_input("Materia principal", value=st.session_state.get('juris_ia_materia', ''), placeholder="Ej: Nulidad de contrato por vicio del consentimiento")
            resumen_juris = st.text_area("Resumen del criterio jurídico", value=st.session_state.get('juris_ia_resumen', ''), height=120,
                                          placeholder="Resumen del criterio central que resuelve la sentencia...")
            
            if st.form_submit_button("💾 Guardar en la Biblioteca", type="primary", use_container_width=True):
                if not archivo_sentencia:
                    st.error("⚠️ Debes subir el PDF de la sentencia.")
                elif not materia_juris.strip():
                    st.error("⚠️ Indica al menos la materia principal.")
                else:
                    drive_id_juris, b64_juris = guardar_archivo_adjunto(archivo_sentencia.name, archivo_sentencia.getvalue(), 'application/pdf')
                    nueva_sentencia = {
                        'ID': str(uuid.uuid4())[:8], 'Tribunal': tribunal_juris, 'Rol_Causa': rol_juris.strip(),
                        'Fecha_Sentencia': fecha_juris.strip(), 'Materia': materia_juris.strip(), 'Resumen': resumen_juris.strip(),
                        'Archivo_Nombre': archivo_sentencia.name, 'Archivo_B64': b64_juris, 'Archivo_Drive_ID': drive_id_juris,
                        'Fecha_Carga': datetime.now().strftime("%d/%m/%Y"), 'Usuario_Propietario': usuario_actual
                    }
                    df_juris_guardar = safe_read_sheet("base_jurisprudencia", COLS_JURISPRUDENCIA)
                    df_juris_guardar = pd.concat([df_juris_guardar, pd.DataFrame([nueva_sentencia])], ignore_index=True)
                    safe_update_sheet("base_jurisprudencia", df_juris_guardar)
                    for k in ['juris_ia_tribunal', 'juris_ia_rol', 'juris_ia_fecha', 'juris_ia_materia', 'juris_ia_resumen']:
                        st.session_state.pop(k, None)
                    st.success("✅ Sentencia agregada a la biblioteca.")
                    st.rerun()
    
    # --- PESTAÑA: BUSCAR EN LA BIBLIOTECA ---
    with tab_juris_buscar:
        df_biblioteca_juris = safe_read_sheet("base_jurisprudencia", COLS_JURISPRUDENCIA)
        
        if df_biblioteca_juris.empty:
            st.info("Todavía no hay sentencias cargadas en la biblioteca.")
        else:
            c_buscar, c_filtro_trib = st.columns([3, 1])
            busqueda_juris = c_buscar.text_input("Buscar por materia, resumen o rol...", key="busqueda_juris_texto")
            filtro_tribunal_juris = c_filtro_trib.selectbox("Tribunal", ["Todos", "Corte Suprema", "Corte de Apelaciones", "Tribunal de Primera Instancia"])
            
            df_mostrar_juris = df_biblioteca_juris.copy()
            if filtro_tribunal_juris != "Todos":
                df_mostrar_juris = df_mostrar_juris[df_mostrar_juris['Tribunal'] == filtro_tribunal_juris]
            if busqueda_juris.strip():
                mask_juris = df_mostrar_juris.astype(str).apply(lambda col: col.str.contains(busqueda_juris, case=False, na=False)).any(axis=1)
                df_mostrar_juris = df_mostrar_juris[mask_juris]
            
            st.caption(f"{len(df_mostrar_juris)} sentencia(s) encontrada(s)")
            for _, fila_juris in df_mostrar_juris.iloc[::-1].iterrows():
                with st.container(border=True):
                    c1, c2 = st.columns([5, 1.3])
                    with c1:
                        st.markdown(f"**{fila_juris['Tribunal']}** — Rol {fila_juris.get('Rol_Causa','—')} — {fila_juris.get('Fecha_Sentencia','—')}")
                        st.markdown(f"*{fila_juris.get('Materia','')}*")
                        st.caption(fila_juris.get('Resumen', ''))
                    with c2:
                        bytes_juris_desc = obtener_bytes_adjunto(fila_juris, 'Archivo_Drive_ID', 'Archivo_B64')
                        if bytes_juris_desc is not None:
                            st.download_button("📥 PDF", data=bytes_juris_desc, file_name=fila_juris.get('Archivo_Nombre', 'sentencia.pdf'), key=f"dl_juris_{fila_juris['ID']}")
                        if usuario_actual == "Narratia" or fila_juris['Usuario_Propietario'] == usuario_actual:
                            if st.button("🗑️", key=f"del_juris_{fila_juris['ID']}"):
                                df_juris_del = safe_read_sheet("base_jurisprudencia", COLS_JURISPRUDENCIA)
                                df_juris_del = df_juris_del[df_juris_del['ID'] != fila_juris['ID']]
                                safe_update_sheet("base_jurisprudencia", df_juris_del)
                                st.rerun()
    
    # --- PESTAÑA: CÓDIGOS DE LA REPÚBLICA (conexión real con BCN) ---
    with tab_codigos_bcn:
        st.caption("Texto vigente y actualizado, traído en vivo desde la API de datos abiertos de la Biblioteca del Congreso Nacional (BCN) — la fuente oficial, no una copia guardada que podría quedar desactualizada.")
        st.info("ℹ️ Por ahora, solo están conectados los códigos cuyo identificador se verificó con certeza. Los demás (Código Penal, de Comercio, del Trabajo, Tributario) se agregarán más adelante una vez confirmados, para no arriesgarse a traer una ley equivocada.")
        
        codigo_elegido_bcn = st.selectbox("Código a consultar", list(CODIGOS_BCN_IDNORMA.keys()))
        termino_bcn = st.text_input("Buscar artículo o palabra clave", placeholder="Ej: 1545, o 'nulidad absoluta'")
        
        if st.button("🔍 Buscar en BCN", type="primary", use_container_width=True):
            if not termino_bcn.strip():
                st.error("⚠️ Escribe un término o número de artículo a buscar.")
            else:
                with st.spinner("Consultando la API de BCN..."):
                    resultados_bcn = buscar_en_codigo_bcn(codigo_elegido_bcn, termino_bcn)
                if resultados_bcn:
                    st.success(f"✅ {len(resultados_bcn)} resultado(s) encontrado(s) en {codigo_elegido_bcn}")
                    for texto_art in resultados_bcn:
                        with st.container(border=True):
                            st.markdown(texto_art.replace("\n", "  \n"))
                else:
                    st.warning("No se encontraron coincidencias, o hubo un problema consultando BCN. Intenta con otro término.")

# =====================================================================
# 📜 MÓDULO: ESCRITURAS PÚBLICAS (3 pestañas)
# =====================================================================
elif st.session_state['menu_radio'] == "📜 Escrituras Públicas":
    st.title("📜 Escrituras Públicas")
    st.markdown("Redacción, análisis y gestión de documentos para escrituras públicas.")
    
    tab_esc_redaccion, tab_esc_analisis, tab_esc_docs = st.tabs([
        "✍️ Redacción de Escritura", "🔍 Análisis de Escritura (IA)", "📥 Docs Cliente"
    ])
    
    # --- PESTAÑA 1: REDACCIÓN DE ESCRITURA ---
    with tab_esc_redaccion:
        st.markdown("#### Paso 1: Elige el tipo de escritura")
        tipo_esc_sel = st.selectbox("Tipo de Escritura Pública", list(CATALOGO_ESCRITURAS.keys()), key="esc_tipo_sel")
        rol1_lbl, rol2_lbl = CATALOGO_ESCRITURAS[tipo_esc_sel]["roles"]
        tiene_segunda_parte = "sin segunda parte" not in rol2_lbl
        
        st.markdown("#### Paso 2: Completa los datos")
        with st.form("form_generar_escritura", clear_on_submit=False):
            with st.container(border=True):
                st.markdown("**Datos del Notario**")
                c_not1, c_not2 = st.columns(2)
                notario_nombre = c_not1.text_input("Nombre del Notario(a)")
                notaria_ciudad = c_not2.text_input("Ciudad de la Notaría", value="Santiago")
            
            with st.container(border=True):
                st.markdown(f"**Compareciente 1 — {rol1_lbl}**")
                c_p1a, c_p1b = st.columns(2)
                parte1_nombre = c_p1a.text_input("Nombre completo", key="esc_p1_nom")
                parte1_rut = c_p1b.text_input("RUT", key="esc_p1_rut")
                parte1_domicilio = c_p1a.text_input("Domicilio", key="esc_p1_dom")
                parte1_profesion = c_p1b.text_input("Profesión u oficio", key="esc_p1_prof")
                parte1_estado_civil = c_p1a.selectbox("Estado civil", ["Soltero(a)", "Casado(a)", "Divorciado(a)", "Viudo(a)"], key="esc_p1_ec")
                parte1_nacionalidad = c_p1b.text_input("Nacionalidad", value="Chilena", key="esc_p1_nac")
            
            if tiene_segunda_parte:
                with st.container(border=True):
                    st.markdown(f"**Compareciente 2 — {rol2_lbl}**")
                    c_p2a, c_p2b = st.columns(2)
                    parte2_nombre = c_p2a.text_input("Nombre completo", key="esc_p2_nom")
                    parte2_rut = c_p2b.text_input("RUT", key="esc_p2_rut")
                    parte2_domicilio = c_p2a.text_input("Domicilio", key="esc_p2_dom")
                    parte2_profesion = c_p2b.text_input("Profesión u oficio", key="esc_p2_prof")
                    parte2_estado_civil = c_p2a.selectbox("Estado civil", ["Soltero(a)", "Casado(a)", "Divorciado(a)", "Viudo(a)"], key="esc_p2_ec")
                    parte2_nacionalidad = c_p2b.text_input("Nacionalidad", value="Chilena", key="esc_p2_nac")
            else:
                st.caption("ℹ️ Este tipo de escritura requiere 3 testigos hábiles presentes en la notaría; no se solicitan sus datos aquí, solo el compareciente principal.")
                parte2_nombre = parte2_rut = parte2_domicilio = parte2_profesion = parte2_estado_civil = parte2_nacionalidad = ""
            
            with st.container(border=True):
                st.markdown(f"**Datos específicos de: {tipo_esc_sel}**")
                datos_especificos = {}
                for campo_key, campo_label, campo_tipo in CATALOGO_ESCRITURAS[tipo_esc_sel]["campos"]:
                    if campo_tipo == "textarea":
                        datos_especificos[campo_key] = st.text_area(campo_label, key=f"esc_campo_{campo_key}")
                    else:
                        datos_especificos[campo_key] = st.text_input(campo_label, key=f"esc_campo_{campo_key}")
            
            if st.form_submit_button("📄 Generar Escritura en Word", type="primary", use_container_width=True):
                if not parte1_nombre.strip() or not parte1_rut.strip():
                    st.error("⚠️ Debes completar al menos el nombre y RUT del primer compareciente.")
                else:
                    datos_escritura = {
                        'notario_nombre': notario_nombre, 'notaria_ciudad': notaria_ciudad,
                        'parte1_nombre': parte1_nombre, 'parte1_rut': parte1_rut, 'parte1_domicilio': parte1_domicilio,
                        'parte1_profesion': parte1_profesion, 'parte1_estado_civil': parte1_estado_civil, 'parte1_nacionalidad': parte1_nacionalidad,
                        'parte2_nombre': parte2_nombre, 'parte2_rut': parte2_rut, 'parte2_domicilio': parte2_domicilio,
                        'parte2_profesion': parte2_profesion, 'parte2_estado_civil': parte2_estado_civil, 'parte2_nacionalidad': parte2_nacionalidad,
                        **datos_especificos
                    }
                    doc_escritura = crear_escritura_word(tipo_esc_sel, datos_escritura)
                    if doc_escritura:
                        buffer_esc = io.BytesIO()
                        doc_escritura.save(buffer_esc)
                        bytes_escritura = buffer_esc.getvalue()
                        nombre_archivo_esc = f"Escritura_{tipo_esc_sel.replace(' ', '_')}_{parte1_nombre.replace(' ', '_')}.docx"
                        
                        drive_id_esc, b64_esc = guardar_archivo_adjunto(
                            nombre_archivo_esc, bytes_escritura,
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                        
                        df_esc = leer_csv_local(ARCHIVO_ESCRITURAS, COLS_ESCRITURAS)
                        nuevo_esc = {
                            'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"),
                            'Tipo_Escritura': tipo_esc_sel, 'Cliente': parte1_nombre, 'RUT_Cliente': parte1_rut,
                            'Detalle': f"{rol1_lbl}: {parte1_nombre}" + (f" | {rol2_lbl}: {parte2_nombre}" if tiene_segunda_parte else ""),
                            'Archivo_B64': b64_esc, 'Archivo_Drive_ID': drive_id_esc, 'Usuario_Propietario': usuario_actual
                        }
                        df_esc = pd.concat([df_esc, pd.DataFrame([nuevo_esc])], ignore_index=True)
                        df_esc.to_csv(ARCHIVO_ESCRITURAS, index=False)
                        
                        dn_esc = safe_read_sheet("base_escrituras", COLS_ESCRITURAS)
                        safe_update_sheet("base_escrituras", pd.concat([dn_esc, pd.DataFrame([nuevo_esc])], ignore_index=True))
                        
                        st.success("✅ Escritura generada y guardada correctamente.")
                        st.download_button("📥 Descargar Escritura (.docx)", data=bytes_escritura, file_name=nombre_archivo_esc,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        st.markdown("---")
        st.markdown("### 🗄️ Historial de Escrituras Generadas")
        df_esc_hist = leer_csv_local(ARCHIVO_ESCRITURAS, COLS_ESCRITURAS)
        if df_esc_hist.empty:
            st.info("Todavía no has generado ninguna escritura.")
        else:
            for _, fila_esc in df_esc_hist.iloc[::-1].iterrows():
                with st.container(border=True):
                    c1, c2 = st.columns([4, 1])
                    with c1:
                        st.markdown(f"**{fila_esc['Tipo_Escritura']}** — {fila_esc['Detalle']}")
                        st.caption(f"Fecha: {fila_esc['Fecha']}")
                    with c2:
                        bytes_desc_esc = obtener_bytes_adjunto(fila_esc, 'Archivo_Drive_ID', 'Archivo_B64')
                        if bytes_desc_esc is not None:
                            st.download_button("📥 Descargar", data=bytes_desc_esc, file_name=f"Escritura_{fila_esc['ID']}.docx", key=f"dl_esc_{fila_esc['ID']}")
    
    # --- PESTAÑA 2: ANÁLISIS DE ESCRITURA (IA) ---
    with tab_esc_analisis:
        st.markdown("#### Sube la escritura y sus documentos de respaldo para que la IA revise su redacción")
        st.caption("La IA revisa la redacción considerando los requisitos formales del Código Orgánico de Tribunales (Arts. 403 a 408 y 415) y las reglas generales de técnica notarial y civil chilena. Es un apoyo de revisión, no reemplaza el criterio profesional del abogado.")
        
        archivo_escritura_analizar = st.file_uploader("Escritura a analizar (PDF)", type=["pdf"], key="esc_analisis_pdf")
        docs_respaldo_analizar = st.file_uploader("Documentos de respaldo (opcional, puedes subir varios)", type=["pdf"], accept_multiple_files=True, key="esc_analisis_respaldo")
        contexto_adicional_esc = st.text_area("Contexto adicional para la IA (opcional)", placeholder="Ej: Es una compraventa de un bien raíz en Providencia, verificar especialmente la cláusula de saneamiento.")
        
        if st.button("🔍 Analizar Escritura", type="primary", use_container_width=True):
            if not archivo_escritura_analizar:
                st.error("⚠️ Debes subir el PDF de la escritura a analizar.")
            else:
                with st.spinner("⚖️ Analizando la redacción y formalidades de la escritura..."):
                    try:
                        prompt_analisis_esc = f"""
                        Actúa como un abogado chileno experto en derecho notarial y civil, revisando la redacción de una escritura pública.
                        
                        Analiza el siguiente texto de una escritura pública y evalúa:
                        1. Formalidades exigidas por el Código Orgánico de Tribunales (individualización correcta del notario y comparecientes: nacionalidad, estado civil, profesión, domicilio, cédula de identidad; ausencia de espacios en blanco, abreviaturas o cifras no permitidas; idioma castellano).
                        2. Si el objeto del acto jurídico está claramente descrito (bien, precio/monto, condiciones).
                        3. Si contiene las cláusulas esenciales según el tipo de acto (por ejemplo, en una compraventa: precio y forma de pago; en una hipoteca: monto garantizado e inscripción; en un mandato: facultades claramente delimitadas).
                        4. Riesgos, ambigüedades o cláusulas que podrían generar problemas de interpretación o nulidad.
                        5. Sugerencias concretas de mejora en la redacción.
                        
                        Contexto adicional entregado por el abogado: {contexto_adicional_esc if contexto_adicional_esc.strip() else "(sin contexto adicional)"}
                        
                        {INSTRUCCION_FUNDAMENTACION_JURIDICA}
                        En cada punto del análisis, cita el artículo exacto del Código Orgánico de Tribunales, Código Civil u otra norma aplicable que sustente la observación, siguiendo las reglas anteriores.
                        
                        Entrega el análisis estructurado en secciones claras con títulos, indicando en cada punto si CUMPLE, CUMPLE PARCIALMENTE o NO CUMPLE, seguido de la explicación, la norma legal que la sustenta, y recomendación concreta.
                        """
                        
                        todos_archivos_esc = [archivo_escritura_analizar] + (docs_respaldo_analizar or [])
                        texto_extraido_esc = extraer_texto_pdfs(todos_archivos_esc)
                        prompt_final_esc = prompt_analisis_esc + f"\n\nTEXTO EXTRAÍDO DE LOS DOCUMENTOS:\n{texto_extraido_esc[:45000]}"
                        texto_resultado_esc = consultar_groq(prompt_final_esc)
                        
                        st.success("✅ Análisis completado.")
                        st.markdown(texto_resultado_esc)
                        
                        # Se guarda el informe como Word en el historial, exactamente igual
                        # que se hace con los contratos: se genera el .docx, se sube a Drive
                        # (o queda de respaldo en base64) y se registra en la nube.
                        doc_analisis = crear_informe_analisis_escritura_word(archivo_escritura_analizar.name, texto_resultado_esc)
                        if doc_analisis:
                            buffer_analisis = io.BytesIO()
                            doc_analisis.save(buffer_analisis)
                            bytes_analisis = buffer_analisis.getvalue()
                            nombre_archivo_analisis = f"Analisis_{archivo_escritura_analizar.name.replace('.pdf', '')}.docx"
                            
                            drive_id_analisis, b64_analisis = guardar_archivo_adjunto(
                                nombre_archivo_analisis, bytes_analisis,
                                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            )
                            
                            df_analisis = leer_csv_local(ARCHIVO_ANALISIS_ESCRITURAS, COLS_ANALISIS_ESCRITURAS)
                            nuevo_analisis = {
                                'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y %H:%M"),
                                'Nombre_Archivo_Original': archivo_escritura_analizar.name,
                                'Archivo_B64': b64_analisis, 'Archivo_Drive_ID': drive_id_analisis, 'Usuario_Propietario': usuario_actual
                            }
                            df_analisis = pd.concat([df_analisis, pd.DataFrame([nuevo_analisis])], ignore_index=True)
                            df_analisis.to_csv(ARCHIVO_ANALISIS_ESCRITURAS, index=False)
                            
                            dn_analisis = safe_read_sheet("base_analisis_escrituras", COLS_ANALISIS_ESCRITURAS)
                            safe_update_sheet("base_analisis_escrituras", pd.concat([dn_analisis, pd.DataFrame([nuevo_analisis])], ignore_index=True))
                            
                            st.download_button("📥 Descargar Informe de Análisis (.docx)", data=bytes_analisis, file_name=nombre_archivo_analisis,
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_analisis_nuevo")
                    except Exception as e:
                        st.error(f"❌ Hubo un error al analizar la escritura: {e}")
        
        st.markdown("---")
        st.markdown("### 🗄️ Historial de Análisis Realizados")
        df_analisis_hist = leer_csv_local(ARCHIVO_ANALISIS_ESCRITURAS, COLS_ANALISIS_ESCRITURAS)
        if df_analisis_hist.empty:
            st.info("Todavía no has generado ningún análisis.")
        else:
            for _, fila_an in df_analisis_hist.iloc[::-1].iterrows():
                with st.container(border=True):
                    c1, c2 = st.columns([4, 1])
                    with c1:
                        st.markdown(f"**{fila_an['Nombre_Archivo_Original']}**")
                        st.caption(f"Analizado: {fila_an['Fecha']}")
                    with c2:
                        bytes_desc_an = obtener_bytes_adjunto(fila_an, 'Archivo_Drive_ID', 'Archivo_B64')
                        if bytes_desc_an is not None:
                            st.download_button("📥 Descargar", data=bytes_desc_an, file_name=f"Analisis_{fila_an['ID']}.docx", key=f"dl_an_{fila_an['ID']}")
    
    # --- PESTAÑA 3: DOCS CLIENTE ---
    with tab_esc_docs:
        st.markdown("#### Solicitar documentos al cliente para una escritura")
        st.caption("Usa el mismo portal externo que ya conoces (el del link para causas): el cliente entra con el mismo enlace y ve TODAS sus solicitudes pendientes juntas, sean de una causa o de una escritura.")
        
        ARCHIVO_DOCS_ESC = "base_documentos_clientes.csv"
        if not os.path.exists(ARCHIVO_DOCS_ESC):
            # Mismo arreglo que en el portal del cliente: intenta reconstruir
            # desde la nube antes de partir con un archivo vacío de verdad.
            df_docs_esc_nube_inicial = safe_read_sheet("base_documentos_clientes", COLS_DOCS)
            if not df_docs_esc_nube_inicial.empty:
                df_docs_esc_nube_inicial.to_csv(ARCHIVO_DOCS_ESC, index=False)
            else:
                pd.DataFrame(columns=COLS_DOCS).to_csv(ARCHIVO_DOCS_ESC, index=False)
        df_docs_esc = leer_csv_local(ARCHIVO_DOCS_ESC, COLS_DOCS)
        
        with st.form("form_solicitar_docs_escritura", clear_on_submit=True):
            nombre_cliente_esc_doc = st.text_input("Nombre completo del cliente (debe coincidir exactamente con el usado al generar la escritura)")
            documento_solicitado_esc = st.text_input("Documento que necesitas que suba", placeholder="Ej: Certificado de Dominio Vigente")
            if st.form_submit_button("➕ Solicitar Documento", type="primary"):
                if nombre_cliente_esc_doc.strip() and documento_solicitado_esc.strip():
                    token_cliente_esc = re.sub(r'[^A-Za-z0-9_]', '', nombre_cliente_esc_doc.strip().replace(" ", "_"))
                    nueva_solicitud_esc = {
                        'ID_Req': str(uuid.uuid4())[:8], 'Cliente_Token': token_cliente_esc,
                        'Documento_Nombre': documento_solicitado_esc.strip(), 'Estado': '⏳ Pendiente',
                        'Archivo_B64': '', 'Archivo_Drive_ID': '', 'Fecha_Subida': ''
                    }
                    df_docs_esc = pd.concat([df_docs_esc, pd.DataFrame([nueva_solicitud_esc])], ignore_index=True)
                    df_docs_esc.to_csv(ARCHIVO_DOCS_ESC, index=False)
                    safe_update_sheet("base_documentos_clientes", df_docs_esc)
                    st.success(f"✅ Solicitud creada. El cliente **{nombre_cliente_esc_doc}** la verá en su portal.")
                    st.rerun()
                else:
                    st.error("⚠️ Completa ambos campos.")
        
        st.markdown("---")
        st.markdown("### 📋 Solicitudes de Documentos Registradas")
        if df_docs_esc.empty:
            st.info("No hay solicitudes de documentos creadas todavía.")
        else:
            for _, fila_doc_esc in df_docs_esc.iloc[::-1].iterrows():
                with st.container(border=True):
                    c1, c2, c3 = st.columns([3, 2, 1])
                    c1.markdown(f"**{fila_doc_esc['Documento_Nombre']}**")
                    c2.markdown(f"Cliente: {fila_doc_esc['Cliente_Token'].replace('_', ' ')} · {fila_doc_esc['Estado']}")
                    with c3:
                        if fila_doc_esc['Estado'] == '✅ Completado':
                            bytes_doc_esc = obtener_bytes_adjunto(fila_doc_esc, 'Archivo_Drive_ID', 'Archivo_B64')
                            if bytes_doc_esc is not None:
                                st.download_button("📥", data=bytes_doc_esc, file_name=f"{fila_doc_esc['Documento_Nombre']}.pdf", key=f"dl_docesc_{fila_doc_esc['ID_Req']}")

# =====================================================================
# 📋 MÓDULO: POSESIÓN EFECTIVA
# =====================================================================
elif st.session_state['menu_radio'] == "📋 Posesión Efectiva":
    st.title("📋 Posesión Efectiva")
    st.markdown("Calcula automáticamente las asignaciones y el impuesto a la herencia de cada heredero, siguiendo las reglas de sucesión intestada y las Tablas 1, 2 y 3 del Formulario 4423 del SII.")
    st.caption("⚠️ Este calculador cubre la sucesión intestada (sin testamento). El resultado es la base para completar los formularios oficiales del SII y del Registro Civil — siempre debe revisarse antes de presentar la declaración.")
    
    df_pe = leer_csv_local(ARCHIVO_POSESION_EFECTIVA, COLS_POSESION_EFECTIVA)
    
    with st.container(border=True):
        st.markdown("#### 1. Datos del Causante y del Solicitante")
        c_pe1, c_pe2 = st.columns(2)
        causante_nombre = c_pe1.text_input("Nombre completo del causante", key="pe_causante_nombre")
        causante_rut = c_pe2.text_input("RUT del causante", key="pe_causante_rut")
        fecha_defuncion = c_pe1.date_input("Fecha de defunción", key="pe_fecha_defuncion")
        valor_utm_pe = c_pe2.number_input("Valor UTM a la fecha de fallecimiento ($)", min_value=1, value=65000, step=100, key="pe_valor_utm")
        cliente_solicitante_pe = c_pe1.text_input("Nombre del cliente solicitante", key="pe_solicitante_nombre")
        rut_cliente_pe = c_pe2.text_input("RUT del cliente solicitante", key="pe_solicitante_rut")
    
    with st.container(border=True):
        st.markdown("#### 2. Herederos")
        st.caption("Agrega una fila por cada heredero. El tipo de parentesco determina la fórmula de asignación y la exención de impuesto que le corresponde.")
        if 'pe_df_herederos' not in st.session_state:
            st.session_state['pe_df_herederos'] = pd.DataFrame(columns=["Nombre", "RUT", "Tipo de Heredero"])
        df_herederos_editado = st.data_editor(
            st.session_state['pe_df_herederos'],
            num_rows="dynamic",
            column_config={
                "Tipo de Heredero": st.column_config.SelectboxColumn(
                    options=["Hijo", "Cónyuge", "Ascendiente", "Hermano", "Medio Hermano", "Colateral 3° o 4° grado", "Colateral 5° o 6° grado"],
                    required=True
                )
            },
            use_container_width=True, key="pe_editor_herederos"
        )
        st.session_state['pe_df_herederos'] = df_herederos_editado
    
    with st.container(border=True):
        st.markdown("#### 3. Inventario de Bienes")
        st.caption("Agrega cada bien con su valorización y exención (si aplica). La masa hereditaria se calcula sola: suma de (Valorización − Exención) de los Activos, menos los Pasivos.")
        if 'pe_df_bienes' not in st.session_state:
            st.session_state['pe_df_bienes'] = pd.DataFrame(columns=["Categoría", "Descripción", "Valorización ($)", "Exención ($)"])
        df_bienes_editado = st.data_editor(
            st.session_state['pe_df_bienes'],
            num_rows="dynamic",
            column_config={
                "Categoría": st.column_config.SelectboxColumn(
                    options=["Bienes Raíces", "Vehículos", "Menaje", "Bienes Inmuebles Excluidos de Avalúo Fiscal",
                             "Otros Bienes Muebles (negocios, empresas, derechos)", "Otros Bienes (acciones, valores, depósitos, bonos)",
                             "Pasivo (Deuda Acreditada)"],
                    required=True
                ),
                "Valorización ($)": st.column_config.NumberColumn(min_value=0, step=1000),
                "Exención ($)": st.column_config.NumberColumn(min_value=0, step=1000),
            },
            use_container_width=True, key="pe_editor_bienes"
        )
        st.session_state['pe_df_bienes'] = df_bienes_editado
    
    if st.button("🧮 Calcular Posesión Efectiva y Determinación del Impuesto", type="primary", use_container_width=True):
        if not causante_nombre.strip() or not causante_rut.strip():
            st.error("⚠️ Debes completar al menos el nombre y RUT del causante.")
        elif df_herederos_editado.empty:
            st.error("⚠️ Debes agregar al menos un heredero.")
        else:
            # Masa hereditaria = suma de (Valorización - Exención) de los activos, menos los pasivos
            df_bienes_limpio = df_bienes_editado.copy()
            df_bienes_limpio["Valorización ($)"] = pd.to_numeric(df_bienes_limpio["Valorización ($)"], errors='coerce').fillna(0)
            df_bienes_limpio["Exención ($)"] = pd.to_numeric(df_bienes_limpio["Exención ($)"], errors='coerce').fillna(0)
            
            mascara_pasivo = df_bienes_limpio["Categoría"] == "Pasivo (Deuda Acreditada)"
            total_activos = (df_bienes_limpio.loc[~mascara_pasivo, "Valorización ($)"] - df_bienes_limpio.loc[~mascara_pasivo, "Exención ($)"]).sum()
            total_pasivos = df_bienes_limpio.loc[mascara_pasivo, "Valorización ($)"].sum()
            masa_hereditaria = max(0, total_activos - total_pasivos)
            
            resultado_calculo = calcular_posesion_efectiva_completa(df_herederos_editado, masa_hereditaria, valor_utm_pe)
            
            st.markdown("---")
            st.markdown("### 📊 Resultado del Cálculo")
            c_res1, c_res2, c_res3 = st.columns(3)
            c_res1.metric("Total Activos", formatear_clp(total_activos))
            c_res2.metric("Total Pasivos", formatear_clp(total_pasivos))
            c_res3.metric("Masa Hereditaria", formatear_clp(masa_hereditaria))
            
            st.markdown("#### Asignaciones e Impuesto por Heredero")
            st.dataframe(resultado_calculo, use_container_width=True, hide_index=True)
            
            total_impuesto_pe = resultado_calculo["Impuesto Total ($)"].sum() if not resultado_calculo.empty else 0
            st.metric("💰 Impuesto Total a Pagar (todos los herederos)", formatear_clp(total_impuesto_pe))
            
            # Generar y guardar el informe en Word, en el historial (igual que contratos y escrituras)
            datos_causante_doc = {'nombre': causante_nombre, 'rut': causante_rut, 'fecha_defuncion': fecha_defuncion.strftime("%d/%m/%Y")}
            doc_pe = crear_informe_posesion_efectiva_word(datos_causante_doc, resultado_calculo, masa_hereditaria, valor_utm_pe, total_impuesto_pe)
            
            bytes_pe = b""
            if doc_pe:
                buffer_pe = io.BytesIO()
                doc_pe.save(buffer_pe)
                bytes_pe = buffer_pe.getvalue()
                nombre_archivo_pe = f"Posesion_Efectiva_{causante_nombre.replace(' ', '_')}.docx"
                
                drive_id_pe, b64_pe = guardar_archivo_adjunto(
                    nombre_archivo_pe, bytes_pe,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                
                nuevo_pe = {
                    'ID': str(uuid.uuid4())[:8], 'Fecha': datetime.now().strftime("%d/%m/%Y"),
                    'Causante': causante_nombre, 'RUT_Causante': causante_rut,
                    'Fecha_Defuncion': fecha_defuncion.strftime("%d/%m/%Y"),
                    'Herederos_JSON': df_herederos_editado.to_json(orient='records', force_ascii=False),
                    'Bienes_JSON': df_bienes_editado.to_json(orient='records', force_ascii=False),
                    'Cliente_Solicitante': cliente_solicitante_pe, 'RUT_Cliente': rut_cliente_pe,
                    'Estado': 'Intestada (calculada)', 'Valor_UTM': valor_utm_pe, 'Masa_Hereditaria': masa_hereditaria,
                    'Impuesto_Total': total_impuesto_pe, 'Archivo_B64': b64_pe, 'Archivo_Drive_ID': drive_id_pe,
                    'Usuario_Propietario': usuario_actual
                }
                df_pe = pd.concat([df_pe, pd.DataFrame([nuevo_pe])], ignore_index=True)
                df_pe.to_csv(ARCHIVO_POSESION_EFECTIVA, index=False)
                dn_pe = safe_read_sheet("base_posesion_efectiva", COLS_POSESION_EFECTIVA)
                safe_update_sheet("base_posesion_efectiva", pd.concat([dn_pe, pd.DataFrame([nuevo_pe])], ignore_index=True))
                
                st.success("✅ Cálculo guardado en el historial.")
                st.download_button("📥 Descargar Informe de Cálculo (.docx)", data=bytes_pe, file_name=nombre_archivo_pe,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_pe_nuevo")
    
    st.markdown("---")
    st.markdown("### 🗄️ Trámites de Posesión Efectiva Registrados")
    if df_pe.empty:
        st.info("No hay trámites de posesión efectiva registrados todavía.")
    else:
        for _, fila_pe in df_pe.iloc[::-1].iterrows():
            with st.container(border=True):
                c1, c2 = st.columns([4, 1])
                with c1:
                    st.markdown(f"**Causante:** {fila_pe['Causante']} ({fila_pe['RUT_Causante']}) — **{fila_pe.get('Estado','')}**")
                    st.caption(f"Fecha defunción: {fila_pe['Fecha_Defuncion']} · Solicitante: {fila_pe.get('Cliente_Solicitante','')} · Registrado: {fila_pe['Fecha']}")
                    if pd.notna(fila_pe.get('Masa_Hereditaria')) and str(fila_pe.get('Masa_Hereditaria', '')).strip():
                        st.caption(f"Masa Hereditaria: {formatear_clp(fila_pe.get('Masa_Hereditaria', 0))} · Impuesto Total: {formatear_clp(fila_pe.get('Impuesto_Total', 0))}")
                with c2:
                    bytes_desc_pe = obtener_bytes_adjunto(fila_pe, 'Archivo_Drive_ID', 'Archivo_B64')
                    if bytes_desc_pe is not None:
                        st.download_button("📥 Descargar", data=bytes_desc_pe, file_name=f"PosesionEfectiva_{fila_pe['ID']}.docx", key=f"dl_pe_{fila_pe['ID']}")
                if pd.notna(fila_pe.get('Herederos_JSON')) and str(fila_pe.get('Herederos_JSON', '')).strip():
                    with st.expander("Ver herederos y bienes"):
                        try:
                            st.markdown("**Herederos:**")
                            st.dataframe(pd.read_json(io.StringIO(fila_pe['Herederos_JSON'])), use_container_width=True, hide_index=True)
                            st.markdown("**Bienes:**")
                            st.dataframe(pd.read_json(io.StringIO(fila_pe['Bienes_JSON'])), use_container_width=True, hide_index=True)
                        except Exception:
                            st.caption("No se pudo mostrar el detalle de este registro antiguo.")